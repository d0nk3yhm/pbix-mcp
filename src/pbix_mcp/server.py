"""
Power BI PBIX MCP Editor Server
================================
Full read/write MCP server for .pbix and .pbit files.

Capabilities:
  READ  — Report layout, visuals, pages, filters, DataMashup (M queries),
          DataModel schema/measures/relationships, settings, metadata
  WRITE — Report layout/visuals/pages/filters, DataMashup M code, settings,
          metadata. DataModel metadata via XPress9 round-trip.

Architecture:
  - PBIX files are ZIP archives
  - We extract components, allow granular inspection/editing, and repack
  - DataModel reading uses native ABF/VertiPaq decoder (XPress9 decompression)
  - DataModel writing works via ABF round-trip (decompress → modify → recompress)
"""

import atexit
import copy
import difflib
import io
import json
import os
import re
import shutil
import sqlite3
import struct
import tempfile
import time
import traceback
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

try:
    from mcp.server.fastmcp import FastMCP
except ModuleNotFoundError as _exc:  # pragma: no cover - import-time guard
    # mcp 2.0 dropped mcp.server.fastmcp (the server API is now
    # mcp.server.mcpserver). pyproject pins mcp<2, so this only fires when the
    # bound is overridden -- say a bare `pip install mcp -U` in an existing
    # environment. Say which package is wrong instead of a bare
    # "No module named 'mcp.server.fastmcp'".
    raise ImportError(
        "pbix-mcp requires the mcp 1.x server API (mcp.server.fastmcp), which "
        "mcp 2.0 removed. Install a compatible version with "
        "`pip install 'mcp>=1.0.0,<2'`."
    ) from _exc

from pbix_mcp.errors import (
    ABFRebuildError,
    DataModelCompressionError,
    DAXMeasureNotFoundError,
    FileAlreadyOpenError,
    FileNotOpenError,
    InvalidPBIXError,
    LayoutParseError,
    PBIXMCPError,
    SessionError,
    UnsafeWriteError,
    UnsupportedFormatError,
    UnsupportedModelEditError,
)
from pbix_mcp.logging_config import logger
from pbix_mcp.models import responses as _responses
from pbix_mcp.models.requests import DimensionRef, FilterContext
from pbix_mcp.models.responses import DAXEvalResponse, DAXResult, ToolResponse

# ---------------------------------------------------------------------------
# MCP Server
# ---------------------------------------------------------------------------
mcp = FastMCP(
    "PowerBI-PBIX-Editor",
    instructions="Full read/write editor for Power BI .pbix/.pbit files",
)

# ---------------------------------------------------------------------------
# State: track open files
# ---------------------------------------------------------------------------
_open_files: dict[str, dict] = {}

# Work directories created by THIS process. pbix_close removes an entry; the
# atexit hook below removes whatever is left, so a process that exits without
# closing (a test run, a script, a crashed sweep) does not strand its
# extractions. Thousands of pbix_mcp_* directories were found accumulated in
# %TEMP% from exactly that class of caller.
_work_dirs: set[str] = set()


def _cleanup_own_work_dirs() -> None:
    for d in list(_work_dirs):
        shutil.rmtree(d, ignore_errors=True)
        _work_dirs.discard(d)


atexit.register(_cleanup_own_work_dirs)

_scavenged = False


def _pid_alive(pid: int) -> bool:
    """Liveness without psutil. NEVER os.kill(pid, 0) on Windows -- any signal
    value other than the CTRL events unconditionally TerminateProcess-es the
    target."""
    if os.name == "nt":
        import ctypes
        PROCESS_QUERY_LIMITED_INFORMATION = 0x1000
        h = ctypes.windll.kernel32.OpenProcess(
            PROCESS_QUERY_LIMITED_INFORMATION, False, int(pid))
        if not h:
            return False
        ctypes.windll.kernel32.CloseHandle(h)
        return True
    try:
        os.kill(int(pid), 0)
    except ProcessLookupError:
        return False
    except PermissionError:
        return True
    except (OverflowError, ValueError):
        # A PID outside the OS's valid range (the test's synthetic dead pid
        # exceeds Linux's signed-int pid space) cannot belong to a live
        # process. Windows never hits this: OpenProcess just fails.
        return False
    return True


def _scavenge_stale_work_dirs() -> None:
    """Once per process, on the first pbix_open: delete sibling pbix_mcp_*
    directories whose owning process is gone.

    The directory name ends in _<pid>_<uuid8> (parsed from the END, because an
    alias may itself contain underscores). A dead pid is deleted immediately; a
    live or unparseable name is deleted only past a 7-day backstop, so another
    process's ACTIVE extraction is never touched even under pid reuse.
    """
    global _scavenged
    if _scavenged:
        return
    _scavenged = True
    root = tempfile.gettempdir()
    now = time.time()
    try:
        entries = os.listdir(root)
    except OSError:
        return
    for name in entries:
        if not name.startswith("pbix_mcp_"):
            continue
        full = os.path.join(root, name)
        if full in _work_dirs or not os.path.isdir(full):
            continue
        parts = name.split("_")
        pid = parts[-2] if len(parts) >= 4 and parts[-2].isdigit() else None
        stale = False
        if pid is not None and not _pid_alive(int(pid)):
            stale = True
        else:
            try:
                stale = (now - os.path.getmtime(full)) > 7 * 86400
            except OSError:
                stale = False
        if stale:
            shutil.rmtree(full, ignore_errors=True)
            logger.debug("Scavenged stale work dir %s", full)
# key = alias (user-chosen or auto), value = {
#   "path": str,              # original file path
#   "work_dir": str,          # temp extraction directory
#   "is_pbit": bool,
#   "modified": bool,
# }


# ---------------------------------------------------------------------------
# ZIP extraction safety limits (guard against decompression bombs and
# path-traversal in untrusted .pbix / .pbit / .pbiviz archives).
#   - Python's own extractall() already strips leading "/" and ".." components
#     and never materialises symlink entries, so classic Zip-Slip is mitigated;
#     these limits add the missing size caps and a defence-in-depth containment
#     check so a small crafted archive cannot fill the disk or escape work_dir.
# ---------------------------------------------------------------------------
_ZIP_MAX_TOTAL_UNCOMPRESSED = 4 * 1024**3   # 4 GiB across all members
_ZIP_MAX_FILE_UNCOMPRESSED = 2 * 1024**3    # 2 GiB for any single member
_ZIP_MAX_MEMBERS = 20000                    # PBIX has ~dozens; this is headroom
_ZIP_MAX_RATIO = 100                        # uncompressed:compressed per member
_ZIP_RATIO_MIN_SIZE = 1 << 16               # only ratio-check members >= 64 KiB


def _safe_join(base_dir: str, *parts: str) -> str:
    """Join *parts onto base_dir, guaranteeing the result stays inside base_dir.

    Prevents path traversal / arbitrary file write when any part is influenced by
    untrusted input — tool arguments (e.g. a theme ``filename``), ZIP member
    names, or manifest fields. Raises UnsafeWriteError if the resolved path would
    escape base_dir via ``..`` components, an absolute path, or (on Windows) a
    drive/UNC prefix.
    """
    base_real = os.path.realpath(base_dir)
    target = os.path.realpath(os.path.join(base_dir, *parts))
    if target != base_real and not target.startswith(base_real + os.sep):
        raise UnsafeWriteError(
            "Refusing to write outside the working directory (path traversal): "
            f"{os.path.join(*parts)!r}"
        )
    return target


def _validate_zip_members(zf: "zipfile.ZipFile", dest: str) -> None:
    """Reject decompression bombs and path-traversal before extracting.

    Raises InvalidPBIXError on any violation. Validates the whole archive up
    front so a malicious file is refused before a single byte is written.
    """
    infos = zf.infolist()
    if len(infos) > _ZIP_MAX_MEMBERS:
        raise InvalidPBIXError(
            f"Archive has too many entries ({len(infos)} > {_ZIP_MAX_MEMBERS}); "
            "refusing to extract (possible zip bomb)."
        )
    dest_real = os.path.realpath(dest)
    total = 0
    for zi in infos:
        # --- size / decompression-bomb caps ---
        total += zi.file_size
        if zi.file_size > _ZIP_MAX_FILE_UNCOMPRESSED:
            raise InvalidPBIXError(
                f"Archive member '{zi.filename}' is too large "
                f"({zi.file_size} bytes > {_ZIP_MAX_FILE_UNCOMPRESSED}); "
                "refusing to extract (possible zip bomb)."
            )
        if total > _ZIP_MAX_TOTAL_UNCOMPRESSED:
            raise InvalidPBIXError(
                f"Archive expands to more than {_ZIP_MAX_TOTAL_UNCOMPRESSED} "
                "bytes uncompressed; refusing to extract (possible zip bomb)."
            )
        if (
            zi.file_size >= _ZIP_RATIO_MIN_SIZE
            and zi.compress_size > 0
            and zi.file_size / zi.compress_size > _ZIP_MAX_RATIO
        ):
            raise InvalidPBIXError(
                f"Archive member '{zi.filename}' has a suspicious compression "
                f"ratio ({zi.file_size // max(zi.compress_size, 1)}:1); "
                "refusing to extract (possible zip bomb)."
            )
        # --- path containment (defence in depth over extractall's own sanitising) ---
        target = os.path.realpath(os.path.join(dest, zi.filename))
        if target != dest_real and not target.startswith(dest_real + os.sep):
            raise InvalidPBIXError(
                f"Archive member '{zi.filename}' resolves outside the "
                "extraction directory; refusing to extract (path traversal)."
            )
        # --- reject symlink entries (belt-and-suspenders: CPython never
        #     creates them, but do not even consider such an archive valid) ---
        if (zi.external_attr >> 16) & 0o170000 == 0o120000:
            raise InvalidPBIXError(
                f"Archive member '{zi.filename}' is a symlink; "
                "refusing to extract."
            )


# ============================= HELPERS =====================================

def _ensure_open(alias: str) -> dict:
    if alias not in _open_files:
        raise FileNotOpenError(
            f"No file open with alias '{alias}'. "
            f"Open files: {list(_open_files.keys()) or '(none)'}"
        )
    return _open_files[alias]


def _extract_pbix(pbix_path: str, work_dir: str) -> None:
    """Extract a PBIX/PBIT ZIP to work_dir (with bomb / traversal guards)."""
    with zipfile.ZipFile(pbix_path, "r") as zf:
        _validate_zip_members(zf, work_dir)
        zf.extractall(work_dir)


def _repack_pbix(work_dir: str, output_path: str, strip_sensitivity_label: bool = False) -> None:
    """Repack work_dir into a PBIX/PBIT ZIP file."""
    # Delete SecurityBindings — Power BI Desktop rejects modified files
    # that still have the original SecurityBindings
    sec_path = os.path.join(work_dir, "SecurityBindings")
    sec_removed = False
    if os.path.exists(sec_path):
        os.remove(sec_path)
        sec_removed = True

    # Update [Content_Types].xml to remove SecurityBindings reference
    if sec_removed:
        ct_path = os.path.join(work_dir, "[Content_Types].xml")
        if os.path.exists(ct_path):
            with open(ct_path, "r", encoding="utf-8") as f:
                ct_xml = f.read()
            ct_xml = ct_xml.replace(
                '<Override PartName="/SecurityBindings" ContentType=""/>',
                ""
            )
            with open(ct_path, "w", encoding="utf-8") as f:
                f.write(ct_xml)

    # Strip MSIP sensitivity label from docProps/custom.xml
    if strip_sensitivity_label:
        custom_path = os.path.join(work_dir, "docProps", "custom.xml")
        if os.path.exists(custom_path):
            import re
            with open(custom_path, "r", encoding="utf-8") as f:
                content = f.read()
            # Remove all MSIP_Label properties
            content = re.sub(
                r'<property[^>]*name="MSIP_Label_[^"]*"[^>]*>.*?</property>',
                "", content
            )
            with open(custom_path, "w", encoding="utf-8") as f:
                f.write(content)

    # Files that must NOT be included in the final ZIP
    _EXCLUDE_FILES = {
        "DataModel.abf",     # temp file from pbix_datamodel_decompress
        "metadata.sqlitedb", # extracted by ModelReader / tools — stale, causes PBI crash
    }
    # Suffixes that are temp artifacts
    _EXCLUDE_SUFFIXES = (".abf", ".tmp", ".bak", ".sqlitedb")

    with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(work_dir):
            for file in files:
                # Skip temp/artifact files
                if file in _EXCLUDE_FILES or file.endswith(_EXCLUDE_SUFFIXES):
                    continue
                file_path = os.path.join(root, file)
                arcname = os.path.relpath(file_path, work_dir).replace("\\", "/")
                # DataModel should be stored, not deflated (it's XPress9 compressed)
                if file == "DataModel":
                    zf.write(file_path, arcname, compress_type=zipfile.ZIP_STORED)
                else:
                    zf.write(file_path, arcname)


def _read_json_component(work_dir: str, rel_path: str) -> Any:
    """Read a JSON component from the extracted work dir."""
    full = os.path.join(work_dir, rel_path)
    if not os.path.exists(full):
        return None
    enc = _detect_encoding(full)
    with open(full, "r", encoding=enc) as f:
        return json.load(f)


def _write_json_component(work_dir: str, rel_path: str, data: Any) -> None:
    """Write a JSON component back, preserving original encoding."""
    full = os.path.join(work_dir, rel_path)
    os.makedirs(os.path.dirname(full), exist_ok=True)
    enc = _detect_encoding(full) if os.path.exists(full) else "utf-16-le"
    text = json.dumps(data, indent=2, ensure_ascii=False)
    with open(full, "wb") as f:
        f.write(text.encode(enc))


def _read_datamashup_m_code(work_dir: str) -> str | None:
    """Extract M code from the DataMashup binary.

    The DataMashup is a binary stream that embeds a ZIP archive.
    We scan for the PK signature to find the inner ZIP, then read
    Formulas/Section1.m from it.
    """
    dm_path = os.path.join(work_dir, "DataMashup")
    if not os.path.exists(dm_path):
        return None

    with open(dm_path, "rb") as f:
        data = f.read()

    # Find the inner ZIP (PK\x03\x04 signature)
    pk_offset = data.find(b"PK\x03\x04")
    if pk_offset == -1:
        return None

    # Find the end of the ZIP (scan for end-of-central-directory)
    eocd_sig = b"PK\x05\x06"
    eocd_pos = data.rfind(eocd_sig)
    if eocd_pos == -1:
        return None

    # EOCD is 22 bytes minimum, but may have a comment
    eocd_comment_len = struct.unpack_from("<H", data, eocd_pos + 20)[0]
    zip_end = eocd_pos + 22 + eocd_comment_len

    zip_data = data[pk_offset:zip_end]

    try:
        with zipfile.ZipFile(io.BytesIO(zip_data), "r") as inner_zf:
            for candidate in [
                "Formulas/Section1.m",
                "formulas/Section1.m",
                "Section1.m",
            ]:
                if candidate in inner_zf.namelist():
                    return inner_zf.read(candidate).decode("utf-8-sig")

            return f"[No Section1.m found. Archive contains: {inner_zf.namelist()}]"
    except zipfile.BadZipFile:
        return "[Could not parse inner DataMashup ZIP]"


def _write_datamashup_m_code(work_dir: str, new_m_code: str) -> bool:
    """Replace M code inside the DataMashup binary.

    Strategy: locate the inner ZIP, extract it, replace Section1.m,
    rebuild the inner ZIP, splice it back into the binary stream.
    """
    dm_path = os.path.join(work_dir, "DataMashup")
    if not os.path.exists(dm_path):
        return False

    with open(dm_path, "rb") as f:
        data = f.read()

    pk_offset = data.find(b"PK\x03\x04")
    if pk_offset == -1:
        return False

    eocd_sig = b"PK\x05\x06"
    eocd_pos = data.rfind(eocd_sig)
    if eocd_pos == -1:
        return False

    eocd_comment_len = struct.unpack_from("<H", data, eocd_pos + 20)[0]
    zip_end = eocd_pos + 22 + eocd_comment_len

    old_zip_data = data[pk_offset:zip_end]

    # Rebuild inner ZIP with new M code
    new_zip_buf = io.BytesIO()
    try:
        with zipfile.ZipFile(io.BytesIO(old_zip_data), "r") as old_zf:
            with zipfile.ZipFile(new_zip_buf, "w", zipfile.ZIP_DEFLATED) as new_zf:
                for item in old_zf.namelist():
                    if item.endswith("Section1.m"):
                        new_zf.writestr(item, new_m_code.encode("utf-8"))
                    else:
                        new_zf.writestr(item, old_zf.read(item))
    except zipfile.BadZipFile:
        return False

    new_zip_bytes = new_zip_buf.getvalue()

    # Splice: prefix + new_zip + suffix
    prefix = data[:pk_offset]
    suffix = data[zip_end:]

    new_data = prefix + new_zip_bytes + suffix

    # If there's a size field at pk_offset - 4, update it
    if pk_offset >= 4:
        old_size = struct.unpack_from("<I", prefix, pk_offset - 4)[0]
        old_zip_len = zip_end - pk_offset
        if old_size == old_zip_len:
            new_data = bytearray(new_data)
            struct.pack_into("<I", new_data, pk_offset - 4, len(new_zip_bytes))
            new_data = bytes(new_data)

    with open(dm_path, "wb") as f:
        f.write(new_data)

    return True


def _detect_encoding(file_path: str) -> str:
    """Detect if a file is UTF-16-LE, UTF-8 BOM, or plain UTF-8."""
    with open(file_path, "rb") as f:
        header = f.read(4)
    if header[:2] == b"\xff\xfe":
        return "utf-16-le"
    if header[:3] == b"\xef\xbb\xbf":
        return "utf-8-sig"
    if len(header) >= 2 and header[1:2] == b"\x00":
        return "utf-16-le"
    return "utf-8"


def _is_pbir(work_dir: str) -> bool:
    """True when the report is stored in PBIR (Report/definition) format.

    Every report authored in the Power BI SERVICE downloads in this shape: the
    classic single ``Report/Layout`` file is absent and the report is instead a
    tree of per-page / per-visual JSON files.
    """
    return os.path.exists(
        os.path.join(work_dir, "Report", "definition", "pages", "pages.json"))


def _get_layout(work_dir: str) -> dict | None:
    """Read the report layout.

    Returns the classic ``Report/Layout`` document when present. For a PBIR
    report (which has no Report/Layout at all) the on-disk tree is converted to
    the same legacy shape, so every consumer reads service-authored reports
    through this one entry point instead of seeing "no layout".

    The converted document is READ-ONLY: it is marked ``__pbir__`` and
    ``_set_layout`` refuses to write it back — doing so would plant a classic
    Report/Layout inside a PBIR file and leave two conflicting definitions.
    """
    layout_path = os.path.join(work_dir, "Report", "Layout")
    if not os.path.exists(layout_path):
        return _get_layout_pbir(work_dir)
    enc = _detect_encoding(layout_path)
    with open(layout_path, "r", encoding=enc) as f:
        return json.load(f)


# Bumped by every layout write, and folded into _layout_stamp. File metadata
# alone cannot see two edits that land in the same filesystem timestamp tick AND
# produce the same byte length -- changing a slicer's selected value from "A" to
# "B" is exactly that. On Windows, whose timestamp granularity is coarse, the
# next evaluate then served the PREVIOUS slicer's result from cache. Counting
# our own writes is exact and free; mtime+size still covers external edits.
_layout_writes: dict[str, int] = {}


def _note_layout_write(work_dir: str) -> None:
    _layout_writes[work_dir] = _layout_writes.get(work_dir, 0) + 1


def _set_layout(work_dir: str, layout: dict) -> None:
    """Persist a layout — classic ``Report/Layout`` or a PBIR tree.

    Every layout mutation funnels through here, so this is where the format is
    decided. A classic report is written back as the single UTF-16-LE
    Report/Layout document. A PBIR report is edited IN PLACE in its
    Report/definition tree, each page/visual patched onto the original file it
    was read from so unmodelled fields survive.
    """
    _note_layout_write(work_dir)
    if _is_pbir(work_dir):
        # A PBIR report is edited IN PLACE in its Report/definition tree — each
        # page/visual is patched onto the original file it was read from. Never
        # write a classic Report/Layout here: that would leave the file with two
        # conflicting report definitions.
        _set_layout_pbir(work_dir, layout)
        return
    if layout.get("__pbir__"):
        raise UnsupportedFormatError(
            "This layout was synthesized from a PBIR report but the target has "
            "no Report/definition tree to write back into. Save it to its own "
            "file, or build a classic report with pbix_create."
        )
    layout_path = os.path.join(work_dir, "Report", "Layout")
    os.makedirs(os.path.dirname(layout_path), exist_ok=True)
    text = json.dumps(layout, ensure_ascii=False)
    with open(layout_path, "wb") as f:
        f.write(text.encode("utf-16-le"))


def _report_config_path(work_dir: str) -> str:
    return os.path.join(work_dir, "Report", "definition", "report.json")


def _get_report_config(work_dir: str) -> dict | None:
    """The document that carries ``resourcePackages`` / ``publicCustomVisuals``.

    Classic reports keep those at the top level of Report/Layout; PBIR reports
    keep them in Report/definition/report.json. Resource and custom-visual
    registration is otherwise identical between the two — the files themselves
    live under the same Report/StaticResources and Report/CustomVisuals paths.
    """
    if _is_pbir(work_dir):
        path = _report_config_path(work_dir)
        if not os.path.exists(path):
            return None
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return None
    return _get_layout(work_dir)


def _set_report_config(work_dir: str, cfg: dict) -> None:
    """Persist a document obtained from :func:`_get_report_config`."""
    if _is_pbir(work_dir):
        path = _report_config_path(work_dir)
        os.makedirs(os.path.dirname(path), exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2, ensure_ascii=False)
        return
    _set_layout(work_dir, cfg)


def _parse_visual_config(vc: dict) -> dict:
    """Parse the 'config' JSON string inside a visual container."""
    config_str = vc.get("config", "{}")
    if isinstance(config_str, str):
        try:
            return json.loads(config_str)
        except json.JSONDecodeError:
            return {}
    return config_str if isinstance(config_str, dict) else {}


def _get_visual_type(config: dict) -> str:
    """Extract visual type from parsed config."""
    sc = config.get("singleVisual", config.get("singleVisualGroup", {}))
    if sc:
        return sc.get("visualType", "unknown")
    return "unknown"


def _get_visual_name(config: dict) -> str:
    """Extract the visual name from config."""
    return config.get("name", "")


def _set_value_by_dot_path(obj: Any, path: str, value: Any) -> None:
    """Set a nested value using a dot-separated path like 'a.b.c'.

    A NUMERIC segment means a list index. Power BI's formatting structures are
    arrays of ``{properties: ...}``, so creating a missing level as a dict keyed
    "0" produced JSON the schema rejects (``title: {"0": ...}`` instead of
    ``title: [...]``). The next segment decides which container to create.
    """
    keys = path.split(".")
    for i, key in enumerate(keys[:-1]):
        want_list = keys[i + 1].isdigit()
        if isinstance(obj, dict):
            if obj.get(key) is None:
                obj[key] = [] if want_list else {}
            obj = obj[key]
        elif isinstance(obj, list):
            if not key.isdigit():
                raise ValueError(
                    f"Path segment '{key}' indexes a list — it must be a number")
            idx = int(key)
            while len(obj) <= idx:
                obj.append([] if want_list else {})
            if obj[idx] is None:
                obj[idx] = [] if want_list else {}
            obj = obj[idx]
        else:
            raise ValueError(f"Cannot traverse into {type(obj)} at key '{key}'")

    final_key = keys[-1]
    if isinstance(obj, dict):
        obj[final_key] = value
    elif isinstance(obj, list):
        if not final_key.isdigit():
            raise ValueError(
                f"Path segment '{final_key}' indexes a list — it must be a number")
        idx = int(final_key)
        while len(obj) <= idx:
            obj.append(None)
        obj[idx] = value
    else:
        raise ValueError(f"Cannot set key '{final_key}' on {type(obj)}")


# ---- Visual formatting helpers ----

_DISPLAY_UNITS = {
    "none": "1D", "thousands": "1000D", "millions": "1000000D",
    "billions": "1000000000D", "trillions": "1000000000000D", "auto": "0D",
}
_LEGEND_POSITIONS = {
    "top": "'Top'", "bottom": "'Bottom'", "left": "'Left'",
    "right": "'Right'", "topCenter": "'TopCenter'",
    "bottomCenter": "'BottomCenter'", "leftCenter": "'LeftCenter'",
    "rightCenter": "'RightCenter'",
}
_ALIGNMENTS = {"left": "'Left'", "center": "'Center'", "right": "'Right'"}


def _pbi_lit(value) -> dict:
    """Convert a Python value to PBI Literal expression wrapper."""
    if isinstance(value, bool):
        raw = "true" if value else "false"
    elif isinstance(value, int):
        raw = f"{value}L"
    elif isinstance(value, float):
        raw = f"{value}D"
    elif isinstance(value, str):
        raw = f"'{value}'"
    else:
        raw = str(value)
    return {"expr": {"Literal": {"Value": raw}}}


def _pbi_props(mapping: dict, src: dict) -> dict:
    """Build PBI properties dict from a key mapping and source values.

    mapping: {pbi_property_name: (src_key, transform_fn_or_None)}
    src: the user-provided dict for this formatting category
    """
    props = {}
    for pbi_key, (src_key, transform) in mapping.items():
        if src_key in src:
            val = src[src_key]
            if transform:
                val = transform(val)
            props[pbi_key] = _pbi_lit(val)
    return props


def _hex_luminance(hex_color: str) -> float:
    """Relative luminance of a hex color (WCAG 2.0 formula). Range 0..1."""
    h = hex_color.lstrip("#")
    if len(h) != 6:
        return 0.5
    r, g, b = int(h[0:2], 16) / 255, int(h[2:4], 16) / 255, int(h[4:6], 16) / 255
    # sRGB linearization
    r = r / 12.92 if r <= 0.03928 else ((r + 0.055) / 1.055) ** 2.4
    g = g / 12.92 if g <= 0.03928 else ((g + 0.055) / 1.055) ** 2.4
    b = b / 12.92 if b <= 0.03928 else ((b + 0.055) / 1.055) ** 2.4
    return 0.2126 * r + 0.7152 * g + 0.0722 * b


def _contrast_ratio(lum1: float, lum2: float) -> float:
    """WCAG contrast ratio between two luminances."""
    lighter = max(lum1, lum2)
    darker = min(lum1, lum2)
    return (lighter + 0.05) / (darker + 0.05)


def _readable_text_color(bg_hex: str) -> str:
    """Return '#FFFFFF' or a dark color for readable text on a background."""
    lum = _hex_luminance(bg_hex)
    return "#FFFFFF" if lum < 0.35 else "#1A1A1A"


def _solid_color(hex_color: str) -> dict:
    """Wrap a hex color in PBI's solid.color structure."""
    return {"solid": {"color": _pbi_lit(hex_color)}}


#: Every format-card name :func:`_build_format_objects` reads out of ``fmt``.
#: Used ONLY to tell "card I don't know" apart from "card I know whose
#: properties I all dropped" when nothing was applied (issue #51: the old
#: message called ``valueAxis`` unrecognised while listing it as supported,
#: sending the caller after the wrong bug). Kept honest by
#: ``tests/test_report_editing.py::TestFormatCardsConstantMatchesMapper``,
#: which scans the mapper for the cards it actually reads.
_FORMAT_CARDS = frozenset({
    "title", "subtitle", "background", "border", "dropShadow", "padding",
    "spacing", "divider", "visualHeader", "visualTooltip", "stylePreset",
    "altText", "lockAspect", "legend", "dataLabels", "labels",
    "categoryAxis", "valueAxis", "dataColors", "grid", "columnHeaders",
    "values", "total", "outline", "shape", "fill", "line", "categoryLabels",
    "slices", "action", "text", "smallMultiples", "rowHeaders", "subTotals",
    "referenceLine", "donut", "bubbles", "markers", "imageScaling", "card",
    "cardTitle", "columnFormatting", "zoom", "general", "visualLink",
    "visualHeaderTooltip",
})


def _build_format_objects(fmt: dict, visual_type: str = "") -> dict:
    """Convert human-readable format dict to PBI objects structures.

    Returns dict with two keys:
      - "_objects": data formatting (labels, legend, axis, dataPoint, grid, etc.)
      - "_vcObjects": visual container formatting (title, subtitle, background,
        border, dropShadow, padding, spacing, visualHeader, divider, etc.)

    All property names and value formats match PBI Desktop March 2026 ground truth.
    Colors use {"solid": {"color": {"expr": {"Literal": {"Value": "'#hex'"}}}}}

    ``visual_type`` disambiguates the properties whose NAME depends on the
    target visual (issue #45: a plain card reads categoryLabels.color, a
    multi-row card reads categoryLabels.categoryLabelFontColor).
    """
    objects: dict[str, list] = {}
    vc_objects: dict[str, list] = {}

    def _add(category: str, props: dict):
        if props:
            objects[category] = [{"properties": props}]

    def _add_vc(category: str, props: dict):
        if props:
            vc_objects[category] = [{"properties": props}]

    # ================================================================
    # vcObjects — visual container formatting
    # ================================================================

    # --- title ---
    if "title" in fmt:
        t = fmt["title"]
        props = {}
        if "show" in t: props["show"] = _pbi_lit(t["show"])
        if "text" in t: props["text"] = _pbi_lit(t["text"])
        if "fontSize" in t: props["fontSize"] = _pbi_lit(float(t["fontSize"]))
        if "color" in t: props["fontColor"] = _solid_color(t["color"])
        if "fontFamily" in t: props["fontFamily"] = _pbi_lit(t["fontFamily"])
        if "bold" in t: props["bold"] = _pbi_lit(t["bold"])
        if "italic" in t: props["italic"] = _pbi_lit(t["italic"])
        if "alignment" in t:
            raw = _ALIGNMENTS.get(t["alignment"], f"'{t['alignment']}'")
            props["alignment"] = {"expr": {"Literal": {"Value": raw}}}
        if "heading" in t: props["heading"] = _pbi_lit(t["heading"])
        if "titleWrap" in t: props["titleWrap"] = _pbi_lit(t["titleWrap"])
        if "background" in t: props["background"] = _solid_color(t["background"])
        _add_vc("title", props)

    # --- subtitle ---
    if "subtitle" in fmt:
        s = fmt["subtitle"]
        props = {}
        if "show" in s: props["show"] = _pbi_lit(s["show"])
        if "text" in s: props["text"] = _pbi_lit(s["text"])
        if "fontSize" in s: props["fontSize"] = _pbi_lit(float(s["fontSize"]))
        if "color" in s: props["fontColor"] = _solid_color(s["color"])
        if "fontFamily" in s: props["fontFamily"] = _pbi_lit(s["fontFamily"])
        if "titleWrap" in s: props["titleWrap"] = _pbi_lit(s["titleWrap"])
        _add_vc("subTitle", props)

    # --- background ---
    if "background" in fmt:
        bg = fmt["background"]
        props = {}
        if "color" in bg: props["color"] = _solid_color(bg["color"])
        if "transparency" in bg: props["transparency"] = _pbi_lit(float(bg["transparency"]))
        if "show" in bg: props["show"] = _pbi_lit(bg["show"])
        else: props["show"] = _pbi_lit(True)
        _add_vc("background", props)

    # --- border ---
    if "border" in fmt:
        bd = fmt["border"]
        props = {}
        if "show" in bd: props["show"] = _pbi_lit(bd["show"])
        if "color" in bd: props["color"] = _solid_color(bd["color"])
        if "radius" in bd: props["radius"] = _pbi_lit(float(bd["radius"]))
        if "width" in bd: props["width"] = _pbi_lit(float(bd["width"]))
        _add_vc("border", props)

    # --- dropShadow ---
    if "dropShadow" in fmt:
        ds = fmt["dropShadow"]
        props = {}
        if "show" in ds: props["show"] = _pbi_lit(ds["show"])
        if "color" in ds: props["color"] = _solid_color(ds["color"])
        if "position" in ds: props["position"] = _pbi_lit(ds["position"])
        if "preset" in ds: props["preset"] = _pbi_lit(ds["preset"])
        if "angle" in ds: props["angle"] = _pbi_lit(float(ds["angle"]))
        if "blur" in ds: props["shadowBlur"] = _pbi_lit(float(ds["blur"]))
        if "distance" in ds: props["shadowDistance"] = _pbi_lit(float(ds["distance"]))
        if "spread" in ds: props["shadowSpread"] = _pbi_lit(float(ds["spread"]))
        if "transparency" in ds: props["transparency"] = _pbi_lit(float(ds["transparency"]))
        _add_vc("dropShadow", props)

    # --- padding ---
    if "padding" in fmt:
        pd = fmt["padding"]
        props = {}
        if isinstance(pd, (int, float)):
            for side in ("top", "bottom", "left", "right"):
                props[side] = _pbi_lit(int(pd))
        else:
            if "top" in pd: props["top"] = _pbi_lit(int(pd["top"]))
            if "bottom" in pd: props["bottom"] = _pbi_lit(int(pd["bottom"]))
            if "left" in pd: props["left"] = _pbi_lit(int(pd["left"]))
            if "right" in pd: props["right"] = _pbi_lit(int(pd["right"]))
        _add_vc("padding", props)

    # --- spacing ---
    if "spacing" in fmt:
        sp = fmt["spacing"]
        props = {}
        props["customizeSpacing"] = _pbi_lit(True)
        if "belowTitle" in sp: props["spaceBelowTitle"] = _pbi_lit(int(sp["belowTitle"]))
        if "belowSubTitle" in sp: props["spaceBelowSubTitle"] = _pbi_lit(int(sp["belowSubTitle"]))
        if "belowTitleArea" in sp: props["spaceBelowTitleArea"] = _pbi_lit(int(sp["belowTitleArea"]))
        if "vertical" in sp: props["verticalSpacing"] = _pbi_lit(int(sp["vertical"]))
        _add_vc("spacing", props)

    # --- divider ---
    if "divider" in fmt:
        dv = fmt["divider"]
        props = {}
        if "show" in dv: props["show"] = _pbi_lit(dv["show"])
        if "color" in dv: props["color"] = _solid_color(dv["color"])
        if "width" in dv: props["width"] = _pbi_lit(float(dv["width"]))
        if "style" in dv: props["style"] = _pbi_lit(dv["style"])
        if "ignorePadding" in dv: props["ignorePadding"] = _pbi_lit(dv["ignorePadding"])
        _add_vc("divider", props)

    # --- visualHeader ---
    if "visualHeader" in fmt:
        vh = fmt["visualHeader"]
        props = {}
        if "show" in vh: props["show"] = _pbi_lit(vh["show"])
        for btn in ("showOptionsMenu", "showFocusModeButton", "showPinButton",
                     "showFilterRestatementButton", "showTooltipButton",
                     "showDrillUpButton", "showDrillDownLevelButton",
                     "showDrillDownExpandButton", "showDrillToggleButton",
                     "showDrillRoleSelector", "showVisualErrorButton",
                     "showVisualWarningButton", "showVisualInformationButton",
                     "showSeeDataLayoutToggleButton"):
            if btn in vh: props[btn] = _pbi_lit(vh[btn])
        _add_vc("visualHeader", props)

    # --- visualTooltip ---
    if "visualTooltip" in fmt:
        vt = fmt["visualTooltip"]
        props = {}
        if "show" in vt: props["show"] = _pbi_lit(vt["show"])
        if "type" in vt: props["type"] = _pbi_lit(vt["type"])
        if "fontSize" in vt: props["fontSize"] = _pbi_lit(float(vt["fontSize"]))
        if "titleFontColor" in vt: props["titleFontColor"] = _solid_color(vt["titleFontColor"])
        if "valueFontColor" in vt: props["valueFontColor"] = _solid_color(vt["valueFontColor"])
        if "actionFontColor" in vt: props["actionFontColor"] = _solid_color(vt["actionFontColor"])
        if "background" in vt: props["background"] = _solid_color(vt["background"])
        _add_vc("visualTooltip", props)

    # --- stylePreset ---
    if "stylePreset" in fmt:
        _add_vc("stylePreset", {"name": _pbi_lit(fmt["stylePreset"])})

    # --- general (vcObjects) ---
    if "altText" in fmt:
        _add_vc("general", {"altText": _pbi_lit(fmt["altText"])})

    # --- lockAspect ---
    if "lockAspect" in fmt:
        _add_vc("lockAspect", {"show": _pbi_lit(fmt["lockAspect"])})

    # ================================================================
    # objects — data formatting
    # ================================================================

    # --- legend ---
    if "legend" in fmt:
        lg = fmt["legend"]
        props = {}
        if "show" in lg: props["show"] = _pbi_lit(lg["show"])
        if "fontSize" in lg: props["fontSize"] = _pbi_lit(float(lg["fontSize"]))
        if "color" in lg: props["fontColor"] = _solid_color(lg["color"])
        if "fontFamily" in lg: props["fontFamily"] = _pbi_lit(lg["fontFamily"])
        if "position" in lg:
            raw = _LEGEND_POSITIONS.get(lg["position"], f"'{lg['position']}'")
            props["position"] = {"expr": {"Literal": {"Value": raw}}}
        _add("legend", props)

    # --- dataLabels / labels (both -> the PBI `labels` object) ---
    # "dataLabels" is the friendly name for a chart's data labels; "labels" is
    # the raw PBI object name and is what a Card's "Callout value" (the big
    # number's colour/size/units) lives under (objects.labels.*). Accept either
    # so `{"labels": {"color": "#.."}}` isn't silently dropped (OpenBI #1 gap).
    _dl = fmt.get("dataLabels", fmt.get("labels"))
    if _dl is not None:
        dl = _dl
        props = {}
        if "show" in dl: props["show"] = _pbi_lit(dl["show"])
        if "fontSize" in dl: props["fontSize"] = _pbi_lit(float(dl["fontSize"]))
        if "color" in dl: props["color"] = _solid_color(dl["color"])
        if "fontFamily" in dl: props["fontFamily"] = _pbi_lit(dl["fontFamily"])
        if "bold" in dl: props["bold"] = _pbi_lit(dl["bold"])
        if "italic" in dl: props["italic"] = _pbi_lit(dl["italic"])
        if "displayUnits" in dl:
            raw = _DISPLAY_UNITS.get(dl["displayUnits"], f"{dl['displayUnits']}D")
            props["labelDisplayUnits"] = {"expr": {"Literal": {"Value": raw}}}
        if "decimalPlaces" in dl: props["labelPrecision"] = _pbi_lit(int(dl["decimalPlaces"]))
        _add("labels", props)

    # --- categoryAxis ---
    if "categoryAxis" in fmt:
        ca = fmt["categoryAxis"]
        props = {}
        if "show" in ca: props["show"] = _pbi_lit(ca["show"])
        if "fontSize" in ca: props["fontSize"] = _pbi_lit(float(ca["fontSize"]))
        if "color" in ca: props["labelColor"] = _solid_color(ca["color"])
        if "fontFamily" in ca: props["fontFamily"] = _pbi_lit(ca["fontFamily"])
        if "title" in ca:
            props["showAxisTitle"] = _pbi_lit(True)
            props["axisTitle"] = _pbi_lit(ca["title"])
        if "titleFontSize" in ca: props["titleFontSize"] = _pbi_lit(float(ca["titleFontSize"]))
        if "gridlineShow" in ca: props["gridlineShow"] = _pbi_lit(ca["gridlineShow"])
        if "innerPadding" in ca: props["innerPadding"] = _pbi_lit(int(ca["innerPadding"]))
        if "invertAxis" in ca: props["invertAxis"] = _pbi_lit(ca["invertAxis"])
        if "concatenateLabels" in ca: props["concatenateLabels"] = _pbi_lit(ca["concatenateLabels"])
        if "axisType" in ca: props["axisType"] = _pbi_lit(ca["axisType"])
        if "start" in ca: props["start"] = _pbi_lit(float(ca["start"]))
        if "end" in ca: props["end"] = _pbi_lit(float(ca["end"]))
        if "switchAxisPosition" in ca: props["switchAxisPosition"] = _pbi_lit(ca["switchAxisPosition"])
        if "preferredCategoryWidth" in ca: props["preferredCategoryWidth"] = _pbi_lit(float(ca["preferredCategoryWidth"]))
        _add("categoryAxis", props)

    # --- valueAxis ---
    if "valueAxis" in fmt:
        va = fmt["valueAxis"]
        props = {}
        if "show" in va: props["show"] = _pbi_lit(va["show"])
        if "fontSize" in va: props["fontSize"] = _pbi_lit(float(va["fontSize"]))
        if "color" in va: props["labelColor"] = _solid_color(va["color"])
        if "fontFamily" in va: props["fontFamily"] = _pbi_lit(va["fontFamily"])
        if "displayUnits" in va:
            raw = _DISPLAY_UNITS.get(va["displayUnits"], f"{va['displayUnits']}D")
            props["labelDisplayUnits"] = {"expr": {"Literal": {"Value": raw}}}
        if "title" in va:
            props["showAxisTitle"] = _pbi_lit(True)
            props["axisTitle"] = _pbi_lit(va["title"])
        if "titleFontSize" in va: props["titleFontSize"] = _pbi_lit(float(va["titleFontSize"]))
        if "gridlineShow" in va: props["gridlineShow"] = _pbi_lit(va["gridlineShow"])
        if "start" in va: props["start"] = _pbi_lit(float(va["start"]))
        if "end" in va: props["end"] = _pbi_lit(float(va["end"]))
        if "switchAxisPosition" in va: props["switchAxisPosition"] = _pbi_lit(va["switchAxisPosition"])
        if "decimalPlaces" in va: props["labelPrecision"] = _pbi_lit(int(va["decimalPlaces"]))
        # Secondary axis of a combo chart (issue #51). Power BI keeps BOTH
        # axes on this one card, the second under `sec`-prefixed property
        # names, so a combo's secondary axis was unreachable through the
        # humanized mapper -- every sec* key was dropped and the visual came
        # back unchanged. Only the measured property names are written; the
        # friendly aliases below map onto those same names.
        if "secShow" in va: props["secShow"] = _pbi_lit(va["secShow"])
        _sec_fs = va.get("secFontSize")
        if _sec_fs is not None: props["secFontSize"] = _pbi_lit(float(_sec_fs))
        _sec_color = va.get("secLabelColor", va.get("secColor"))
        if _sec_color is not None: props["secLabelColor"] = _solid_color(_sec_color)
        _sec_units = va.get("secLabelDisplayUnits", va.get("secDisplayUnits"))
        if _sec_units is not None:
            raw = _DISPLAY_UNITS.get(_sec_units, f"{_sec_units}D")
            props["secLabelDisplayUnits"] = {"expr": {"Literal": {"Value": raw}}}
        # `secTitle` mirrors the primary `title`: naming it turns it on.
        # `secAxisTitle` / `secShowAxisTitle` are the raw pair, and an
        # explicit secShowAxisTitle wins over the alias's implied True.
        if "secTitle" in va:
            props["secShowAxisTitle"] = _pbi_lit(True)
            props["secAxisTitle"] = _pbi_lit(va["secTitle"])
        if "secAxisTitle" in va: props["secAxisTitle"] = _pbi_lit(va["secAxisTitle"])
        if "secShowAxisTitle" in va:
            props["secShowAxisTitle"] = _pbi_lit(va["secShowAxisTitle"])
        if "secStart" in va: props["secStart"] = _pbi_lit(float(va["secStart"]))
        if "secEnd" in va: props["secEnd"] = _pbi_lit(float(va["secEnd"]))
        # alignZeros pins the two axes' zero lines together — a whole-card
        # property, not a per-axis one.
        if "alignZeros" in va: props["alignZeros"] = _pbi_lit(va["alignZeros"])
        _add("valueAxis", props)

    # --- dataColors (dataPoint) ---
    # NOTE: Multi-color dataColors with per-series/per-category selectors
    # are handled in pbix_format_visual directly (needs visual projections).
    # This only handles the single-color fallback.
    if "dataColors" in fmt and not fmt.get("_skip_datacolors"):
        colors = fmt["dataColors"]
        if isinstance(colors, list) and colors:
            props = {"fill": _solid_color(colors[0])}
            _add("dataPoint", props)

    # --- grid (table/matrix) ---
    if "grid" in fmt:
        gr = fmt["grid"]
        props = {}
        if "gridVertical" in gr: props["gridVertical"] = _pbi_lit(gr["gridVertical"])
        if "gridHorizontal" in gr: props["gridHorizontal"] = _pbi_lit(gr["gridHorizontal"])
        if "rowPadding" in gr: props["rowPadding"] = _pbi_lit(int(gr["rowPadding"]))
        if "outlineColor" in gr: props["outlineColor"] = _solid_color(gr["outlineColor"])
        if "outlineWeight" in gr: props["outlineWeight"] = _pbi_lit(int(gr["outlineWeight"]))
        if "textSize" in gr: props["textSize"] = _pbi_lit(float(gr["textSize"]))
        if "gridHorizontalColor" in gr: props["gridHorizontalColor"] = _solid_color(gr["gridHorizontalColor"])
        if "gridVerticalColor" in gr: props["gridVerticalColor"] = _solid_color(gr["gridVerticalColor"])
        _add("grid", props)

    # --- columnHeaders (table/matrix) ---
    if "columnHeaders" in fmt:
        ch = fmt["columnHeaders"]
        props = {}
        if "bold" in ch: props["bold"] = _pbi_lit(ch["bold"])
        if "fontSize" in ch: props["fontSize"] = _pbi_lit(float(ch["fontSize"]))
        if "fontFamily" in ch: props["fontFamily"] = _pbi_lit(ch["fontFamily"])
        if "fontColor" in ch: props["fontColor"] = _solid_color(ch["fontColor"])
        if "backColor" in ch: props["backColor"] = _solid_color(ch["backColor"])
        if "alignment" in ch: props["alignment"] = _pbi_lit(ch["alignment"])
        if "autoSizeColumnWidth" in ch: props["autoSizeColumnWidth"] = _pbi_lit(ch["autoSizeColumnWidth"])
        if "wordWrap" in ch: props["wordWrap"] = _pbi_lit(ch["wordWrap"])
        _add("columnHeaders", props)

    # --- values (table rows) ---
    if "values" in fmt:
        vl = fmt["values"]
        props = {}
        if "bold" in vl: props["bold"] = _pbi_lit(vl["bold"])
        if "fontSize" in vl: props["fontSize"] = _pbi_lit(float(vl["fontSize"]))
        if "fontFamily" in vl: props["fontFamily"] = _pbi_lit(vl["fontFamily"])
        if "fontColor" in vl: props["fontColor"] = _solid_color(vl["fontColor"])
        if "backColor" in vl: props["backColor"] = _solid_color(vl["backColor"])
        if "wordWrap" in vl: props["wordWrap"] = _pbi_lit(vl["wordWrap"])
        # Alternating row colors
        if "backColorPrimary" in vl: props["backColorPrimary"] = _solid_color(vl["backColorPrimary"])
        if "backColorSecondary" in vl: props["backColorSecondary"] = _solid_color(vl["backColorSecondary"])
        if "fontColorPrimary" in vl: props["fontColorPrimary"] = _solid_color(vl["fontColorPrimary"])
        if "fontColorSecondary" in vl: props["fontColorSecondary"] = _solid_color(vl["fontColorSecondary"])
        _add("values", props)

    # --- total (table/matrix totals row) ---
    if "total" in fmt:
        tt = fmt["total"]
        props = {}
        if "show" in tt: props["show"] = _pbi_lit(tt["show"])
        if "bold" in tt: props["bold"] = _pbi_lit(tt["bold"])
        if "fontSize" in tt: props["fontSize"] = _pbi_lit(float(tt["fontSize"]))
        if "fontColor" in tt: props["fontColor"] = _solid_color(tt["fontColor"])
        if "backColor" in tt: props["backColor"] = _solid_color(tt["backColor"])
        _add("total", props)

    # --- outline ---
    if "outline" in fmt:
        ol = fmt["outline"]
        props = {}
        if "show" in ol: props["show"] = _pbi_lit(ol["show"])
        if "weight" in ol: props["weight"] = _pbi_lit(int(ol["weight"]))
        if "color" in ol:
            # Stroke colour naming is per-visual (issue #47, measured on
            # Desktop-authored shapes): the `shape` vintage carries it as
            # outline.lineColor, basicShape as line.lineColor — the generic
            # outline.color matched no measured shape, so the authored
            # colour rendered as Desktop's default stroke.
            if visual_type == "shape":
                props["lineColor"] = _solid_color(ol["color"])
            elif visual_type != "basicShape":
                props["color"] = _solid_color(ol["color"])
        if visual_type == "basicShape":
            line_props = {}
            if "color" in ol: line_props["lineColor"] = _solid_color(ol["color"])
            if "weight" in ol: line_props["weight"] = _pbi_lit(int(ol["weight"]))
            if "transparency" in ol:
                line_props["transparency"] = _pbi_lit(float(ol["transparency"]))
            _add("line", line_props)
            props.pop("weight", None)
        _add("outline", props)

    # --- shape (buttons, shapes) ---
    if "shape" in fmt:
        sh = fmt["shape"]
        props = {}
        if "map" in sh: props["map"] = _pbi_lit(sh["map"])
        # Geometry (issue #47): line vs rectangle vs oval vs arrow. Measured
        # Desktop shapes carry shape.tileShape (`shape` vintage) or
        # general.shapeType (basicShape).
        geometry = sh.get("tileShape", sh.get("geometry"))
        if geometry is not None:
            if visual_type == "basicShape":
                _add("general", {"shapeType": _pbi_lit(geometry)})
            else:
                props["tileShape"] = _pbi_lit(geometry)
        if "rotation" in sh:
            # Rotation lives on its OWN card in every measured Desktop shape
            # (issue #47): rotation.shapeAngle on `shape`,
            # rotation.angle on basicShape — never shape.rotation, which
            # Desktop ignores. Non-shape visuals keep the legacy spelling.
            if visual_type == "shape":
                _add("rotation",
                     {"shapeAngle": _pbi_lit(float(sh["rotation"]))})
            elif visual_type == "basicShape":
                _add("rotation", {"angle": _pbi_lit(float(sh["rotation"]))})
            else:
                props["rotation"] = _pbi_lit(int(sh["rotation"]))
        _add("shape", props)

    # --- fill (shape fill) ---
    if "fill" in fmt:
        fl = fmt["fill"]
        props = {}
        if "color" in fl: props["fillColor"] = _solid_color(fl["color"])
        if "transparency" in fl: props["transparency"] = _pbi_lit(float(fl["transparency"]))
        if "show" in fl: props["show"] = _pbi_lit(fl["show"])
        _add("fill", props)

    # --- line (line charts) ---
    if "line" in fmt:
        ln = fmt["line"]
        props = {}
        if "lineStyle" in ln: props["lineStyle"] = _pbi_lit(ln["lineStyle"])
        if "strokeWidth" in ln: props["strokeWidth"] = _pbi_lit(float(ln["strokeWidth"]))
        if "joinType" in ln: props["joinType"] = _pbi_lit(int(ln["joinType"]))
        if "showMarker" in ln: props["showMarker"] = _pbi_lit(ln["showMarker"])
        if "markerShape" in ln: props["markerShape"] = _pbi_lit(ln["markerShape"])
        if "markerSize" in ln: props["markerSize"] = _pbi_lit(int(ln["markerSize"]))
        _add("lineStyles", props)

    # --- categoryLabels (card / multi-row card / pie / donut) ---
    if "categoryLabels" in fmt:
        cl = fmt["categoryLabels"]
        props = {}
        if "show" in cl: props["show"] = _pbi_lit(cl["show"])
        if "fontSize" in cl: props["fontSize"] = _pbi_lit(float(cl["fontSize"]))
        if "color" in cl:
            # The colour property's NAME depends on the visual (issue #45,
            # measured by write-save-reopen-readback): a plain card renders
            # categoryLabels.color; a multi-row card renders
            # categoryLabels.categoryLabelFontColor. Writing the multi-row
            # name onto a card was accepted, persisted, and rendered as
            # nothing.
            color_prop = ("color" if visual_type == "card"
                          else "categoryLabelFontColor")
            props[color_prop] = _solid_color(cl["color"])
        if "fontFamily" in cl: props["fontFamily"] = _pbi_lit(cl["fontFamily"])
        _add("categoryLabels", props)

    # --- slices (pie/donut) ---
    if "slices" in fmt:
        sl = fmt["slices"]
        props = {}
        if "innerRadius" in sl: props["innerRadiusRatio"] = _pbi_lit(int(sl["innerRadius"]))
        _add("slices", props)

    # --- action (vcObjects.visualLink — action buttons) ---
    if "action" in fmt:
        ac = fmt["action"]
        props = {}
        # Desktop-authored buttons carry the action in vcObjects.visualLink
        # (issue #48, measured): this key used to write objects.visualLink,
        # a bucket no Desktop button reads — the documented `action` key
        # produced a persisted-but-dead action.
        if "show" in ac: props["show"] = _pbi_lit(ac["show"])
        if "type" in ac: props["type"] = _pbi_lit(ac["type"])
        if "navigationSection" in ac: props["navigationSection"] = _pbi_lit(ac["navigationSection"])
        if "bookmark" in ac: props["bookmark"] = _pbi_lit(ac["bookmark"])
        if "webUrl" in ac: props["webUrl"] = _pbi_lit(ac["webUrl"])
        if "tooltip" in ac: props["tooltip"] = _pbi_lit(ac["tooltip"])
        _add_vc("visualLink", props)

    # --- text (actionButton label; objects.text with the default-state
    # selector, matching Desktop-authored buttons — issue #48) ---
    if "text" in fmt and visual_type == "actionButton":
        tx = fmt["text"]
        if isinstance(tx, str):
            tx = {"text": tx, "show": True}
        props = {}
        if "text" in tx: props["text"] = _pbi_lit(tx["text"])
        if "show" in tx: props["show"] = _pbi_lit(tx["show"])
        if "fontSize" in tx: props["fontSize"] = _pbi_lit(float(tx["fontSize"]))
        if "fontColor" in tx: props["fontColor"] = _solid_color(tx["fontColor"])
        if props:
            objects["text"] = [{"properties": props,
                                "selector": {"id": "default"}}]

    # --- smallMultiples ---
    if "smallMultiples" in fmt:
        sm = fmt["smallMultiples"]
        props = {}
        if "minWidth" in sm: props["minWidth"] = _pbi_lit(int(sm["minWidth"]))
        if "maxWidth" in sm: props["maxWidth"] = _pbi_lit(int(sm["maxWidth"]))
        if "minHeight" in sm: props["minHeight"] = _pbi_lit(int(sm["minHeight"]))
        _add("smallMultiplesLayout", props)

    # --- rowHeaders (matrix) ---
    if "rowHeaders" in fmt:
        rh = fmt["rowHeaders"]
        props = {}
        if "bold" in rh: props["bold"] = _pbi_lit(rh["bold"])
        if "fontSize" in rh: props["fontSize"] = _pbi_lit(float(rh["fontSize"]))
        if "fontFamily" in rh: props["fontFamily"] = _pbi_lit(rh["fontFamily"])
        if "fontColor" in rh: props["fontColor"] = _solid_color(rh["fontColor"])
        if "alignment" in rh: props["alignment"] = _pbi_lit(rh["alignment"])
        _add("rowHeaders", props)

    # --- subTotals (matrix) ---
    if "subTotals" in fmt:
        st = fmt["subTotals"]
        props = {}
        if "bold" in st: props["bold"] = _pbi_lit(st["bold"])
        if "fontSize" in st: props["fontSize"] = _pbi_lit(float(st["fontSize"]))
        if "fontColor" in st: props["fontColor"] = _solid_color(st["fontColor"])
        if "backColor" in st: props["backColor"] = _solid_color(st["backColor"])
        if "columnSubtotals" in st: props["columnSubtotals"] = _pbi_lit(st["columnSubtotals"])
        if "rowSubtotals" in st: props["rowSubtotals"] = _pbi_lit(st["rowSubtotals"])
        _add("subTotals", props)

    # --- referenceLine ---
    if "referenceLine" in fmt:
        rl = fmt["referenceLine"]
        props = {}
        if "show" in rl: props["show"] = _pbi_lit(rl["show"])
        if "displayName" in rl: props["displayName"] = _pbi_lit(rl["displayName"])
        if "color" in rl: props["lineColor"] = _solid_color(rl["color"])
        if "style" in rl: props["style"] = _pbi_lit(rl["style"])
        if "width" in rl: props["width"] = _pbi_lit(float(rl["width"]))
        if "transparency" in rl: props["transparency"] = _pbi_lit(float(rl["transparency"]))
        if "position" in rl: props["position"] = _pbi_lit(rl["position"])
        _add("y1AxisReferenceLine", props)

    # --- donut ---
    if "donut" in fmt:
        dn = fmt["donut"]
        props = {}
        if "innerRadius" in dn: props["innerRadius"] = _pbi_lit(int(dn["innerRadius"]))
        if "radius" in dn: props["radius"] = _pbi_lit(int(dn["radius"]))
        if "maxSlices" in dn: props["maxSlicesVisible"] = _pbi_lit(int(dn["maxSlices"]))
        _add("donut", props)

    # --- bubbles (scatter chart) ---
    if "bubbles" in fmt:
        bb = fmt["bubbles"]
        props = {}
        if "size" in bb: props["bubbleSize"] = _pbi_lit(int(bb["size"]))
        if "shape" in bb: props["markerShape"] = _pbi_lit(bb["shape"])
        if "rangeType" in bb: props["markerRangeType"] = _pbi_lit(bb["rangeType"])
        _add("bubbles", props)

    # --- markers (scatter/line) ---
    if "markers" in fmt:
        mk = fmt["markers"]
        props = {}
        if "borderWidth" in mk: props["borderWidth"] = _pbi_lit(float(mk["borderWidth"]))
        if "transparency" in mk: props["transparency"] = _pbi_lit(float(mk["transparency"]))
        _add("markers", props)

    # --- imageScaling ---
    if "imageScaling" in fmt:
        props = {"imageScalingType": _pbi_lit(fmt["imageScaling"])}
        _add("imageScaling", props)

    # --- card (new card visual styling) ---
    if "card" in fmt and isinstance(fmt["card"], dict):
        cd = fmt["card"]
        props = {}
        if "barShow" in cd: props["barShow"] = _pbi_lit(cd["barShow"])
        if "barColor" in cd: props["barColor"] = _solid_color(cd["barColor"])
        if "barWeight" in cd: props["barWeight"] = _pbi_lit(float(cd["barWeight"]))
        if "cardPadding" in cd: props["cardPadding"] = _pbi_lit(float(cd["cardPadding"]))
        if "outlineStyle" in cd: props["outlineStyle"] = _pbi_lit(float(cd["outlineStyle"]))
        _add("card", props)

    # --- cardTitle ---
    if "cardTitle" in fmt:
        ct = fmt["cardTitle"]
        props = {}
        if "color" in ct: props["color"] = _solid_color(ct["color"])
        if "fontSize" in ct: props["fontSize"] = _pbi_lit(float(ct["fontSize"]))
        _add("cardTitle", props)

    # --- columnFormatting (table/matrix) ---
    if "columnFormatting" in fmt:
        cf = fmt["columnFormatting"]
        props = {}
        if "alignment" in cf: props["alignment"] = _pbi_lit(cf["alignment"])
        if "displayUnits" in cf:
            raw = _DISPLAY_UNITS.get(cf["displayUnits"], f"{cf['displayUnits']}D")
            props["labelDisplayUnits"] = {"expr": {"Literal": {"Value": raw}}}
        if "decimalPlaces" in cf: props["labelPrecision"] = _pbi_lit(int(cf["decimalPlaces"]))
        if "styleHeader" in cf: props["styleHeader"] = _pbi_lit(cf["styleHeader"])
        if "styleTotal" in cf: props["styleTotal"] = _pbi_lit(cf["styleTotal"])
        _add("columnFormatting", props)

    # --- zoom (scatter chart zoom slider) ---
    if "zoom" in fmt:
        _add("zoom", {"show": _pbi_lit(fmt["zoom"])})

    # --- general.objects (image URL, layout, orientation) ---
    if "general" in fmt and isinstance(fmt["general"], dict):
        gn = fmt["general"]
        props = {}
        if "layout" in gn: props["layout"] = _pbi_lit(gn["layout"])
        if "orientation" in gn: props["orientation"] = _pbi_lit(float(gn["orientation"]))
        _add("general", props)

    # --- visualLink (vcObjects — action buttons navigation) ---
    if "visualLink" in fmt:
        vl = fmt["visualLink"]
        props = {}
        if "show" in vl: props["show"] = _pbi_lit(vl["show"])
        if "type" in vl: props["type"] = _pbi_lit(vl["type"])
        if "tooltip" in vl: props["tooltip"] = _pbi_lit(vl["tooltip"])
        if "showDefaultTooltip" in vl: props["showDefaultTooltip"] = _pbi_lit(vl["showDefaultTooltip"])
        if "navigationSection" in vl: props["navigationSection"] = _pbi_lit(vl["navigationSection"])
        if "bookmark" in vl: props["bookmark"] = _pbi_lit(vl["bookmark"])
        if "webUrl" in vl: props["webUrl"] = _pbi_lit(vl["webUrl"])
        _add_vc("visualLink", props)

    # --- visualHeaderTooltip (vcObjects) ---
    if "visualHeaderTooltip" in fmt:
        vht = fmt["visualHeaderTooltip"]
        props = {}
        if "text" in vht: props["text"] = _pbi_lit(vht["text"])
        if "type" in vht: props["type"] = _pbi_lit(vht["type"])
        if "bold" in vht: props["bold"] = _pbi_lit(vht["bold"])
        if "fontSize" in vht: props["fontSize"] = _pbi_lit(float(vht["fontSize"]))
        if "fontFamily" in vht: props["fontFamily"] = _pbi_lit(vht["fontFamily"])
        if "transparency" in vht: props["transparency"] = _pbi_lit(float(vht["transparency"]))
        if "background" in vht: props["themedBackground"] = _solid_color(vht["background"])
        if "titleFontColor" in vht: props["themedTitleFontColor"] = _solid_color(vht["titleFontColor"])
        _add_vc("visualHeaderTooltip", props)

    return {"_objects": objects, "_vcObjects": vc_objects}


# ============================= MCP TOOLS ===================================

# ---- Section 3: File Management ----

@mcp.tool()
def pbix_open(file_path: str, alias: str = "") -> str:
    """Open a PBIX or PBIT file for editing.

    Args:
        file_path: Full path to the .pbix or .pbit file
        alias: Short name to reference this file (auto-generated if empty)
    """
    file_path = os.path.abspath(file_path)
    if not os.path.exists(file_path):
        raise InvalidPBIXError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in (".pbix", ".pbit"):
        raise InvalidPBIXError(f"Expected .pbix or .pbit file, got '{ext}'")

    if not alias:
        alias = Path(file_path).stem

    if alias in _open_files:
        raise FileAlreadyOpenError(f"Alias '{alias}' is already in use. Close it first or choose a different alias.")

    # Create work directory. The timestamp alone is only second-granular, so
    # two same-alias opens within one second -- e.g. parallel PROCESSES each
    # opening a file under the same alias -- landed in the SAME directory and
    # silently overwrote each other's extracted model: one process then read
    # the other's tables, and the concurrent extract/validate race surfaced as
    # a bogus "path traversal" refusal. The pid+uuid suffix makes the
    # directory unique per open.
    work_dir = os.path.join(
        tempfile.gettempdir(),
        f"pbix_mcp_{alias}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        f"_{os.getpid()}_{uuid.uuid4().hex[:8]}",
    )
    os.makedirs(work_dir, exist_ok=True)
    _work_dirs.add(work_dir)
    _scavenge_stale_work_dirs()

    try:
        logger.info("Opening %s as '%s'", file_path, alias)
        _extract_pbix(file_path, work_dir)
        logger.debug("Extracted to %s", work_dir)
    except PBIXMCPError:
        raise
    except Exception as e:
        logger.error("Failed to extract %s: %s", file_path, e)
        shutil.rmtree(work_dir, ignore_errors=True)
        _work_dirs.discard(work_dir)
        raise InvalidPBIXError(f"Failed to extract: {e}")

    # Detect DirectQuery / composite models by checking for connections in DataModel
    _dq_flag = False
    dm_path = os.path.join(work_dir, "DataModel")
    if os.path.exists(dm_path):
        try:
            from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
            from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
            dm_bytes = open(dm_path, "rb").read()
            abf = decompress_datamodel(dm_bytes)
            db_bytes = read_metadata_sqlite(abf)
            import sqlite3
            tmp_db = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
            tmp_db.write(db_bytes)
            tmp_db.close()
            conn = sqlite3.connect(tmp_db.name)
            # Check for DirectQuery partitions (Mode=1 is DirectQuery, Mode=0 is Import)
            # Note: Type=4 for both Import and DirectQuery; Mode distinguishes them
            dq_partitions = conn.execute(
                "SELECT COUNT(*) FROM [Partition] WHERE Mode = 1"
            ).fetchone()[0]
            conn.close()
            os.unlink(tmp_db.name)
            if dq_partitions > 0:
                _dq_flag = True
                logger.warning(
                    "DirectQuery detected: %d DirectQuery partition(s). "
                    "Data operations (table reads, DAX evaluation) will not work. "
                    "Layout, measures, and metadata operations are still available.",
                    dq_partitions,
                )
        except Exception:
            pass  # If detection fails, continue — the file might still be usable

    _open_files[alias] = {
        "path": file_path,
        "work_dir": work_dir,
        "is_pbit": ext == ".pbit",
        "modified": False,
        "is_directquery": _dq_flag,
    }

    # Inventory
    components = []
    for root, dirs, files in os.walk(work_dir):
        for f in files:
            rel = os.path.relpath(os.path.join(root, f), work_dir)
            size = os.path.getsize(os.path.join(root, f))
            components.append(f"  {rel} ({size:,} bytes)")

    return ToolResponse.ok(
        f"Opened '{file_path}' as '{alias}'\n"
        f"Type: {'PBIT template' if ext == '.pbit' else 'PBIX report'}"
        f"{' ⚠️ DirectQuery detected — data operations unavailable, layout/measures/metadata OK' if _dq_flag else ''}\n"
        f"Components:\n" + "\n".join(sorted(components))
    ).to_text()


@mcp.tool()
def pbix_save(alias: str, output_path: str = "", overwrite: bool = False, backup: bool = True,
              strip_sensitivity_label: bool = False) -> str:
    """Save/repack the modified PBIX/PBIT file.

    Creates an automatic .bak backup before overwriting (unless backup=False).
    Set overwrite=False to refuse overwriting an existing file.

    Args:
        alias: The alias of the open file
        output_path: Where to save. Empty = overwrite original.
        overwrite: If False (default), refuse to overwrite an existing file
        backup: If True (default), create a .bak backup before overwriting
        strip_sensitivity_label: If True, remove MSIP sensitivity labels from the file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        # A PBIP-opened session saves back into its project folder unless an
        # explicit .pbix/.pbit output_path asks for a converted copy.
        if info.get("pbip_dir") and not output_path:
            summary = _save_pbip(info)
            info["modified"] = False
            _dax_cache.pop(alias, None)
            return ToolResponse.ok(summary).to_text()

        target = output_path or info["path"]
        target = os.path.abspath(target)
        logger.info("Saving '%s' to %s (overwrite=%s, backup=%s)", alias, target, overwrite, backup)

        # Safety: refuse overwrite if explicitly disabled
        if not overwrite and os.path.exists(target) and target != info["path"]:
            raise UnsafeWriteError(f"'{target}' already exists and overwrite=False. Use overwrite=True or choose a different path.")

        # If overwriting original, create backup
        if backup and target == info["path"] and os.path.exists(target):
            backup_path = target + ".bak"
            shutil.copy2(target, backup_path)

        _repack_pbix(work_dir, target, strip_sensitivity_label=strip_sensitivity_label)
        # Only mark the session clean when we wrote back to the ORIGINAL file.
        # Exporting a copy elsewhere must leave it dirty, otherwise a later
        # pbix_close (without force) would silently discard the work-dir edits
        # and the original on disk would never receive them.
        if target == info["path"]:
            info["modified"] = False
        # Clear DAX cache since data may have changed
        _dax_cache.pop(alias, None)
        size = os.path.getsize(target)
        return ToolResponse.ok(f"Saved '{alias}' to {target} ({size:,} bytes)").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise SessionError(f"Save failed: {e}")


@mcp.tool()
def pbix_close(alias: str, force: bool = False) -> str:
    """Close an open file and clean up temporary files.

    Refuses to close files with unsaved changes unless force=True.

    Args:
        alias: The alias of the open file
        force: If False (default), refuse to close files with unsaved changes
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        if info.get("modified") and not force:
            raise UnsafeWriteError(
                f"'{alias}' has unsaved changes. Use pbix_save first, or pbix_close with force=True to discard changes."
            )

        shutil.rmtree(work_dir, ignore_errors=True)
        _work_dirs.discard(work_dir)
        logger.info("Closed '%s'", alias)
        del _open_files[alias]
        # Clear DAX cache to avoid stale data on reopen
        _dax_cache.pop(alias, None)
        return ToolResponse.ok(f"Closed '{alias}'.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise SessionError(f"Close failed: {e}")


@mcp.tool()
def pbix_report_format(alias: str) -> str:
    """Report which layout format the open file uses, and whether it is writable.

    Two formats exist in the wild:

    * **classic** — a single ``Report/Layout`` document. Fully readable AND
      writable by every layout tool.
    * **PBIR** — a ``Report/definition/`` tree of per-page / per-visual JSON.
      This is what every report authored in the Power BI SERVICE downloads as.
      pbix-mcp reads AND writes it: the tree is converted to the classic shape
      for reading, and an edit is patched back onto the original page/visual
      files, so fields this converter doesn't model survive untouched and no
      classic ``Report/Layout`` is ever planted alongside the tree.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        is_pbir = _is_pbir(work_dir)
        layout = _get_layout(work_dir)
        sections = (layout or {}).get("sections", [])
        pages = len(sections)
        visuals = sum(len(s.get("visualContainers", [])) for s in sections)
        fmt = "PBIR" if is_pbir else ("classic" if layout else "none")
        writable = "yes" if layout else "no"
        return ToolResponse.ok(
            f"Report format: {fmt}\n"
            f"  Pages: {pages}\n"
            f"  Visuals: {visuals}\n"
            f"  Layout readable: {'yes' if layout else 'no'}\n"
            f"  Layout writable: {writable}",
            data={
                "format": fmt,
                "is_pbir": is_pbir,
                "readable": layout is not None,
                "writable": bool(layout),
                "pages": pages,
                "visuals": visuals,
            },
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_list_open() -> str:
    """List all currently open PBIX/PBIT files."""
    if not _open_files:
        return ToolResponse.ok("No files currently open.").to_text()
    lines = []
    for alias, info in _open_files.items():
        status = "modified" if info.get("modified") else "clean"
        ftype = "PBIT" if info.get("is_pbit") else "PBIX"
        lines.append(f"  {alias}: {info['path']} [{ftype}, {status}]")
    return ToolResponse.ok("Open files:\n" + "\n".join(lines)).to_text()


# ---- Section 4: Report Layout tools ----

@mcp.tool()
def pbix_get_pages(alias: str) -> str:
    """List all pages in the report with visual counts.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found in this file")

        sections = layout.get("sections", [])
        lines = [f"Report has {len(sections)} page(s):\n"]
        for i, sec in enumerate(sections):
            name = sec.get("displayName", f"Page {i}")
            vis_count = len(sec.get("visualContainers", []))
            width = sec.get("width", "?")
            height = sec.get("height", "?")
            hidden = " [HIDDEN]" if sec.get("config", "").find('"visibility":1') >= 0 else ""
            lines.append(f"  [{i}] {name} — {vis_count} visuals, {width}x{height}{hidden}")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_page_visuals(alias: str, page_index: int = 0) -> str:
    """List all visuals on a specific page.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range (0-{len(sections)-1})")

        page = sections[page_index]
        page_name = page.get("displayName", f"Page {page_index}")
        containers = page.get("visualContainers", [])

        lines = [f"Page '{page_name}' has {len(containers)} visual(s):\n"]
        for i, vc in enumerate(containers):
            config = _parse_visual_config(vc)
            vtype = _get_visual_type(config)
            vname = _get_visual_name(config)
            x = vc.get("x", 0)
            y = vc.get("y", 0)
            w = vc.get("width", 0)
            h = vc.get("height", 0)
            lines.append(f"  [{i}] {vtype} (name={vname}) at ({x},{y}) size {w}x{h}")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_visual_detail(alias: str, page_index: int, visual_index: int) -> str:
    """Get the full configuration JSON for a specific visual.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        vc = containers[visual_index]
        config = _parse_visual_config(vc)
        result = {
            "x": vc.get("x", 0),
            "y": vc.get("y", 0),
            "width": vc.get("width", 0),
            "height": vc.get("height", 0),
            "z": vc.get("z", 0),
            "config": config,
        }
        # Include query and dataTransforms if present
        for key in ("query", "dataTransforms", "filters"):
            raw = vc.get(key)
            if raw:
                if isinstance(raw, str):
                    try:
                        result[key] = json.loads(raw)
                    except json.JSONDecodeError:
                        result[key] = raw
                else:
                    result[key] = raw

        return ToolResponse.ok(json.dumps(result, indent=2, ensure_ascii=False)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_visual_property(
    alias: str, page_index: int, visual_index: int,
    property_path: str, value: str
) -> str:
    """Set a property on a visual using a dot-path (e.g. 'singleVisual.title.text').

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
        property_path: Dot-separated path into the config JSON
        value: New value (JSON-encoded string, e.g. '"hello"' or '42' or 'true')
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        vc = containers[visual_index]
        config = _parse_visual_config(vc)

        # Parse the value as JSON
        try:
            parsed_value = json.loads(value)
        except json.JSONDecodeError:
            parsed_value = value  # treat as raw string

        _set_value_by_dot_path(config, property_path, parsed_value)

        # Write config back
        vc["config"] = json.dumps(config, ensure_ascii=False)
        rebound = _recompile_classic_binding(info, vc, config)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        note = " (data binding recompiled)" if rebound else ""
        return ToolResponse.ok(
            f"Set {property_path} = {value} on page {page_index}, "
            f"visual {visual_index}{note}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_update_visual_json(
    alias: str, page_index: int, visual_index: int, config_json: str
) -> str:
    """Replace the entire config JSON for a visual.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
        config_json: Complete config JSON string to replace
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        # Validate JSON
        try:
            new_config = json.loads(config_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid JSON: {e}")

        containers[visual_index]["config"] = json.dumps(new_config, ensure_ascii=False)
        rebound = _recompile_classic_binding(
            info, containers[visual_index], new_config)
        _set_layout(info["work_dir"], layout)
        _warn_unbound_field_parameters(info, new_config)
        info["modified"] = True
        note = " (data binding recompiled)" if rebound else ""
        return ToolResponse.ok(
            f"Updated visual config on page {page_index}, "
            f"visual {visual_index}{note}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


def _recompile_classic_binding(info: dict, vc: dict, config: dict) -> bool:
    """Re-derive query + dataTransforms after a CLASSIC layout config edit.

    Power BI Desktop does not re-derive a visual's data binding from
    ``config.singleVisual`` when it loads a classic ``Report/Layout``; it reads
    the sibling ``query`` / ``dataTransforms`` blobs. Editing projections or
    prototypeQuery through pbix_set_visual_property / pbix_update_visual_json
    therefore changed the config and left the COMPILED binding pointing at the
    old field, so Desktop kept rendering the previous column with no error
    anywhere. The creation paths already compile; these two did not.

    Returns True when a binding was written. Best-effort by design: a config
    edit must never fail because the binding could not be compiled.
    """
    sv = (config or {}).get("singleVisual") or {}
    if not (sv.get("prototypeQuery") and sv.get("projections")):
        return False            # textbox / shape / image / button: no binding
    try:
        from pbix_mcp.report_binding import compile_visual_binding
        q, dt = compile_visual_binding(sv, _report_type_resolver(info))
        if q is None:
            return False
        vc["query"] = json.dumps(q, ensure_ascii=False)
        vc["dataTransforms"] = json.dumps(dt, ensure_ascii=False)
        vc.setdefault("filters", "[]")
        # compile_visual_binding rewrites bare value-role columns to implicit
        # Aggregations inside `sv` itself, so the config has to be re-serialized
        # from the MUTATED object or the prototype and the compiled query
        # disagree.
        vc["config"] = json.dumps(config, ensure_ascii=False)
        return True
    except Exception:
        return False


@mcp.tool()
def pbix_add_page(alias: str, display_name: str, width: int = 1280, height: int = 720) -> str:
    """Add a new blank page to the report.

    Args:
        alias: The alias of the open file
        display_name: Name for the new page
        width: Page width in pixels (default 1280)
        height: Page height in pixels (default 720)
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        import uuid
        new_section = {
            "displayName": display_name,
            # 1 = FitToPage, what Desktop stamps on a new page (0 is the
            # deprecated dynamic mode).
            "displayOption": 1,
            "name": str(uuid.uuid4()).replace("-", ""),
            "width": width,
            "height": height,
            "visualContainers": [],
            "config": json.dumps({"visibility": 0}),
            "filters": "[]",
            "ordinal": len(layout.get("sections", [])),
        }

        layout.setdefault("sections", []).append(new_section)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        idx = len(layout["sections"]) - 1
        return ToolResponse.ok(f"Added page '{display_name}' at index {idx} ({width}x{height})").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_remove_page(alias: str, page_index: int) -> str:
    """Remove a page from the report.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index to remove
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        removed = sections.pop(page_index)
        name = removed.get("displayName", f"Page {page_index}")
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(f"Removed page '{name}' (was index {page_index}). {len(sections)} pages remain.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


def _resolve_page(sections: list, page: str) -> int:
    """Resolve a page reference (index or displayName) to an index."""
    try:
        idx = int(page)
    except (TypeError, ValueError):
        for i, sec in enumerate(sections):
            if (sec.get("displayName") or "").lower() == str(page).lower():
                return i
        raise LayoutParseError(
            f"Page '{page}' not found. Available: "
            f"{[s.get('displayName') for s in sections]}")
    if idx < 0 or idx >= len(sections):
        raise LayoutParseError(
            f"Page index {idx} out of range (0..{len(sections) - 1})")
    return idx


@mcp.tool()
def pbix_rename_page(alias: str, page: str, new_name: str) -> str:
    """Rename a report page.

    Args:
        alias: The alias of the open file
        page: Page index (e.g. "0") or current displayName
        new_name: New display name for the page
    """
    logger.info(f"pbix_rename_page: {alias} {page!r} -> {new_name!r}")
    try:
        if not new_name or not new_name.strip():
            raise LayoutParseError("new_name must not be empty")
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        idx = _resolve_page(sections, page)
        old = sections[idx].get("displayName")
        # The internal `name` is an identity other objects reference (bookmarks,
        # drillthrough, page navigation), so only displayName changes.
        sections[idx]["displayName"] = new_name

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Renamed page {idx} from '{old}' to '{new_name}'").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_reorder_pages(alias: str, page_order: str) -> str:
    """Reorder the report's pages.

    Args:
        alias: The alias of the open file
        page_order: Comma-separated page references (indices or displayNames)
                    in the desired order, e.g. "Summary,Detail,0". Pages left
                    out keep their relative order after the ones listed.
    """
    logger.info(f"pbix_reorder_pages: {alias} {page_order!r}")
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        refs = [p.strip() for p in page_order.split(",") if p.strip()]
        if not refs:
            raise LayoutParseError("page_order must name at least one page")

        wanted: list = []
        for ref in refs:
            idx = _resolve_page(sections, ref)
            if idx in wanted:
                raise LayoutParseError(
                    f"Page '{ref}' listed more than once in page_order")
            wanted.append(idx)

        rest = [i for i in range(len(sections)) if i not in wanted]
        layout["sections"] = [sections[i] for i in wanted + rest]
        # Classic Report/Layout carries an explicit ordinal per section; PBIR
        # takes the order from the list itself.
        for i, sec in enumerate(layout["sections"]):
            if "ordinal" in sec:
                sec["ordinal"] = i

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        names = [s.get("displayName") for s in layout["sections"]]
        return ToolResponse.ok(f"Page order is now: {names}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_page_visibility(alias: str, page: str, hidden: bool) -> str:
    """Show or hide a report page (hidden pages are not shown in view mode).

    Args:
        alias: The alias of the open file
        page: Page index or displayName
        hidden: True to hide the page, False to show it
    """
    logger.info(f"pbix_set_page_visibility: {alias} {page!r} hidden={hidden}")
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        idx = _resolve_page(sections, page)
        sec = sections[idx]

        # Both formats are driven through the classic shape: the section
        # `config` JSON with 0 = AlwaysVisible, 1 = HiddenInViewMode. The PBIR
        # writer converts it to the enum name on the way out.
        raw = sec.get("config", "{}")
        try:
            cfg = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except json.JSONDecodeError:
            cfg = {}
        cfg["visibility"] = 1 if hidden else 0
        sec["config"] = json.dumps(cfg, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        state = "hidden" if hidden else "visible"
        return ToolResponse.ok(
            f"Page '{sec.get('displayName')}' is now {state}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_duplicate_page(alias: str, page: str, new_name: str = "") -> str:
    """Duplicate a page and all of its visuals.

    Args:
        alias: The alias of the open file
        page: Page index or displayName to copy
        new_name: Display name for the copy (default: "<name> (copy)")
    """
    logger.info(f"pbix_duplicate_page: {alias} {page!r}")
    try:
        import uuid

        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        idx = _resolve_page(sections, page)
        src = sections[idx]
        copy_sec = copy.deepcopy(src)

        # Every identity in the copy must be fresh, or the two pages collide:
        # page name, PBIR folder id, and every visual name.
        copy_sec["name"] = uuid.uuid4().hex[:20]
        copy_sec.pop("__pbir_page__", None)
        copy_sec.pop("isActive", None)
        copy_sec["displayName"] = (
            new_name.strip() or f"{src.get('displayName', 'Page')} (copy)")

        renamed = 0
        for vc in copy_sec.get("visualContainers", []) or []:
            vc.pop("__pbir_visual__", None)
            raw = vc.get("config", "{}")
            try:
                cfg = json.loads(raw) if isinstance(raw, str) else (raw or {})
            except json.JSONDecodeError:
                continue
            if cfg.get("name"):
                cfg["name"] = uuid.uuid4().hex[:20]
                renamed += 1
            vc["config"] = json.dumps(cfg, ensure_ascii=False)

        sections.insert(idx + 1, copy_sec)
        for i, sec in enumerate(sections):
            if "ordinal" in sec:
                sec["ordinal"] = i

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Duplicated page '{src.get('displayName')}' as "
            f"'{copy_sec['displayName']}' at index {idx + 1} "
            f"({renamed} visual(s) copied with new identities)").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_move_visual(alias: str, page: str, visual_index: int,
                     x: float = -1, y: float = -1,
                     width: float = -1, height: float = -1,
                     z: float = -1) -> str:
    """Move and/or resize an existing visual.

    Position lives on the visual CONTAINER, not in its config JSON, so
    pbix_set_visual_property cannot reach it.

    Args:
        alias: The alias of the open file
        page: Page index or displayName
        visual_index: Zero-based visual index on that page
        x: New left position in px (-1 = leave unchanged)
        y: New top position in px (-1 = leave unchanged)
        width: New width in px (-1 = leave unchanged)
        height: New height in px (-1 = leave unchanged)
        z: New z-order (-1 = leave unchanged)
    """
    logger.info(f"pbix_move_visual: {alias} {page!r}[{visual_index}]")
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        idx = _resolve_page(sections, page)
        containers = sections[idx].get("visualContainers", []) or []
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(
                f"Visual index {visual_index} out of range on page "
                f"'{sections[idx].get('displayName')}' "
                f"(0..{len(containers) - 1})")

        vc = containers[visual_index]
        changes = {}
        for key, val in (("x", x), ("y", y), ("width", width),
                         ("height", height), ("z", z)):
            if val is None or val < 0:
                continue
            if key in ("width", "height") and val == 0:
                raise LayoutParseError(f"{key} must be greater than 0")
            vc[key] = val
            changes[key] = val
        if not changes:
            raise LayoutParseError(
                "Nothing to change — pass at least one of x/y/width/height/z")

        # Classic containers repeat the geometry inside config.layouts; keep the
        # two copies consistent or Desktop renders the stale one.
        raw = vc.get("config", "{}")
        try:
            cfg = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except json.JSONDecodeError:
            cfg = {}
        for lay in cfg.get("layouts", []) or []:
            pos = lay.get("position")
            if isinstance(pos, dict):
                pos.update(changes)
        if cfg:
            vc["config"] = json.dumps(cfg, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Moved visual {visual_index} on page "
            f"'{sections[idx].get('displayName')}': {changes}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_duplicate_visual(alias: str, page: str, visual_index: int,
                          target_page: str = "",
                          offset_x: float = 20, offset_y: float = 20) -> str:
    """Copy a visual, onto the same page or another one.

    Args:
        alias: The alias of the open file
        page: Source page index or displayName
        visual_index: Zero-based visual index on the source page
        target_page: Destination page (default: same page)
        offset_x: Horizontal offset for the copy, px (same-page copies only)
        offset_y: Vertical offset for the copy, px (same-page copies only)
    """
    logger.info(f"pbix_duplicate_visual: {alias} {page!r}[{visual_index}]")
    try:
        import uuid

        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        src_idx = _resolve_page(sections, page)
        containers = sections[src_idx].get("visualContainers", []) or []
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(
                f"Visual index {visual_index} out of range "
                f"(0..{len(containers) - 1})")

        dst_idx = (_resolve_page(sections, target_page)
                   if target_page else src_idx)
        vc = copy.deepcopy(containers[visual_index])
        vc.pop("__pbir_visual__", None)

        if dst_idx == src_idx:
            # A copy landing exactly on top of the original looks like nothing
            # happened, so nudge it.
            vc["x"] = (vc.get("x") or 0) + offset_x
            vc["y"] = (vc.get("y") or 0) + offset_y

        raw = vc.get("config", "{}")
        try:
            cfg = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except json.JSONDecodeError:
            cfg = {}
        new_name = uuid.uuid4().hex[:20]
        cfg["name"] = new_name
        if dst_idx == src_idx:
            for lay in cfg.get("layouts", []) or []:
                pos = lay.get("position")
                if isinstance(pos, dict):
                    pos["x"] = (pos.get("x") or 0) + offset_x
                    pos["y"] = (pos.get("y") or 0) + offset_y
        vc["config"] = json.dumps(cfg, ensure_ascii=False)

        sections[dst_idx].setdefault("visualContainers", []).append(vc)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Copied visual {visual_index} from page "
            f"'{sections[src_idx].get('displayName')}' to "
            f"'{sections[dst_idx].get('displayName')}' as '{new_name}'"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


def _read_ndjson_rows(path: str) -> list[dict]:
    """Stream an NDJSON file (one JSON object per line) into a row list.

    The STREAMING row source (issue #46): a caller converting a large table
    writes batches to the file and frees them as it goes, then hands over
    the path — instead of holding source rows + a row-dict list + the whole
    serialized JSON text simultaneously (measured at 3x the data size).
    Blank lines are skipped; a malformed line errors with its line number.
    """
    rows: list[dict] = []
    with open(path, "r", encoding="utf-8-sig") as f:
        for ln, line in enumerate(f, 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = json.loads(line)
            except json.JSONDecodeError as e:
                raise ValueError(
                    f"{path}:{ln}: not valid JSON ({e.msg}). NDJSON needs "
                    f"exactly one JSON object per line.") from e
            if not isinstance(obj, dict):
                raise ValueError(
                    f"{path}:{ln}: expected a JSON object per line, got "
                    f"{type(obj).__name__}.")
            rows.append(obj)
    return rows


@mcp.tool()
def pbix_create(
    file_path: str,
    alias: str = "",
    tables_json: str = "",
    measures_json: str = "",
    relationships_json: str = "",
) -> str:
    """Create a new PBIX file and open it for editing.

    Builds a valid PBIX entirely from scratch — no templates or skeletons.
    Every layer is generated from code: PBIX ZIP shell, ABF binary container,
    db.xml, metadata SQLite, VertiPaq column data, and report layout.

    Args:
        file_path: Where to save the new file (e.g. "my_report.pbix")
        alias: Alias for the opened file (auto-generated if empty)
        tables_json: Optional JSON array of tables with columns and rows, e.g.
            '[{"name": "Sales", "columns": [{"name": "Amount", "data_type": "Double"},
              {"name": "Product", "data_type": "String"}],
              "rows": [{"Amount": 100.0, "Product": "Widget"}]}]'
            Supported data_type values: String, Int64, Double, DateTime, Decimal, Boolean
            Optional per-table fields:
            - "rows_path": path to an NDJSON file (one JSON row object per
              line) used INSTEAD of inline "rows" — the streaming row source
              for large tables: write batches to the file and free them,
              then pass the path, instead of serializing the whole dataset
              into this one string (which costs ~3x the data size in RAM).
              Mutually exclusive with "rows".
            - "source_csv": "/path/to/data.csv" — M expression references CSV for Refresh
            - "source_db": {"type": "sqlserver", "server": "localhost", "database": "mydb",
              "table": "orders"} — M expression references database for Refresh/DirectQuery.
              Supported types: "sqlserver", "mysql", "sqlite", "postgresql",
              "mariadb" (MySQL DirectQuery via MariaDB adapter),
              "excel" (needs path+sheet), "json"/"web"/"api" (needs url),
              "azuresql"/"azure" (same as sqlserver for Azure SQL)
            - "mode": "directquery" — live database queries (default: "import").
              DirectQuery requires source_db and a running database server.
        measures_json: Optional JSON array of measures, e.g.
            '[{"table": "Sales", "name": "Total", "expression": "SUM(Sales[Amount])",
              "format_string": "$#,0.00"}]'
            Optional per-measure fields: "description", "format_string"
            (display format code, e.g. "$#,0.00", "0.0%", "#,0")
        relationships_json: Optional JSON array of relationships, e.g.
            '[{"from_table": "Sales", "from_column": "ProductID",
              "to_table": "Products", "to_column": "ProductID"}]'
    """
    try:
        from pbix_mcp.builder import PBIXBuilder

        builder = PBIXBuilder()

        if tables_json:
            for tdef in json.loads(tables_json):
                rows = tdef.get("rows")
                rows_path = tdef.get("rows_path")
                if rows_path:
                    if rows:
                        return ToolResponse.error(
                            f"Table '{tdef.get('name')}': 'rows' and "
                            f"'rows_path' are mutually exclusive — pass one "
                            f"row source.", "INVALID_INPUT").to_text()
                    if not os.path.exists(rows_path):
                        return ToolResponse.error(
                            f"Table '{tdef.get('name')}': rows_path not "
                            f"found: {rows_path}", "INVALID_INPUT").to_text()
                    rows = _read_ndjson_rows(rows_path)
                builder.add_table(
                    tdef["name"],
                    tdef.get("columns", []),
                    rows=rows,
                    hidden=tdef.get("hidden", False),
                    source_csv=tdef.get("source_csv"),
                    source_db=tdef.get("source_db"),
                    mode=tdef.get("mode", "import"),
                )

        if measures_json:
            for mdef in json.loads(measures_json):
                builder.add_measure(
                    mdef["table"],
                    mdef["name"],
                    mdef["expression"],
                    mdef.get("description", ""),
                    format_string=mdef.get("format_string"),
                )

        if relationships_json:
            for rdef in json.loads(relationships_json):
                builder.add_relationship(
                    rdef["from_table"],
                    rdef["from_column"],
                    rdef["to_table"],
                    rdef["to_column"],
                )

        builder.add_page("Page 1")

        abs_path = builder.save(file_path)
        size = os.path.getsize(abs_path)

        # Auto-open the created file
        result = pbix_open(abs_path, alias)
        return ToolResponse.ok(f"Created '{abs_path}' ({size:,} bytes) and opened it.\n{result}").to_text()

    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# Built-in visual types with no data roles at all. Deliberately narrow: an
# unrecognized type (any custom visual) is assumed to HAVE wells, so only the
# provable cases are refused.
_NO_DATA_ROLE_VISUALS = frozenset({
    "textbox", "image", "shape", "basicShape", "actionButton",
})


def _field_parameter_tables(info: dict) -> set:
    """Names of tables that are FIELD PARAMETERS (ParameterMetadata
    ExtendedProperty on a column). Best-effort: empty set on any failure."""
    out: set = set()
    try:
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        dm_path = os.path.join(info["work_dir"], "DataModel")
        with open(dm_path, "rb") as f:
            meta_bytes = read_metadata_sqlite(decompress_datamodel(f.read()))
        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        os.write(fd, meta_bytes)
        os.close(fd)
        try:
            conn = sqlite3.connect(tmp_path)
            for row in conn.execute(
                "SELECT DISTINCT t.Name FROM ExtendedProperty ep "
                "JOIN [Column] c ON ep.ObjectID = c.ID AND ep.ObjectType = 4 "
                "JOIN [Table] t ON c.TableID = t.ID "
                "WHERE ep.Name = 'ParameterMetadata'"):
                out.add(row[0])
            conn.close()
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except Exception:  # noqa: BLE001 - detection is advisory only
        pass
    return out


def _warn_unbound_field_parameters(info: dict, config: dict) -> None:
    """Warn when a projection targets a field-parameter display column but the
    visual carries no queryFieldParametersByRole.

    Desktop does not error on that shape -- it silently treats the display
    column as ordinary text and degrades the well to an implicit Count
    (OpenBI findings #19: six equal "Count of Metric" bars instead of
    field-swapping). A warning is the difference between a five-minute fix
    and a silently wrong chart."""
    try:
        sv = config.get("singleVisual") or {}
        projections = sv.get("projections") or {}
        if not projections:
            return
        qfp = sv.get("queryFieldParametersByRole") or {}
        param_tables = _field_parameter_tables(info)
        if not param_tables:
            return
        for role, items in projections.items():
            if role in qfp:
                continue
            for it in items or []:
                ref = str((it or {}).get("queryRef") or "")
                # The binding compiler may already have wrapped the bare
                # column in an implicit aggregation -- the very degradation
                # being warned about -- so the ref can arrive as
                # "CountNonNull(Metric.Metric)". Unwrap before matching.
                m = re.match(r"^[A-Za-z]+\(([^)]+)\)$", ref)
                if m:
                    ref = m.group(1)
                table = ref.split(".", 1)[0] if "." in ref else ""
                if table in param_tables:
                    _responses.add_pending_warning(
                        f"Projection '{ref}' ({role}) targets field-parameter "
                        f"table '{table}' without queryFieldParametersByRole "
                        f"-- Power BI will silently degrade it to an implicit "
                        f"Count. Use pbix_bind_field_parameter to bind it.")
    except Exception:  # noqa: BLE001 - advisory only
        pass


def _report_type_resolver(info: dict):
    """Build ``resolve_type(entity, prop, is_measure) -> data_type`` from the
    open model's metadata, for report-binding type codes (best-effort).

    Plain internal helper — NOT an MCP tool (must stay above pbix_add_visual's
    ``@mcp.tool()`` without displacing it)."""
    col_types: dict = {}
    try:
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        dm_path = os.path.join(info["work_dir"], "DataModel")
        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta_bytes = read_metadata_sqlite(abf)
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta_bytes)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            _AMO = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                    10: "Decimal", 11: "Boolean"}
            for r in conn.execute(
                # COALESCE: Type 4 (calculated-table column) keeps its name
                # in InferredName with ExplicitName NULL, so the key would be
                # (table, None) and the type lookup would silently miss.
                "SELECT t.Name tn, COALESCE(c.ExplicitName, c.InferredName) cn, "
                "       c.ExplicitDataType edt "
                "FROM [Column] c JOIN [Table] t ON t.ID = c.TableID "
                "WHERE c.Type IN (1, 2, 4)"
            ):
                col_types[(r["tn"], r["cn"])] = _AMO.get(r["edt"], "String")
            conn.close()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass
    except Exception:
        col_types = {}

    def resolve(entity, prop, is_measure):
        # Measures: Measure.DataType is unreliable (currency measures are stored
        # as Int64), so fall through to the compiler's numeric/decimal default
        # (259/1) — matching Desktop for the common case.
        if is_measure:
            return None
        return col_types.get((entity, prop))

    return resolve


@mcp.tool()
def pbix_add_visual(
    alias: str,
    page_index: int,
    visual_type: str,
    x: int = 0,
    y: int = 0,
    width: int = 300,
    height: int = 200,
    config_json: str = "",
    sort_by: str = "",
    sort_direction: str = "desc",
) -> str:
    """Add a new visual to a report page.

    Supports all Power BI visual types: card, table, clusteredBarChart,
    clusteredColumnChart, lineChart, pieChart, donutChart, shape (buttons),
    image, slicer, textbox, and any custom visual type.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_type: Visual type (e.g. "card", "clusteredBarChart", "shape", "image", "textbox")
        x: X position in pixels
        y: Y position in pixels
        width: Width in pixels
        height: Height in pixels
        config_json: Optional full config JSON to merge (for advanced properties)
        sort_by: Optional visual-level sort field — one of the visual's own
            fields, as a bare name ("Pipeline Value"), DAX-style reference
            ("[Pipeline Value]", "'Table'[Col]"), or queryRef ("Table.Field").
            Authors the Desktop-style prototypeQuery.OrderBy clause (and the
            same clause in the compiled query); without it the Power BI
            service falls back to category-ascending query order. Requires a
            data binding in config_json (prototypeQuery + projections).
        sort_direction: "asc"/"ascending" or "desc"/"descending" (default
            "desc", matching Desktop's usual value-descending chart default)
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        import uuid
        visual_name = str(uuid.uuid4()).replace("-", "")[:16]

        config: dict = {
            "name": visual_name,
            "singleVisual": {
                "visualType": visual_type,
            },
        }

        # Merge custom config if provided
        if config_json:
            try:
                custom = json.loads(config_json)
                if isinstance(custom, dict):
                    for key, val in custom.items():
                        if key == "singleVisual" and isinstance(val, dict):
                            config["singleVisual"].update(val)
                        else:
                            config[key] = val
            except json.JSONDecodeError:
                raise LayoutParseError("Invalid config_json")

        # Textbox visuals: normalize paragraphs structure for Fabric compatibility
        if visual_type == "textbox":
            sv = config["singleVisual"]
            sv.setdefault("drillFilterOtherVisuals", True)
            config.setdefault("layouts", [{"id": 0, "position": {
                "x": float(x), "y": float(y), "z": 0,
                "width": float(width), "height": float(height),
            }}])

            # Normalize paragraphs: fix double-nesting, px→pt, strip unsupported keys
            gen = sv.get("objects", {}).get("general", [])
            for entry in gen:
                props = entry.get("properties", {})
                paras = props.get("paragraphs", [])
                # Fix double-nested {"paragraphs": [...]} → [...]
                if isinstance(paras, dict) and "paragraphs" in paras:
                    paras = paras["paragraphs"]
                    props["paragraphs"] = paras
                if isinstance(paras, list):
                    for para in paras:
                        # Remove horizontalTextAlignment (Fabric rejects it)
                        para.pop("horizontalTextAlignment", None)
                        for tr in para.get("textRuns", []):
                            ts = tr.get("textStyle", {})
                            # Convert px to pt for Fabric
                            fs = ts.get("fontSize", "")
                            if isinstance(fs, str) and fs.endswith("px"):
                                ts["fontSize"] = fs.replace("px", "pt")

        # Image visuals: embed the image file and set up ResourcePackageItem
        if visual_type == "image":
            sv = config["singleVisual"]
            sv.setdefault("drillFilterOtherVisuals", True)
            config.setdefault("layouts", [{"id": 0, "position": {
                "x": float(x), "y": float(y), "z": 0,
                "width": float(width), "height": float(height),
            }}])

            # Check if config has a local file path to embed
            img_url = (sv.get("objects", {}).get("general", [{}])[0]
                       .get("properties", {}).get("imageUrl", {}))
            src_path = img_url.get("sourcePath", "")  # custom field for local files
            if src_path:
                # NEVER persist the private key: it is not Power BI schema and
                # would leak the author's local filesystem path into the saved
                # report. A path that cannot be read fails LOUD rather than
                # silently shipping a visual with no image.
                img_url.pop("sourcePath", None)
                if not os.path.isfile(src_path):
                    raise LayoutParseError(
                        f"Image file not found: {src_path} (pbix_add_image is "
                        "the supported API and also accepts base64 bytes)")
            if src_path:
                # LEGACY private hook. pbix_add_image is the supported API —
                # it does this registration plus Desktop's full container
                # (howCreated / z / tabOrder / padding) and accepts bytes as
                # well as a path. Kept working for existing callers, and now
                # sharing the same registration helpers, so it inherits the
                # [Content_Types].xml `</Types>` fallback (the old
                # json-anchored replace silently no-opped on documents with no
                # json Default) and the no-clobber item naming.
                data, ext = _resolve_image_source(src_path, "")
                item_name = _register_resource(
                    info, layout, data, visual_name, ext, 100)
                sv.setdefault("objects", {})["general"] = \
                    _image_url_object(item_name)

        page = sections[page_index]
        # If this visual belongs to a singleVisualGroup, its container x/y are
        # stored RELATIVE to the group origin — pbix_get_visual_positions adds
        # the group origin back on read. Convert the caller's ABSOLUTE coords to
        # group-relative so a grouped child is not corrupted; only top-level
        # visuals are page-clamped.
        group_origin = None
        parent_group = config.get("parentGroupName")
        if parent_group:
            for _vc in page.get("visualContainers", []):
                _cfg = _parse_visual_config(_vc)
                if _cfg.get("name") == parent_group and _cfg.get("singleVisualGroup"):
                    group_origin = (float(_vc.get("x", 0)), float(_vc.get("y", 0)))
                    break

        if group_origin is not None:
            x = float(x) - group_origin[0]
            y = float(y) - group_origin[1]
        else:
            # Clamp position to stay within page bounds.
            page_w = page.get("width", 1280)
            page_h = page.get("height", 720)
            x = min(float(x), page_w - width)
            y = min(float(y), page_h - height)
            x = max(0.0, x)
            y = max(0.0, y)

        # Desktop stamps every container with a 1000-step z and
        # tabOrder = z + 1000 (ground truth: add_image / Desktop-authored
        # corpus). Ledger issues-3: add_visual used to write z=0 and no
        # tabOrder at all.
        z_new = _next_layer_z(page)
        tab_order = z_new + 1000

        # Desktop-authored visuals of EVERY type carry config.layouts and
        # default drillFilterOtherVisuals: true (ledger issues-8 field
        # audit against the GeoSales tableEx); both were previously only
        # written for image visuals.
        config.setdefault("layouts", [{"id": 0, "position": {
            "x": float(x), "y": float(y), "z": z_new,
            "width": float(width), "height": float(height),
            "tabOrder": tab_order,
        }}])
        if config.get("singleVisual", {}).get("prototypeQuery"):
            config["singleVisual"].setdefault("drillFilterOtherVisuals", True)

        for lay in config["layouts"]:
            pos = lay.get("position", {})
            pos["x"] = x
            pos["y"] = y
            pos["z"] = z_new
            pos["tabOrder"] = tab_order

        # Opt-in visual-level sort: author prototypeQuery.OrderBy BEFORE the
        # config is serialized and the binding compiled (the compiler deep-
        # copies the prototype, carrying the clause into the compiled query).
        # Unlike the best-effort binding compile below, a bad sort fails LOUD —
        # silently dropping a requested sort would be another silent-wrong.
        if sort_by:
            sv_sort = config.get("singleVisual", {})
            if not (sv_sort.get("prototypeQuery") or {}).get("Select"):
                raise LayoutParseError(
                    "sort_by requires a data binding — provide config_json with "
                    "singleVisual.prototypeQuery + projections.")
            from pbix_mcp.report_binding import attach_order_by
            try:
                attach_order_by(sv_sort, sort_by, sort_direction)
            except ValueError as e:
                raise LayoutParseError(str(e))

        container = {
            "x": x,
            "y": y,
            "z": z_new,
            "width": float(width),
            "height": float(height),
            "tabOrder": tab_order,
            "config": json.dumps(config, ensure_ascii=False),
        }
        if visual_type == "image":
            container["filters"] = "[]"

        # Compile the data binding (query + dataTransforms) Power BI Desktop's
        # report loader requires on data visuals. Without it a report carrying
        # report-level config / visual objects fails to load with "Failed to
        # load the report", even though the model opens fine. Non-data visuals
        # (textbox / shape / image / button) have no projections and are skipped.
        sv_final = config.get("singleVisual", {})
        if sv_final.get("prototypeQuery") and sv_final.get("projections"):
            try:
                from pbix_mcp.report_binding import compile_visual_binding
                q, dt = compile_visual_binding(sv_final, _report_type_resolver(info))
                if q is not None:
                    container["query"] = json.dumps(q, ensure_ascii=False)
                    container["dataTransforms"] = json.dumps(dt, ensure_ascii=False)
                    container.setdefault("filters", "[]")
                # compile_visual_binding rewrites bare value-role columns to
                # implicit Aggregations IN the prototypeQuery/projections (as
                # Desktop's field well does) — re-serialize the config so the
                # prototype Desktop re-derives the live query from is aggregated
                # too; the compiled query alone leaves the chart empty.
                container["config"] = json.dumps(config, ensure_ascii=False)
            except Exception:
                pass  # best-effort: never block visual creation on binding

        page.setdefault("visualContainers", []).append(container)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True

        _warn_unbound_field_parameters(info, config)
        idx = len(page["visualContainers"]) - 1
        page_name = page.get("displayName", f"Page {page_index}")
        return ToolResponse.ok(f"Added {visual_type} visual at ({x},{y}) {width}x{height} on '{page_name}' (index {idx})").to_text()

    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_visual_sort(
    alias: str,
    page_index: int,
    visual_index: int,
    sort_by: str = "",
    sort_direction: str = "desc",
) -> str:
    """Set (or clear) the visual-level sort on an existing data visual.

    Authors the Desktop-style ``prototypeQuery.OrderBy`` clause and recompiles
    the visual's query/dataTransforms binding so the compiled query carries the
    same clause. Without an OrderBy, the Power BI service falls back to
    category-ascending query order — Desktop's usual value-descending bar/
    column default comes from Desktop AUTHORING an OrderBy, not the renderer.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
        sort_by: Field to sort by — one of the visual's own fields, as a bare
            name ("Pipeline Value"), DAX-style reference ("[Pipeline Value]",
            "'Table'[Col]", "Table[Col]"), or queryRef ("Table.Field").
            Empty string clears any existing sort.
        sort_direction: "asc"/"ascending" or "desc"/"descending" (default "desc")
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")
        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        vc = containers[visual_index]
        config = _parse_visual_config(vc)
        sv = config.get("singleVisual", {})
        proto = sv.get("prototypeQuery") or {}
        if not proto.get("Select"):
            raise LayoutParseError(
                f"Visual {visual_index} ({sv.get('visualType', 'unknown')}) has "
                "no data binding (prototypeQuery) — nothing to sort.")

        from pbix_mcp.report_binding import SORT_ASCENDING, attach_order_by, compile_visual_binding
        if sort_by:
            try:
                matched_ref, dcode = attach_order_by(sv, sort_by, sort_direction)
            except ValueError as e:
                raise LayoutParseError(str(e))
            word = "ascending" if dcode == SORT_ASCENDING else "descending"
            msg = f"Sort set: {matched_ref} {word}."
        else:
            proto.pop("OrderBy", None)
            msg = "Sort cleared (service will use its default query order)."

        # Recompile so the compiled query/dataTransforms match the prototype
        # (compile also re-applies implicit value-role aggregations, mutating
        # the config — serialize it AFTER, like pbix_add_visual does).
        q, dt = compile_visual_binding(sv, _report_type_resolver(info))
        if q is not None:
            vc["query"] = json.dumps(q, ensure_ascii=False)
            vc["dataTransforms"] = json.dumps(dt, ensure_ascii=False)
            vc.setdefault("filters", "[]")
        vc["config"] = json.dumps(config, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(msg).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}",
                                  getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_bind_field_parameter(
    alias: str,
    page_index: int,
    visual_index: int,
    role: str,
    parameter_name: str,
    initial_field: str = "",
) -> str:
    """Bind a field parameter into a visual's field well (real field-swapping).

    Putting the parameter's display column straight into a projection does NOT
    work: Desktop silently treats it as ordinary text and degrades the well to
    an implicit Count -- plausible-looking bars, wrong semantics, no error.
    The working shape (diffed against a Desktop-authored binding) keeps the
    currently-RESOLVED field in the projection and expresses the parameter
    linkage separately. This tool authors all five pieces:

    1. ``projections.<role>`` holds the resolved field's queryRef, with the
       matching prototypeQuery Select entry carrying ``NativeReferenceName``;
    2. ``queryFieldParametersByRole`` on singleVisual carries the parameter
       linkage (index/length/display-column expr);
    3. ``columnProperties`` restates the parameter's display label for the
       resolved field;
    4. the compiled query joins the parameter table and gains a Where clause
       selecting the resolved field through the hidden "<name> Fields" column
       (NAMEOF-style triple-quoted literal);
    5. the resolved field's dataTransforms select carries
       ``sourceFieldParameters`` provenance.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
        role: The field well to bind (e.g. "Y", "Values", "Category")
        parameter_name: Name of the field-parameter table (as created by
            pbix_datamodel_add_field_parameter)
        initial_field: Which of the parameter's fields the visual shows before
            any slicer selection -- a display name ("Revenue") or a field ref
            ("Sales[Total Revenue]" / "'Sales'[Total Revenue]"). Defaults to
            the parameter's first field.
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")
        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")
        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        # A visual with no field wells cannot host a field parameter. Binding
        # one anyway used to "succeed", leaving a textbox carrying a query,
        # dataTransforms and a Y projection -- structurally incoherent, and
        # pbix_doctor does not flag it. Refuse only the types that provably
        # have no data roles, so custom and unrecognized types still bind.
        target_type = _parse_visual_config(
            containers[visual_index]).get("singleVisual", {}).get("visualType", "")
        if target_type in _NO_DATA_ROLE_VISUALS:
            raise LayoutParseError(
                f"Visual {visual_index} on page {page_index} is a "
                f"'{target_type}', which has no field wells -- there is no "
                f"role for a field parameter to swap. Bind it to a data "
                f"visual (chart, card, table, matrix, slicer, ...) instead.")

        # --- the parameter's definition, from the model -------------------
        if parameter_name not in _field_parameter_tables(info):
            raise LayoutParseError(
                f"'{parameter_name}' is not a field parameter in this model "
                f"(no ParameterMetadata). Create one with "
                f"pbix_datamodel_add_field_parameter first.")
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf
        dm_path = os.path.join(info["work_dir"], "DataModel")
        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta_bytes = read_metadata_sqlite(abf)
        td = read_table_from_abf(abf, parameter_name, meta_bytes,
                                 include_calculated=True)
        cols = td.get("columns") or []
        fields_col = f"{parameter_name} Fields"
        order_col = f"{parameter_name} Order"
        if parameter_name not in cols or fields_col not in cols:
            raise LayoutParseError(
                f"'{parameter_name}' does not have the expected field-"
                f"parameter columns ('{parameter_name}', '{fields_col}').")
        di, fi = cols.index(parameter_name), cols.index(fields_col)
        oi = cols.index(order_col) if order_col in cols else None
        tuples = [(r[di], r[fi], (r[oi] if oi is not None else n))
                  for n, r in enumerate(td.get("rows") or [])]
        if not tuples:
            raise LayoutParseError(f"'{parameter_name}' has no fields.")
        tuples.sort(key=lambda t: t[2])

        # --- resolve initial_field ----------------------------------------
        chosen = None
        if initial_field:
            want = initial_field.strip()
            canonical = None
            try:
                _t, _n, canonical = _normalize_field_ref(want)
            except Exception:  # noqa: BLE001 - maybe it is a display name
                pass
            for disp, ref, _o in tuples:
                if disp == want or (canonical and ref == canonical):
                    chosen = (disp, ref)
                    break
            if chosen is None:
                raise LayoutParseError(
                    f"initial_field {initial_field!r} is not one of "
                    f"'{parameter_name}''s fields: "
                    + ", ".join(d for d, _r, _o in tuples))
        else:
            chosen = (tuples[0][0], tuples[0][1])
        display, canonical_ref = chosen
        f_table, f_name, _canon = _normalize_field_ref(canonical_ref)

        # Column or measure? The Select node kind must match the model.
        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        os.write(fd, meta_bytes)
        os.close(fd)
        try:
            conn = sqlite3.connect(tmp_path)
            conn.row_factory = sqlite3.Row

            def _field_is_measure(ref: str) -> bool | None:
                """True=measure, False=column, None=missing from the model."""
                try:
                    t_, n_, _c = _normalize_field_ref(ref)
                except Exception:  # noqa: BLE001
                    return None
                tr = conn.execute(
                    "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                    (t_,)).fetchone()
                if not tr:
                    return None
                if conn.execute(
                        "SELECT 1 FROM Measure WHERE TableID = ? AND Name = ?",
                        (tr["ID"], n_)).fetchone():
                    return True
                if conn.execute(
                        "SELECT 1 FROM [Column] WHERE TableID = ? AND "
                        "COALESCE(ExplicitName, InferredName) = ?",
                        (tr["ID"], n_)).fetchone():
                    return False
                return None

            # An AGGREGATING value role can only host measure-backed
            # parameter fields. A raw column there renders EMPTY whichever
            # way it is projected — bare, Desktop drops the un-aggregated
            # well; Sum-wrapped, the queryRef no longer matches any NAMEOF
            # ref (issue #36) — so both silent-empty shapes are refused
            # here, loudly, at author time (issue #37). Columns stay valid
            # for grouping roles (Category/Axis/Rows — the documented
            # dimension-swap use) and for table/slicer wells, which show
            # raw values.
            from pbix_mcp.report_binding import (
                _SLICER_TYPES,
                _TABLE_TYPES,
                _VALUE_ROLES,
            )
            aggregating = (target_type not in _TABLE_TYPES
                           and target_type not in _SLICER_TYPES)
            if aggregating and role in _VALUE_ROLES:
                col_fields = [(d, r) for d, r, _o in tuples
                              if _field_is_measure(r) is False]
                if col_fields:
                    listing = ", ".join(
                        f"'{d}' ({r})" for d, r in col_fields)
                    example_d, example_r = col_fields[0]
                    et, en, _ec = _normalize_field_ref(example_r)
                    raise LayoutParseError(
                        f"A field parameter bound to the aggregating value "
                        f"role '{role}' of a '{target_type}' must list "
                        f"MEASURES, but '{parameter_name}' contains column "
                        f"field(s): {listing}. A raw column in a value role "
                        f"renders an EMPTY visual in Power BI (no bars, no "
                        f"value axis). Add measures — e.g. "
                        f"pbix_datamodel_add_measure(..., '{et}', "
                        f"'Total {en}', 'SUM({et}[{en}])') — and point the "
                        f"parameter at those, or bind this parameter to a "
                        f"grouping role (Category/Axis/Rows) instead.")

            trow = conn.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (f_table,)).fetchone()
            if not trow:
                raise LayoutParseError(
                    f"Parameter field {canonical_ref} points at a table that "
                    f"no longer exists ('{f_table}').")
            is_measure = bool(conn.execute(
                "SELECT 1 FROM Measure WHERE TableID = ? AND Name = ?",
                (trow["ID"], f_name)).fetchone())
            if not is_measure and not conn.execute(
                    "SELECT 1 FROM [Column] WHERE TableID = ? AND "
                    "COALESCE(ExplicitName, InferredName) = ?",
                    (trow["ID"], f_name)).fetchone():
                raise LayoutParseError(
                    f"Parameter field {canonical_ref} no longer exists in "
                    f"the model.")
            conn.close()
        finally:
            # Close BEFORE unlinking: Windows refuses to delete a file with
            # an open handle (WinError 32), and the refusal paths above raise
            # while the connection is still open.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        # --- author the visual config -------------------------------------
        vc = containers[visual_index]
        config = _parse_visual_config(vc)
        sv = config.setdefault("singleVisual", {})
        proto = sv.setdefault("prototypeQuery", {"Version": 2, "From": [],
                                                 "Select": []})
        from_list = proto.setdefault("From", [])
        selects = proto.setdefault("Select", [])
        query_ref = f"{f_table}.{f_name}"

        # Alias for the resolved field's table in the PROTOTYPE (the
        # parameter table itself joins only in the COMPILED query).
        alias_name = None
        used = {fr.get("Name") for fr in from_list}
        for fr in from_list:
            if fr.get("Entity") == f_table:
                alias_name = fr.get("Name")
                break
        if alias_name is None:
            base = (f_table[:1].lower() or "t")
            alias_name = base
            n = 0
            while alias_name in used:
                n += 1
                alias_name = f"{base}{n}"
            from_list.append({"Name": alias_name, "Entity": f_table,
                              "Type": 0})

        # Drop any select previously projected into this role (including the
        # naive parameter-display-column shape this tool replaces) -- unless
        # another role still projects it.
        old_refs = {it.get("queryRef")
                    for it in (sv.get("projections") or {}).get(role, [])}
        old_refs.add(f"{parameter_name}.{parameter_name}")
        still_projected = set()
        for r2, items in (sv.get("projections") or {}).items():
            if r2 == role:
                continue
            for it in items or []:
                still_projected.add(it.get("queryRef"))
        dropped_sel = [s2 for s2 in selects
                       if s2.get("Name") in old_refs
                       and s2.get("Name") not in still_projected]
        proto["Select"] = [s2 for s2 in selects if s2 not in dropped_sel]
        selects = proto["Select"]

        # A prior pbix_set_visual_sort OrderBy may point at a select we just
        # dropped -- rebinding Y from [TR] to [TC] otherwise leaves the
        # compiled query ordering by a field it no longer selects, a DANGLING
        # reference. The sort was on THIS role's field and the role still has
        # one, so re-point it at the newly bound field (preserving the user's
        # intent to sort by the value axis) rather than silently losing it.
        order_by = proto.get("OrderBy") or []
        if order_by and dropped_sel:
            def _inner(node):
                if "Aggregation" in node:
                    node = node["Aggregation"].get("Expression", {})
                return node.get("Measure") or node.get("Column") or {}

            dropped_props = {(_inner(d).get("Property") or "") for d in dropped_sel}
            kept = []
            for ob in order_by:
                prop = _inner(ob.get("Expression") or {}).get("Property") or ""
                if prop and prop in dropped_props:
                    ob = copy.deepcopy(ob)
                    ob["Expression"] = {
                        ("Measure" if is_measure else "Column"): {
                            "Expression": {"SourceRef": {"Source": alias_name}},
                            "Property": f_name,
                        }}
                kept.append(ob)
            proto["OrderBy"] = kept

        node_kind = "Measure" if is_measure else "Column"
        if not any(s2.get("Name") == query_ref for s2 in selects):
            selects.append({
                node_kind: {
                    "Expression": {"SourceRef": {"Source": alias_name}},
                    "Property": f_name,
                },
                "Name": query_ref,
                "NativeReferenceName": display,
            })

        sv.setdefault("projections", {})[role] = [{"queryRef": query_ref}]
        sv.setdefault("queryFieldParametersByRole", {})[role] = [{
            "index": 0,
            "length": 1,
            "expr": {"Column": {
                "Expression": {"SourceRef": {"Entity": parameter_name}},
                "Property": parameter_name,
            }},
        }]
        sv.setdefault("columnProperties", {})[query_ref] = {
            "displayName": display}

        # --- recompile the caches (query + dataTransforms) ----------------
        from pbix_mcp.report_binding import compile_visual_binding
        q, dt = compile_visual_binding(sv, _report_type_resolver(info))
        if q is not None:
            vc["query"] = json.dumps(q, ensure_ascii=False)
            vc["dataTransforms"] = json.dumps(dt, ensure_ascii=False)
            vc.setdefault("filters", "[]")
        # Serialize AFTER compile: implicit value-role aggregation mutates
        # the prototype, and the persisted config must match the query.
        vc["config"] = json.dumps(config, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Field parameter '{parameter_name}' bound to role '{role}' of "
            f"visual {visual_index} on page {page_index}; initial field "
            f"'{display}' ({canonical_ref}). A slicer over "
            f"'{parameter_name}' now swaps this well's field.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}",
                                  str(getattr(e, "code", "") or
                                      "PBIX_MCP_ERROR")).to_text()


@mcp.tool()
def pbix_format_visual(
    alias: str, page_index: int, visual_index: int, format_json: str
) -> str:
    """Format a visual with human-readable properties (colors, titles, fonts).

    Converts simple formatting options to PBI's internal objects structure.
    Merges with existing formatting — only specified properties are changed.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
        format_json: JSON object with formatting options. Supported keys:

            title: {text, show, fontSize, color, fontFamily, bold, italic, alignment}
            subtitle: {text, show, fontSize, color, fontFamily}
            dataLabels: {show, fontSize, color, displayUnits, decimalPlaces}
            legend: {show, position, fontSize, color}
                position: "top", "bottom", "left", "right", "topCenter"
            categoryAxis: {show, fontSize, color, title, gridlineShow, innerPadding,
                invertAxis, axisType, start, end, switchAxisPosition}
            valueAxis: {show, fontSize, color, displayUnits, title, titleFontSize,
                gridlineShow, start, end, decimalPlaces, switchAxisPosition,
                secShow, secFontSize, secColor, secDisplayUnits, secTitle,
                secAxisTitle, secShowAxisTitle, secStart, secEnd, alignZeros}
                displayUnits: "none", "thousands", "millions", "billions", "auto"
                sec* = a combo chart's SECONDARY value axis (both axes share
                this one card); alignZeros pins the two zero lines together
            background: {color, transparency}
            border: {show, color, radius, width}
            dropShadow: {show, color, angle, blur, distance, spread, transparency,
                position, preset}
            padding: number | {top, bottom, left, right}
            spacing: {belowTitle, belowSubTitle, belowTitleArea, vertical}
            divider: {show, color, width, style, ignorePadding}
            visualHeader: {show, showOptionsMenu, showFocusModeButton, showPinButton,
                showFilterRestatementButton, showTooltipButton, showDrillUpButton, ...}
            visualTooltip: {show, type, fontSize, titleFontColor, valueFontColor,
                actionFontColor, background}
            dataColors: ["#hex1", "#hex2", ...]
            grid: {gridVertical, gridHorizontal, rowPadding, outlineColor, outlineWeight,
                gridHorizontalColor, gridVerticalColor}
            columnHeaders: {bold, fontSize, fontFamily, fontColor, backColor, alignment}
            values: {bold, fontSize, fontFamily, fontColor, backColor, wordWrap,
                backColorPrimary, backColorSecondary, fontColorPrimary, fontColorSecondary}
            total: {show, bold, fontSize, fontColor, backColor}
            outline: {show, weight, color}
            fill: {color, transparency, show}
            line: {lineStyle, strokeWidth, showMarker, markerShape, markerSize}
            categoryLabels: {show, fontSize, color, fontFamily}
            slices: {innerRadius}
            smallMultiples: {minWidth, maxWidth, minHeight}
            stylePreset: "name"
            altText: "description"
            lockAspect: true/false

            Example: {"title": {"text": "Sales", "fontSize": 16}, "dataLabels": {"show": true}}
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        try:
            fmt = json.loads(format_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid format_json: {e}")

        vc = containers[visual_index]
        config = _parse_visual_config(vc)
        sv = config.setdefault("singleVisual", {})
        existing_objects = sv.setdefault("objects", {})

        # Handle dataColors with per-series/per-category selectors
        colors = fmt.get("dataColors") if isinstance(fmt, dict) else None
        if isinstance(colors, list) and len(colors) > 1:
            projections = sv.get("projections", {})
            proto = sv.get("prototypeQuery", {})
            selects = proto.get("Select", [])

            # Build dataPoint entries with selectors
            dp_entries = []

            # Determine selector type based on visual projections
            y_refs = [p.get("queryRef", "") for p in projections.get("Y", [])]
            cat_refs = [p.get("queryRef", "") for p in projections.get("Category", [])]

            if len(y_refs) > 1:
                # Multi-measure chart (column chart with Revenue + Cost)
                # Selector: {"metadata": "Table.Measure"}
                for i, y_ref in enumerate(y_refs):
                    if i < len(colors):
                        entry = {"properties": {"fill": _solid_color(colors[i])}}
                        entry["selector"] = {"metadata": y_ref}
                        dp_entries.append(entry)

            elif len(cat_refs) >= 1 and len(colors) > 1:
                # Single-measure chart with category axis (donut, bar by category)
                # Need to find unique category values from the data
                # Selector: {"data": [{scopeId: {Comparison: ...}}]}
                cat_ref = cat_refs[0]  # e.g. "Products.Category"
                parts = cat_ref.split(".")
                if len(parts) == 2:
                    entity, prop = parts
                    # Try to get unique values from the table data
                    try:
                        from pbix_mcp.formats.model_reader import ModelReader
                        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
                        td = model.get_table(entity)
                        col_idx = td["columns"].index(prop)
                        unique_vals = sorted(set(
                            row[col_idx] for row in td["rows"]
                            if row[col_idx] is not None
                        ))
                        for i, val in enumerate(unique_vals):
                            if i < len(colors):
                                entry = {"properties": {"fill": _solid_color(colors[i])}}
                                entry["selector"] = {
                                    "data": [{
                                        "scopeId": {
                                            "Comparison": {
                                                "ComparisonKind": 0,
                                                "Left": {
                                                    "Column": {
                                                        "Expression": {
                                                            "SourceRef": {"Entity": entity}
                                                        },
                                                        "Property": prop,
                                                    }
                                                },
                                                "Right": {
                                                    "Literal": {"Value": f"'{val}'"}
                                                },
                                            }
                                        }
                                    }]
                                }
                                dp_entries.append(entry)
                    except Exception:
                        pass  # Fall through to single-color

            if dp_entries:
                existing_objects["dataPoint"] = dp_entries
                fmt["_skip_datacolors"] = True  # Skip the single-color fallback

        visual_type = sv.get("visualType", "")
        result = _build_format_objects(fmt, visual_type=visual_type)
        new_objects = result.get("_objects", {})
        new_vc_objects = result.get("_vcObjects", {})

        # Deep-merge so a partial update keeps the unspecified sibling properties
        # of an existing object entry (matched by selector). Replacing the whole
        # category array used to silently drop e.g. border width/radius when only
        # the colour was sent.
        def _merge_obj_entries(existing_list, entries):
            out = list(existing_list or [])
            for ne in entries:
                sel = ne.get("selector")
                match = next((e for e in out if e.get("selector") == sel), None)
                if match is not None:
                    match.setdefault("properties", {}).update(ne.get("properties", {}))
                else:
                    out.append(ne)
            return out

        # Merge data formatting into singleVisual.objects
        for category, entries in new_objects.items():
            existing_objects[category] = _merge_obj_entries(existing_objects.get(category), entries)

        # Merge container formatting into singleVisual.vcObjects
        if new_vc_objects:
            existing_vc_objects = sv.setdefault("vcObjects", {})
            for category, entries in new_vc_objects.items():
                existing_vc_objects[category] = _merge_obj_entries(existing_vc_objects.get(category), entries)

        applied = list(new_objects.keys()) + list(new_vc_objects.keys())
        if not applied:
            # Reporting success for a no-op is the worst failure shape: the
            # caller believes the formatting landed. Name what was ignored --
            # and name it ACCURATELY. Blaming the card while listing that same
            # card as supported (issue #51: "none of ['valueAxis'] is a
            # recognised key ... Supported keys include: ... valueAxis") sent
            # the caller after the wrong bug; the real fault was that every
            # PROPERTY inside the recognised card was dropped.
            ignored = (sorted(k for k in fmt if not k.startswith("_"))
                       if isinstance(fmt, dict) else [])
            known = [k for k in ignored if k in _FORMAT_CARDS]
            unknown = [k for k in ignored if k not in _FORMAT_CARDS]
            parts = []
            for k in known:
                inner = fmt.get(k)
                props = (sorted(inner) if isinstance(inner, dict)
                         else [repr(inner)])
                parts.append(f"{k}: no recognised properties in {props}")
            if unknown:
                parts.append(f"unrecognised card(s): {unknown}")
            detail = "; ".join(parts) or "(empty input)"
            vt = f" on visual type '{visual_type}'" if visual_type else ""
            raise LayoutParseError(
                f"No formatting was applied{vt}, so the visual is unchanged "
                f"— {detail}. Some properties are only recognised on the "
                f"visual type that owns them. Supported cards include: "
                f"title, subtitle, background, border, padding, dataLabels, "
                f"legend, categoryAxis, valueAxis, dataColors, visualHeader, "
                f"altText. Values are human-readable, e.g. "
                f'{{"title": {{"text": "Sales", "show": true}}}} — not raw '
                f"Power BI object descriptors.")

        vc["config"] = json.dumps(config, ensure_ascii=False)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True

        return ToolResponse.ok(
            f"Formatted visual {visual_index} on page {page_index}: {', '.join(applied)}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_remove_visual(alias: str, page_index: int, visual_index: int) -> str:
    """Remove a visual from a report page.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        removed = containers.pop(visual_index)
        config = _parse_visual_config(removed)
        vtype = _get_visual_type(config)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(f"Removed {vtype} visual (was index {visual_index}). {len(containers)} visuals remain.").to_text()

    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_layout_raw(alias: str) -> str:
    """Get the raw Report/Layout JSON.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")
        return ToolResponse.ok(json.dumps(layout, indent=2, ensure_ascii=False)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_layout_raw(alias: str, layout_json: str) -> str:
    """Write raw layout JSON back to Report/Layout.

    Args:
        alias: The alias of the open file
        layout_json: Complete layout JSON string
    """
    try:
        info = _ensure_open(alias)
        try:
            layout = json.loads(layout_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid JSON: {e}")
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok("Layout updated.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_filters(alias: str, page_index: int = -1) -> str:
    """Get report-level or page-level filters.

    Args:
        alias: The alias of the open file
        page_index: Page index for page filters, or -1 for report-level filters
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        if page_index == -1:
            # Report-level filters
            filters_raw = layout.get("filters", "[]")
        else:
            sections = layout.get("sections", [])
            if page_index < 0 or page_index >= len(sections):
                raise LayoutParseError(f"Page index {page_index} out of range")
            filters_raw = sections[page_index].get("filters", "[]")

        if isinstance(filters_raw, str):
            try:
                filters = json.loads(filters_raw)
            except json.JSONDecodeError:
                filters = filters_raw
        else:
            filters = filters_raw

        level = f"page {page_index}" if page_index >= 0 else "report"
        return ToolResponse.ok(f"Filters ({level}):\n{json.dumps(filters, indent=2, ensure_ascii=False)}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_filters(alias: str, filters_json: str, page_index: int = -1) -> str:
    """Set report-level or page-level filters.

    Args:
        alias: The alias of the open file
        filters_json: JSON array of filter definitions
        page_index: Page index for page filters, or -1 for report-level filters
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        # Validate JSON
        try:
            json.loads(filters_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid JSON: {e}")

        if page_index == -1:
            layout["filters"] = filters_json
        else:
            sections = layout.get("sections", [])
            if page_index < 0 or page_index >= len(sections):
                raise LayoutParseError(f"Page index {page_index} out of range")
            sections[page_index]["filters"] = filters_json

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        level = f"page {page_index}" if page_index >= 0 else "report"
        return ToolResponse.ok(f"Filters updated ({level}).").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_settings(alias: str) -> str:
    """Get report settings.

    Reads ``Report/definition/report.json`` ``settings`` on a PBIR report and
    the legacy ``Report/Settings`` part on a classic one — the authoritative
    location differs by format, and reading the classic part on a PBIR file
    reports "no settings" for a report that has them.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        if _is_pbir(work_dir):
            cfg = _get_report_config(work_dir) or {}
            settings = cfg.get("settings")
            if settings is None:
                return ToolResponse.ok(
                    "No settings found in Report/definition/report.json."
                ).to_text()
            return ToolResponse.ok(
                json.dumps(settings, indent=2, ensure_ascii=False),
                data={"settings": settings, "source": "report.json"}).to_text()
        settings = _read_json_component(work_dir, os.path.join("Report", "Settings"))
        if settings is None:
            return ToolResponse.ok("No Settings found.").to_text()
        return ToolResponse.ok(json.dumps(settings, indent=2, ensure_ascii=False)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_settings(alias: str, settings_json: str) -> str:
    """Write report settings.

    Writes ``Report/definition/report.json`` ``settings`` on a PBIR report and
    the legacy ``Report/Settings`` part on a classic one. Writing the classic
    part into a PBIR report created a second, conflicting settings document
    that Power BI ignores while the real one stayed stale.

    Args:
        alias: The alias of the open file
        settings_json: Complete settings JSON string
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        try:
            settings = json.loads(settings_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid JSON: {e}")
        if _is_pbir(work_dir):
            cfg = _get_report_config(work_dir)
            if cfg is None:
                raise LayoutParseError(
                    "No Report/definition/report.json found.")
            cfg["settings"] = settings
            _set_report_config(work_dir, cfg)
            info["modified"] = True
            return ToolResponse.ok(
                "Settings updated in Report/definition/report.json.").to_text()
        _write_json_component(work_dir, os.path.join("Report", "Settings"), settings)
        info["modified"] = True
        return ToolResponse.ok("Settings updated.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_get_bookmarks(alias: str) -> str:
    """Get report bookmarks.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        config_str = layout.get("config", "{}")
        if isinstance(config_str, str):
            try:
                config = json.loads(config_str)
            except json.JSONDecodeError:
                config = {}
        else:
            config = config_str

        bookmarks = config.get("bookmarks", [])
        if not bookmarks:
            return ToolResponse.ok("No bookmarks found.").to_text()

        lines = [f"Report has {len(bookmarks)} bookmark(s):\n"]
        for i, bm in enumerate(bookmarks):
            name = bm.get("displayName", bm.get("name", f"Bookmark {i}"))
            lines.append(f"  [{i}] {name}")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_add_bookmark(
    alias: str,
    display_name: str,
    target_page: str = "",
    hidden_visuals: str = "",
    report_filter_json: str = "",
) -> str:
    """Create a report bookmark that captures page and visual state.

    Args:
        alias: The alias of the open file
        display_name: Name for the bookmark (e.g. "Sales Overview", "Q4 Filter")
        target_page: Optional page displayName or index to navigate to when bookmark is applied.
                     If empty, bookmark targets the first page.
        hidden_visuals: Optional comma-separated list of visual names to hide when
                        bookmark is applied (e.g. "visual_0,visual_2"). Other visuals
                        stay visible.
        report_filter_json: Optional JSON array of report-level filters to apply
                            when bookmark is activated, e.g.
                            '[{"target":{"table":"Sales","column":"Region"},"operator":"In","values":["West"]}]'
    """
    import uuid as _uuid

    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if not sections:
            raise LayoutParseError("Report has no pages")

        # Resolve target page
        target_section = None
        if target_page:
            # Try numeric index first
            try:
                idx = int(target_page)
                if 0 <= idx < len(sections):
                    target_section = sections[idx]
            except ValueError:
                pass
            # Try display name match
            if not target_section:
                for sec in sections:
                    if sec.get("displayName", "").lower() == target_page.lower():
                        target_section = sec
                        break
            if not target_section:
                raise LayoutParseError(
                    f"Page '{target_page}' not found. "
                    f"Available: {[s.get('displayName') for s in sections]}"
                )
        else:
            target_section = sections[0]

        section_name = target_section.get("name", "ReportSection1")

        # Build visual state — all visuals visible unless in hidden list
        hidden_set = set()
        if hidden_visuals:
            hidden_set = {v.strip() for v in hidden_visuals.split(",") if v.strip()}

        visual_states = {}
        for vc in target_section.get("visualContainers", []):
            vc_config_str = vc.get("config", "{}")
            try:
                vc_config = json.loads(vc_config_str) if isinstance(vc_config_str, str) else vc_config_str
            except json.JSONDecodeError:
                continue
            vname = vc_config.get("name", "")
            if vname:
                visual_states[vname] = {
                    "visualType": vc_config.get("singleVisual", {}).get("visualType", "unknown"),
                    "hidden": vname in hidden_set,
                }

        # Build bookmark object
        bookmark_id = str(_uuid.uuid4()).replace("-", "")[:20]
        bookmark = {
            "displayName": display_name,
            "name": f"Bookmark{bookmark_id}",
            "explorationState": {
                "version": "1.2",
                "activeSection": section_name,
            },
            "options": {
                "targetVisualNames": list(visual_states.keys()) if visual_states else [],
            },
        }

        # Add visual display states if any visuals hidden.
        # Power BI's display.mode enum is "hidden"|"maximize"|"spotlight"|
        # "elevation" — there is NO "visible". A visible visual is expressed by
        # OMITTING mode (applying the bookmark with no display override shows
        # it). Writing mode:"visible" made Desktop ignore the block / mishandle
        # Selection-pane state, so hidden visuals get {"display":{"mode":
        # "hidden"}} and visible ones get a bare {"singleVisual":{}} (no mode).
        # `sections` is required by the PBIR bookmark schema and is present in
        # every Desktop-authored bookmark, so it is written unconditionally.
        bookmark["explorationState"]["sections"] = {
            section_name: {
                "visualContainers": {
                    vname: (
                        {"singleVisual": {"display": {"mode": "hidden"}}}
                        if state["hidden"]
                        else {"singleVisual": {}}
                    )
                    for vname, state in visual_states.items()
                }
            }
        }

        # Add report-level filters if provided
        if report_filter_json:
            try:
                filters = json.loads(report_filter_json)
                bookmark["explorationState"]["filters"] = {"byExpr": filters}
            except json.JSONDecodeError:
                raise LayoutParseError("Invalid report_filter_json — must be valid JSON array")

        # Insert into layout config
        config_str = layout.get("config", "{}")
        if isinstance(config_str, str):
            try:
                config = json.loads(config_str)
            except json.JSONDecodeError:
                config = {}
        else:
            config = config_str

        config.setdefault("bookmarks", []).append(bookmark)
        layout["config"] = json.dumps(config, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True

        hidden_msg = f", hiding: {hidden_visuals}" if hidden_visuals else ""
        return ToolResponse.ok(
            f"Created bookmark '{display_name}' → page '{target_section.get('displayName')}'"
            f"{hidden_msg}. Total bookmarks: {len(config['bookmarks'])}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_remove_bookmark(alias: str, bookmark_index: int) -> str:
    """Remove a bookmark by index.

    Args:
        alias: The alias of the open file
        bookmark_index: Zero-based index of the bookmark to remove (from pbix_get_bookmarks)
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        config_str = layout.get("config", "{}")
        if isinstance(config_str, str):
            try:
                config = json.loads(config_str)
            except json.JSONDecodeError:
                config = {}
        else:
            config = config_str

        bookmarks = config.get("bookmarks", [])
        if not bookmarks:
            return ToolResponse.ok("No bookmarks to remove.").to_text()

        if bookmark_index < 0 or bookmark_index >= len(bookmarks):
            raise LayoutParseError(
                f"Index {bookmark_index} out of range (0–{len(bookmarks) - 1})"
            )

        removed = bookmarks.pop(bookmark_index)
        name = removed.get("displayName", removed.get("name", "?"))
        config["bookmarks"] = bookmarks
        layout["config"] = json.dumps(config, ensure_ascii=False)

        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Removed bookmark '{name}'. Remaining: {len(bookmarks)}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_metadata(alias: str) -> str:
    """Get file metadata — component inventory and sizes.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        lines = [f"Metadata for '{alias}' ({info['path']}):\n"]
        total = 0
        for root, dirs, files in os.walk(work_dir):
            for f in files:
                fp = os.path.join(root, f)
                rel = os.path.relpath(fp, work_dir)
                size = os.path.getsize(fp)
                total += size
                lines.append(f"  {rel}: {size:,} bytes")
        lines.append(f"\nTotal: {total:,} bytes")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 5: Resources & Theme tools ----

@mcp.tool()
def pbix_list_resources(alias: str) -> str:
    """List all static resources (images, custom visuals, themes).

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        resource_dirs = [
            "Report/StaticResources",
            "Report/CustomVisuals",
        ]
        lines = ["Resources:\n"]
        found = False
        for rd in resource_dirs:
            rd_full = os.path.join(work_dir, rd)
            if os.path.isdir(rd_full):
                for root, dirs, files in os.walk(rd_full):
                    for f in files:
                        fp = os.path.join(root, f)
                        rel = os.path.relpath(fp, work_dir)
                        size = os.path.getsize(fp)
                        lines.append(f"  {rel} ({size:,} bytes)")
                        found = True
        if not found:
            return ToolResponse.ok("No resources found.").to_text()
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# --- Registered resources (images & friends) -------------------------------
#
# A file-backed report resource is registered by three touchpoints (all
# verified against Desktop-authored corpus files — GeoSales_Dashboard,
# Agents_Performance):
#   1. the bytes at Report/StaticResources/RegisteredResources/<item>
#   2. a `<Default Extension="<ext>" ContentType=""/>` in [Content_Types].xml
#   3. a type-tagged entry in the layout's top-level `resourcePackages`
#      RegisteredResources package
# An image visual then points at it with an `ImageUrl` ResourcePackageItem
# expr (PackageType 1). Item types: 100 = image, 200 = shape map,
# 201 = custom theme, 202 = base theme.

_IMAGE_MAX_BYTES = 5 * 1024 * 1024
_RESOURCE_ITEM_TYPES = {"image": 100, "shapemap": 200,
                        "customtheme": 201, "basetheme": 202}
# Scaling literals live under objects.imageScaling (NOT objects.general) as
# single-quoted PBI string literals. 'Fit' is corpus-verified.
_IMAGE_SCALING = {"fit": "Fit", "fill": "Fill", "normal": "Normal"}


def _sniff_image_ext(data: bytes) -> str | None:
    """Canonical extension from the file's MAGIC BYTES — never from a caller's
    filename or content-type claim. Returns None for anything unrecognized.

    Raster formats are exact magic-byte matches. SVG is text, so its detection
    skips a UTF-8/UTF-16 BOM, whitespace, the XML declaration, DOCTYPE, and
    comments, then requires the ROOT ELEMENT to be <svg> — a substring search
    would classify any XML that merely mentions svg somewhere.
    """
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if data[:3] == b"\xff\xd8\xff":
        return "jpg"
    if data[:6] in (b"GIF87a", b"GIF89a"):
        return "gif"
    if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
        return "webp"
    if data[:2] == b"BM":
        return "bmp"
    if data[:4] in (b"II*\x00", b"MM\x00*"):
        return "tiff"
    if data[:4] == b"\x00\x00\x01\x00":
        return "ico"

    # --- SVG (text) ---
    head = data[:8192]
    for bom, enc in ((b"\xef\xbb\xbf", "utf-8"),
                     (b"\xff\xfe", "utf-16-le"), (b"\xfe\xff", "utf-16-be")):
        if head.startswith(bom):
            head = head[len(bom):]
            try:
                head = head.decode(enc, errors="ignore").encode("utf-8")
            except Exception:
                return None
            break
    i, n = 0, len(head)
    while i < n:
        while i < n and head[i:i + 1].isspace():
            i += 1
        if head[i:i + 4] == b"<!--":                      # comment
            end = head.find(b"-->", i)
            if end < 0:
                return None
            i = end + 3
        elif head[i:i + 5] == b"<?xml" or head[i:i + 2] == b"<?":   # decl / PI
            end = head.find(b"?>", i)
            if end < 0:
                return None
            i = end + 2
        elif head[i:i + 9].upper() == b"<!DOCTYPE":       # doctype (may nest [])
            depth = 0
            while i < n:
                c = head[i:i + 1]
                if c == b"[":
                    depth += 1
                elif c == b"]":
                    depth -= 1
                elif c == b">" and depth <= 0:
                    i += 1
                    break
                i += 1
        else:
            break
    tag = head[i:i + 4].lower()
    if tag == b"<svg" and head[i + 4:i + 5] in (b"", b" ", b"\t", b"\r", b"\n", b">", b"/"):
        return "svg"
    return None


def _sniff_resource_ext(data: bytes, item_type: int) -> str | None:
    """Extension for a registered resource of ``item_type``.

    Images (100) must be images. Shape maps (200) and themes (201/202) are
    JSON in Desktop-authored files (corpus-verified: a type-200
    ``us-states….json`` TopoJSON, a type-202 ``BaseThemes/….json``), so those
    types accept JSON — and NOT an image, which Power BI could not consume.
    """
    if item_type == 100:
        return _sniff_image_ext(data)
    text = data.lstrip(b"\xef\xbb\xbf").lstrip()
    if text[:1] in (b"{", b"["):
        try:
            json.loads(data.decode("utf-8-sig"))
        except Exception:
            return None
        return "json"
    return None


def _sanitize_item_name(name: str, ext: str) -> str:
    """Restrict an item name to [A-Za-z0-9._-] and force the sniffed
    extension, so the name always agrees with [Content_Types].xml."""
    stem = os.path.basename(name or "").strip()
    stem = os.path.splitext(stem)[0]
    stem = re.sub(r"[^A-Za-z0-9._-]", "_", stem).strip("._-")
    if not stem:
        stem = "image"
    return f"{stem[:80]}.{ext}"


def _load_resource_bytes(image_path: str, image_base64: str) -> bytes:
    """Read resource bytes from a local path or a base64 payload (exactly one).

    No remote fetching: a URL would make the engine download untrusted content
    on the caller's behalf. Callers holding a URL fetch it themselves and pass
    the bytes.
    """
    if bool(image_path) == bool(image_base64):
        raise LayoutParseError(
            "Provide exactly one image source: image_path or image_base64.")
    if image_path:
        if not os.path.isfile(image_path):
            raise LayoutParseError(f"Image file not found: {image_path}")
        if os.path.getsize(image_path) > _IMAGE_MAX_BYTES:
            raise LayoutParseError(
                f"Image exceeds the {_IMAGE_MAX_BYTES // (1024 * 1024)} MB limit.")
        with open(image_path, "rb") as f:
            return f.read()
    import base64 as _b64
    payload = (image_base64 or "").strip()
    if payload.startswith("data:"):          # tolerate a full data: URI
        payload = payload.split(",", 1)[-1]
    # Strip line wrapping first: MIME / `base64` / openssl output is wrapped at
    # 64-76 columns, which validate=True would reject outright.
    payload = re.sub(r"\s+", "", payload)
    try:
        data = _b64.b64decode(payload, validate=True)
    except Exception as e:
        raise LayoutParseError(f"image_base64 is not valid base64: {e}")
    if len(data) > _IMAGE_MAX_BYTES:
        raise LayoutParseError(
            f"Image exceeds the {_IMAGE_MAX_BYTES // (1024 * 1024)} MB limit.")
    return data


def _ensure_content_type_default(work_dir: str, ext: str) -> bool:
    """Ensure [Content_Types].xml declares ``ext`` (empty ContentType, as
    Desktop writes). Returns True when the file was changed.

    Inserts before the json Default when present, else before ``</Types>`` —
    the old json-anchored string replace SILENTLY no-opped on documents with
    no json Default (e.g. the repo's own fixtures), shipping a pbix whose
    image part had no declared extension.
    """
    ct_path = os.path.join(work_dir, "[Content_Types].xml")
    if not os.path.exists(ct_path):
        return False
    with open(ct_path, "r", encoding="utf-8") as f:
        ct_xml = f.read()
    if re.search(rf'Extension\s*=\s*"{re.escape(ext)}"', ct_xml, re.IGNORECASE):
        return False
    entry = f'<Default Extension="{ext}" ContentType=""/>'
    if '<Default Extension="json"' in ct_xml:
        ct_xml = ct_xml.replace('<Default Extension="json"',
                                entry + '<Default Extension="json"', 1)
    elif "</Types>" in ct_xml:
        ct_xml = ct_xml.replace("</Types>", entry + "</Types>", 1)
    else:
        return False
    with open(ct_path, "w", encoding="utf-8") as f:
        f.write(ct_xml)
    return True


# A PBIR report.json declares resource items with STRING types in a FLAT
# package, where classic Report/Layout nests the package and uses numeric
# codes. Verified against test_corpus/{Ecommerce_Conversion,IT_Support}.pbix.
_PBIR_ITEM_TYPE = {100: "Image", 200: "ShapeMap", 201: "CustomTheme",
                   202: "BaseTheme"}
_PBIR_ITEM_TYPE_INV = {v: k for k, v in _PBIR_ITEM_TYPE.items()}

# Package-level type. Verified against the corpus: every classic Layout writes
# 1 for RegisteredResources and 2 for SharedResources, and the matching PBIR
# report.json writes the enum names.
_PBIR_PACKAGE_TYPE = {1: "RegisteredResources", 2: "SharedResources"}
_PBIR_PACKAGE_TYPE_INV = {v: k for k, v in _PBIR_PACKAGE_TYPE.items()}


def _classic_theme_to_pbir(theme: dict, original: dict | None = None) -> dict:
    """Classic ``config.themeCollection`` -> the PBIR report.json shape.

    Only `type` differs structurally (int vs enum name). PBIR also carries
    `reportVersionAtImport`, which the classic shape cannot express, so it is
    preserved from ``original`` per theme slot rather than dropped.
    """
    original = original or {}
    out: dict = {}
    # PBIR requires reportVersionAtImport on every theme slot. A newly added
    # customTheme has no original to inherit it from, so it borrows the one the
    # report already declares.
    fallback_version = None
    for slot in ("baseTheme", "customTheme"):
        v = (original.get(slot) or {}).get("reportVersionAtImport")
        if v is not None:
            fallback_version = v
            break

    for slot, entry in (theme or {}).items():
        if not isinstance(entry, dict):
            continue
        base = copy.deepcopy(original.get(slot, {}))
        # The report's real baseTheme is authoritative on PBIR; the classic
        # code substitutes a built-in default that would then contradict
        # resourcePackages. Only the customTheme overlay is taken from the
        # caller.
        if slot == "baseTheme" and base:
            out[slot] = base
            continue
        base.update({k: v for k, v in entry.items() if k != "type"})
        ttype = entry.get("type", base.get("type"))
        base["type"] = (ttype if isinstance(ttype, str)
                        else _PBIR_PACKAGE_TYPE.get(ttype, "SharedResources"))
        # `version` is the classic spelling; PBIR uses reportVersionAtImport.
        base.pop("version", None)
        if base.get("reportVersionAtImport") is None and fallback_version:
            base["reportVersionAtImport"] = copy.deepcopy(fallback_version)
        out[slot] = base
    # Never drop a theme slot the report already had.
    for slot, entry in original.items():
        out.setdefault(slot, entry)
    return out


def _pbir_packages_to_classic(pkgs: list) -> list:
    """PBIR report.json resourcePackages -> the nested classic Layout shape."""
    out = []
    for pkg in pkgs or []:
        if "resourcePackage" in pkg:      # already classic
            out.append(copy.deepcopy(pkg))
            continue
        inner = {
            "name": pkg.get("name", ""),
            "type": _PBIR_PACKAGE_TYPE_INV.get(pkg.get("type"), 1),
            "items": [{"type": _PBIR_ITEM_TYPE_INV.get(it.get("type"), 100),
                       "path": it.get("path", it.get("name", "")),
                       "name": it.get("name", "")}
                      for it in (pkg.get("items") or [])],
            "disabled": False,
        }
        out.append({"resourcePackage": inner})
    return out


def _classic_packages_to_pbir(pkgs: list, original: list | None = None) -> list:
    """Inverse of :func:`_pbir_packages_to_classic`.

    ``original`` is the report.json list this came from; keys PBIR carries that
    the classic shape cannot express are preserved per package/item name, so a
    read/write cycle does not strip them.
    """
    orig_pkgs = {p.get("name"): p for p in (original or [])
                 if isinstance(p, dict) and "resourcePackage" not in p}
    out = []
    for pkg in pkgs or []:
        inner = pkg.get("resourcePackage", pkg)
        name = inner.get("name", "")
        base = copy.deepcopy(orig_pkgs.get(name, {}))
        orig_items = {i.get("name"): i for i in (base.get("items") or [])}
        ptype = inner.get("type")
        base["name"] = name
        base["type"] = (ptype if isinstance(ptype, str)
                        else _PBIR_PACKAGE_TYPE.get(ptype, "RegisteredResources"))
        items = []
        for it in (inner.get("items") or []):
            iname = it.get("name", "")
            item = copy.deepcopy(orig_items.get(iname, {}))
            itype = it.get("type")
            item["name"] = iname
            item["path"] = it.get("path", iname)
            item["type"] = (itype if isinstance(itype, str)
                            else _PBIR_ITEM_TYPE.get(itype, "Image"))
            items.append(item)
        base["items"] = items
        base.pop("disabled", None)   # classic-only
        out.append(base)
    return out

# Classic Report/Layout stores a page's scaling mode as an int; PBIR page.json
# stores the enum name. Ordering is the ordinal order of the anyOf list in
# .../report/definition/page/2.1.0/schema.json, corroborated by the corpus:
# every Desktop-authored classic page uses 1 and every service-authored PBIR
# page uses "FitToPage".
_PBIR_DISPLAY_OPTION = {
    0: "DeprecatedDynamic",
    1: "FitToPage",
    2: "FitToWidth",
    3: "ActualSize",
    4: "ActualSizeTopLeft",
}
_PBIR_DISPLAY_OPTION_INV = {v: k for k, v in _PBIR_DISPLAY_OPTION.items()}


def _is_pbir_report_config(cfg: dict) -> bool:
    """True when this document is a PBIR report.json rather than a Layout."""
    if "sections" in cfg:
        return False
    for pkg in cfg.get("resourcePackages", []) or []:
        if "resourcePackage" in pkg:
            return False
        if isinstance(pkg.get("type"), str):
            return True
    return "$schema" in cfg and "sections" not in cfg


def _registered_resource_items(layout: dict) -> list:
    """The RegisteredResources package's item list, created if absent."""
    rp = layout.setdefault("resourcePackages", [])
    for pkg in rp:
        inner = pkg.get("resourcePackage", pkg)
        if inner.get("name") == "RegisteredResources":
            return inner.setdefault("items", [])
    if _is_pbir_report_config(layout):
        reg = {"name": "RegisteredResources",
               "type": "RegisteredResources", "items": []}
        rp.append(reg)
    else:
        reg = {"name": "RegisteredResources", "type": 1, "items": [],
               "disabled": False}
        rp.append({"resourcePackage": reg})
    return reg["items"]


def _ensure_resource_item(layout: dict, item_name: str, item_type: int = 100) -> None:
    """Idempotently register an item in RegisteredResources (path == name ==
    the bare filename, exactly as Desktop writes it)."""
    items = _registered_resource_items(layout)
    if not any(i.get("name") == item_name for i in items):
        if _is_pbir_report_config(layout):
            items.append({"name": item_name, "path": item_name,
                          "type": _PBIR_ITEM_TYPE.get(item_type, "Image")})
        else:
            items.append({"type": item_type, "path": item_name,
                          "name": item_name})


def _register_resource(info: dict, layout: dict, data: bytes, name: str,
                       ext: str, item_type: int = 100) -> str:
    """Write the bytes + declare the extension + register the layout item.

    Returns the final (possibly uniquified) item name; the caller persists the
    mutated ``layout``.
    """
    work_dir = info["work_dir"]
    item_name = _sanitize_item_name(name, ext)
    res_dir = os.path.join(work_dir, "Report", "StaticResources",
                           "RegisteredResources")
    os.makedirs(res_dir, exist_ok=True)

    # Never clobber a DIFFERENT resource that already uses this name; reuse the
    # name when the bytes are identical. Name matching is CASE-INSENSITIVE and
    # adopts the existing casing: on a case-insensitive filesystem (macOS,
    # Windows — where Desktop runs) "logo.png" and "Logo.png" are one file, so
    # keeping the caller's casing would register a layout item + visual
    # reference for a part that never lands in the .pbix.
    existing: dict = {}
    for it in _registered_resource_items(layout):
        nm = it.get("name")
        if nm:
            existing.setdefault(nm.lower(), nm)
    try:
        for entry in os.listdir(res_dir):
            existing.setdefault(entry.lower(), entry)
    except OSError:
        pass

    stem, dot_ext = os.path.splitext(item_name)
    n = 0
    while True:
        match = existing.get(item_name.lower())
        if match is None:
            break                                  # free name
        target = _safe_join(res_dir, match)
        if os.path.isfile(target):
            with open(target, "rb") as f:
                if f.read() == data:
                    item_name = match              # identical bytes — reuse
                    break
        elif not os.path.exists(target):
            item_name = match      # registered but absent: rewrite that item
            break
        n += 1                     # different bytes (or a directory) — uniquify
        item_name = f"{stem}_{n}{dot_ext}"

    with open(_safe_join(res_dir, item_name), "wb") as f:
        f.write(data)
    _ensure_content_type_default(work_dir, ext)
    _ensure_resource_item(layout, item_name, item_type)
    return item_name


def _image_url_object(item_name: str) -> list:
    """objects.general holding the ImageUrl ResourcePackageItem expr."""
    return [{"properties": {"imageUrl": {"expr": {"ResourcePackageItem": {
        "PackageName": "RegisteredResources",
        "PackageType": 1,
        "ItemName": item_name,
    }}}}}]


def _image_scaling_object(scaling: str) -> list:
    """objects.imageScaling for 'Fit' | 'Fill' | 'Normal'."""
    return [{"properties": {"imageScalingType": {
        "expr": {"Literal": {"Value": f"'{scaling}'"}}}}}]


def _next_layer_z(page: dict) -> int:
    """Desktop assigns z in 1000-steps per insert (corpus-verified)."""
    zs = []
    for vc in page.get("visualContainers", []):
        try:
            zs.append(float(vc.get("z", 0) or 0))
        except (TypeError, ValueError):
            continue
    return int(max(zs) + 1000) if zs else 0


def _resolve_image_source(image_path: str, image_base64: str,
                         item_type: int = 100) -> tuple[bytes, str]:
    """Bytes + sniffed extension, with a typed error on unusable payloads."""
    data = _load_resource_bytes(image_path, image_base64)
    ext = _sniff_resource_ext(data, item_type)
    if ext is None:
        if item_type == 100:
            raise LayoutParseError(
                "Unrecognized image data — supported: PNG, JPEG, GIF, WebP, "
                "BMP, TIFF, ICO, SVG (detected from the file contents, not "
                "the name).")
        raise LayoutParseError(
            "Shape maps and themes must be JSON (Desktop stores them as "
            "JSON resources) — the supplied data is not valid JSON.")
    return data, ext


@mcp.tool()
def pbix_register_resource(
    alias: str, name: str, image_path: str = "", image_base64: str = "",
    resource_type: str = "image",
) -> str:
    """Register a file resource (image, shape map, theme) in the report.

    Writes the bytes to ``Report/StaticResources/RegisteredResources/``,
    declares the extension in ``[Content_Types].xml``, and adds the item to
    the layout's ``resourcePackages`` — the three touchpoints Power BI Desktop
    uses. Returns the final item name; point an image visual at it with
    pbix_set_image (or use pbix_add_image, which registers AND places in one
    call).

    The file TYPE comes from the CONTENT, never from the name or a caller's
    claim: images (type 100) must be PNG/JPEG/GIF/WebP/BMP/TIFF/ICO/SVG, and
    shape maps / themes (200/201/202) must be JSON — the form Desktop stores
    them in. 5 MB max. Item names are sanitized to
    ``[A-Za-z0-9._-]``, forced to the sniffed extension, and uniquified rather
    than overwriting a different existing resource.

    Args:
        alias: The alias of the open file
        name: Desired item name (its extension is replaced with the sniffed one)
        image_path: Local file to read (exactly one of image_path/image_base64)
        image_base64: Base64-encoded bytes — for callers holding an upload or a
            data URI (a full ``data:...;base64,...`` string is accepted). The
            engine never fetches remote URLs itself.
        resource_type: image (100, default), shapeMap (200), customTheme (201),
            or baseTheme (202)
    """
    try:
        logger.info("pbix_register_resource name=%r type=%r", name, resource_type)
        info = _ensure_open(alias)
        item_type = _RESOURCE_ITEM_TYPES.get(
            (resource_type or "image").replace("_", "").replace(" ", "").lower())
        if item_type is None:
            raise LayoutParseError(
                f"Unknown resource_type {resource_type!r} — use one of: "
                "image, shapeMap, customTheme, baseTheme")

        data, ext = _resolve_image_source(image_path, image_base64, item_type)
        # resourcePackages lives in Report/Layout (classic) or
        # Report/definition/report.json (PBIR); the FILES land under the same
        # Report/StaticResources path either way.
        cfg = _get_report_config(info["work_dir"])
        if cfg is None:
            raise LayoutParseError(
                "No report definition found (neither Report/Layout nor "
                "Report/definition/report.json).")
        item_name = _register_resource(info, cfg, data, name or image_path,
                                       ext, item_type)
        _set_report_config(info["work_dir"], cfg)
        info["modified"] = True
        return ToolResponse.ok(
            f"Registered '{item_name}' ({len(data):,} bytes, {ext}) in "
            f"RegisteredResources (item type {item_type}).",
            data={"item_name": item_name, "bytes": len(data), "format": ext},
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_image(
    alias: str, page_index: int = 0, image_path: str = "",
    image_base64: str = "", name: str = "", x: int = 40, y: int = 40,
    width: int = 300, height: int = 200, scaling: str = "Fit",
) -> str:
    """Add an image visual to a page — registration + placement in one call.

    Registers the bytes as a report resource (see pbix_register_resource) and
    adds a visual container matching Desktop's own image insert field-for-field
    (verified against Desktop-authored corpus reports): ``howCreated``,
    1000-step ``z`` / ``tabOrder``, ``layouts[0].position``,
    ``drillFilterOtherVisuals``, the ``ImageUrl`` ResourcePackageItem expr,
    ``objects.imageScaling``, and ``vcObjects.padding`` 0D on all four sides.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        image_path: Local image file (exactly one of image_path/image_base64)
        image_base64: Base64-encoded image bytes (a full data: URI is accepted)
        name: Optional item name (default: derived from the file name)
        x: X position in pixels
        y: Y position in pixels
        width: Width in pixels
        height: Height in pixels
        scaling: "Fit" (default), "Fill", or "Normal"; empty omits the
            imageScaling object (Desktop also writes image visuals without it)
    """
    try:
        logger.info("pbix_add_image page=%d name=%r", page_index, name)
        info = _ensure_open(alias)
        scale = None
        if scaling:
            scale = _IMAGE_SCALING.get(scaling.strip().lower())
            if scale is None:
                raise LayoutParseError(
                    f"Invalid scaling {scaling!r} — use Fit, Fill, or Normal "
                    "(or empty to omit).")

        data, ext = _resolve_image_source(image_path, image_base64)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError(
                "No legacy Report/Layout found. Image authoring for the PBIR "
                "(Report/definition) format is not yet supported.")
        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")
        page = sections[page_index]

        item_name = _register_resource(info, layout, data,
                                       name or image_path, ext, 100)

        # Clamp to the page, like pbix_add_visual.
        xf = max(0.0, min(float(x), page.get("width", 1280) - float(width)))
        yf = max(0.0, min(float(y), page.get("height", 720) - float(height)))
        z = _next_layer_z(page)
        tab_order = z + 1000

        import uuid as _uuid
        objects: dict = {"general": _image_url_object(item_name)}
        if scale:
            objects["imageScaling"] = _image_scaling_object(scale)
        pad_lit = {"expr": {"Literal": {"Value": "0D"}}}
        config = {
            "name": _uuid.uuid4().hex[:20],
            "layouts": [{"id": 0, "position": {
                "x": xf, "y": yf, "z": z,
                "width": float(width), "height": float(height),
                "tabOrder": tab_order,
            }}],
            "singleVisual": {
                "visualType": "image",
                "drillFilterOtherVisuals": True,
                "objects": objects,
                "vcObjects": {"padding": [{"properties": {
                    "left": pad_lit, "top": pad_lit,
                    "right": pad_lit, "bottom": pad_lit}}]},
            },
            "howCreated": "InsertVisualButton",
        }
        page.setdefault("visualContainers", []).append({
            "x": xf, "y": yf, "z": z,
            "width": float(width), "height": float(height),
            "tabOrder": tab_order,
            "config": json.dumps(config, ensure_ascii=False),
            "filters": "[]",
        })
        _set_layout(info["work_dir"], layout)
        info["modified"] = True

        idx = len(page["visualContainers"]) - 1
        page_name = page.get("displayName", f"Page {page_index}")
        return ToolResponse.ok(
            f"Image '{item_name}' ({len(data):,} bytes, {ext}) added to "
            f"'{page_name}' at ({xf:.0f},{yf:.0f}) {width}x{height}"
            + (f", scaling {scale}" if scale else "")
            + f" (visual index {idx}).",
            data={"item_name": item_name, "visual_index": idx,
                  "visual_name": config["name"], "format": ext},
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_image(
    alias: str, page_index: int, visual_index: int, image_path: str = "",
    image_base64: str = "", item_name: str = "", name: str = "",
    scaling: str = "",
) -> str:
    """Repoint or restyle an EXISTING image visual.

    Supply new bytes (image_path / image_base64) to register a fresh resource
    and point the visual at it, and/or ``item_name`` to point at an
    already-registered item, and/or ``scaling`` to change the fit. The
    previously referenced resource is left in place — another visual may
    reference the same item.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
        visual_index: Zero-based visual index on the page (must be an image visual)
        image_path: New local image file (optional)
        image_base64: New image bytes, base64 (optional)
        item_name: Point at this already-registered item instead (optional)
        name: Item name to use when registering new bytes (optional)
        scaling: "Fit", "Fill", or "Normal" (optional; unchanged when empty)
    """
    try:
        logger.info("pbix_set_image page=%d visual=%d", page_index, visual_index)
        info = _ensure_open(alias)
        if not (image_path or image_base64 or item_name or scaling):
            raise LayoutParseError(
                "Nothing to change — provide image_path/image_base64, "
                "item_name, and/or scaling.")
        if (image_path or image_base64) and item_name:
            raise LayoutParseError(
                "Provide either new bytes (image_path/image_base64) or an "
                "existing item_name, not both — the registered name of new "
                "bytes is chosen by `name`.")
        scale = None
        if scaling:
            scale = _IMAGE_SCALING.get(scaling.strip().lower())
            if scale is None:
                raise LayoutParseError(
                    f"Invalid scaling {scaling!r} — use Fit, Fill, or Normal.")

        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No legacy Report/Layout found.")
        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")
        containers = sections[page_index].get("visualContainers", [])
        if visual_index < 0 or visual_index >= len(containers):
            raise LayoutParseError(f"Visual index {visual_index} out of range")

        vc = containers[visual_index]
        config = _parse_visual_config(vc)
        sv = config.get("singleVisual", {})
        if sv.get("visualType") != "image":
            raise LayoutParseError(
                f"Visual {visual_index} is a "
                f"'{sv.get('visualType', 'unknown')}', not an image visual.")

        changes = []
        if image_path or image_base64:
            data, ext = _resolve_image_source(image_path, image_base64)
            item_name = _register_resource(info, layout, data,
                                           name or image_path, ext, 100)
            changes.append(f"image -> '{item_name}' ({len(data):,} bytes, {ext})")
        elif item_name:
            known = {i.get("name") for i in _registered_resource_items(layout)}
            if item_name not in known:
                raise LayoutParseError(
                    f"Item '{item_name}' is not registered in "
                    "RegisteredResources. Known items: "
                    f"{', '.join(sorted(k for k in known if k)) or '(none)'}")
            changes.append(f"image -> '{item_name}'")

        objects = sv.setdefault("objects", {})
        if item_name:
            objects["general"] = _image_url_object(item_name)
        if scale:
            objects["imageScaling"] = _image_scaling_object(scale)
            changes.append(f"scaling {scale}")

        vc["config"] = json.dumps(config, ensure_ascii=False)
        _set_layout(info["work_dir"], layout)
        info["modified"] = True
        return ToolResponse.ok(
            f"Image visual {visual_index} updated: " + "; ".join(changes) + ".",
            data={"item_name": item_name or None, "visual_index": visual_index},
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()

# --- Custom visual (.pbiviz) embedding -------------------------------------
#
# A Power BI custom visual is registered in a legacy Report/Layout by three
# byte-identical touchpoints, all equal to the visual's GUID (verified against
# Desktop-authored reports and our own toolchain-built visual):
#   1. the folder    Report/CustomVisuals/<guid>/ (holding the .pbiviz verbatim)
#   2. the top-level  Layout["publicCustomVisuals"] array (a list of guid strings)
#   3. each visualContainer's singleVisual.visualType == <guid>
# The GUID MUST come from the .pbiviz manifest — never fabricated — or the three
# touchpoints diverge and Desktop silently drops the visual. `resourcePackages`
# is NOT used for custom visuals (that is for images/themes); the earlier
# type-7 registration here was non-canonical and Desktop ignored it.

# Our own bundled HTML-rendering custom visual (source in
# assets/pbix_html_visual/visual_src). Renders a `content` data role — an
# HTML/CSS/SVG string produced by a DAX measure — inside Power BI's sandboxed
# visual iframe. Built with the powerbi-visuals-tools toolchain; Desktop-verified.
_HTML_VISUAL_GUID = "pbixHtml5C3A2F1E9B7D46A8C0E1D2F3A4B5C6D7"
_HTML_VISUAL_ROLE = "content"
# AnalysisServices silently truncates a text cell past this; keep authored HTML
# comfortably under it and warn the caller before we cross the line.
_DAX_STRING_MAX = 32000


def _bundled_html_pbiviz() -> str:
    """Absolute path to the bundled PBIX HTML custom visual (.pbiviz)."""
    asset_dir = os.path.join(os.path.dirname(__file__), "assets", "pbix_html_visual")
    if os.path.isdir(asset_dir):
        for fn in sorted(os.listdir(asset_dir)):
            if fn.endswith(".pbiviz"):
                return os.path.join(asset_dir, fn)
    raise LayoutParseError(
        "Bundled PBIX HTML visual not found under assets/pbix_html_visual/. "
        "Pass pbiviz_path=... with your own .pbiviz instead."
    )


def _read_pbiviz_manifest(pbiviz_path: str) -> dict:
    """Read visual metadata (guid/name/displayName/version/apiVersion) from a
    .pbiviz archive. The GUID is authoritative — read from the package manifest,
    never generated."""
    import zipfile as _zf

    if not os.path.exists(pbiviz_path):
        raise LayoutParseError(f"File not found: {pbiviz_path}")
    if not _zf.is_zipfile(pbiviz_path):
        raise LayoutParseError("Not a valid .pbiviz file (not a ZIP archive)")

    with _zf.ZipFile(pbiviz_path, "r") as zf:
        names = zf.namelist()
        manifest = None
        # package.json carries the canonical visual{} block in a toolchain build;
        # fall back to the resources/<guid>.pbiviz.json bundle if absent.
        for cand in ["package.json"] + [n for n in names if n.endswith("pbiviz.json")]:
            if cand in names:
                try:
                    m = json.loads(zf.read(cand))
                except (json.JSONDecodeError, KeyError):
                    continue
                if isinstance(m, dict) and isinstance(m.get("visual"), dict) and m["visual"].get("guid"):
                    manifest = m
                    break
        if not manifest:
            raise LayoutParseError(
                "No visual.guid found in the .pbiviz manifest "
                f"(package.json / *.pbiviz.json). Contents: {names[:10]}"
            )

    v = manifest["visual"]
    guid = v["guid"]
    if not guid or not guid.replace("_", "").isalnum():
        raise LayoutParseError(f"Invalid custom visual GUID in manifest: {guid!r}")
    return {
        "guid": guid,
        "name": v.get("name", guid),
        "display_name": v.get("displayName", v.get("name", guid)),
        "version": v.get("version", manifest.get("version", "1.0.0.0")),
        "api_version": manifest.get("apiVersion", "5.11.0"),
    }


def _embed_custom_visual(work_dir: str, layout: dict, pbiviz_path: str) -> dict:
    """Embed a .pbiviz verbatim into a legacy Report/Layout PBIX and register it.

    Extracts the archive into ``Report/CustomVisuals/<guid>/`` and appends
    ``<guid>`` to top-level ``layout["publicCustomVisuals"]`` (deduped). Returns
    the manifest metadata dict (incl. ``guid``). Mutates ``layout`` in place;
    the caller persists it with ``_set_layout``.
    """
    import shutil
    import zipfile as _zf

    meta = _read_pbiviz_manifest(pbiviz_path)
    guid = meta["guid"]

    # Folder == guid. `guid` is validated alnum above, but contain every write
    # to work_dir anyway (Zip-Slip / CWE-22 defence on untrusted archives).
    cv_dir = _safe_join(work_dir, "Report", "CustomVisuals", guid)
    if os.path.isdir(cv_dir):
        shutil.rmtree(cv_dir)          # replace on re-import (idempotent)
    os.makedirs(cv_dir, exist_ok=True)

    with _zf.ZipFile(pbiviz_path, "r") as zf:
        n_files = 0
        for name in zf.namelist():
            if name.endswith("/"):
                continue
            target = _safe_join(cv_dir, name)
            os.makedirs(os.path.dirname(target), exist_ok=True)
            with zf.open(name) as src, open(target, "wb") as dst:
                shutil.copyfileobj(src, dst)
            n_files += 1

    # Register the guid — the linchpin. Top-level array of guid strings.
    pcv = layout.get("publicCustomVisuals")
    if not isinstance(pcv, list):
        pcv = []
    if guid not in pcv:
        pcv.append(guid)
    layout["publicCustomVisuals"] = pcv

    meta["files"] = n_files
    return meta


@mcp.tool()
def pbix_add_custom_visual(alias: str, pbiviz_path: str) -> str:
    """Import a custom visual (.pbiviz) into the report.

    Extracts the .pbiviz package verbatim into ``Report/CustomVisuals/<guid>/``
    and registers its GUID in the layout's ``publicCustomVisuals`` — exactly how
    Power BI Desktop embeds a custom visual. The GUID is read from the .pbiviz
    manifest (never fabricated). After importing, place it with pbix_add_visual
    using the returned GUID as ``visual_type`` (or use pbix_add_html_visual for
    the bundled PBIX HTML visual, which does everything in one call).

    Args:
        alias: The alias of the open file
        pbiviz_path: Absolute path to the .pbiviz file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        cfg = _get_report_config(work_dir)
        if cfg is None:
            raise LayoutParseError(
                "No report definition found (neither Report/Layout nor "
                "Report/definition/report.json).")

        meta = _embed_custom_visual(work_dir, cfg, pbiviz_path)
        _set_report_config(work_dir, cfg)
        info["modified"] = True

        guid = meta["guid"]
        return ToolResponse.ok(
            f"Custom visual '{meta['display_name']}' imported successfully!\n"
            f"  GUID: {guid}\n"
            f"  Version: {meta['version']}  (apiVersion {meta['api_version']})\n"
            f"  Files: {meta['files']} extracted to Report/CustomVisuals/{guid}/\n"
            f"  Registered in publicCustomVisuals.\n\n"
            f"To place on a page, use:\n"
            f"  pbix_add_visual(alias, page_index, visual_type=\"{guid}\", ...)"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_reference_public_visual(alias: str, guid: str) -> str:
    """Reference a public (AppSource) custom visual by GUID — no file payload.

    Certified AppSource visuals (e.g. Deneb, GUID
    ``deneb7E15AEF80B9E4D4F8E12924291ECE89A``) are resolved by the Power BI
    service FROM APPSOURCE for report consumers — referencing one needs only
    its GUID in the layout's top-level ``publicCustomVisuals`` array. Unlike
    pbix_add_custom_visual, NOTHING is extracted into
    ``Report/CustomVisuals/`` and no .pbiviz is required: zero file parts,
    zero [Content_Types].xml changes, resourcePackages untouched
    (service-verified against app.powerbi.com on a certified-only tenant).

    OFFLINE: Desktop fetches the visual from AppSource at report open and
    caches it per-machine (ExtensionCache under Desktop's LocalAppData).
    With no network and a cold cache the report still opens — the container
    shows Desktop's unavailable-visual placeholder while everything else
    works; after one online open the cache serves the visual offline.

    After registering, place the visual with
    ``pbix_add_visual(alias, page_index, visual_type="<guid>", ...)`` — for
    Deneb, put the Vega-Lite/Vega spec in
    ``config.singleVisual.objects.vega`` (string-Literal properties
    ``provider``/``version``/``jsonSpec``/``jsonConfig``) and bind fields to
    its single ``dataset`` role. De-register with pbix_remove_custom_visual
    (the folder branch is a no-op for reference-only registrations).

    Args:
        alias: The alias of the open file
        guid: The visual's marketplace GUID, used VERBATIM (the service
            resolves certified visuals by exact GUID — no normalization).
            Letters/digits/underscores/hyphens, e.g.
            "deneb7E15AEF80B9E4D4F8E12924291ECE89A" or the legacy
            "PBI_CV_23E12E97-..." hyphenated form.
    """
    try:
        logger.info("pbix_reference_public_visual guid=%r", guid)
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        # The manifest reader's rule widened with '-': legacy PBI_CV_<GUID>
        # marketplace ids carry hyphenated GUID segments, and a
        # publicCustomVisuals entry is a plain string, so hyphens are
        # structurally safe. The GUID is registered verbatim — never
        # normalized — because the service resolves by exact GUID.
        guid = (guid or "").strip()
        if not guid or not guid.replace("_", "").replace("-", "").isalnum():
            raise LayoutParseError(
                f"Invalid custom visual GUID: {guid!r} "
                "(letters, digits, underscores, and hyphens only)")

        cfg = _get_report_config(work_dir)
        if cfg is None:
            raise LayoutParseError(
                "No report definition found (neither Report/Layout nor "
                "Report/definition/report.json).")

        pcv = cfg.get("publicCustomVisuals")
        if not isinstance(pcv, list):
            pcv = []
        already = guid in pcv
        if not already:
            pcv.append(guid)
        cfg["publicCustomVisuals"] = pcv
        _set_report_config(work_dir, cfg)
        info["modified"] = True

        return ToolResponse.ok(
            (f"GUID '{guid}' was already registered." if already else
             f"Public visual '{guid}' registered in publicCustomVisuals.") +
            "\nThe service auto-loads certified visuals from AppSource; "
            "place one with:\n"
            f"  pbix_add_visual(alias, page_index, visual_type=\"{guid}\", ...)",
            data={"publicCustomVisuals": pcv},
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_remove_custom_visual(alias: str, visual_name: str) -> str:
    """Remove a custom visual package from the report.

    Args:
        alias: The alias of the open file
        visual_name: Name of the custom visual (from pbix_list_resources)
    """
    import shutil

    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        # `visual_name` is the GUID (folder name / publicCustomVisuals entry).
        # Remove the embedded files.
        cv_dir = os.path.join(work_dir, "Report", "CustomVisuals", visual_name)
        if os.path.isdir(cv_dir):
            shutil.rmtree(cv_dir)

        layout = _get_layout(work_dir)
        if layout:
            # De-register from publicCustomVisuals (the canonical registration).
            pcv = layout.get("publicCustomVisuals")
            if isinstance(pcv, list):
                layout["publicCustomVisuals"] = [g for g in pcv if g != visual_name]
            # Also strip any legacy non-canonical resourcePackages entry keyed on
            # the name (older embeds registered a type-7 package).
            resource_packages = layout.get("resourcePackages", [])
            if resource_packages:
                layout["resourcePackages"] = [
                    rp for rp in resource_packages
                    if rp.get("resourcePackage", rp).get("name") != visual_name
                ]
            _set_layout(work_dir, layout)

        info["modified"] = True
        return ToolResponse.ok(
            f"Custom visual '{visual_name}' removed from report "
            f"(files + publicCustomVisuals registration)."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


# --- HTML custom-visual authoring (turnkey create / view / edit) -----------
#
# The bundled PBIX HTML visual renders a `content` data-role string as HTML in
# Power BI's sandboxed visual iframe. The string is produced by a DAX measure —
# either a static HTML literal or a data-driven expression (FORMAT()/& concat).
# These tools embed the visual, author/edit the measure, and place a fully
# String-bound container in one call.


def _html_to_dax_literal(html: str) -> str:
    """Wrap a raw HTML string as a DAX string-literal expression.

    DAX string literals are double-quoted; an embedded ``"`` is escaped by
    doubling it. The caller's HTML can therefore use normal double-quoted
    attributes (``class="x"``) with no manual escaping.
    """
    return '"' + html.replace('"', '""') + '"'


def _decode_html_dax_literal(expr: str) -> str | None:
    """Inverse of :func:`_html_to_dax_literal`. Returns the HTML for a pure
    string-literal measure, or ``None`` if the expression is data-driven DAX
    (not a single quoted literal) and can't be losslessly decoded to plain HTML.
    """
    s = (expr or "").strip()
    if len(s) >= 2 and s[0] == '"' and s[-1] == '"':
        inner = s[1:-1]
        # A lone (un-doubled) quote would mean the literal ended early -> not a
        # single literal, so treat as data-driven.
        if '"' in inner.replace('""', ""):
            return None
        return inner.replace('""', '"')
    return None


def _first_model_table(info: dict) -> str:
    """First data-model table name (home for an auto-created HTML measure)."""
    from pbix_mcp.formats.model_reader import ModelReader
    model = ModelReader(info["path"], work_dir=info.get("work_dir"))
    stats = model.statistics or []
    if not stats:
        raise LayoutParseError("Model has no tables to attach the HTML measure to.")
    return str(stats[0]["TableName"])


def _resolve_model_field(info: dict, field: str) -> tuple[str, str]:
    """Resolve ``field`` to a (table, column) pair. Accepts ``Table[Column]``,
    ``Table.Column``, or a bare ``Column`` (unique match across the model)."""
    from pbix_mcp.formats.model_reader import ModelReader
    schema = ModelReader(info["path"], work_dir=info.get("work_dir")).schema or []
    raw = field.strip()
    table_hint = None
    col = raw
    if "[" in raw and raw.endswith("]"):
        table_hint, col = raw[:raw.index("[")].strip().strip("'"), raw[raw.index("[") + 1:-1].strip()
    elif "." in raw:
        table_hint, col = raw.rsplit(".", 1)
        table_hint = table_hint.strip().strip("'")
        col = col.strip()
    matches = [(r["TableName"], r["ColumnName"]) for r in schema
               if r["ColumnName"] == col and (table_hint is None or r["TableName"] == table_hint)]
    if not matches:
        raise LayoutParseError(
            f"category_field '{field}' not found. Use Table[Column] / Table.Column / Column.")
    if len({m[0] for m in matches}) > 1:
        tables = ", ".join(sorted({m[0] for m in matches}))
        raise LayoutParseError(
            f"category_field '{field}' is ambiguous across tables ({tables}); "
            "qualify it as Table[Column].")
    return matches[0]


def _iter_html_visuals(layout: dict):
    """Yield (page_index, visual_index, container, single_visual) for every
    embedded PBIX HTML visual in a legacy Report/Layout."""
    for pi, sec in enumerate(layout.get("sections", [])):
        for vi, vc in enumerate(sec.get("visualContainers", [])):
            try:
                cfg = json.loads(vc.get("config", "{}"))
            except (json.JSONDecodeError, TypeError):
                continue
            sv = cfg.get("singleVisual", {})
            if sv.get("visualType") == _HTML_VISUAL_GUID:
                yield pi, vi, vc, sv


def _html_measure_ref(single_visual: dict) -> tuple[str | None, str | None]:
    """Extract (table, measure_name) bound to the content role of an HTML visual."""
    proto = single_visual.get("prototypeQuery", {})
    from_map = {f.get("Name"): f.get("Entity") for f in proto.get("From", [])}
    for sel in proto.get("Select", []):
        m = sel.get("Measure")
        if m:
            src = m.get("Expression", {}).get("SourceRef", {})
            entity = from_map.get(src.get("Source"), src.get("Entity"))
            return entity, m.get("Property")
    return None, None


@mcp.tool()
def pbix_add_html_visual(
    alias: str,
    page_index: int = 0,
    html: str = "",
    dax: str = "",
    x: float = 40,
    y: float = 40,
    width: float = 480,
    height: float = 320,
    measure_name: str = "",
    measure_table: str = "",
    css: str = "",
    pbiviz_path: str = "",
    template: str = "",
    template_spec_json: str = "",
    category_field: str = "",
) -> str:
    """Create a custom HTML / CSS / SVG visual on a report page (turnkey).

    The one-call path to a fully-rendered custom visual: it (1) embeds the
    bundled "PBIX HTML" custom visual (or your own ``pbiviz_path``), (2) authors a
    DAX measure whose string value IS the HTML, and (3) places a fully data-bound
    visual container that renders it. Power BI Desktop renders arbitrary HTML +
    inline ``<style>`` CSS + inline ``<svg>`` + inline ``<script>`` inside its
    sandboxed visual iframe — build KPI cards, SVG charts / gauges / maps, badges,
    custom tables, etc. External URLs are blocked by the sandbox; embed images as
    base64 ``data:`` URIs.

    Provide EITHER ``html`` (a raw HTML string — double-quotes are escaped for
    you) OR ``dax`` (a full DAX string expression, for data-driven HTML via
    ``FORMAT()`` / ``&`` concatenation and ``SELECTEDVALUE`` context). ``css``,
    when given with ``html``, is inlined as a leading ``<style>`` block.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index to place the visual on
        html: Raw HTML/CSS/SVG string to render (mutually exclusive with ``dax``)
        dax: Full DAX string expression producing the HTML (for data injection)
        x: X position in report px
        y: Y position in report px
        width: Width in report px
        height: Height in report px
        measure_name: Name for the created content measure (auto-named if empty)
        measure_table: Table to hold the measure (first model table if empty)
        css: Optional CSS inlined as a leading ``<style>`` block (used with ``html``)
        pbiviz_path: Optional path to your own HTML-rendering .pbiviz to embed
        template: Optional built-in template name (kpi_card / bar_chart / gauge /
            table / progress / badge) — rendered into ``html`` for you. List them
            with pbix_html_template().
        template_spec_json: JSON spec for ``template`` (e.g.
            ``{"title":"Sales","value":"1.2M","spark":[3,5,4,8]}``)
        category_field: Optional column (``Table[Column]`` / ``Table.Column`` /
            ``Column``) that makes the visual CROSS-FILTER the rest of the report,
            like a native visual. Tag clickable elements in your HTML/SVG with
            ``data-pbix-select="<value>"`` (the category value); clicking one
            selects that value and filters every other visual bound to the same
            field. Ctrl/Cmd-click multi-selects; clicking the background clears.
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        # A named template renders into `html` (escaping-safe, professional).
        if template:
            from pbix_mcp import html_templates
            try:
                spec = json.loads(template_spec_json) if template_spec_json else {}
            except json.JSONDecodeError as e:
                raise LayoutParseError(f"Invalid template_spec_json: {e}")
            if not isinstance(spec, dict):
                raise LayoutParseError("template_spec_json must be a JSON object.")
            try:
                html = html_templates.render(template, spec)
            except ValueError as e:
                raise LayoutParseError(str(e))

        if bool(html) == bool(dax):
            raise LayoutParseError(
                "Provide exactly one content source: `html` (or `template`) or `dax`.")

        # Validate the target page up front — authoring a measure is an
        # expensive ABF rebuild, so never orphan one on a bad page index. The
        # measure-add below touches only the DataModel, not Report/Layout, so
        # this `layout` object stays valid across it.
        layout = _get_layout(work_dir)   # classic or PBIR
        if not layout:
            raise LayoutParseError("No report layout found.")
        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(
                f"Page index {page_index} out of range (report has "
                f"{len(sections)} page(s))")

        # 1. DAX expression for the content measure.
        if dax:
            expr = dax
        else:
            full_html = (f"<style>{css}</style>" if css else "") + html
            if len(full_html) > _DAX_STRING_MAX:
                raise LayoutParseError(
                    f"HTML is {len(full_html)} chars; Power BI silently truncates a "
                    f"text cell past ~{_DAX_STRING_MAX}. Trim it or split the content."
                )
            expr = _html_to_dax_literal(full_html)

        table = measure_table or _first_model_table(info)

        # Auto-name the measure uniquely if not supplied.
        if not measure_name:
            n_existing = sum(1 for _ in _iter_html_visuals(layout))
            measure_name = f"HTML Visual {n_existing + 1}"

        # 2. Author the measure (full ABF rebuild -> any length supported).
        add_res = json.loads(pbix_datamodel_add_measure(
            alias, table, measure_name, expr,
            description="HTML content for a PBIX HTML custom visual"))
        if not add_res.get("success", False):
            return ToolResponse.error(
                f"Failed to author HTML measure '{measure_name}' on table '{table}': "
                f"{add_res.get('message') or add_res.get('error')}",
                "MEASURE_ADD_FAILED").to_text()

        # 3. Embed the custom visual + place the bound container.
        meta = _embed_custom_visual(
            work_dir, layout, pbiviz_path or _bundled_html_pbiviz())
        guid = meta["guid"]

        page = sections[page_index]
        page_w = page.get("width", 1280)
        page_h = page.get("height", 720)
        xf = max(0.0, min(float(x), page_w - float(width)))
        yf = max(0.0, min(float(y), page_h - float(height)))

        import uuid as _uuid
        proto_from = [{"Name": "t", "Entity": table, "Type": 0}]
        proto_select = [{
            "Measure": {
                "Expression": {"SourceRef": {"Source": "t"}},
                "Property": measure_name,
            },
            "Name": "C",
        }]
        projections = {_HTML_VISUAL_ROLE: [{"queryRef": "C"}]}

        # Optional cross-filter category: bind a column so the visual receives
        # per-value selection identities (wired in the visual to data-pbix-select).
        cat_table = cat_col = None
        if category_field:
            cat_table, cat_col = _resolve_model_field(info, category_field)
            src = "t" if cat_table == table else "c"
            if src == "c":
                proto_from.append({"Name": "c", "Entity": cat_table, "Type": 0})
            proto_select.append({
                "Column": {
                    "Expression": {"SourceRef": {"Source": src}},
                    "Property": cat_col,
                },
                "Name": "cat",
            })
            projections["category"] = [{"queryRef": "cat"}]

        single_visual = {
            "visualType": guid,
            "projections": projections,
            "prototypeQuery": {
                "Version": 2,
                "From": proto_from,
                "Select": proto_select,
            },
            "drillFilterOtherVisuals": True,
        }
        # Carry the position in config.layouts too — modern Power BI Desktop reads
        # the visual position from here for custom visuals; a container without it
        # can fault the whole report load. (Matches Desktop-authored custom visuals.)
        # Desktop's 1000-step z / tabOrder stamping (ledger issues-3; same
        # convention as pbix_add_image and Desktop-authored corpus files).
        z_new = _next_layer_z(page)
        tab_order = z_new + 1000
        config = {
            "name": _uuid.uuid4().hex[:16],
            "layouts": [{"id": 0, "position": {
                "x": xf, "y": yf, "z": z_new,
                "width": float(width), "height": float(height),
                "tabOrder": tab_order,
            }}],
            "singleVisual": single_visual,
        }

        # The content measure MUST bind as String (underlyingType 1 /
        # queryMetadata.Type 2048) — Desktop-verified. Override just this measure;
        # defer everything else to the model resolver.
        base_resolver = _report_type_resolver(info)

        def _res(entity, prop, is_measure):
            if is_measure and prop == measure_name:
                return "String"
            return base_resolver(entity, prop, is_measure)

        from pbix_mcp.report_binding import compile_visual_binding
        q, dt = compile_visual_binding(single_visual, _res)

        container = {
            "x": xf, "y": yf, "z": z_new,
            "width": float(width), "height": float(height),
            "tabOrder": tab_order,
            "config": json.dumps(config, ensure_ascii=False),
            "filters": "[]",
        }
        if q is not None:
            container["query"] = json.dumps(q, ensure_ascii=False)
            container["dataTransforms"] = json.dumps(dt, ensure_ascii=False)

        page.setdefault("visualContainers", []).append(container)
        _set_layout(work_dir, layout)
        info["modified"] = True

        idx = len(page["visualContainers"]) - 1
        page_name = page.get("displayName", f"Page {page_index}")
        return ToolResponse.ok(
            f"HTML visual placed on '{page_name}' (visual index {idx}).\n"
            f"  Custom visual: {meta['display_name']} ({guid})\n"
            f"  Content measure: '{measure_name}' on table '{table}'\n"
            f"  Position: ({xf:.0f},{yf:.0f}) {float(width):.0f}x{float(height):.0f}\n"
            f"  View with pbix_get_html_visual; edit with pbix_set_html_visual."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_get_html_visual(alias: str, page_index: int = -1) -> str:
    """List the HTML custom visuals in the report and their content.

    Returns each PBIX HTML visual with its page/visual index, position, the DAX
    measure that feeds it (table + name), the raw measure expression, and — when
    the measure is a plain HTML literal — the decoded HTML string.

    Args:
        alias: The alias of the open file
        page_index: Restrict to one zero-based page (default -1 = all pages)
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No legacy Report/Layout found.")

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        expr_by_key = {(m["TableName"], m["Name"]): m.get("Expression", "")
                       for m in (model.dax_measures or [])}
        expr_by_name = {m["Name"]: m.get("Expression", "")
                        for m in (model.dax_measures or [])}

        visuals = []
        for pi, vi, vc, sv in _iter_html_visuals(layout):
            if page_index >= 0 and pi != page_index:
                continue
            table, mname = _html_measure_ref(sv)
            expr = expr_by_key.get((table, mname)) or expr_by_name.get(mname, "")
            decoded = _decode_html_dax_literal(expr)
            visuals.append({
                "page_index": pi,
                "visual_index": vi,
                "position": {"x": vc.get("x"), "y": vc.get("y"),
                             "width": vc.get("width"), "height": vc.get("height")},
                "measure_table": table,
                "measure_name": mname,
                "dax_expression": expr,
                "html": decoded,
                "data_driven": decoded is None,
            })

        return ToolResponse.ok(
            message=f"{len(visuals)} HTML visual(s) in the report.",
            data={"count": len(visuals), "visuals": visuals},
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_set_html_visual(
    alias: str,
    page_index: int = 0,
    visual_index: int = -1,
    html: str = "",
    dax: str = "",
    css: str = "",
    measure_name: str = "",
) -> str:
    """Edit an existing HTML custom visual's content (updates its DAX measure).

    Locate the visual either by ``page_index`` + ``visual_index`` (as reported by
    pbix_get_html_visual), or by the ``measure_name`` it is bound to. Provide the
    new content as EITHER ``html`` (raw string, escaped for you) OR ``dax`` (full
    DAX string expression). The visual container itself is untouched — only the
    measure's expression changes, so position/size/binding are preserved.

    Args:
        alias: The alias of the open file
        page_index: Page of the target visual (used with ``visual_index``)
        visual_index: Visual index on the page (default -1 = first HTML visual on
            the page, or the one matching ``measure_name``)
        html: New raw HTML/CSS/SVG (mutually exclusive with ``dax``)
        dax: New full DAX string expression (for data injection)
        css: Optional CSS inlined as a leading ``<style>`` block (used with ``html``)
        measure_name: Target by bound measure name instead of visual index
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No legacy Report/Layout found.")

        if bool(html) == bool(dax):
            raise LayoutParseError("Provide exactly one of `html` or `dax`.")

        # Resolve the target measure.
        target_measure = None
        target_table = None
        candidates = list(_iter_html_visuals(layout))
        if measure_name:
            for pi, vi, vc, sv in candidates:
                t, m = _html_measure_ref(sv)
                if m == measure_name:
                    target_table, target_measure = t, m
                    break
        else:
            for pi, vi, vc, sv in candidates:
                if pi != page_index:
                    continue
                if visual_index < 0 or vi == visual_index:
                    target_table, target_measure = _html_measure_ref(sv)
                    break

        if not target_measure:
            return ToolResponse.error(
                "No matching HTML visual found. Use pbix_get_html_visual to list "
                "the page_index/visual_index/measure_name of existing HTML visuals.",
                "HTML_VISUAL_NOT_FOUND").to_text()

        # New expression.
        if dax:
            expr = dax
        else:
            full_html = (f"<style>{css}</style>" if css else "") + html
            if len(full_html) > _DAX_STRING_MAX:
                raise LayoutParseError(
                    f"HTML is {len(full_html)} chars; Power BI silently truncates a "
                    f"text cell past ~{_DAX_STRING_MAX}. Trim it or split the content."
                )
            expr = _html_to_dax_literal(full_html)

        mod_res = json.loads(pbix_datamodel_modify_measure(
            alias, target_measure, expr))
        if not mod_res.get("success", False):
            return ToolResponse.error(
                f"Failed to update HTML measure '{target_measure}': "
                f"{mod_res.get('message') or mod_res.get('error')}",
                "MEASURE_MODIFY_FAILED").to_text()

        info["modified"] = True
        return ToolResponse.ok(
            f"HTML visual content updated (measure '{target_measure}' on table "
            f"'{target_table}')."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(str(e))


@mcp.tool()
def pbix_html_template(kind: str = "", spec_json: str = "") -> str:
    """Render a professional, escaping-safe HTML/SVG snippet for an HTML visual.

    Call with no ``kind`` to list the available templates and their spec keys.
    Otherwise returns the ready HTML in ``data.html`` — pass it straight to
    ``pbix_add_html_visual(html=...)`` (or use ``pbix_add_html_visual(template=...,
    template_spec_json=...)`` to do both in one call). All user text is HTML-escaped.

    Templates:
      - kpi_card   {title, value, subtitle?, accent?, spark?[numbers]}
      - bar_chart  {title, items:[[label,value],...], accent?, value_suffix?}
      - gauge      {title, percent, accent?, center_label?}
      - table      {headers:[...], rows:[[...],...], accent?, align_right_from?}
      - progress   {title, items:[[label,percent],...], accent?}
      - badge      {text, color?, filled?}

    Args:
        kind: Template name (empty = list the catalog)
        spec_json: JSON object with the template's parameters
    """
    try:
        from pbix_mcp import html_templates
        if not kind:
            catalog = {k: v[1] for k, v in html_templates.TEMPLATES.items()}
            return ToolResponse.ok(
                message="Available HTML templates (call with kind + spec_json).",
                data={"templates": catalog}).to_text()
        try:
            spec = json.loads(spec_json) if spec_json else {}
        except json.JSONDecodeError as e:
            return ToolResponse.error(f"Invalid spec_json: {e}", "BAD_SPEC").to_text()
        if not isinstance(spec, dict):
            return ToolResponse.error("spec_json must be a JSON object.", "BAD_SPEC").to_text()
        try:
            html = html_templates.render(kind, spec)
        except ValueError as e:
            return ToolResponse.error(str(e), "BAD_TEMPLATE").to_text()
        return ToolResponse.ok(
            message=f"Rendered '{kind}' ({len(html)} chars).",
            data={"html": html}).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_svg_measure(kind: str = "", spec_json: str = "", alias: str = "",
                     measure_table: str = "", measure_name: str = "") -> str:
    """Generate DAX for an SVG data-URI image measure — optionally add it.

    Emits a DAX expression evaluating to ``data:image/svg+xml;utf8,<svg …>``.
    With ``DataCategory='ImageUrl'`` such a measure renders as a live vector
    image in table/matrix cells — in Power BI Desktop AND the service
    (service-verified), PDF export, and subscriptions — with zero custom
    visuals; the SVG recomputes under any filter context. Colors are
    percent-encoded (%23…), utf8 (never base64), and numeric interpolation is
    locale-proof.

    Call with no ``kind`` to list the available templates. Templates:
    data_bar (proportional bar), bullet (bar + target tick), pill (badge with
    a DAX text expression), icon_updown (arrow by sign), sparkline (polyline
    of a measure per category column). Dynamic parts of a spec are DAX
    sub-expressions (e.g. "value": "[Total Revenue]"); styling parts are
    plain values.

    Args:
        kind: Template name (empty = list templates)
        spec_json: JSON object with the template's parameters, e.g.
            '{"value": "[Total Revenue]", "max_value": "CALCULATE([Total Revenue], ALL(Sales))"}'
        alias: Optional — with measure_name, also ADD the measure to this open
            file (with DataCategory='ImageUrl') instead of just returning DAX
        measure_table: Optional home table for the added measure
            (default: first table)
        measure_name: Optional name for the added measure
    """
    try:
        logger.info("pbix_svg_measure kind=%r add=%s", kind, bool(measure_name))
        from pbix_mcp import svg_measures
        if not kind:
            return ToolResponse.ok(
                message=("SVG measure templates. Provide `kind` + `spec_json`; "
                         "add `alias` + `measure_name` to author the measure "
                         "directly (DataCategory=ImageUrl)."),
                data={"templates": {k: v[1] for k, v in svg_measures.TEMPLATES.items()}},
            ).to_text()
        try:
            spec = json.loads(spec_json) if spec_json else {}
        except json.JSONDecodeError as e:
            return ToolResponse.error(f"Invalid spec_json: {e}", "BAD_SPEC").to_text()
        if not isinstance(spec, dict):
            return ToolResponse.error("spec_json must be a JSON object.", "BAD_SPEC").to_text()
        try:
            dax = svg_measures.render(kind, spec)
        except (ValueError, TypeError) as e:
            return ToolResponse.error(str(e), "BAD_TEMPLATE").to_text()
        if len(dax) > _DAX_STRING_MAX:
            return ToolResponse.error(
                f"Generated DAX is {len(dax)} chars; Power BI silently "
                f"truncates text past ~{_DAX_STRING_MAX}.", "BAD_TEMPLATE").to_text()

        added = False
        if alias or measure_name:
            if not (alias and measure_name):
                return ToolResponse.error(
                    "To add the measure, provide BOTH alias and measure_name.",
                    "BAD_SPEC").to_text()
            info = _ensure_open(alias)
            home = measure_table or _first_model_table(info)
            add_res = json.loads(pbix_datamodel_add_measure(
                alias, home, measure_name, dax,
                description=f"SVG image measure ({kind})",
                data_category="ImageUrl"))
            if not add_res.get("success", False):
                return ToolResponse.error(
                    f"Measure add failed: {add_res.get('message')}",
                    add_res.get("error_code") or "MEASURE_ADD_FAILED").to_text()
            added = True
            measure_table = home

        msg = f"Rendered '{kind}' DAX ({len(dax)} chars)."
        if added:
            msg += (f" Added measure '{measure_name}' on '{measure_table}' "
                    f"with DataCategory=ImageUrl.")
        return ToolResponse.ok(
            message=msg,
            data={"dax": dax, "chars": len(dax), "added": added}).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# A custom report theme is applied as a customTheme OVERLAY on a valid built-in
# baseTheme. Power BI resolves the built-in by name (no base-theme JSON needs to
# ship); the overlay in RegisteredResources supplies the palette. Verified
# against real Power BI Desktop 2.152.
_BUILTIN_BASE_THEME = "CY24SU10"
_BUILTIN_BASE_VERSION = "5.63"


@mcp.tool()
def pbix_get_theme(alias: str) -> str:
    """Get the current report theme JSON.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        # Custom themes live in RegisteredResources; the built-in base (if any) in
        # SharedResources/BaseThemes. Read both, custom first.
        theme_dirs = [
            os.path.join(work_dir, "Report", "StaticResources", "RegisteredResources"),
            os.path.join(work_dir, "Report", "StaticResources", "SharedResources", "BaseThemes"),
        ]
        themes = []
        seen = set()
        for theme_dir in theme_dirs:
            if not os.path.isdir(theme_dir):
                continue
            for f in sorted(os.listdir(theme_dir)):
                if not f.endswith(".json") or f in seen:
                    continue
                fp = os.path.join(theme_dir, f)
                try:
                    with open(fp, "r", encoding="utf-8") as fh:
                        theme = json.load(fh)
                except (json.JSONDecodeError, OSError):
                    continue
                if "dataColors" not in theme and "visualStyles" not in theme:
                    continue  # skip non-theme JSON in RegisteredResources
                seen.add(f)
                themes.append(f"Theme file: {f}\n{json.dumps(theme, indent=2, ensure_ascii=False)}")
        if not themes:
            return ToolResponse.ok("No theme JSON files found.").to_text()
        return ToolResponse.ok("\n\n".join(themes)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_theme(alias: str, theme_json: str, filename: str = "CY24SU11.json") -> str:
    """Set the report theme JSON.

    Writes to both BaseThemes and RegisteredResources if the theme file
    exists in RegisteredResources (custom themes used by the report).

    Args:
        alias: The alias of the open file
        theme_json: Complete theme JSON string
        filename: Theme filename (default: CY24SU11.json)
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        try:
            theme = json.loads(theme_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid JSON: {e}")

        written_to = []

        # A custom theme must be registered as a customTheme OVERLAY on a valid
        # built-in baseTheme — NOT by overwriting baseTheme.name. Power BI reads
        # baseTheme.name as a BUILT-IN theme id, so a custom name (e.g. "Modern
        # Blue") fails to resolve and Desktop silently falls back to its default
        # palette (the report loads but chart colors are wrong). The custom theme
        # JSON belongs in RegisteredResources (item type 201); SharedResources/
        # BaseThemes is only for the built-in base. Verified against real Power BI
        # Desktop (Cars Sales / Briqlab): the theme's dataColors now apply.
        # `filename` is caller-controlled; _safe_join contains it to work_dir
        # (CWE-22/CWE-73).
        reg_dir = os.path.join(work_dir, "Report", "StaticResources", "RegisteredResources")
        os.makedirs(reg_dir, exist_ok=True)
        with open(_safe_join(reg_dir, filename), "w", encoding="utf-8") as fh:
            json.dump(theme, fh, indent=2, ensure_ascii=False)
        written_to.append("RegisteredResources")

        # Drop the stale copy an older pbix-mcp wrote into BaseThemes (it made the
        # custom theme masquerade as the base theme).
        stale = _safe_join(
            os.path.join(work_dir, "Report", "StaticResources", "SharedResources", "BaseThemes"),
            filename)
        if os.path.exists(stale):
            try:
                os.remove(stale)
            except OSError:
                pass

        # Update layout config to reference the theme
        layout = _get_layout(work_dir)
        if layout:
            rp = layout.get("resourcePackages", []) or []

            # Remove any stale SharedResources item that registered this custom
            # theme as a base theme (BaseThemes/<filename>).
            for pkg in rp:
                inner = pkg.get("resourcePackage", pkg)
                if inner.get("name") == "SharedResources":
                    inner["items"] = [
                        it for it in inner.get("items", [])
                        if it.get("path") != f"BaseThemes/{filename}"]

            # Ensure a RegisteredResources package (type 1) with the theme item
            # (type 201). The item name/path is the filename, matching customTheme.
            reg_pkg = None
            for pkg in rp:
                inner = pkg.get("resourcePackage", pkg)
                if inner.get("name") == "RegisteredResources":
                    reg_pkg = inner
                    break
            if reg_pkg is None:
                reg_pkg = {"name": "RegisteredResources", "type": 1, "items": [], "disabled": False}
                rp.append({"resourcePackage": reg_pkg})
            items = reg_pkg.setdefault("items", [])
            theme_item = next(
                (it for it in items if it.get("path") == filename or it.get("name") == filename), None)
            if theme_item is None:
                items.append({"type": 201, "path": filename, "name": filename})
            else:
                theme_item.update({"type": 201, "path": filename, "name": filename})
            layout["resourcePackages"] = rp

            # config.themeCollection: a valid built-in baseTheme + the custom
            # overlay. Power BI resolves the built-in by name (no base file
            # shipped). Also fill the report-level config keys Desktop expects.
            config_str = layout.get("config", "{}")
            try:
                config = json.loads(config_str) if isinstance(config_str, str) else (config_str or {})
            except json.JSONDecodeError:
                config = {}
            config.setdefault("version", "5.37")
            config.setdefault("activeSectionIndex", 0)
            config.setdefault("linguisticSchemaSyncVersion", 2)
            config["themeCollection"] = {
                "baseTheme": {"name": _BUILTIN_BASE_THEME, "version": _BUILTIN_BASE_VERSION, "type": 2},
                "customTheme": {"name": filename, "version": _BUILTIN_BASE_VERSION, "type": 1},
            }
            layout["config"] = json.dumps(config, ensure_ascii=False)

            _set_layout(work_dir, layout)
            written_to.append("layout config")

        info["modified"] = True
        return ToolResponse.ok(f"Theme saved to {filename} ({', '.join(written_to)})").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


def _load_theme_data_colors(work_dir: str) -> list[str]:
    """Load dataColors from the active theme (RegisteredResources first, then BaseThemes)."""
    for subdir in ("RegisteredResources", "SharedResources/BaseThemes"):
        theme_dir = os.path.join(work_dir, "Report", "StaticResources", subdir)
        if os.path.isdir(theme_dir):
            for f in os.listdir(theme_dir):
                if f.endswith(".json"):
                    with open(os.path.join(theme_dir, f)) as fh:
                        try:
                            theme = json.load(fh)
                            if "dataColors" in theme:
                                return [c.upper() for c in theme["dataColors"]]
                        except (json.JSONDecodeError, KeyError):
                            pass
    return []


def _resolve_theme_color(data_colors: list[str], color_id: int, percent: float) -> str:
    """Resolve a ThemeDataColor reference to a hex color string."""
    if color_id < len(data_colors):
        base = data_colors[color_id]
    else:
        base = "#808080"
    r, g, b = int(base[1:3], 16), int(base[3:5], 16), int(base[5:7], 16)
    if percent > 0:
        r = int(r + (255 - r) * percent)
        g = int(g + (255 - g) * percent)
        b = int(b + (255 - b) * percent)
    elif percent < 0:
        r = int(r * (1 + percent))
        g = int(g * (1 + percent))
        b = int(b * (1 + percent))
    return f"#{max(0,min(255,r)):02X}{max(0,min(255,g)):02X}{max(0,min(255,b)):02X}"


@mcp.tool()
def pbix_extract_colors(alias: str) -> str:
    """Extract all colors from the report — theme, visuals, and page backgrounds.

    Scans the theme JSON and every visual's objects/vcObjects for hex color
    values. Also resolves ThemeDataColor references (ColorId + Percent) to
    their actual rendered hex values. Returns a deduplicated list with
    locations so you know what to change for a complete recolor.

    Args:
        alias: The alias of the open file
    """
    import re
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        colors: dict[str, list[str]] = {}  # hex -> [locations]
        data_colors = _load_theme_data_colors(work_dir)

        def _add(hex_color: str, location: str):
            h = hex_color.upper()
            colors.setdefault(h, []).append(location)

        # Scan theme files
        for subdir in ("SharedResources/BaseThemes", "RegisteredResources"):
            theme_dir = os.path.join(work_dir, "Report", "StaticResources", subdir)
            if os.path.isdir(theme_dir):
                for f in os.listdir(theme_dir):
                    if f.endswith(".json"):
                        with open(os.path.join(theme_dir, f)) as fh:
                            text = fh.read()
                        for m in re.finditer(r'#[0-9A-Fa-f]{6}\b', text):
                            _add(m.group(), f"theme:{f}")

        # Scan layout — both hex literals AND ThemeDataColor refs
        layout = _get_layout(work_dir)
        if layout:
            for si, sec in enumerate(layout.get("sections", [])):
                page_name = sec.get("displayName", f"Page {si}")
                # Page-level config
                page_cfg_str = sec.get("config", "{}")
                if isinstance(page_cfg_str, str):
                    for m in re.finditer(r"'(#[0-9A-Fa-f]{6})'", page_cfg_str):
                        _add(m.group(1), f"{page_name}:pageConfig")

                for vi, vc in enumerate(sec.get("visualContainers", [])):
                    config_str = vc.get("config", "{}")
                    if isinstance(config_str, dict):
                        config_str = json.dumps(config_str)
                    config = json.loads(config_str) if isinstance(config_str, str) else config_str
                    sv = config.get("singleVisual", {}) if isinstance(config, dict) else {}
                    vtype = sv.get("visualType", "?") if isinstance(sv, dict) else "?"
                    loc = f"{page_name}[{vi}]:{vtype}"

                    # Find hex literals
                    for m in re.finditer(r"'(#[0-9A-Fa-f]{6})'", config_str):
                        _add(m.group(1), loc)

                    # Find ThemeDataColor refs (escaped JSON inside config strings)
                    for m in re.finditer(
                        r'"ThemeDataColor"\s*:\s*\{\s*"ColorId"\s*:\s*(\d+)\s*,\s*"Percent"\s*:\s*([-\d.]+)\s*\}',
                        config_str
                    ):
                        cid, pct = int(m.group(1)), float(m.group(2))
                        resolved = _resolve_theme_color(data_colors, cid, pct)
                        _add(resolved, f"{loc} [ThemeDataColor:{cid},{pct}]")

                    # Also check escaped variants (config stored as JSON string in JSON)
                    for m in re.finditer(
                        r'\\"ThemeDataColor\\"\s*:\s*\{\s*\\"ColorId\\"\s*:\s*(\d+)\s*,\s*\\"Percent\\"\s*:\s*([-\d.]+)\s*\}',
                        config_str
                    ):
                        cid, pct = int(m.group(1)), float(m.group(2))
                        resolved = _resolve_theme_color(data_colors, cid, pct)
                        _add(resolved, f"{loc} [ThemeDataColor:{cid},{pct}]")

        lines = []
        for hex_c in sorted(colors.keys()):
            locs = sorted(set(colors[hex_c]))
            lines.append(f"  {hex_c}  ({len(locs)} refs): {', '.join(locs[:8])}")

        return ToolResponse.ok(
            f"Found {len(colors)} unique colors:\n" + "\n".join(lines)
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_recolor(alias: str, color_map_json: str) -> str:
    """Global find-and-replace of colors across the entire report.

    Replaces hex colors in the theme (BaseThemes + RegisteredResources),
    the report layout (all visuals, pages, config), and page configs.
    Case-insensitive matching.

    Also converts ThemeDataColor references to direct hex Literal values
    when the resolved theme color is in the replacement map. This ensures
    ALL color references are replaced, not just hex literals.

    Args:
        alias: The alias of the open file
        color_map_json: JSON object mapping old hex colors to new ones, e.g.
            {"#0F7C7B": "#C2185B", "#1AA6A5": "#E91E63", "#E7E4D8": "#F5E6F0"}
    """
    import re
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        try:
            color_map = json.loads(color_map_json)
        except json.JSONDecodeError as e:
            raise LayoutParseError(f"Invalid color_map_json: {e}")

        # Normalize keys to uppercase
        cmap = {k.upper(): v for k, v in color_map.items()}
        total_replacements = 0
        contrast_fixes = 0

        # Load theme dataColors for resolving ThemeDataColor refs
        data_colors = _load_theme_data_colors(work_dir)

        # Auto-extend: map unmapped OLD theme palette colors to new palette.
        # Only applies when the user is doing a real recolor (not identity map).
        # Skip colors that are already target values (already correct).
        if data_colors:
            # Check if this is a real recolor (at least one key != value)
            is_real_recolor = any(k != v.upper() for k, v in cmap.items())
            if is_real_recolor:
                target_values = {v.upper() for v in cmap.values()}
                new_palette = list(dict.fromkeys(cmap.values()))
                if new_palette:
                    for i, dc in enumerate(data_colors):
                        dc_upper = dc.upper()
                        if dc_upper not in cmap and dc_upper not in target_values:
                            cmap[dc_upper] = new_palette[i % len(new_palette)]

        def _replace_hex(text: str) -> tuple[str, int]:
            count = 0
            for old, new in cmap.items():
                pattern = re.compile(re.escape(old), re.IGNORECASE)
                text, n = pattern.subn(new, text)
                count += n
            return text, count

        def _replace_theme_ref(m) -> str:
            """Replace a ThemeDataColor ref with a Literal hex if it matches the color map."""
            cid, pct = int(m.group(1)), float(m.group(2))
            resolved = _resolve_theme_color(data_colors, cid, pct).upper()
            # Check if this resolved color is in our replacement map
            new_color = cmap.get(resolved)
            if not new_color:
                # Also check close matches (ThemeDataColor percent shifts
                # produce slightly different hex than exact theme colors)
                for old_c, new_c in cmap.items():
                    if cid < len(data_colors) and data_colors[cid].upper() == old_c:
                        new_color = new_c
                        break
            if new_color:
                return f'"Literal":{{"Value":"\'{new_color}\'"}}'
            return m.group(0)  # no match, keep original

        def _replace_theme_ref_escaped(m) -> str:
            """Same but for escaped JSON (config strings inside JSON)."""
            cid, pct = int(m.group(1)), float(m.group(2))
            resolved = _resolve_theme_color(data_colors, cid, pct).upper()
            new_color = cmap.get(resolved)
            if not new_color:
                for old_c, new_c in cmap.items():
                    if cid < len(data_colors) and data_colors[cid].upper() == old_c:
                        new_color = new_c
                        break
            if new_color:
                return f'\\"Literal\\":{{\\"Value\\":\\"\'{new_color}\'\\"}}'
            return m.group(0)

        # Replace in theme files
        for subdir in ("SharedResources/BaseThemes", "RegisteredResources"):
            theme_dir = os.path.join(work_dir, "Report", "StaticResources", subdir)
            if os.path.isdir(theme_dir):
                for f in os.listdir(theme_dir):
                    if f.endswith(".json"):
                        fp = os.path.join(theme_dir, f)
                        with open(fp, "r", encoding="utf-8") as fh:
                            text = fh.read()
                        new_text, n = _replace_hex(text)
                        if n > 0:
                            # Fix theme foreground/textClasses contrast vs background
                            try:
                                theme_obj = json.loads(new_text)
                                bg = theme_obj.get("background", "#FFFFFF")
                                fg = theme_obj.get("foreground", "#000000")
                                if isinstance(bg, str) and isinstance(fg, str):
                                    bg_lum = _hex_luminance(bg)
                                    fg_lum = _hex_luminance(fg)
                                    if _contrast_ratio(bg_lum, fg_lum) < 3.0:
                                        ideal = _readable_text_color(bg)
                                        theme_obj["foreground"] = ideal
                                        # Also fix textClasses
                                        tc = theme_obj.get("textClasses", {})
                                        for cls in tc.values():
                                            if isinstance(cls, dict) and "color" in cls:
                                                cls["color"] = ideal
                                        new_text = json.dumps(theme_obj, indent=2, ensure_ascii=False)
                                        contrast_fixes += 1
                            except (json.JSONDecodeError, ValueError):
                                pass
                            with open(fp, "w", encoding="utf-8") as fh:
                                fh.write(new_text)
                            total_replacements += n

        # Replace in layout — hex colors + ThemeDataColor refs
        layout = _get_layout(work_dir)
        if layout:
            layout_str = json.dumps(layout, ensure_ascii=False)

            # Replace hex literals
            new_str, n = _replace_hex(layout_str)
            total_replacements += n

            # Replace ThemeDataColor references (non-escaped)
            prev = new_str
            new_str = re.sub(
                r'"ThemeDataColor"\s*:\s*\{\s*"ColorId"\s*:\s*(\d+)\s*,\s*"Percent"\s*:\s*([-\d.]+)\s*\}',
                _replace_theme_ref, new_str
            )
            total_replacements += (len(prev) - len(new_str)) // 10 if len(new_str) != len(prev) else 0

            # Replace ThemeDataColor references (escaped — config strings)
            prev = new_str
            new_str = re.sub(
                r'\\"ThemeDataColor\\"\s*:\s*\{\s*\\"ColorId\\"\s*:\s*(\d+)\s*,\s*\\"Percent\\"\s*:\s*([-\d.]+)\s*\}',
                _replace_theme_ref_escaped, new_str
            )
            total_replacements += (len(prev) - len(new_str)) // 10 if len(new_str) != len(prev) else 0

            # Count actual ThemeDataColor replacements properly
            theme_refs_before = len(re.findall(r'ThemeDataColor', layout_str))
            theme_refs_after = len(re.findall(r'ThemeDataColor', new_str))
            theme_replaced = theme_refs_before - theme_refs_after

            new_layout = json.loads(new_str)

            # --- Inject per-selector dataPoint entries for chart visuals ---
            # After text replacement, walk each visual and ensure chart series/categories
            # get explicit color assignments from the new theme palette.
            new_data_colors = _load_theme_data_colors(work_dir)
            # If theme palette lacks diversity, try rebuilding from original + cmap
            if new_data_colors and len(set(c.upper() for c in new_data_colors)) <= 2 and data_colors:
                rebuilt = [cmap.get(dc.upper(), dc) for dc in data_colors]
                if len(set(c.upper() for c in rebuilt)) > len(set(c.upper() for c in new_data_colors)):
                    new_data_colors = rebuilt
            # If still empty or single-color, use cmap values spread, or a default palette
            if not new_data_colors or len(set(c.upper() for c in new_data_colors)) <= 1:
                unique_targets = list(dict.fromkeys(cmap.values()))
                if len(unique_targets) >= 3:
                    new_data_colors = unique_targets
                else:
                    # Generate a gradient from the primary color
                    primary = unique_targets[0] if unique_targets else "#4E79A7"
                    pr, pg, pb = int(primary[1:3], 16), int(primary[3:5], 16), int(primary[5:7], 16)
                    new_data_colors = []
                    for pct in [0.0, 0.15, 0.30, 0.45, 0.60, 0.75, 0.85, 0.93]:
                        r = int(pr + (255 - pr) * pct)
                        g = int(pg + (255 - pg) * pct)
                        b = int(pb + (255 - pb) * pct)
                        new_data_colors.append(f"#{r:02X}{g:02X}{b:02X}")

            visuals_colored = 0
            for section in new_layout.get("sections", []):
                for vc in section.get("visualContainers", []):
                    try:
                        config = _parse_visual_config(vc)
                        sv = config.get("singleVisual", {})
                        vtype = sv.get("visualType", "")

                        # Apply themed row colors to table/matrix visuals
                        if vtype in ("tableEx", "pivotTable") and new_data_colors:
                            tbl_objects = sv.setdefault("objects", {})
                            primary = new_data_colors[0]

                            # Grid styling
                            if "grid" not in tbl_objects:
                                tbl_objects["grid"] = [{"properties": {
                                    "outlineColor": _solid_color(primary),
                                    "gridHorizontalColor": _solid_color(primary),
                                }}]

                            # Alternating row colors: primary tint (25%) and lighter tint (10%)
                            if "values" not in tbl_objects:
                                pr, pg, pb = int(primary[1:3], 16), int(primary[3:5], 16), int(primary[5:7], 16)
                                # Row primary: 25% tint of palette color
                                rp = f"#{int(pr + (255-pr)*0.75):02X}{int(pg + (255-pg)*0.75):02X}{int(pb + (255-pb)*0.75):02X}"
                                # Row secondary: 10% tint (very light)
                                rs = f"#{int(pr + (255-pr)*0.90):02X}{int(pg + (255-pg)*0.90):02X}{int(pb + (255-pb)*0.90):02X}"
                                # Text colors: readable against each row bg
                                fp = _readable_text_color(rp)
                                fs = _readable_text_color(rs)
                                tbl_objects["values"] = [{"properties": {
                                    "backColorPrimary": _solid_color(rp),
                                    "backColorSecondary": _solid_color(rs),
                                    "fontColorPrimary": _solid_color(fp),
                                    "fontColorSecondary": _solid_color(fs),
                                }}]

                            # Column headers: palette primary bg, readable text
                            if "columnHeaders" not in tbl_objects:
                                tbl_objects["columnHeaders"] = [{"properties": {
                                    "fontColor": _solid_color(_readable_text_color(primary)),
                                    "backColor": _solid_color(primary),
                                    "bold": _pbi_lit(True),
                                }}]

                            vc["config"] = json.dumps(config, ensure_ascii=False)
                            visuals_colored += 1
                            continue

                        # Only process chart visuals that render data series
                        chart_types = {
                            "clusteredBarChart", "clusteredColumnChart", "stackedBarChart",
                            "stackedColumnChart", "hundredPercentStackedBarChart",
                            "hundredPercentStackedColumnChart", "lineChart", "areaChart",
                            "stackedAreaChart", "lineClusteredColumnComboChart",
                            "lineStackedColumnComboChart", "pieChart", "donutChart",
                            "treemap", "waterfallChart", "funnel", "scatterChart",
                            "ribbonChart",
                        }
                        if vtype not in chart_types:
                            continue

                        projections = sv.get("projections", {})
                        if not projections:
                            continue

                        # Check if visual already has per-selector dataPoint entries
                        objects = sv.get("objects", {})
                        existing_dp = objects.get("dataPoint", [])
                        has_selectors = any(e.get("selector") for e in existing_dp)
                        if has_selectors:
                            continue  # Already has explicit per-selector colors

                        dp_entries = []
                        y_refs = [p.get("queryRef", "") for p in projections.get("Y", []) if p.get("queryRef")]
                        cat_refs = [p.get("queryRef", "") for p in projections.get("Category", []) if p.get("queryRef")]
                        series_refs = [p.get("queryRef", "") for p in projections.get("Series", []) if p.get("queryRef")]

                        if len(y_refs) > 1:
                            # Multi-measure chart: spread colors evenly across palette
                            n = len(y_refs)
                            spread = max(len(new_data_colors) // max(n, 1), 1) if new_data_colors else 1
                            for i, y_ref in enumerate(y_refs):
                                color = new_data_colors[(i * spread) % len(new_data_colors)] if new_data_colors else "#808080"
                                entry = {
                                    "properties": {"fill": _solid_color(color)},
                                    "selector": {"metadata": y_ref},
                                }
                                dp_entries.append(entry)

                        elif cat_refs and len(y_refs) <= 1:
                            # Category-based chart: assign colors per category value
                            cat_ref = cat_refs[0]
                            parts = cat_ref.split(".")
                            if len(parts) == 2:
                                entity, prop = parts
                                try:
                                    from pbix_mcp.formats.model_reader import ModelReader
                                    model = ModelReader(info["path"], work_dir=work_dir)
                                    td = model.get_table(entity)
                                    col_idx = td["columns"].index(prop)
                                    unique_vals = sorted(set(
                                        row[col_idx] for row in td["rows"]
                                        if row[col_idx] is not None
                                    ))
                                    n_vals = len(unique_vals)
                                    spread = max(len(new_data_colors) // max(n_vals, 1), 1) if new_data_colors else 1
                                    for i, val in enumerate(unique_vals):
                                        color = new_data_colors[(i * spread) % len(new_data_colors)] if new_data_colors else "#808080"
                                        entry = {
                                            "properties": {"fill": _solid_color(color)},
                                            "selector": {
                                                "data": [{
                                                    "scopeId": {
                                                        "Comparison": {
                                                            "ComparisonKind": 0,
                                                            "Left": {
                                                                "Column": {
                                                                    "Expression": {"SourceRef": {"Entity": entity}},
                                                                    "Property": prop,
                                                                }
                                                            },
                                                            "Right": {
                                                                "Literal": {"Value": f"'{val}'"}
                                                            },
                                                        }
                                                    }
                                                }]
                                            },
                                        }
                                        dp_entries.append(entry)
                                except Exception:
                                    pass

                        if dp_entries:
                            sv.setdefault("objects", {})["dataPoint"] = dp_entries
                            vc["config"] = json.dumps(config, ensure_ascii=False)
                            visuals_colored += 1
                    except Exception:
                        continue  # Skip visuals with parse errors

            # --- Contrast readability pass ---
            # Walk each visual and fix text vs background contrast issues.
            # If a card/chart has a dark background, text should be light; vice versa.
            for section in new_layout.get("sections", []):
                # Page background contrast
                pg_config_str = section.get("config", "{}")
                try:
                    pg_config = json.loads(pg_config_str) if isinstance(pg_config_str, str) else pg_config_str
                except Exception:
                    pg_config = {}

                for vc in section.get("visualContainers", []):
                    try:
                        config = _parse_visual_config(vc)
                        sv = config.get("singleVisual", {})
                        vtype = sv.get("visualType", "")
                        vc_objs = sv.get("vcObjects", {})
                        objects = sv.setdefault("objects", {})
                        changed = False

                        # Extract background color
                        bg_hex = None
                        bg_entries = vc_objs.get("background", [])
                        for be in bg_entries:
                            try:
                                bg_hex = be["properties"]["color"]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                            except (KeyError, TypeError, AttributeError):
                                pass

                        # Strip borders from all visuals by default during recolor.
                        # Borders are a design choice — if users want them, they
                        # can set them explicitly via pbix_format_visual.
                        border_entries = vc_objs.get("border", [])
                        for brd in border_entries:
                            brd_props = brd.get("properties", {})
                            if "show" in brd_props:
                                brd_props["show"] = _pbi_lit(False)
                                changed = True

                        # Pie/donut charts have hardcoded gray leader lines that
                        # clash with dark backgrounds. Remove background from
                        # these charts — their slices provide all the color.
                        if vtype in ("pieChart", "donutChart") and bg_entries:
                            for be in bg_entries:
                                be_props = be.get("properties", {})
                                be_props["show"] = _pbi_lit(False)
                            # Fix labels and title to be dark (page bg is typically light)
                            for te in vc_objs.get("title", []):
                                te_props = te.get("properties", {})
                                if "fontColor" in te_props:
                                    te_props["fontColor"] = _solid_color("#1A1A1A")
                            for le in objects.get("labels", []):
                                le_props = le.get("properties", {})
                                if "color" in le_props:
                                    le_props["color"] = _solid_color("#1A1A1A")
                            for le in objects.get("legend", []):
                                le_props = le.get("properties", {})
                                if "labelColor" in le_props:
                                    le_props["labelColor"] = _solid_color("#1A1A1A")
                            bg_hex = None
                            changed = True

                        if bg_hex and re.match(r'^#[0-9A-Fa-f]{6}$', bg_hex):
                            ideal_text = _readable_text_color(bg_hex)
                            bg_lum = _hex_luminance(bg_hex)

                            # Fix title font color
                            title_entries = vc_objs.get("title", [])
                            for te in title_entries:
                                props = te.get("properties", {})
                                fc = props.get("fontColor", {})
                                try:
                                    cur = fc["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                    if re.match(r'^#[0-9A-Fa-f]{6}$', cur):
                                        text_lum = _hex_luminance(cur)
                                        ratio = _contrast_ratio(bg_lum, text_lum)
                                        if ratio < 3.0:  # WCAG AA minimum for large text
                                            props["fontColor"] = _solid_color(ideal_text)
                                            changed = True
                                            contrast_fixes += 1
                                except (KeyError, TypeError, AttributeError):
                                    pass

                            # Fix subtitle font color
                            sub_entries = vc_objs.get("subTitle", [])
                            for se in sub_entries:
                                props = se.get("properties", {})
                                fc = props.get("fontColor", {})
                                try:
                                    cur = fc["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                    if re.match(r'^#[0-9A-Fa-f]{6}$', cur):
                                        text_lum = _hex_luminance(cur)
                                        ratio = _contrast_ratio(bg_lum, text_lum)
                                        if ratio < 3.0:
                                            props["fontColor"] = _solid_color(ideal_text)
                                            changed = True
                                            contrast_fixes += 1
                                except (KeyError, TypeError, AttributeError):
                                    pass

                            # Fix card calloutValue / label colors (card visuals)
                            for obj_cat in ("labels", "calloutValue", "categoryLabels"):
                                for entry in objects.get(obj_cat, []):
                                    props = entry.get("properties", {})
                                    for color_key in ("color", "fontColor", "labelColor"):
                                        fc = props.get(color_key, {})
                                        try:
                                            cur = fc["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                            if re.match(r'^#[0-9A-Fa-f]{6}$', cur):
                                                text_lum = _hex_luminance(cur)
                                                ratio = _contrast_ratio(bg_lum, text_lum)
                                                if ratio < 3.0:
                                                    props[color_key] = _solid_color(ideal_text)
                                                    changed = True
                                                    contrast_fixes += 1
                                        except (KeyError, TypeError, AttributeError):
                                            pass

                            # Card defaults: show categoryLabels, hide title (less
                            # redundant). Set readable colors for both labels and
                            # calloutValue on dark backgrounds.
                            if vtype == "card":
                                # Hide title by default (categoryLabels shows the name)
                                title_entries = vc_objs.get("title", [])
                                for te in title_entries:
                                    te_props = te.get("properties", {})
                                    if "show" not in te_props or te_props.get("show", {}).get("expr", {}).get("Literal", {}).get("Value") == "true":
                                        te_props["show"] = _pbi_lit(False)
                                        changed = True
                                # Set categoryLabels color for readability
                                if "categoryLabels" not in objects:
                                    cat_props = {"show": _pbi_lit(True)}
                                    if bg_lum < 0.25:
                                        cat_props["color"] = _solid_color(ideal_text)
                                        contrast_fixes += 1
                                    objects["categoryLabels"] = [{"properties": cat_props}]
                                    changed = True
                                if bg_lum < 0.25 and "calloutValue" not in objects:
                                    objects["calloutValue"] = [{"properties": {
                                        "color": _solid_color(ideal_text),
                                    }}]
                                    changed = True
                                    contrast_fixes += 1

                        # Also check chart axis/label colors vs chart background
                        # Skip for pie/donut — their bg was stripped above
                        chart_bg = None
                        if vtype not in ("pieChart", "donutChart"):
                            for be in bg_entries:
                                try:
                                    chart_bg = be["properties"]["color"]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                except (KeyError, TypeError, AttributeError):
                                    pass
                        if chart_bg and re.match(r'^#[0-9A-Fa-f]{6}$', chart_bg):
                            chart_bg_lum = _hex_luminance(chart_bg)
                            ideal = _readable_text_color(chart_bg)

                            # Fix existing axis/label text colors
                            for obj_cat in ("categoryAxis", "valueAxis", "legend", "dataLabels"):
                                for entry in objects.get(obj_cat, []):
                                    props = entry.get("properties", {})
                                    for color_key in ("labelColor", "fontColor", "color"):
                                        fc = props.get(color_key, {})
                                        try:
                                            cur = fc["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                            if re.match(r'^#[0-9A-Fa-f]{6}$', cur):
                                                text_lum = _hex_luminance(cur)
                                                ratio = _contrast_ratio(chart_bg_lum, text_lum)
                                                if ratio < 3.0:
                                                    props[color_key] = _solid_color(ideal)
                                                    changed = True
                                                    contrast_fixes += 1
                                        except (KeyError, TypeError, AttributeError):
                                            pass

                            # Inject text color for chart elements that have NO explicit
                            # color but would inherit unreadable theme foreground.
                            # Only for chart visuals (not cards, tables, textboxes).
                            _chart_vis_types = {
                                "clusteredBarChart", "clusteredColumnChart", "stackedBarChart",
                                "stackedColumnChart", "hundredPercentStackedBarChart",
                                "hundredPercentStackedColumnChart", "lineChart", "areaChart",
                                "stackedAreaChart", "lineClusteredColumnComboChart",
                                "lineStackedColumnComboChart", "pieChart", "donutChart",
                                "treemap", "waterfallChart", "funnel", "scatterChart",
                                "ribbonChart",
                            }
                            if vtype in _chart_vis_types and (chart_bg_lum < 0.15 or _contrast_ratio(chart_bg_lum, _hex_luminance(new_data_colors[0] if new_data_colors else "#000000")) < 3.0):
                                color_map_items = {
                                    "categoryAxis": "labelColor",
                                    "valueAxis": "labelColor",
                                    "legend": "labelColor",
                                    "labels": "color",
                                }
                                for obj_cat, color_key in color_map_items.items():
                                    if obj_cat not in objects:
                                        objects[obj_cat] = [{"properties": {
                                            color_key: _solid_color(ideal),
                                        }}]
                                        changed = True
                                        contrast_fixes += 1
                                    else:
                                        # Exists but may lack a color — add it if missing or unreadable
                                        for entry in objects[obj_cat]:
                                            eprops = entry.get("properties", {})
                                            if color_key not in eprops:
                                                eprops[color_key] = _solid_color(ideal)
                                                changed = True
                                                contrast_fixes += 1
                                            else:
                                                try:
                                                    cur = eprops[color_key]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                                    if re.match(r'^#[0-9A-Fa-f]{6}$', cur):
                                                        ratio = _contrast_ratio(chart_bg_lum, _hex_luminance(cur))
                                                        if ratio < 3.0:
                                                            eprops[color_key] = _solid_color(ideal)
                                                            changed = True
                                                            contrast_fixes += 1
                                                except (KeyError, TypeError, AttributeError):
                                                    pass

                        # Check table alternating row contrast (backColor vs fontColor pairs)
                        for val_entry in objects.get("values", []):
                            vprops = val_entry.get("properties", {})
                            for bg_key, fg_key in (
                                ("backColorPrimary", "fontColorPrimary"),
                                ("backColorSecondary", "fontColorSecondary"),
                            ):
                                try:
                                    bg_val = vprops[bg_key]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                    fg_val = vprops[fg_key]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                    if re.match(r'^#[0-9A-Fa-f]{6}$', bg_val) and re.match(r'^#[0-9A-Fa-f]{6}$', fg_val):
                                        ratio = _contrast_ratio(_hex_luminance(bg_val), _hex_luminance(fg_val))
                                        if ratio < 3.0:
                                            vprops[fg_key] = _solid_color(_readable_text_color(bg_val))
                                            changed = True
                                            contrast_fixes += 1
                                except (KeyError, TypeError, AttributeError):
                                    pass

                        # Check column header contrast (backColor vs fontColor)
                        for hdr_entry in objects.get("columnHeaders", []):
                            hprops = hdr_entry.get("properties", {})
                            try:
                                hbg = hprops["backColor"]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                hfg = hprops["fontColor"]["solid"]["color"]["expr"]["Literal"]["Value"].strip("'")
                                if re.match(r'^#[0-9A-Fa-f]{6}$', hbg) and re.match(r'^#[0-9A-Fa-f]{6}$', hfg):
                                    ratio = _contrast_ratio(_hex_luminance(hbg), _hex_luminance(hfg))
                                    if ratio < 3.0:
                                        hprops["fontColor"] = _solid_color(_readable_text_color(hbg))
                                        changed = True
                                        contrast_fixes += 1
                            except (KeyError, TypeError, AttributeError):
                                pass

                        if changed:
                            vc["config"] = json.dumps(config, ensure_ascii=False)
                    except Exception:
                        continue

            _set_layout(work_dir, new_layout)

        info["modified"] = True

        parts = [f"Replaced {total_replacements} hex color occurrences"]
        if theme_replaced > 0:
            parts.append(f"{theme_replaced} ThemeDataColor refs converted to hex")
        if visuals_colored > 0:
            parts.append(f"{visuals_colored} chart(s) got per-series/category colors")
        if contrast_fixes > 0:
            parts.append(f"{contrast_fixes} text contrast fix(es)")
        parts.append(f"({len(cmap)} colors mapped)")

        return ToolResponse.ok(" + ".join(parts)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_get_linguistic_schema(alias: str) -> str:
    """Get the Q&A linguistic schema XML.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        ls_path = os.path.join(work_dir, "Report", "LinguisticSchema")
        if not os.path.exists(ls_path):
            return ToolResponse.ok("No linguistic schema found.").to_text()
        enc = _detect_encoding(ls_path)
        with open(ls_path, "r", encoding=enc) as f:
            return ToolResponse.ok(f.read()).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_linguistic_schema(alias: str, schema_xml: str) -> str:
    """Set (replace) the Q&A linguistic schema XML.

    Args:
        alias: The alias of the open file
        schema_xml: The new linguistic schema XML content
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        ls_path = os.path.join(work_dir, "Report", "LinguisticSchema")
        os.makedirs(os.path.dirname(ls_path), exist_ok=True)
        # Write in UTF-16-LE (Power BI native)
        with open(ls_path, "wb") as f:
            f.write(schema_xml.encode("utf-16-le"))
        info["modified"] = True
        return ToolResponse.ok("Linguistic schema updated.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 6: DataMashup (M Code) tools ----

@mcp.tool()
def pbix_get_m_code(alias: str) -> str:
    """Get the Power Query M code from the DataMashup.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        m_code = _read_datamashup_m_code(info["work_dir"])
        if m_code is None:
            return ToolResponse.ok("No DataMashup found in this file.").to_text()
        return ToolResponse.ok(m_code).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_m_code(alias: str, m_code: str) -> str:
    """Set the Power Query M code in the DataMashup.

    Args:
        alias: The alias of the open file
        m_code: New M code to write into the DataMashup
    """
    try:
        info = _ensure_open(alias)
        ok = _write_datamashup_m_code(info["work_dir"], m_code)
        if not ok:
            return ToolResponse.error("Failed to write M code. DataMashup may not exist or be corrupt.", PBIXMCPError.code).to_text()
        info["modified"] = True
        return ToolResponse.ok("M code updated in DataMashup.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 7: DataModel READ tools (native ABF/VertiPaq) ----

@mcp.tool()
def pbix_get_model_schema(alias: str) -> str:
    """Get the data model schema — all tables, columns, and data types.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_schema_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        schema = model.schema
        return ToolResponse.ok(format_schema_table(schema)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_get_model_measures(alias: str) -> str:
    """Get all DAX measures from the data model.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_measures_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        measures = model.dax_measures
        return ToolResponse.ok(format_measures_table(measures)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_get_model_relationships(alias: str) -> str:
    """Get all relationships in the data model.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_relationships_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        rels = model.relationships
        return ToolResponse.ok(format_relationships_table(rels)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_get_model_power_query(alias: str) -> str:
    """Get Power Query expressions from the model.

    This reads M expressions as stored in the DataModel itself
    (different from the DataMashup M code).

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_power_query_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        pq = model.power_query
        return ToolResponse.ok(format_power_query_table(pq)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_list_data_sources(alias: str) -> str:
    """List all data sources with connection details for each table.

    Parses M expressions from Partition.QueryDefinition to extract
    connection type, server, database, table, mode, and file paths.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        pq = model.power_query

        import re
        mode_names = {0: "Import", 1: "DirectQuery"}
        lines = []
        for entry in pq:
            tname = entry.get("TableName", "")
            expr = entry.get("Expression", "")
            if not expr:
                continue

            # Parse connection type and parameters from M expression
            source_type = "Embedded"
            details = {}

            if "Sql.Database(" in expr:
                source_type = "SQL Server"
                m = re.search(r'Sql\.Database\("([^"]*)",\s*"([^"]*)"', expr)
                if m:
                    details["server"] = m.group(1)
                    details["database"] = m.group(2)
                m2 = re.search(r'Schema="([^"]*)".*?Item="([^"]*)"', expr)
                if m2:
                    details["schema"] = m2.group(1)
                    details["table"] = m2.group(2)
            elif "PostgreSQL.Database(" in expr:
                source_type = "PostgreSQL"
                m = re.search(r'PostgreSQL\.Database\("([^"]*)",\s*"([^"]*)"', expr)
                if m:
                    details["server"] = m.group(1)
                    details["database"] = m.group(2)
                m2 = re.search(r'Schema="([^"]*)".*?Item="([^"]*)"', expr)
                if m2:
                    details["schema"] = m2.group(1)
                    details["table"] = m2.group(2)
            elif "MySQL.Database(" in expr:
                source_type = "MySQL"
                m = re.search(r'MySQL\.Database\("([^"]*)",\s*"([^"]*)"', expr)
                if m:
                    details["server"] = m.group(1)
                    details["database"] = m.group(2)
                m2 = re.search(r'Schema="([^"]*)".*?Item="([^"]*)"', expr)
                if m2:
                    details["schema"] = m2.group(1)
                    details["table"] = m2.group(2)
            elif "MariaDB.Contents(" in expr:
                source_type = "MariaDB"
                m = re.search(r'MariaDB\.Contents\("([^"]*)",\s*"([^"]*)"', expr)
                if m:
                    details["server"] = m.group(1)
                    details["database"] = m.group(2)
            elif "Odbc.DataSource(" in expr and "SQLite" in expr:
                source_type = "SQLite"
                m = re.search(r'Database=([^;"\}]+)', expr)
                if m:
                    details["path"] = m.group(1)
            elif "Csv.Document(" in expr:
                source_type = "CSV"
                m = re.search(r'File\.Contents\("([^"]*)"', expr)
                if m:
                    details["path"] = m.group(1)
            elif "Excel.Workbook(" in expr:
                source_type = "Excel"
                m = re.search(r'File\.Contents\("([^"]*)"', expr)
                if m:
                    details["path"] = m.group(1)
                m2 = re.search(r'Item="([^"]*)"', expr)
                if m2:
                    details["sheet"] = m2.group(1)
            elif "Json.Document(" in expr or "Web.Contents(" in expr:
                source_type = "JSON/Web"
                m = re.search(r'Web\.Contents\("([^"]*)"', expr)
                if m:
                    details["url"] = m.group(1)
            elif "#table(" in expr:
                source_type = "Embedded"

            # Get mode from metadata
            mode_str = "Import"
            try:
                mode_rows = model._query_metadata(
                    "SELECT p.Mode FROM Partition p JOIN [Table] t ON p.TableID = t.ID "
                    "WHERE t.Name = ? AND t.ModelID = 1 "
                    "AND t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'",
                    (tname,)
                )
                if mode_rows:
                    mode_str = mode_names.get(mode_rows[0].get("Mode", 0), "Import")
            except Exception:
                pass

            detail_str = ", ".join(f"{k}={v}" for k, v in details.items())
            lines.append(f"  {tname}: {source_type} ({mode_str}){' — ' + detail_str if detail_str else ''}")

        if not lines:
            return ToolResponse.ok("No data sources found.").to_text()
        return ToolResponse.ok(f"Data sources ({len(lines)} tables):\n\n" + "\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_update_data_source(
    alias: str, table_name: str, new_source_json: str
) -> str:
    """Update a table's data source connection without full DataModel rebuild.

    Changes the M expression (Partition.QueryDefinition) and optionally the
    mode (Import/DirectQuery). This is a lightweight metadata-only operation
    that does NOT regenerate VertiPaq data.

    Args:
        alias: The alias of the open file
        table_name: Table to update
        new_source_json: JSON with new connection parameters. Examples:
            '{"server": "new-server.example.com", "database": "prod_db"}'
            '{"type": "postgresql", "server": "pg.local", "port": 5432, "database": "analytics", "table": "orders"}'
            '{"type": "csv", "path": "C:/data/sales.csv"}'
            '{"mode": "directquery"}'
            Supported types: sqlserver, postgresql, mysql, mariadb, sqlite, csv, excel, json, azuresql
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        new_source = json.loads(new_source_json)
        from pbix_mcp.builder import _build_m_expression

        def _do_update(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row

            # Find the partition for this table
            row = conn.execute(
                "SELECT p.ID, p.QueryDefinition, p.Mode, t.ID as TableID "
                "FROM Partition p JOIN [Table] t ON p.TableID = t.ID "
                "WHERE t.Name = ? AND t.ModelID = 1 "
                "AND t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'",
                (table_name,)
            ).fetchone()
            if not row:
                raise ValueError(f"Table '{table_name}' not found")

            part_id = row["ID"]
            current_mode = row["Mode"] or 0

            # Read column definitions for M expression generation
            cols = [{"name": c["ExplicitName"],
                     "data_type": {2: "String", 6: "Int64", 8: "Double",
                                   9: "DateTime", 10: "Decimal", 11: "Boolean"
                                   }.get(c["ExplicitDataType"], "String")}
                    for c in conn.execute(
                        "SELECT ExplicitName, ExplicitDataType FROM [Column] "
                        "WHERE TableID = ? AND Type = 1 ORDER BY ID",
                        (row["TableID"],)
                    )]

            # Determine new mode
            new_mode = current_mode
            if "mode" in new_source:
                new_mode = 1 if new_source["mode"] == "directquery" else 0

            # Build source_db dict for M expression generator
            source_db = None
            source_csv = None
            is_dq = new_mode == 1

            src_type = new_source.get("type", "").lower()
            if src_type in ("sqlserver", "azuresql", "azure"):
                source_db = {
                    "type": src_type if src_type != "azure" else "azuresql",
                    "server": new_source.get("server", "localhost"),
                    "database": new_source.get("database", ""),
                    "table": new_source.get("table", table_name),
                    "schema": new_source.get("schema", "dbo"),
                }
            elif src_type == "postgresql":
                source_db = {
                    "type": "postgresql",
                    "server": new_source.get("server", "localhost"),
                    "port": new_source.get("port", 5432),
                    "database": new_source.get("database", ""),
                    "table": new_source.get("table", table_name),
                    "schema": new_source.get("schema", "public"),
                }
            elif src_type == "mysql":
                source_db = {
                    "type": "mysql",
                    "server": new_source.get("server", "localhost"),
                    "port": new_source.get("port", 3306),
                    "database": new_source.get("database", ""),
                    "table": new_source.get("table", table_name),
                }
            elif src_type == "mariadb":
                source_db = {
                    "type": "mariadb",
                    "server": new_source.get("server", "localhost"),
                    "port": new_source.get("port", 3306),
                    "database": new_source.get("database", ""),
                    "table": new_source.get("table", table_name),
                }
            elif src_type == "sqlite":
                source_db = {
                    "type": "sqlite",
                    "path": new_source.get("path", ""),
                    "table": new_source.get("table", table_name),
                }
            elif src_type == "csv":
                source_csv = new_source.get("path", "")
            elif src_type == "excel":
                source_db = {
                    "type": "excel",
                    "path": new_source.get("path", ""),
                    "sheet": new_source.get("sheet", "Sheet1"),
                }
            elif src_type in ("json", "web", "api"):
                source_db = {
                    "type": "json",
                    "url": new_source.get("url", ""),
                }
            elif not src_type and ("server" in new_source or "database" in new_source):
                # Partial update — rewrite with same type, infer from current M expression
                current_qd = row["QueryDefinition"] or ""
                if "Sql.Database(" in current_qd:
                    source_db = {"type": "sqlserver"}
                elif "PostgreSQL.Database(" in current_qd:
                    source_db = {"type": "postgresql", "port": 5432, "schema": "public"}
                elif "MySQL.Database(" in current_qd:
                    source_db = {"type": "mysql", "port": 3306}
                else:
                    source_db = {"type": "sqlserver"}
                # Merge new params
                for k, v in new_source.items():
                    if k != "mode":
                        source_db[k] = v
                if "table" not in source_db:
                    source_db["table"] = table_name

            if source_db or source_csv:
                new_m = _build_m_expression(
                    table_name, cols,
                    source_csv=source_csv,
                    source_db=source_db,
                    is_directquery=is_dq,
                )
                conn.execute(
                    "UPDATE Partition SET QueryDefinition = ?, Mode = ? WHERE ID = ?",
                    (new_m, new_mode, part_id),
                )
            elif "mode" in new_source:
                # Mode-only change
                conn.execute(
                    "UPDATE Partition SET Mode = ? WHERE ID = ?",
                    (new_mode, part_id),
                )
            else:
                raise ValueError("No recognized connection parameters in new_source_json")

            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_update)
        info["modified"] = True

        src_type = new_source.get("type", "connection")
        return ToolResponse.ok(
            f"Data source updated for '{table_name}': {src_type}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes (lightweight, no rebuild)"
        ).to_text()
    except json.JSONDecodeError as e:
        return ToolResponse.error(f"Invalid JSON: {e}", "INVALID_INPUT").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_partition_m(alias: str, table_name: str, m_expression: str) -> str:
    """Set a table partition's raw Power Query M expression.

    The table-scoped complement to pbix_set_m_code (which replaces the whole
    DataMashup) and pbix_update_data_source (which builds the M from
    structured connection parameters): the expression is written to
    ``Partition.QueryDefinition`` verbatim (ledger issues-12). Metadata-only —
    the cached VertiPaq rows are untouched, so the file keeps opening with
    its current data, and Power BI runs the new M on the next Refresh.

    Args:
        alias: The alias of the open file
        table_name: The table whose partition to update
        m_expression: The complete M expression (e.g. ``let Source = ... in
            Source``); written as-is, no validation of the M itself
    """
    try:
        info = _ensure_open(alias)
        if not (m_expression or "").strip():
            return ToolResponse.error(
                "m_expression is empty — pass the complete M expression.",
                "INVALID_INPUT").to_text()
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.", DataModelCompressionError.code).to_text()

        def _do_update(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            row = conn.execute(
                "SELECT p.ID FROM Partition p JOIN [Table] t ON p.TableID = t.ID "
                "WHERE t.Name = ? AND t.ModelID = 1 "
                "AND t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'",
                (table_name,)).fetchone()
            if not row:
                raise ValueError(f"Table '{table_name}' not found")
            conn.execute(
                "UPDATE Partition SET QueryDefinition = ? WHERE ID = ?",
                (m_expression, row["ID"]))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_update)
        info["modified"] = True
        return ToolResponse.ok(
            f"Partition M set for '{table_name}' ({len(m_expression)} chars)\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes "
            "(metadata-only, cached data untouched)").to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "INVALID_INPUT").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_get_model_columns(alias: str) -> str:
    """Get all DAX calculated columns from the model.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_dax_columns_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        cols = model.dax_columns
        return ToolResponse.ok(format_dax_columns_table(cols)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_get_table_data(alias: str, table_name: str, max_rows: int = 50) -> str:
    """Get sample data from a table in the data model.

    Args:
        alias: The alias of the open file
        table_name: Name of the table to query
        max_rows: Maximum rows to return (default 50)
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally. "
                "Use layout, measure, and metadata tools instead.",
                UnsupportedFormatError.code,
            ).to_text()
        from pbix_mcp.formats.model_reader import ModelReader, format_table_data
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        table_data = model.get_table(table_name, max_rows=max_rows)
        if not table_data["columns"] or not table_data["rows"]:
            return ToolResponse.ok(f"No data found in table '{table_name}'.").to_text()
        return ToolResponse.ok(format_table_data(table_data, max_rows=max_rows)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


def _format_csv_value(val, delimiter: str = ",") -> str:
    """Format a single value for CSV output."""
    if val is None:
        return ""
    if isinstance(val, (int, float)):
        return str(val)
    from datetime import date, datetime
    if isinstance(val, (datetime, date)):
        return val.isoformat()
    s = str(val)
    # Quote if contains delimiter, quote, or newline
    if delimiter in s or '"' in s or "\n" in s or "\r" in s:
        s = s.replace('"', '""')
        return f'"{s}"'
    return s


def _is_system_table(name: str) -> bool:
    """Check if a table is a system/internal table (hidden from users)."""
    return name.startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate"))


@mcp.tool()
def pbix_export_table_csv(
    alias: str, table_name: str, output_path: str, delimiter: str = ","
) -> str:
    """Export a table's data to a CSV file.

    Writes all rows of the table (no row limit) with headers. Strings are
    quoted when they contain the delimiter, quotes, or newlines. Dates are
    formatted as ISO 8601. Works on Import files only (not DirectQuery).

    Args:
        alias: The alias of the open file
        table_name: Name of the table to export
        output_path: Absolute path for the CSV file
        delimiter: Field delimiter (default ',')
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        table_data = model.get_table(table_name, max_rows=0)

        if not table_data["columns"]:
            return ToolResponse.error(
                f"Table '{table_name}' not found.", "TABLE_NOT_FOUND"
            ).to_text()

        cols = table_data["columns"]
        rows = table_data["rows"]

        with open(output_path, "w", encoding="utf-8-sig", newline="") as f:
            f.write(delimiter.join(_format_csv_value(c, delimiter) for c in cols) + "\n")
            for row in rows:
                f.write(delimiter.join(_format_csv_value(v, delimiter) for v in row) + "\n")

        file_size = os.path.getsize(output_path)
        return ToolResponse.ok(
            f"Exported '{table_name}' to {output_path}\n"
            f"  {len(rows):,} rows × {len(cols)} columns ({file_size:,} bytes)"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "CSV_EXPORT_ERROR").to_text()


@mcp.tool()
def pbix_export_all_tables_csv(alias: str, output_dir: str) -> str:
    """Export every data table in the model to separate CSV files.

    Creates one CSV per table in the output directory. Skips system tables
    (H$, R$, U$, LocalDateTable, DateTableTemplate). Works on Import files only.

    Args:
        alias: The alias of the open file
        output_dir: Absolute path for the output directory (created if missing)
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        os.makedirs(output_dir, exist_ok=True)

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        stats = model.statistics
        data_tables = [t for t in stats if not _is_system_table(t["TableName"])]

        exported = []
        errors = []
        for t in data_tables:
            tname = t["TableName"]
            try:
                safe_name = "".join(c if c.isalnum() or c in "-_. " else "_" for c in tname)
                csv_path = os.path.join(output_dir, f"{safe_name}.csv")
                tdata = model.get_table(tname, max_rows=0)
                if not tdata["columns"]:
                    continue
                with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
                    f.write(",".join(_format_csv_value(c) for c in tdata["columns"]) + "\n")
                    for row in tdata["rows"]:
                        f.write(",".join(_format_csv_value(v) for v in row) + "\n")
                exported.append(f"  {tname}: {len(tdata['rows']):,} rows -> {safe_name}.csv")
            except Exception as e:
                errors.append(f"  {tname}: {e}")

        msg = f"Exported {len(exported)} tables to {output_dir}\n" + "\n".join(exported)
        if errors:
            msg += "\n\nErrors:\n" + "\n".join(errors)
        return ToolResponse.ok(msg).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "CSV_EXPORT_ERROR").to_text()


@mcp.tool()
def pbix_find_value(
    alias: str, search_value: str, case_sensitive: bool = False, max_matches: int = 100
) -> str:
    """Search for a value across all tables and columns in the model.

    Returns all table.column locations where the value appears, with the
    number of matching rows per column. Works on Import files only.

    Args:
        alias: The alias of the open file
        search_value: The value to search for (string comparison)
        case_sensitive: If False (default), compares case-insensitively
        max_matches: Maximum locations to report (default 100)
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        stats = model.statistics
        data_tables = [t for t in stats if not _is_system_table(t["TableName"])]

        needle = search_value if case_sensitive else search_value.lower()
        matches: list[tuple[str, str, int, list]] = []  # (table, col, count, samples)

        for t in data_tables:
            tname = t["TableName"]
            try:
                tdata = model.get_table(tname, max_rows=0)
                cols = tdata["columns"]
                for ci, cname in enumerate(cols):
                    count = 0
                    samples: list = []
                    for row in tdata["rows"]:
                        val = row[ci]
                        if val is None:
                            continue
                        s = str(val) if case_sensitive else str(val).lower()
                        if needle in s:
                            count += 1
                            if len(samples) < 3:
                                samples.append(str(val))
                    if count > 0:
                        matches.append((tname, cname, count, samples))
                        if len(matches) >= max_matches:
                            break
                if len(matches) >= max_matches:
                    break
            except Exception:
                continue

        if not matches:
            return ToolResponse.ok(f"No matches found for '{search_value}'.").to_text()

        lines = [f"Found '{search_value}' in {len(matches)} location(s):\n"]
        for tname, cname, count, samples in matches:
            sample_str = ", ".join(f"'{s}'" for s in samples[:3])
            lines.append(f"  {tname}.{cname}: {count:,} matches (e.g. {sample_str})")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "SEARCH_ERROR").to_text()


def _parse_where_clause(where: str) -> list[dict]:
    """Parse a simple SQL-like WHERE clause into conditions.

    Supports: column = 'value', column != 'value', column > N, column < N,
    column >= N, column <= N, column IN ('a', 'b'), column LIKE '%x%'
    joined by AND/OR. Returns list of {col, op, value, connector} dicts.
    """
    import re as _re
    if not where.strip():
        return []

    # Tokenize by AND/OR (preserving connectors)
    parts = _re.split(r'\s+(AND|OR)\s+', where, flags=_re.IGNORECASE)
    conditions = []
    for i, part in enumerate(parts):
        if part.upper() in ("AND", "OR"):
            continue
        connector = "AND"
        if i > 0:
            connector = parts[i - 1].upper()

        # Match: col OP value
        m = _re.match(r"\s*([a-zA-Z_][\w\s\-\.]*?)\s*(IN|LIKE|!=|>=|<=|>|<|=)\s*(.+)\s*$", part, _re.IGNORECASE)
        if not m:
            raise ValueError(f"Can't parse condition: '{part}'")
        col = m.group(1).strip()
        op = m.group(2).upper()
        val = m.group(3).strip()

        # Parse value
        if op == "IN":
            # ('a', 'b', 'c')
            val = val.strip("()")
            items = [x.strip().strip("'\"") for x in val.split(",")]
            parsed_val = items
        elif val.startswith("'") and val.endswith("'"):
            parsed_val = val[1:-1]
        elif val.startswith('"') and val.endswith('"'):
            parsed_val = val[1:-1]
        else:
            # Try number
            try:
                parsed_val = float(val) if "." in val else int(val)
            except ValueError:
                parsed_val = val
        conditions.append({"col": col, "op": op, "value": parsed_val, "connector": connector})
    return conditions


def _eval_condition(row_val, op: str, value) -> bool:
    """Evaluate a single condition against a row value."""
    if row_val is None:
        return False
    try:
        if op == "=":
            return str(row_val) == str(value)
        if op == "!=":
            return str(row_val) != str(value)
        if op == ">":
            return float(row_val) > float(value)
        if op == ">=":
            return float(row_val) >= float(value)
        if op == "<":
            return float(row_val) < float(value)
        if op == "<=":
            return float(row_val) <= float(value)
        if op == "IN":
            return str(row_val) in [str(x) for x in value]
        if op == "LIKE":
            # Convert SQL LIKE to regex
            import re as _re
            pattern = _re.escape(str(value)).replace("%", ".*").replace("_", ".")
            return bool(_re.match(f"^{pattern}$", str(row_val), _re.IGNORECASE))
    except (ValueError, TypeError):
        return False
    return False


@mcp.tool()
def pbix_query_table(
    alias: str,
    table_name: str,
    where: str = "",
    columns: str = "",
    max_rows: int = 100,
    order_by: str = "",
) -> str:
    """Filter table rows with a SQL-like WHERE clause.

    Supports operators: =, !=, >, >=, <, <=, LIKE, IN. Conditions joined
    by AND/OR. Column values can be strings ('USA'), numbers (42), or
    lists for IN (('USA', 'Canada')).

    Args:
        alias: The alias of the open file
        table_name: Name of the table to query
        where: WHERE clause, e.g. "Country = 'USA' AND Amount > 1000"
        columns: Comma-separated columns to return (empty = all)
        max_rows: Maximum rows to return (default 100)
        order_by: Column name to sort by (optional, append ' DESC' for descending)
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        tdata = model.get_table(table_name, max_rows=0)

        if not tdata["columns"]:
            return ToolResponse.error(
                f"Table '{table_name}' not found.", "TABLE_NOT_FOUND"
            ).to_text()

        all_cols = tdata["columns"]
        col_idx = {c: i for i, c in enumerate(all_cols)}

        # Parse WHERE
        conditions = _parse_where_clause(where) if where else []
        for c in conditions:
            if c["col"] not in col_idx:
                return ToolResponse.error(
                    f"Column '{c['col']}' not found in table '{table_name}'", "COLUMN_NOT_FOUND"
                ).to_text()

        # Filter rows
        filtered = []
        for row in tdata["rows"]:
            if not conditions:
                filtered.append(row)
                continue
            # Eval AND/OR — simple left-to-right
            result = _eval_condition(row[col_idx[conditions[0]["col"]]],
                                     conditions[0]["op"], conditions[0]["value"])
            for cond in conditions[1:]:
                val = row[col_idx[cond["col"]]]
                r = _eval_condition(val, cond["op"], cond["value"])
                if cond["connector"] == "AND":
                    result = result and r
                else:
                    result = result or r
            if result:
                filtered.append(row)

        # Column projection
        if columns.strip():
            proj_cols = [c.strip() for c in columns.split(",")]
            for c in proj_cols:
                if c not in col_idx:
                    return ToolResponse.error(
                        f"Column '{c}' not found", "COLUMN_NOT_FOUND"
                    ).to_text()
            proj_idx = [col_idx[c] for c in proj_cols]
            filtered = [[r[i] for i in proj_idx] for r in filtered]
            out_cols = proj_cols
        else:
            out_cols = all_cols

        # ORDER BY
        if order_by.strip():
            ob = order_by.strip()
            reverse = False
            if ob.upper().endswith(" DESC"):
                ob = ob[:-5].strip()
                reverse = True
            elif ob.upper().endswith(" ASC"):
                ob = ob[:-4].strip()
            if ob not in [c for c in out_cols]:
                return ToolResponse.error(
                    f"ORDER BY column '{ob}' not in output", "COLUMN_NOT_FOUND"
                ).to_text()
            ob_idx = out_cols.index(ob)
            filtered.sort(key=lambda r: (r[ob_idx] is None, r[ob_idx]), reverse=reverse)

        total = len(filtered)
        shown = filtered[:max_rows]

        # Format output
        from pbix_mcp.formats.model_reader import format_table_data
        formatted = format_table_data({"columns": out_cols, "rows": shown}, max_rows=max_rows)
        header = f"Query returned {total:,} rows"
        if total > max_rows:
            header += f" (showing first {max_rows})"
        return ToolResponse.ok(f"{header}\n\n{formatted}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "QUERY_ERROR").to_text()


@mcp.tool()
def pbix_table_stats(alias: str, table_name: str) -> str:
    """Profile a table — per-column stats (min/max/avg/distinct/nulls).

    For strings: distinct count, null count, min/max length, top 5 values.
    For numbers: min/max/avg/sum/null count.
    For dates: min/max/null count.

    Args:
        alias: The alias of the open file
        table_name: Name of the table to profile
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from collections import Counter
        from datetime import date, datetime

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        tdata = model.get_table(table_name, max_rows=0)

        if not tdata["columns"]:
            return ToolResponse.error(
                f"Table '{table_name}' not found.", "TABLE_NOT_FOUND"
            ).to_text()

        cols = tdata["columns"]
        rows = tdata["rows"]
        total_rows = len(rows)

        lines = [f"# Stats for '{table_name}' ({total_rows:,} rows, {len(cols)} columns)\n"]

        for ci, cname in enumerate(cols):
            values = [r[ci] for r in rows]
            nulls = sum(1 for v in values if v is None)
            non_null = [v for v in values if v is not None]

            if not non_null:
                lines.append(f"## {cname}")
                lines.append(f"  All {total_rows:,} values are null")
                lines.append("")
                continue

            # Detect type from first non-null value
            sample = non_null[0]
            if isinstance(sample, (int, float)) and not isinstance(sample, bool):
                vals = [float(v) for v in non_null]
                mn, mx = min(vals), max(vals)
                avg = sum(vals) / len(vals)
                lines.append(f"## {cname} (numeric)")
                lines.append(f"  count={len(non_null):,}, nulls={nulls:,}")
                lines.append(f"  min={mn:g}, max={mx:g}, avg={avg:.2f}, sum={sum(vals):g}")
                lines.append(f"  distinct={len(set(vals)):,}")
            elif isinstance(sample, (datetime, date)):
                lines.append(f"## {cname} (datetime)")
                lines.append(f"  count={len(non_null):,}, nulls={nulls:,}")
                lines.append(f"  min={min(non_null)}, max={max(non_null)}")
                lines.append(f"  distinct={len(set(non_null)):,}")
            else:
                # String
                strs = [str(v) for v in non_null]
                lens = [len(s) for s in strs]
                distinct = set(strs)
                counter = Counter(strs)
                top = counter.most_common(5)
                lines.append(f"## {cname} (string)")
                lines.append(f"  count={len(non_null):,}, nulls={nulls:,}, distinct={len(distinct):,}")
                lines.append(f"  length: min={min(lens)}, max={max(lens)}, avg={sum(lens)/len(lens):.1f}")
                lines.append(f"  top 5: {', '.join(f'{v!r} ({c})' for v, c in top)}")
            lines.append("")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "STATS_ERROR").to_text()


@mcp.tool()
def pbix_data_diff(alias_a: str, alias_b: str, table_name: str, key_columns: str) -> str:
    """Diff row data between the same table in two PBIX files.

    Matches rows by key_columns (comma-separated), then reports:
    - Added rows (in B, not in A)
    - Removed rows (in A, not in B)
    - Changed rows (same key, different values)

    Args:
        alias_a: Alias of the first (old) file
        alias_b: Alias of the second (new) file
        table_name: Table to diff (must exist in both files)
        key_columns: Comma-separated columns to match rows by
    """
    try:
        info_a = _ensure_open(alias_a)
        info_b = _ensure_open(alias_b)
        if info_a.get("is_directquery") or info_b.get("is_directquery"):
            return ToolResponse.error(
                "DirectQuery files don't store data locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.formats.model_reader import ModelReader
        model_a = ModelReader(info_a["path"], work_dir=info_a.get("work_dir"))
        model_b = ModelReader(info_b["path"], work_dir=info_b.get("work_dir"))

        t_a = model_a.get_table(table_name, max_rows=0)
        t_b = model_b.get_table(table_name, max_rows=0)

        if not t_a["columns"]:
            return ToolResponse.error(f"Table '{table_name}' not in file A", "TABLE_NOT_FOUND").to_text()
        if not t_b["columns"]:
            return ToolResponse.error(f"Table '{table_name}' not in file B", "TABLE_NOT_FOUND").to_text()

        keys = [k.strip() for k in key_columns.split(",")]
        cols_a, cols_b = t_a["columns"], t_b["columns"]

        for k in keys:
            if k not in cols_a:
                return ToolResponse.error(f"Key column '{k}' not in table A", "COLUMN_NOT_FOUND").to_text()
            if k not in cols_b:
                return ToolResponse.error(f"Key column '{k}' not in table B", "COLUMN_NOT_FOUND").to_text()

        key_idx_a = [cols_a.index(k) for k in keys]
        key_idx_b = [cols_b.index(k) for k in keys]

        def row_key(row, key_idx):
            return tuple(str(row[i]) for i in key_idx)

        map_a = {row_key(r, key_idx_a): r for r in t_a["rows"]}
        map_b = {row_key(r, key_idx_b): r for r in t_b["rows"]}

        added_keys = set(map_b) - set(map_a)
        removed_keys = set(map_a) - set(map_b)
        common_keys = set(map_a) & set(map_b)

        # Compare common rows by value
        common_cols = [c for c in cols_a if c in cols_b]
        changed = []
        for k in common_keys:
            ra, rb = map_a[k], map_b[k]
            row_changes = []
            for cname in common_cols:
                va = ra[cols_a.index(cname)]
                vb = rb[cols_b.index(cname)]
                if str(va) != str(vb):
                    row_changes.append((cname, va, vb))
            if row_changes:
                changed.append((k, row_changes))

        lines = [
            f"# Data diff: '{table_name}' (key: {key_columns})",
            "",
            f"File A: {len(t_a['rows']):,} rows",
            f"File B: {len(t_b['rows']):,} rows",
            "",
            f"Summary: {len(added_keys)} added, {len(removed_keys)} removed, {len(changed)} changed",
        ]

        if added_keys:
            lines.append(f"\n## Added ({len(added_keys)}):")
            for k in sorted(list(added_keys))[:20]:
                lines.append(f"  + {' / '.join(k)}")
            if len(added_keys) > 20:
                lines.append(f"  ... and {len(added_keys) - 20} more")

        if removed_keys:
            lines.append(f"\n## Removed ({len(removed_keys)}):")
            for k in sorted(list(removed_keys))[:20]:
                lines.append(f"  - {' / '.join(k)}")
            if len(removed_keys) > 20:
                lines.append(f"  ... and {len(removed_keys) - 20} more")

        if changed:
            lines.append(f"\n## Changed ({len(changed)}):")
            for k, row_changes in changed[:20]:
                lines.append(f"  ~ {' / '.join(k)}")
                for cname, va, vb in row_changes:
                    lines.append(f"      {cname}: {va!r} -> {vb!r}")
            if len(changed) > 20:
                lines.append(f"  ... and {len(changed) - 20} more")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "DIFF_ERROR").to_text()


@mcp.tool()
def pbix_replace_value(
    alias: str,
    table_name: str,
    column_name: str,
    old_value: str,
    new_value: str,
    case_sensitive: bool = True,
) -> str:
    """Find and replace ALL occurrences of a value in a column.

    Reads the table, replaces all rows where the column matches old_value,
    and writes the updated data back via DataModel rebuild.

    LIMITATION: Uses the full rebuild pipeline — works on builder-created
    files but may break PBI Desktop files with SQL Server imports (destroys
    M expressions). For PBI Desktop files, use with caution.

    Args:
        alias: The alias of the open file
        table_name: Name of the table
        column_name: Name of the column to modify
        old_value: Value to find (exact match)
        new_value: Value to replace with
        case_sensitive: If False, matches strings case-insensitively (default True)
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — table data is not stored locally.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        tdata = model.get_table(table_name, max_rows=0)

        if not tdata["columns"]:
            return ToolResponse.error(
                f"Table '{table_name}' not found.", "TABLE_NOT_FOUND"
            ).to_text()

        cols = tdata["columns"]
        if column_name not in cols:
            return ToolResponse.error(
                f"Column '{column_name}' not found in table '{table_name}'. "
                f"Available: {', '.join(cols)}",
                "COLUMN_NOT_FOUND"
            ).to_text()

        col_idx = cols.index(column_name)

        # Replace values in rows
        replaced = 0
        if case_sensitive:
            def matches(v):
                return str(v) == old_value
        else:
            needle = old_value.lower()
            def matches(v):
                return str(v).lower() == needle

        # Detect target data type from first non-null sample
        sample = next((r[col_idx] for r in tdata["rows"] if r[col_idx] is not None), None)

        # Coerce new_value to match column type
        def coerce(v):
            if sample is None:
                return v
            if isinstance(sample, bool):
                return v.lower() in ("true", "1", "yes")
            if isinstance(sample, int):
                return int(v)
            if isinstance(sample, float):
                return float(v)
            return v

        try:
            new_val_typed = coerce(new_value)
        except (ValueError, TypeError) as e:
            return ToolResponse.error(
                f"new_value '{new_value}' cannot be converted to column type: {e}",
                "TYPE_MISMATCH"
            ).to_text()

        # Rebuild rows as list-of-dicts (required by _rebuild_datamodel)
        new_rows = []
        for row in tdata["rows"]:
            row_dict = dict(zip(cols, row))
            if matches(row_dict[column_name]):
                row_dict[column_name] = new_val_typed
                replaced += 1
            new_rows.append(row_dict)

        if replaced == 0:
            return ToolResponse.ok(
                f"No matches found — '{old_value}' not in {table_name}.{column_name}"
            ).to_text()

        # Get column definitions from metadata for _rebuild_datamodel
        dm_path = os.path.join(info["work_dir"], "DataModel")
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()
        abf = decompress_datamodel(dm_bytes)
        meta_bytes = read_metadata_sqlite(abf)

        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        os.write(fd, meta_bytes)
        os.close(fd)
        try:
            conn = sqlite3.connect(tmp_path)
            conn.row_factory = sqlite3.Row
            _AMO_TO_TYPE = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                            10: "Decimal", 11: "Boolean"}
            col_rows = conn.execute(
                """SELECT c.ExplicitName, c.ExplicitDataType, c.DataCategory
                   FROM [Column] c
                   JOIN [Table] t ON c.TableID = t.ID
                   WHERE t.Name = ? AND c.Type = 1
                   ORDER BY c.ID""",
                (table_name,)
            ).fetchall()
            conn.close()
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        if not col_rows:
            return ToolResponse.error(
                f"Table '{table_name}' has no user columns.", "TABLE_NOT_FOUND"
            ).to_text()

        columns_def = [
            {"name": cr["ExplicitName"],
             "data_type": _AMO_TO_TYPE.get(cr["ExplicitDataType"], "String"),
             "data_category": cr["DataCategory"]}
            for cr in col_rows
        ]

        # Filter new_rows to only include columns that exist in columns_def
        # (tdata may include RowNumber/hidden cols that columns_def excludes)
        valid_names = {c["name"] for c in columns_def}
        filtered_rows = [
            {k: v for k, v in r.items() if k in valid_names}
            for r in new_rows
        ]

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            table_updates={table_name: {"columns": columns_def, "rows": filtered_rows}},
        )
        info["modified"] = True

        return ToolResponse.ok(
            f"Replaced {replaced:,} occurrences of '{old_value}' with '{new_value}' "
            f"in {table_name}.{column_name}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", "REPLACE_ERROR").to_text()


# --- Field parameters (Desktop-recognized) ---------------------------------
#
# A real Desktop field parameter is a CALCULATED table whose partition holds a
# {("Display", NAMEOF('T'[C]), n), ...} tuple set, whose three columns are
# calc-table columns (Type=4) with the display column sorted by the hidden
# Order column, and whose Fields column carries the ExtendedProperty
# 'ParameterMetadata' = {"version":3,"kind":2}. Every value below was diffed
# against Desktop-authored ground truth (test_corpus/Ecommerce_Conversion.pbix,
# two genuine field parameters).
#
# pbix-mcp authors the table as ordinary static data (full VertiPaq storage —
# exactly what Desktop's own files physically contain, since NAMEOF tuples are
# constants) and then stamps the Desktop metadata shape on top with a
# metadata-only splice. _rebuild_datamodel recognizes the shape, rebuilds the
# static data, and re-stamps — so field parameters survive every rebuild-based
# edit instead of tripping the calculated-table refusal.

_FIELD_PARAM_METADATA_JSON = '{"version":3,"kind":2}'

_FIELD_REF_RE = re.compile(
    r"^\s*(?:'((?:[^']|'')+)'|([^'\[\]]+?))\s*\[\s*([^\[\]]+?)\s*\]\s*$"
)


def _normalize_field_ref(ref: str) -> tuple[str, str, str]:
    """Parse a field reference ("Table[Col]" / "'Table'[Col]") and return
    (table, name, canonical) where canonical is the DAX-quoted form
    ("'Table'[Col]") that NAMEOF() evaluates to — the exact string Desktop
    stores in a field parameter's Fields column."""
    m = _FIELD_REF_RE.match(ref or "")
    if not m:
        raise ValueError(
            f"Invalid field ref {ref!r} — use \"Table[Field]\" or "
            f"\"'Table'[Field]\" (field = column or measure name)")
    table = (m.group(1) or m.group(2) or "").strip()
    if m.group(1):
        table = table.replace("''", "'")
    name = m.group(3).strip()
    canonical = f"'{table.replace(chr(39), chr(39) * 2)}'[{name}]"
    return table, name, canonical


def _field_parameter_query_definition(fields: list[dict]) -> str:
    """Build the calculated-table tuple-set DAX exactly as Desktop writes it.

    ``fields`` entries carry "display" and the CANONICAL "ref" from
    _normalize_field_ref."""
    lines = []
    for i, f in enumerate(fields):
        display = f["display"].replace('"', '""')
        lines.append(f'    ("{display}", NAMEOF({f["ref"]}), {i})')
    return "{\n" + ",\n".join(lines) + "\n}"


def _detect_field_parameter_shape(conn, tid: int) -> dict | None:
    """Return {"columns": [...], "query_definition": str} when table ``tid``
    is a field parameter whose static rows the rebuild can reproduce
    (calculated partition + parseable NAMEOF tuple set + exactly 3 data
    columns with physical VertiPaq storage). None = not a field parameter."""
    prow = conn.execute(
        "SELECT Type, QueryDefinition FROM [Partition] WHERE TableID = ? LIMIT 1",
        (tid,),
    ).fetchone()
    if prow is None or prow["Type"] != 2 or not prow["QueryDefinition"]:
        return None
    crows = conn.execute(
        # COALESCE, and the RowNumber test has to run on the same expression:
        # a Type 4 column has ExplicitName NULL, and `NULL NOT LIKE ...` is NULL
        # (falsy), so the bare form dropped every calculated-table column from
        # the shape check -- which is exactly what a field parameter is made of.
        "SELECT COALESCE(ExplicitName, InferredName) AS ExplicitName, "
        "       ExplicitDataType, InferredDataType FROM [Column] "
        "WHERE TableID = ? AND Type IN (1, 4) "
        "AND COALESCE(ExplicitName, InferredName) NOT LIKE 'RowNumber%' "
        "ORDER BY ID",
        (tid,),
    ).fetchall()
    if len(crows) != 3:
        return None
    col_names = [c["ExplicitName"] for c in crows]
    try:
        from pbix_mcp.dax.calc_tables import _parse_field_parameter
        # Strip DAX comments first (Desktop-authored field parameters carry
        # them), mirroring calc_tables._evaluate_table_expression.
        qd = re.sub(r'--[^\n]*', '', prow["QueryDefinition"])
        qd = re.sub(r'//[^\n]*', '', qd).strip()
        parsed = _parse_field_parameter(qd, {"columns": col_names})
    except Exception:
        parsed = None
    if not parsed:
        return None
    amo_map = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
               10: "Decimal", 11: "Boolean"}
    cols = []
    for c in crows:
        amo = c["ExplicitDataType"]
        if amo == 1:  # automatic (Desktop calc-table column) -> inferred type
            amo = c["InferredDataType"] or 2
        cols.append({"name": c["ExplicitName"],
                     "data_type": amo_map.get(amo, "String")})
    return {"columns": cols, "query_definition": prow["QueryDefinition"]}


def _apply_field_parameter_metadata(dm_path: str, specs: list[dict]) -> tuple[int, int]:
    """Stamp Desktop's field-parameter metadata onto static 3-column tables.

    Each spec = {"table": name, "query_definition": tuple-set DAX}. Applies,
    per table (all values from Desktop-authored ground truth): Table
    SystemFlags=2; the three data columns become calc-table columns (Type=4,
    ExplicitDataType=1 automatic + real InferredDataType, SourceColumn
    [Value1..3]); display column visible + sorted by the hidden Order column;
    Fields column hidden + sorted by Order + ExtendedProperty
    ParameterMetadata; Order column hidden, FormatString '0', SummarizeBy Sum;
    partition flipped to calculated (Type=2) holding the NAMEOF tuple set; and
    the display->Fields group-by wiring (RelatedColumnDetails/GroupByColumn).
    """
    def _do_apply(conn: sqlite3.Connection):
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        # Older pbix-mcp-built files predate these metadata tables.
        c.execute(
            'CREATE TABLE IF NOT EXISTS [ExtendedProperty]( [ID] INTEGER, '
            '[ObjectID] INTEGER, [ObjectType] INTEGER, [Name] TEXT, '
            '[Type] INTEGER, [Value] TEXT, [ModifiedTime] INTEGER, '
            'PRIMARY KEY("ID" ASC) )')
        c.execute(
            'CREATE TABLE IF NOT EXISTS [RelatedColumnDetails]( [ID] INTEGER, '
            '[ColumnID] INTEGER, [ModifiedTime] INTEGER, '
            'PRIMARY KEY("ID" ASC) )')
        c.execute(
            'CREATE TABLE IF NOT EXISTS [GroupByColumn]( [ID] INTEGER, '
            '[RelatedColumnDetailsID] INTEGER, [GroupingColumnID] INTEGER, '
            '[ModifiedTime] INTEGER, PRIMARY KEY("ID" ASC) )')

        maxid_row = c.execute(
            "SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
        max_id = int(maxid_row[0]) if maxid_row else 0
        import datetime
        epoch = datetime.datetime(1601, 1, 1)
        # NAIVE UTC (utcnow() is deprecated): the FILETIME delta subtracts a
        # naive epoch, and an aware operand would raise TypeError.
        _now = datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None)
        filetime = int((_now - epoch).total_seconds() * 10_000_000)

        for spec in specs:
            tname = spec["table"]
            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (tname,)).fetchone()
            if not trow:
                raise ValueError(f"Field parameter table '{tname}' not found")
            tid = trow["ID"]
            crows = c.execute(
                "SELECT ID FROM [Column] WHERE TableID = ? AND Type IN (1, 4) "
                "AND ExplicitName NOT LIKE 'RowNumber%' ORDER BY ID",
                (tid,)).fetchall()
            if len(crows) != 3:
                raise ValueError(
                    f"Field parameter table '{tname}' must have exactly 3 "
                    f"data columns, found {len(crows)}")
            disp_id, fields_id, order_id = (r["ID"] for r in crows)

            c.execute("UPDATE [Table] SET SystemFlags = 2 WHERE ID = ?", (tid,))
            # Desktop stamps SystemFlags=2 on EVERY calc-table object (all
            # columns incl. RowNumber, and the partition) and keeps columns
            # MDX-visible (IsAvailableInMDX=1) — ground truth Ecommerce corpus.
            c.execute("UPDATE [Column] SET SystemFlags = 2 WHERE TableID = ?",
                      (tid,))
            common = ("Type = 4, ExplicitDataType = 1, SourceColumn = ?, "
                      "InferredName = ?, IsAvailableInMDX = 1")
            c.execute(
                f"UPDATE [Column] SET {common}, InferredDataType = 2, "
                f"IsHidden = 0, SortByColumnID = ?, SummarizeBy = 2 "
                f"WHERE ID = ?",
                ("[Value1]", "Value1", order_id, disp_id))
            c.execute(
                f"UPDATE [Column] SET {common}, InferredDataType = 2, "
                f"IsHidden = 1, SortByColumnID = ?, SummarizeBy = 2 "
                f"WHERE ID = ?",
                ("[Value2]", "Value2", order_id, fields_id))
            c.execute(
                f"UPDATE [Column] SET {common}, InferredDataType = 6, "
                f"IsHidden = 1, SortByColumnID = 0, SummarizeBy = 3, "
                f"FormatString = '0' WHERE ID = ?",
                ("[Value3]", "Value3", order_id))
            c.execute(
                "UPDATE [Partition] SET Type = 2, Mode = 0, SystemFlags = 2, "
                "QueryDefinition = ? WHERE TableID = ?",
                (spec["query_definition"], tid))

            # ParameterMetadata marker on the Fields column (idempotent).
            c.execute(
                "DELETE FROM [ExtendedProperty] WHERE ObjectID = ? "
                "AND ObjectType = 4 AND Name = 'ParameterMetadata'",
                (fields_id,))
            max_id += 1
            c.execute(
                "INSERT INTO [ExtendedProperty] (ID, ObjectID, ObjectType, "
                "Name, Type, Value, ModifiedTime) "
                "VALUES (?, ?, 4, 'ParameterMetadata', 1, ?, ?)",
                (max_id, fields_id, _FIELD_PARAM_METADATA_JSON, filetime))

            # Display column groups by the Fields column (Desktop wiring).
            c.execute(
                "DELETE FROM [GroupByColumn] WHERE RelatedColumnDetailsID IN "
                "(SELECT ID FROM [RelatedColumnDetails] WHERE ColumnID = ?)",
                (disp_id,))
            c.execute(
                "DELETE FROM [RelatedColumnDetails] WHERE ColumnID = ?",
                (disp_id,))
            max_id += 1
            rcd_id = max_id
            c.execute(
                "INSERT INTO [RelatedColumnDetails] (ID, ColumnID, "
                "ModifiedTime) VALUES (?, ?, ?)", (rcd_id, disp_id, filetime))
            c.execute(
                "UPDATE [Column] SET RelatedColumnDetailsID = ? WHERE ID = ?",
                (rcd_id, disp_id))
            max_id += 1
            c.execute(
                "INSERT INTO [GroupByColumn] (ID, RelatedColumnDetailsID, "
                "GroupingColumnID, ModifiedTime) VALUES (?, ?, ?, ?)",
                (max_id, rcd_id, fields_id, filetime))

        c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'",
                  (str(max_id),))
        conn.commit()

    return _modify_metadata_only(dm_path, _do_apply)


def _apply_calculated_column_metadata(
    dm_path: str, specs: list[dict]
) -> tuple[int, int]:
    """Flip builder-materialized data columns into calculated columns.

    Each spec = {"table", "column", "expression", "amo_type"}. The builder
    encoded the values as an ordinary data column; a Desktop calculated column
    is physically identical (ColumnStorage / hierarchy / segment data), so only
    the metadata differs — verified field-for-field against a Desktop-authored
    calc column (test_corpus/GeoSales_Dashboard.pbix, fct_Orders[Discount
    Group]): Type=2, the DAX Expression, SourceColumn NULL, ExplicitDataType 1
    (Automatic) with InferredDataType carrying the real AMO type.
    """
    def _do_apply(conn: sqlite3.Connection):
        c = conn.cursor()
        import datetime
        epoch = datetime.datetime(1601, 1, 1)
        # NAIVE UTC (utcnow() is deprecated): the delta's epoch is naive.
        _now = datetime.datetime.now(datetime.timezone.utc).replace(tzinfo=None)
        filetime = int((_now - epoch).total_seconds() * 10_000_000)
        for spec in specs:
            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (spec["table"],)).fetchone()
            if not trow:
                raise ValueError(
                    f"Calculated-column table '{spec['table']}' not found")
            tid = trow[0]
            # SystemFlags=2 on a calculated column that lives on a calculated
            # TABLE — Desktop stamps every object of such a table that way, and
            # a plain 0 there is not the shape it writes (verified across all
            # 57 corpus tables of that shape). On an ordinary table the flag
            # stays as it was.
            sysflags = spec.get("system_flags")
            extra = ", SystemFlags = ?" if sysflags is not None else ""
            params: list = [spec["expression"], spec["amo_type"], filetime,
                            filetime, filetime]
            if sysflags is not None:
                params.append(sysflags)
            params += [tid, spec["column"]]
            n = c.execute(
                "UPDATE [Column] SET Type = 2, Expression = ?, "
                "SourceColumn = NULL, ExplicitDataType = 1, "
                "InferredDataType = ?, ModifiedTime = ?, "
                "StructureModifiedTime = ?, RefreshedTime = ?"
                f"{extra} "
                "WHERE TableID = ? AND ExplicitName = ?",
                params).rowcount
            if n != 1:
                raise ValueError(
                    f"Expected to re-stamp exactly one column "
                    f"'{spec['table']}'[{spec['column']}], updated {n}")
        conn.commit()

    return _modify_metadata_only(dm_path, _do_apply)


def _apply_calculated_table_metadata(
    dm_path: str, specs: list[dict]
) -> tuple[int, int]:
    """Flip builder-materialized static tables into calculated tables.

    Each spec = {"table", "expression"}. The builder emitted the evaluated rows
    as an ordinary imported table; a Desktop calculated table is physically the
    same (VertiPaq storage + hierarchies) and differs only in metadata —
    verified field-for-field against Desktop-authored calc tables in
    test_corpus/GeoSales_Dashboard.pbix ('DiscountGroup'): Table SystemFlags=2;
    partition Type=2 + SystemFlags=2 carrying the DAX as its QueryDefinition;
    and every data column Type=4 with ExplicitName NULL, the name moved to
    InferredName, SourceColumn '[Name]', ExplicitDataType 1 (Automatic) with the
    real type left in InferredDataType, SystemFlags=2, IsAvailableInMDX=1.
    """
    def _do_apply(conn: sqlite3.Connection):
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        for spec in specs:
            tname = spec["table"]
            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (tname,)).fetchone()
            if not trow:
                raise ValueError(f"Calculated table '{tname}' not found")
            tid = trow["ID"]
            c.execute("UPDATE [Table] SET SystemFlags = 2 WHERE ID = ?", (tid,))
            c.execute(
                "UPDATE [Partition] SET Type = 2, SystemFlags = 2, "
                "QueryDefinition = ? WHERE TableID = ?",
                (spec["expression"], tid))
            # Columns this table owns as CALCULATED columns are not the
            # partition's own data columns and must survive untouched for
            # _apply_calculated_column_metadata to find them: it looks a column
            # up by ExplicitName, and rewriting every column to Type=4 with
            # ExplicitName NULL both destroyed the calc columns and made that
            # lookup miss. Power BI's auto date/time tables are exactly this
            # shape — a Date column from the partition expression plus six
            # calculated columns — and used to be refused outright because of it.
            skip = {n for n in (spec.get("calc_columns") or ())}
            # Data columns -> calc-table columns. RowNumber (Type=3) only gets
            # the system flag, exactly as Desktop leaves it.
            for crow in c.execute(
                "SELECT ID, ExplicitName, Type FROM [Column] WHERE TableID = ? "
                "ORDER BY ID", (tid,)).fetchall():
                if crow["Type"] == 3 or (
                        crow["ExplicitName"] or "").startswith("RowNumber"):
                    c.execute("UPDATE [Column] SET SystemFlags = 2 WHERE ID = ?",
                              (crow["ID"],))
                    continue
                if crow["ExplicitName"] in skip:
                    # Only the flags Desktop puts on every object of a calc
                    # table; the calc-column stamp follows and owns the rest.
                    c.execute(
                        "UPDATE [Column] SET SystemFlags = 2, "
                        "IsAvailableInMDX = 1 WHERE ID = ?", (crow["ID"],))
                    continue
                name = crow["ExplicitName"]
                # Keep the model's own SourceColumn when it has one: a calc
                # table copied from another table qualifies its columns
                # ('DateAutoTemplate[Year]'), and a bare '[Year]' does not
                # resolve. Fall back to the bare form only for a table we are
                # authoring from scratch, which has no source to qualify.
                src = (spec.get("source_columns") or {}).get(name) or f"[{name}]"
                c.execute(
                    "UPDATE [Column] SET Type = 4, ExplicitName = NULL, "
                    "InferredName = ?, SourceColumn = ?, ExplicitDataType = 1, "
                    "SystemFlags = 2, IsAvailableInMDX = 1 WHERE ID = ?",
                    (name, src, crow["ID"]))
        conn.commit()

    return _modify_metadata_only(dm_path, _do_apply)


# AMO ObjectType, as used by [Annotation].ObjectType and friends. Derived
# empirically rather than from documentation: across the 24-report corpus, all
# 4,345 (ObjectID, ObjectType) pairs resolved to exactly one entity table each,
# with no ObjectType matching two candidates and none left unmatched.
_AMO_OBJECT_TYPE = {
    1: "Model", 3: "Table", 4: "Column", 7: "Relationship",
    8: "Measure", 9: "Hierarchy", 12: "KPI", 41: "Expression",
}

# Model metadata the from-scratch builder never emits. Each entry is
# (table, {fk_column: kind}, identity_columns). ``kind`` says how to translate
# an ID across the rebuild:
#   "model"/"culture"  — the singleton row
#   "table"/"column"/"measure"/"hierarchy"/"relationship"/"expression"/"kpi"
#                      — resolved by NAME, since IDs are reassigned
#   "object"           — an (ObjectID, ObjectType) pair; the sibling
#                        ObjectType column names the entity
#   "self:<Table>"     — a row carried earlier in this same pass
# ``identity_columns`` are compared against what the rebuild already wrote, so
# a row the builder reproduced on its own is never duplicated.
#
# Order matters: parents precede the children that reference them.
_CARRY_SPEC: list[tuple[str, dict[str, str], tuple[str, ...]]] = [
    ("Expression", {"ModelID": "model",
                    "ParameterValuesColumnID": "column?"}, ("Name",)),
    ("Function", {"ModelID": "model"}, ("Name",)),
    ("DataSource", {"ModelID": "model"}, ("Name",)),
    ("LinguisticMetadata", {"CultureID": "culture"}, ("CultureID",)),
    ("Perspective", {"ModelID": "model"}, ("Name",)),
    ("PerspectiveTable", {"PerspectiveID": "self:Perspective",
                          "TableID": "table"}, ("PerspectiveID", "TableID")),
    ("PerspectiveColumn", {"PerspectiveTableID": "self:PerspectiveTable",
                           "ColumnID": "column"},
     ("PerspectiveTableID", "ColumnID")),
    ("PerspectiveMeasure", {"PerspectiveTableID": "self:PerspectiveTable",
                            "MeasureID": "measure"},
     ("PerspectiveTableID", "MeasureID")),
    ("PerspectiveHierarchy", {"PerspectiveTableID": "self:PerspectiveTable",
                              "HierarchyID": "hierarchy"},
     ("PerspectiveTableID", "HierarchyID")),
    ("KPI", {"MeasureID": "measure"}, ("MeasureID",)),
    ("RelatedColumnDetails", {"ColumnID": "column"}, ("ColumnID",)),
    ("GroupByColumn", {"RelatedColumnDetailsID": "self:RelatedColumnDetails",
                       "GroupingColumnID": "column"},
     ("RelatedColumnDetailsID", "GroupingColumnID")),
    ("Variation", {"ColumnID": "column", "RelationshipID": "relationship",
                   "DefaultHierarchyID": "hierarchy?",
                   "DefaultColumnID": "column?"}, ("ColumnID", "Name")),
    ("FormatStringDefinition", {"ObjectID": "object"},
     ("ObjectID", "ObjectType")),
    ("Annotation", {"ObjectID": "object"}, ("ObjectID", "ObjectType", "Name")),
    ("ExtendedProperty", {"ObjectID": "object"},
     ("ObjectID", "ObjectType", "Name")),
    ("ChangedProperty", {"ObjectID": "object"},
     ("ObjectID", "ObjectType", "Property")),
    # --- features no corpus file exercises (issue #7) -----------------------
    # A rebuild dropped every one of these silently, with success: true and an
    # empty warnings list. Reproduced with a calculation group authored by this
    # project's own pbix_datamodel_add_calculation_group: CalculationGroup,
    # CalculationItem, Table.CalculationGroupID and the Type=7 partition all
    # went to zero.
    #
    # These come AFTER FormatStringDefinition and PerspectiveTable above,
    # because "self:" resolves against rows carried earlier in the same pass.
    ("QueryGroup", {"ModelID": "model"}, ("ModelID", "Folder")),
    ("CalculationGroup", {"TableID": "table"}, ("TableID",)),
    ("CalculationItem",
     {"CalculationGroupID": "self:CalculationGroup",
      "FormatStringDefinitionID": "self:FormatStringDefinition?"},
     ("CalculationGroupID", "Name")),
    ("CalculationExpression",
     {"CalculationGroupID": "self:CalculationGroup",
      "FormatStringDefinitionID": "self:FormatStringDefinition?"},
     ("CalculationGroupID",)),
    ("Set", {"TableID": "table"}, ("TableID", "Name")),
    ("PerspectiveSet", {"PerspectiveTableID": "self:PerspectiveTable",
                        "SetID": "self:Set"},
     ("PerspectiveTableID", "SetID")),
    ("ObjectTranslation", {"CultureID": "culture", "ObjectID": "object"},
     ("CultureID", "ObjectID", "ObjectType", "Property")),
    ("DetailRowsDefinition", {"ObjectID": "object"},
     ("ObjectID", "ObjectType")),
    ("AlternateOf", {"ColumnID": "column", "BaseColumnID": "column?",
                     "BaseTableID": "table?"}, ("ColumnID",)),
    ("RefreshPolicy", {"TableID": "table"}, ("TableID",)),
    ("Calendar", {"TableID": "table"}, ("TableID", "Name")),
    ("TimeUnitColumnAssociation", {"CalendarID": "self:Calendar"},
     ("CalendarID", "TimeUnit")),
    ("CalendarColumnReference",
     {"TimeUnitColumnAssociationID": "self:TimeUnitColumnAssociation",
      "ColumnID": "column"}, ("TimeUnitColumnAssociationID", "ColumnID")),
    ("AnalyticsAIMetadata", {"ModelID": "model"}, ("ModelID", "Name")),
]

# What each carried table means to someone using the report, for the warning
# text. "Annotation: 130 rows" tells a user nothing.
_CARRY_MEANING = {
    "Expression": "shared M expressions / query parameters",
    "DataSource": "declared data sources",
    "LinguisticMetadata": "Q&A synonyms and phrasings",
    "Perspective": "perspectives",
    "PerspectiveTable": "perspectives",
    "PerspectiveColumn": "perspectives",
    "PerspectiveMeasure": "perspectives",
    "PerspectiveHierarchy": "perspectives",
    "KPI": "KPI definitions on measures",
    "RelatedColumnDetails": "column grouping",
    "GroupByColumn": "column grouping",
    "Variation": "auto date/time drill-down wired to a date column",
    "FormatStringDefinition": "dynamic format strings",
    "Annotation": "annotations",
    "ExtendedProperty": "extended properties",
    "ChangedProperty": "changed-property bookkeeping",
    "Function": "user-defined DAX functions",
    "QueryGroup": "query display folders",
    "CalculationGroup": "calculation groups",
    "CalculationItem": "calculation groups",
    "CalculationExpression": "calculation groups",
    "Set": "named sets",
    "PerspectiveSet": "named sets in perspectives",
    "ObjectTranslation": "translated names for other languages",
    "DetailRowsDefinition": "detail-rows / drill-through expressions",
    "AlternateOf": "aggregation-table wiring",
    "RefreshPolicy": "incremental refresh policies",
    "Calendar": "calendars",
    "TimeUnitColumnAssociation": "calendar time-unit associations",
    "CalendarColumnReference": "calendar column references",
    "AnalyticsAIMetadata": "AI analytics metadata",
}


# Authoring properties on [Table] and [Column] rows. The builder DOES create
# those rows, so the carry-over above never covered them — it only moves rows
# the builder omits entirely — and every rebuild-path edit silently reset them
# to defaults. What that costs a user: hidden tables and columns become visible,
# a currency or date column loses its format string, "Month sorted by MonthNo"
# reverts to alphabetical, an ImageUrl column stops rendering as an image, and
# a numeric column like Year gets SummarizeBy=Sum so dragging it into a visual
# adds the years together.
#
# Deliberately excluded: storage and type fields (ColumnStorageID,
# InferredDataType, IsAvailableInMDX, AttributeHierarchyID, the *ModifiedTime
# stamps). Those describe how the rebuilt data is physically stored and MUST
# take their new values.
_TABLE_PROPERTIES = (
    "IsHidden", "IsPrivate", "ShowAsVariationsOnly", "DataCategory",
    "Description", "ExcludeFromModelRefresh",
)

# AMO PartitionSourceType: 1 = Query (names a DataMashup query), 4 = M
# (the expression is the source). The builder emits M; a model authored by an
# older Desktop uses Query and its partitions must keep that shape.
_PARTITION_TYPE_QUERY = 1
_COLUMN_PROPERTIES = (
    "IsHidden", "FormatString", "SummarizeBy", "DataCategory", "DisplayOrdinal",
    "Description", "DisplayFolder", "IsKey", "IsNullable", "IsUnique",
    "IsDefaultLabel", "IsDefaultImage", "EncodingHint", "Alignment",
    "KeepUniqueRows", "ErrorMessage",
)

# Measures and hierarchies survive a rebuild by name, but the builder writes
# defaults for everything else about them. On a 102-measure model that meant 89
# measures losing the DisplayFolder they were organised into, hidden helper
# measures such as Date[_ShowValueForDates] becoming visible, and 18 measures
# silently retyped (Boolean and Int64 both landing on Double). DataType is
# included deliberately: unlike a column's storage type it is the measure's
# declared result type, not a fact about how the rebuilt data is stored.
_MEASURE_PROPERTIES = (
    "IsHidden", "DisplayFolder", "Description", "FormatString", "DataType",
    "DataCategory", "IsSimpleMeasure", "LineageTag",
)
_HIERARCHY_PROPERTIES = (
    "IsHidden", "DisplayFolder", "Description", "HideMembers", "LineageTag",
)


def _snapshot_object_properties(conn: sqlite3.Connection) -> dict:
    """Capture authoring properties of every table and column, keyed by NAME.

    ``SortByColumnID`` is stored as the referenced column's name, because the
    rebuild reassigns every ID and a stale number would point at whatever column
    happens to land on it.
    """
    conn.row_factory = sqlite3.Row
    have_t = {r[1] for r in conn.execute("PRAGMA table_info([Table])")}
    have_c = {r[1] for r in conn.execute("PRAGMA table_info([Column])")}
    tprops = [p for p in _TABLE_PROPERTIES if p in have_t]
    cprops = [p for p in _COLUMN_PROPERTIES if p in have_c]

    col_name_by_id = {
        r["ID"]: r["nm"] for r in conn.execute(
            "SELECT c.ID, COALESCE(c.ExplicitName, c.InferredName) AS nm "
            "FROM [Column] c JOIN [Table] t ON c.TableID = t.ID "
            "WHERE t.ModelID = 1")
    }
    tables = {
        r["Name"]: {p: r[p] for p in tprops}
        for r in conn.execute("SELECT * FROM [Table] WHERE ModelID = 1")
    }
    columns = {}
    for r in conn.execute(
        "SELECT c.*, t.Name AS _tbl FROM [Column] c "
        "JOIN [Table] t ON c.TableID = t.ID WHERE t.ModelID = 1"
    ):
        nm = r["ExplicitName"] or r["InferredName"]
        if not nm or nm.startswith("RowNumber"):
            continue
        d = {p: r[p] for p in cprops}
        if "SortByColumnID" in have_c and r["SortByColumnID"]:
            d["__sort_by__"] = col_name_by_id.get(r["SortByColumnID"])
        columns[f"{r['_tbl']}\x00{nm}"] = d

    # Partition source, keyed by table name. The builder always emits a Type=4
    # (M) partition holding the rows inline, which is right for a table it
    # authored but wrong for one that already had a source: a legacy Type=1
    # (Query) partition names a query in the DataMashup, and replacing it with
    # inline M orphans that query. Power BI Desktop then opens the file, shows
    # "There are pending changes in your queries that haven't been applied" and
    # an EMPTY Data pane — the model loads with no tables at all.
    have_p = {r[1] for r in conn.execute("PRAGMA table_info([Partition])")}
    partitions = {}
    if {"Type", "QueryDefinition"} <= have_p:
        for r in conn.execute(
            "SELECT t.Name AS tbl, p.Type, p.Mode, p.QueryDefinition "
            "FROM [Partition] p JOIN [Table] t ON p.TableID = t.ID "
            "WHERE t.ModelID = 1 AND t.SystemFlags = 0"
        ):
            # Only legacy Query partitions need carrying; an M partition the
            # builder regenerates already matches what the rebuild wrote.
            if r["Type"] == _PARTITION_TYPE_QUERY:
                partitions[r["tbl"]] = {
                    "Type": r["Type"], "Mode": r["Mode"],
                    "QueryDefinition": r["QueryDefinition"],
                }
    # Measures and hierarchies, keyed "<table>\x00<name>" like the columns above.
    measures: dict[str, dict] = {}
    hierarchies: dict[str, dict] = {}
    for tbl, props, out in (("Measure", _MEASURE_PROPERTIES, measures),
                            ("Hierarchy", _HIERARCHY_PROPERTIES, hierarchies)):
        have = {r[1] for r in conn.execute(f"PRAGMA table_info([{tbl}])")}
        keep = [p for p in props if p in have]
        if not keep:
            continue
        for r in conn.execute(
            f"SELECT o.*, t.Name AS _tbl FROM [{tbl}] o "
            "JOIN [Table] t ON o.TableID = t.ID WHERE t.ModelID = 1"
        ):
            if r["Name"]:
                out[f"{r['_tbl']}\x00{r['Name']}"] = {p: r[p] for p in keep}

    # Hierarchy levels, keyed "<table>\x00<hierarchy>\x00<ordinal>". Only the
    # lineage tag is carried: a level's name and column are rebuilt from the
    # hierarchy definition, but the tag is the identity downstream tooling
    # (TMDL/PBIP round-trips, service lineage) matches on across a save.
    levels = {}
    have_l = {r[1] for r in conn.execute("PRAGMA table_info([Level])")}
    if "LineageTag" in have_l:
        for r in conn.execute(
            "SELECT l.Ordinal, l.LineageTag, h.Name AS hname, t.Name AS tname "
            "FROM [Level] l JOIN [Hierarchy] h ON l.HierarchyID = h.ID "
            "JOIN [Table] t ON h.TableID = t.ID WHERE t.ModelID = 1"
        ):
            if r["LineageTag"]:
                levels[f"{r['tname']}\x00{r['hname']}\x00{r['Ordinal']}"] = \
                    r["LineageTag"]

    return {"tables": tables, "columns": columns, "partitions": partitions,
            "measures": measures, "hierarchies": hierarchies, "levels": levels}


def _restore_object_properties(dm_path: str, snap: dict) -> None:
    """Re-apply snapshotted authoring properties after a rebuild.

    Objects that no longer exist are skipped; the rebuild legitimately removed
    them. A sort-by target that did not survive is cleared rather than left
    pointing at an arbitrary column.
    """
    if not snap or not any(snap.get(k) for k in
                           ("tables", "columns", "partitions",
                            "measures", "hierarchies", "levels")):
        return

    def _do(conn: sqlite3.Connection):
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        ids = {}
        for r in c.execute(
            "SELECT c.ID, t.Name AS tbl, "
            "       COALESCE(c.ExplicitName, c.InferredName) AS nm "
            "FROM [Column] c JOIN [Table] t ON c.TableID = t.ID "
            "WHERE t.ModelID = 1"
        ):
            if r["nm"]:
                ids[f"{r['tbl']}\x00{r['nm']}"] = r["ID"]

        for name, props in (snap.get("tables") or {}).items():
            sets = {k: v for k, v in props.items() if v is not None}
            if not sets:
                continue
            c.execute(
                f"UPDATE [Table] SET {', '.join(f'[{k}] = ?' for k in sets)} "
                f"WHERE Name = ? AND ModelID = 1",
                [*sets.values(), name])

        for key, props in (snap.get("columns") or {}).items():
            cid = ids.get(key)
            if cid is None:
                continue
            sets = {k: v for k, v in props.items()
                    if k != "__sort_by__" and v is not None}
            sort_target = props.get("__sort_by__")
            if sort_target is not None:
                tbl = key.split("\x00", 1)[0]
                sets["SortByColumnID"] = ids.get(f"{tbl}\x00{sort_target}", 0)
            if not sets:
                continue
            c.execute(
                f"UPDATE [Column] SET {', '.join(f'[{k}] = ?' for k in sets)} "
                f"WHERE ID = ?", [*sets.values(), cid])

        # Measures and hierarchies (see _snapshot_object_properties). Matched on
        # (table, name); anything the rebuild legitimately dropped is skipped.
        # A snapshotted NULL is applied rather than skipped, because for these
        # the builder's default is the non-NULL value: Description defaults to
        # '' and DisplayFolder to '', so treating NULL as "no opinion" would
        # leave every measure carrying an empty folder it never had.
        for obj, key in (("Measure", "measures"), ("Hierarchy", "hierarchies")):
            for name, props in (snap.get(key) or {}).items():
                tbl, oname = name.split("\x00", 1)
                if not props:
                    continue
                c.execute(
                    f"UPDATE [{obj}] SET "
                    f"{', '.join(f'[{k}] = ?' for k in props)} "
                    f"WHERE Name = ? AND TableID = ("
                    f"  SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1)",
                    [*props.values(), oname, tbl])

        for key, tag in (snap.get("levels") or {}).items():
            tbl, hname, ordinal = key.split("\x00", 2)
            c.execute(
                "UPDATE [Level] SET LineageTag = ? WHERE Ordinal = ? "
                "AND HierarchyID = (SELECT h.ID FROM [Hierarchy] h "
                "  JOIN [Table] t ON h.TableID = t.ID "
                "  WHERE h.Name = ? AND t.Name = ? AND t.ModelID = 1)",
                (tag, int(ordinal), hname, tbl))

        # Restore legacy Query partitions (see _snapshot_object_properties).
        for tname, pprops in (snap.get("partitions") or {}).items():
            c.execute(
                "UPDATE [Partition] SET Type = ?, Mode = ?, QueryDefinition = ? "
                "WHERE TableID = (SELECT ID FROM [Table] "
                "                 WHERE Name = ? AND ModelID = 1)",
                (pprops["Type"], pprops["Mode"], pprops["QueryDefinition"], tname))

        # A table may have at most ONE key column. The builder marks every
        # RowNumber column IsKey (right for a table it authored), but the
        # restore above re-applies IsKey to whichever column the author
        # actually keyed on — a date table's Date column, or a designated
        # business key. Both then claim the key and Power BI Desktop refuses
        # to open the file at all:
        #   "The table 'DateAutoTemplate' has two columns with the IsKey
        #    property set to True."
        # Corpus ground truth (248 RowNumber columns across 24 Desktop-authored
        # files): IsKey is 0 for exactly the 14 whose table has another key
        # column, and 1 for the other 234 — never two per table. IsUnique stays
        # 1 in all 248, so only IsKey is cleared. This is not confined to
        # auto-date system tables: Ecommerce dimDate, IT_Support dim_Date and
        # three MS_Store_Sales tables are ordinary SystemFlags=0 tables.
        c.execute(
            "UPDATE [Column] SET IsKey = 0 "
            "WHERE Type = 3 AND IsKey = 1 AND TableID IN ("
            "  SELECT TableID FROM [Column] WHERE IsKey = 1 AND Type <> 3)")
        conn.commit()

    _modify_metadata_only(dm_path, _do)


def _identity_maps(conn: sqlite3.Connection) -> tuple[dict, dict]:
    """Build (id -> identity, identity -> id) for every referencable entity.

    A rebuild reassigns every primary key, so a carried row's foreign keys are
    meaningless as numbers. Names survive; these maps translate between the two.
    Columns are keyed on ``COALESCE(ExplicitName, InferredName)`` because
    calculated-table and auto-date columns leave ExplicitName NULL.
    """
    fwd: dict[int, tuple] = {}
    rev: dict[tuple, int] = {}

    def _put(oid, ident):
        if oid is None or ident is None:
            return
        fwd[oid] = ident
        rev.setdefault(ident, oid)

    for r in conn.execute("SELECT ID FROM [Model]"):
        _put(r[0], ("Model",))
    for r in conn.execute("SELECT ID FROM [Culture]"):
        _put(r[0], ("Culture",))
    for r in conn.execute("SELECT ID, Name FROM [Table] WHERE ModelID = 1"):
        _put(r[0], ("Table", r[1]))
    for r in conn.execute(
        "SELECT c.ID, t.Name, COALESCE(c.ExplicitName, c.InferredName) "
        "FROM [Column] c JOIN [Table] t ON c.TableID = t.ID "
        "WHERE t.ModelID = 1"
    ):
        _put(r[0], ("Column", r[1], r[2]))
    for r in conn.execute(
        "SELECT m.ID, t.Name, m.Name FROM [Measure] m "
        "JOIN [Table] t ON m.TableID = t.ID WHERE t.ModelID = 1"
    ):
        _put(r[0], ("Measure", r[1], r[2]))
    try:
        for r in conn.execute(
            "SELECT h.ID, t.Name, h.Name FROM [Hierarchy] h "
            "JOIN [Table] t ON h.TableID = t.ID WHERE t.ModelID = 1"
        ):
            _put(r[0], ("Hierarchy", r[1], r[2]))
    except sqlite3.Error:
        pass
    try:
        for r in conn.execute(
            "SELECT k.ID, t.Name, m.Name FROM [KPI] k "
            "JOIN [Measure] m ON k.MeasureID = m.ID "
            "JOIN [Table] t ON m.TableID = t.ID"
        ):
            _put(r[0], ("KPI", r[1], r[2]))
    except sqlite3.Error:
        pass
    try:
        for r in conn.execute("SELECT ID, Name FROM [Expression]"):
            _put(r[0], ("Expression", r[1]))
    except sqlite3.Error:
        pass
    # A relationship has no name; its endpoints are its identity.
    try:
        for r in conn.execute(
            "SELECT r.ID, ft.Name, COALESCE(fc.ExplicitName, fc.InferredName), "
            "       tt.Name, COALESCE(tc.ExplicitName, tc.InferredName) "
            "FROM [Relationship] r "
            "JOIN [Column] fc ON r.FromColumnID = fc.ID "
            "JOIN [Table] ft ON fc.TableID = ft.ID "
            "JOIN [Column] tc ON r.ToColumnID = tc.ID "
            "JOIN [Table] tt ON tc.TableID = tt.ID"
        ):
            _put(r[0], ("Relationship", r[1], r[2], r[3], r[4]))
    except sqlite3.Error:
        pass
    return fwd, rev


def _snapshot_carryable_metadata(conn: sqlite3.Connection) -> dict:
    """Capture the model metadata a from-scratch rebuild would discard.

    The builder writes only [Table], [Column], [Partition], [Relationship] and
    [Measure]; everything else in metadata.sqlitedb is lost. Each row is stored
    with its foreign keys already translated into name-based identities, so it
    can be re-attached after the rebuild has renumbered every ID.
    """
    conn.row_factory = sqlite3.Row
    fwd, _rev = _identity_maps(conn)
    snap: dict[str, list[dict]] = {}
    # Singleton tables are kept verbatim across a rebuild, but two of their
    # columns point INTO tables the rebuild clears, so they would be left
    # dangling. Captured by identity here and repaired after the restore.
    snap["__singletons__"] = {}
    try:
        r = conn.execute("SELECT ID, LinguisticMetadataID FROM [Culture]").fetchone()
        if r and r["LinguisticMetadataID"]:
            snap["__singletons__"]["culture_lm"] = r["LinguisticMetadataID"]
    except sqlite3.Error:
        pass
    try:
        r = conn.execute("SELECT DefaultMeasureID FROM [Model]").fetchone()
        if r and r["DefaultMeasureID"]:
            snap["__singletons__"]["default_measure"] = fwd.get(r["DefaultMeasureID"])
    except sqlite3.Error:
        pass

    for tname, fks, _identity in _CARRY_SPEC:
        try:
            rows = conn.execute(f"SELECT * FROM [{tname}]").fetchall()
        except sqlite3.Error:
            continue
        kept = []
        for row in rows:
            d = dict(row)
            ok = True
            for col, kind in fks.items():
                if col not in d:
                    continue
                val = d[col]
                if val in (None, 0):
                    d[col] = None
                    continue
                if kind.startswith("self:"):
                    # Strip a trailing "?" so the remap key is the real table
                    # name; the "?" only says the reference is OPTIONAL.
                    d[col] = ("#self", kind[5:].rstrip("?"), val)
                    continue
                ident = fwd.get(val)
                if ident is None:
                    if kind.endswith("?"):
                        d[col] = None
                        continue
                    ok = False
                    break
                d[col] = ("#ref",) + ident
            if ok:
                kept.append(d)
        if kept:
            snap[tname] = kept
    return snap


def _restore_carryable_metadata(dm_path: str, snap: dict) -> list[str]:
    """Re-attach snapshotted metadata after a rebuild. Returns what was lost.

    Fail-safe by construction: a row whose owner no longer exists (a deleted
    table, a renamed column) is SKIPPED and reported, never written with a
    dangling foreign key. A half-attached model is worse than a missing
    annotation.
    """
    if not snap:
        return []
    dropped: dict[str, int] = {}

    def _do(conn: sqlite3.Connection):
        conn.row_factory = sqlite3.Row
        c = conn.cursor()
        _fwd, rev = _identity_maps(conn)
        row = c.execute(
            "SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
        max_id = int(row[0]) if row else 0
        remap: dict[str, dict[int, int]] = {}

        for tname, fks, identity in _CARRY_SPEC:
            rows = snap.get(tname)
            if not rows:
                continue
            try:
                cols = [r[1] for r in c.execute(f"PRAGMA table_info([{tname}])")]
            except sqlite3.Error:
                cols = []
            if not cols:
                # The rebuilt schema has no such table, so these rows have
                # nowhere to go. Report it — a table missing from the schema is
                # exactly how [Function] was lost without anyone noticing.
                dropped[tname] = dropped.get(tname, 0) + len(rows)
                continue
            for d in rows:
                resolved = {}
                ok = True
                for col in cols:
                    if col not in d:
                        continue
                    val = d[col]
                    if isinstance(val, tuple) and val and val[0] == "#ref":
                        new = rev.get(tuple(val[1:]))
                        if new is None:
                            ok = fks.get(col, "").endswith("?")
                            if not ok:
                                break
                            new = None
                        resolved[col] = new
                    elif isinstance(val, tuple) and val and val[0] == "#self":
                        new = remap.get(val[1], {}).get(val[2])
                        if new is None:
                            # An OPTIONAL self reference drops to NULL rather
                            # than taking its whole row down with it: a
                            # calculation item should survive losing its format
                            # string, not vanish because of it.
                            ok = fks.get(col, "").endswith("?")
                            if not ok:
                                break
                            new = None
                        resolved[col] = new
                    else:
                        resolved[col] = val
                if not ok:
                    dropped[tname] = dropped.get(tname, 0) + 1
                    continue

                # Never duplicate what the rebuild already reproduced.
                where = " AND ".join(
                    f"[{k}] IS ?" for k in identity if k in resolved)
                if where:
                    dup = c.execute(
                        f"SELECT ID FROM [{tname}] WHERE {where}",
                        [resolved[k] for k in identity if k in resolved],
                    ).fetchone()
                    if dup:
                        remap.setdefault(tname, {})[d["ID"]] = dup["ID"]
                        continue

                max_id += 1
                old_id = d.get("ID")
                resolved["ID"] = max_id
                names = [k for k in cols if k in resolved]
                c.execute(
                    f"INSERT INTO [{tname}] ({','.join('[' + n + ']' for n in names)}) "
                    f"VALUES ({','.join('?' for _ in names)})",
                    [resolved[n] for n in names],
                )
                if old_id is not None:
                    remap.setdefault(tname, {})[old_id] = max_id

        # A calculation group is wired up from BOTH ends, and the rebuild only
        # rewrites one of them. Carrying the CalculationGroup row back is not
        # enough: [Table].CalculationGroupID points the other way, and a
        # calculation group's partition must be Type=7 (CalculationGroup
        # source). Without these two the model loads with the group's table
        # present but inert -- the items never appear.
        try:
            groups = c.execute(
                "SELECT ID, TableID FROM [CalculationGroup]").fetchall()
        except sqlite3.Error:
            groups = []
        for g in groups:
            try:
                c.execute("UPDATE [Table] SET CalculationGroupID = ? "
                          "WHERE ID = ?", (g["ID"], g["TableID"]))
                # QueryDefinition MUST be cleared, not just left behind. The
                # rebuild writes an Enter-data M query for every partition, and
                # Power BI rejects the whole file on open with "Partition 'X'
                # in table 'X' has the QueryDefinition property set which is
                # not a valid field for this partition type." Metadata-level
                # checks all passed -- referential integrity clean, every other
                # field byte-identical to an authored calculation group -- and
                # only opening the file in Desktop surfaced it.
                c.execute("UPDATE [Partition] SET Type = 7, "
                          "QueryDefinition = NULL WHERE TableID = ?",
                          (g["TableID"],))
            except sqlite3.Error:
                pass

        # Re-point the singleton references at the rows just re-created.
        singles = snap.get("__singletons__") or {}
        if "culture_lm" in singles:
            new_lm = remap.get("LinguisticMetadata", {}).get(singles["culture_lm"])
            if new_lm is None:
                row = c.execute("SELECT ID FROM [LinguisticMetadata] "
                                "LIMIT 1").fetchone()
                new_lm = row["ID"] if row else 0
            try:
                c.execute("UPDATE [Culture] SET LinguisticMetadataID = ?", (new_lm,))
            except sqlite3.Error:
                pass
        if singles.get("default_measure"):
            try:
                c.execute("UPDATE [Model] SET DefaultMeasureID = ?",
                          (rev.get(tuple(singles["default_measure"]), 0),))
            except sqlite3.Error:
                pass

        c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'",
                  (str(max_id),))
        conn.commit()

    _modify_metadata_only(dm_path, _do)

    out = []
    for tname, n in sorted(dropped.items(), key=lambda kv: -kv[1]):
        meaning = _CARRY_MEANING.get(tname, tname)
        out.append(f"{n} {tname} row(s) could not be re-attached ({meaning}) — "
                   f"the object they referenced no longer exists after the edit")
    return out


def _rebuild_datamodel(
    info: dict,
    table_updates: dict[str, dict] | None = None,
    extra_tables: list[dict] | None = None,
    extra_measures: list[dict] | None = None,
    extra_relationships: list[dict] | None = None,
    remove_tables: set[str] | None = None,
    remove_relationships: list[tuple[str, str, str, str]] | None = None,
    calc_authoring: bool = False,
    restamp_calc_tables: set[str] | None = None,
    lost_report: list[str] | None = None,
) -> tuple[int, int]:
    """Rebuild the entire DataModel using the builder pipeline.

    Reads all existing tables, measures, relationships, and row data.
    Applies updates/additions/removals, then regenerates the DataModel from scratch.

    Args:
        info: Open file info dict from _ensure_open()
        table_updates: {table_name: {"columns": [...], "rows": [...]}} to replace
        extra_tables: New tables to add: [{"name", "columns", "rows"}, ...]
        extra_measures: New measures: [{"table", "name", "expression", "format_string"}, ...]
        extra_relationships: New rels: [{"from_table", "from_column", "to_table", "to_column"}, ...]
        remove_tables: Set of table names to exclude from rebuild
        remove_relationships: List of (from_table, from_col, to_table, to_col) to exclude

    Returns (old_dm_size, new_dm_size).
    """
    from pbix_mcp.builder import PBIXBuilder
    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
    from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf

    table_updates = table_updates or {}
    extra_tables = extra_tables or []
    extra_measures = extra_measures or []
    extra_relationships = extra_relationships or []
    remove_tables = remove_tables or set()
    remove_relationships = remove_relationships or []
    restamp_calc_tables = restamp_calc_tables or set()
    # Anything the rebuild cannot carry across is collected here and surfaced
    # to the caller. A rebuild used to drop perspectives, Q&A metadata, dynamic
    # format strings and shared M expressions while reporting success with an
    # empty warnings list; silence is the part that made it dangerous.
    lost_metadata: list[str] = []
    _responses.clear_pending_warnings()

    dm_path = os.path.join(info["work_dir"], "DataModel")
    with open(dm_path, "rb") as f:
        dm_bytes = f.read()

    abf = decompress_datamodel(dm_bytes)
    meta_bytes = read_metadata_sqlite(abf)

    # Read structure from metadata
    tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
    tmp.write(meta_bytes)
    tmp.close()
    try:
        conn = sqlite3.connect(tmp.name)
        conn.row_factory = sqlite3.Row

        _AMO_TO_TYPE = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                        10: "Decimal", 11: "Boolean"}

        # Get all existing user tables
        tables = []
        special_tables: list[tuple[str, str]] = []
        field_param_specs: list[dict] = []
        for trow in conn.execute(
            "SELECT ID, Name FROM [Table] WHERE ModelID = 1 "
            "AND Name NOT LIKE 'H$%' AND Name NOT LIKE 'R$%' ORDER BY ID"
        ):
            tid, tname = trow["ID"], trow["Name"]
            cols = []
            for crow in conn.execute(
                "SELECT ExplicitName, ExplicitDataType, DataCategory FROM [Column] "
                "WHERE TableID = ? AND Type = 1 ORDER BY ID", (tid,)
            ):
                dt = _AMO_TO_TYPE.get(crow["ExplicitDataType"], "String")
                cols.append({"name": crow["ExplicitName"], "data_type": dt,
                             "data_category": crow["DataCategory"]})

            # Detect tables the from-scratch rebuild can't faithfully reproduce.
            # A CALCULATED table (Partition.Type=2, e.g. DATATABLE/GENERATESERIES),
            # a table carrying CALCULATED COLUMNS (Column.Type=4 = calc-table
            # column, or Column.Type=2 = DAX column on a normal table), or a
            # DirectQuery/dual partition all have data Power BI derives from a DAX
            # expression / live source — the builder can't reproduce it, so a
            # rebuild would silently open those tables EMPTY or drop the column.
            # These are refused. A MEASURE-ONLY container (no data columns — a
            # "_Measures" table) IS preserved: it re-emits as a RowNumber-only
            # empty table, which the builder now supports.
            #
            # EXCEPTION — field parameters: a calculated table whose partition
            # holds a {("d", NAMEOF('T'[C]), n)} tuple set IS reproducible
            # (its rows are constants with full physical VertiPaq storage), so
            # it is rebuilt as static data and the Desktop metadata shape is
            # re-stamped after the rebuild (_apply_field_parameter_metadata).
            fp = _detect_field_parameter_shape(conn, tid)
            if fp is not None:
                field_param_specs.append(
                    {"table": tname, "query_definition": fp["query_definition"]})
                tables.append({"name": tname, "columns": fp["columns"]})
                continue
            prow = conn.execute(
                "SELECT Type FROM [Partition] WHERE TableID = ? LIMIT 1", (tid,)
            ).fetchone()
            n_type2 = conn.execute(
                "SELECT COUNT(*) FROM [Column] WHERE TableID = ? AND Type = 2",
                (tid,),
            ).fetchone()[0]
            n_type4 = conn.execute(
                "SELECT COUNT(*) FROM [Column] WHERE TableID = ? AND Type = 4",
                (tid,),
            ).fetchone()[0]
            is_calc_table = prow is not None and prow["Type"] == 2
            if (is_calc_table or n_type2 > 0 or n_type4 > 0) \
                    and tname not in remove_tables:
                # A special table being REMOVED needs no reproduction — never
                # let it block the edit (also the escape hatch for any
                # unsupported table: pbix_datamodel_remove_table always works).
                #
                # calc-authoring mode: a NORMAL table whose only special feature
                # is Type=2 calculated columns (not a calc TABLE, not Type=4
                # calc-table columns) is allowed IF the caller re-supplies it via
                # table_updates — pbix_datamodel_add_calculated_column
                # re-materializes its calc columns as data and re-stamps them
                # (Type=2 + Expression) after the rebuild.
                #
                # A calculated TABLE is likewise allowed when the caller names
                # it in restamp_calc_tables — pbix_datamodel_add_calculated_table
                # preserves its rows (Type=4 columns ARE readable) via
                # table_updates and re-stamps the calc-table metadata afterwards.
                calc_ok = calc_authoring and (
                    (not is_calc_table and n_type4 == 0 and n_type2 > 0
                     and tname in table_updates)
                    or (is_calc_table and tname in restamp_calc_tables
                        and tname in table_updates)
                )
                if not calc_ok:
                    kind = ("calculated table" if is_calc_table
                            else "table with calculated columns")
                    special_tables.append((tname, kind))
            tables.append({"name": tname, "columns": cols})

        if special_tables:
            conn.close()
            listing = ", ".join(f"'{n}' ({k})" for n, k in special_tables)
            raise UnsupportedModelEditError(
                "This edit rebuilds the whole DataModel, but the model contains "
                f"tables the rebuild can't preserve: {listing}. Power BI computes "
                "their rows/values from a DAX expression, which the rebuild can't "
                "recompute — rebuilding would reopen them empty or drop the "
                "calculated column. To avoid corrupting the file the edit was "
                "refused. The surgical DataModel tools do NOT rebuild and work on "
                "these models: pbix_datamodel_add_measure / modify_measure / "
                "remove_measure / modify_column. Editing models that contain "
                "calculated tables/columns via the rebuild path is a known "
                "limitation. (Models whose only special table is a measure-only "
                "container ARE supported.)"
            )

        # Get existing measures
        measures = []
        for mrow in conn.execute(
            "SELECT t.Name as tbl, m.Name, m.Expression, m.FormatString, "
            "m.Description, m.DataCategory "
            "FROM Measure m JOIN [Table] t ON m.TableID = t.ID"
        ):
            measures.append({
                "table": mrow["tbl"], "name": mrow["Name"],
                "expression": mrow["Expression"],
                "format_string": mrow["FormatString"] or "",
                "description": mrow["Description"] or "",
                "data_category": mrow["DataCategory"],
            })

        # Get existing relationships — preserve the semantic columns (IsActive,
        # CrossFilteringBehavior, cardinality, ...) so an unrelated datamodel
        # edit doesn't silently reset bidirectional / inactive / many-to-many /
        # one-to-one relationships back to active many-to-one (OpenBI #3).
        rels = []
        for rrow in conn.execute(
            # Calculated-table and auto-date columns carry no ExplicitName —
            # their name lives in InferredName. Reading only ExplicitName gives
            # a relationship endpoint of None, which pre-build validation then
            # rejects as "column does not exist".
            "SELECT ft.Name as ft, "
            "COALESCE(fc.ExplicitName, fc.InferredName) as fc, "
            "tt.Name as tt, "
            "COALESCE(tc.ExplicitName, tc.InferredName) as tc, "
            "r.IsActive, r.CrossFilteringBehavior, "
            "r.FromCardinality, r.ToCardinality, "
            "r.RelyOnReferentialIntegrity, r.JoinOnDateBehavior, "
            "r.SecurityFilteringBehavior, r.Type "
            "FROM Relationship r "
            "JOIN [Table] ft ON r.FromTableID = ft.ID "
            "JOIN [Column] fc ON r.FromColumnID = fc.ID "
            "JOIN [Table] tt ON r.ToTableID = tt.ID "
            "JOIN [Column] tc ON r.ToColumnID = tc.ID"
        ):
            rels.append({
                "from_table": rrow["ft"], "from_column": rrow["fc"],
                "to_table": rrow["tt"], "to_column": rrow["tc"],
                # NULL IsActive (shouldn't occur in Desktop files) defaults to
                # active — 0 is a valid value, so `bool(None)` must not silently
                # deactivate a relationship.
                "is_active": (True if rrow["IsActive"] is None
                              else bool(rrow["IsActive"])),
                "cross_filter_behavior": rrow["CrossFilteringBehavior"] or 1,
                "from_cardinality": rrow["FromCardinality"] or 2,
                "to_cardinality": rrow["ToCardinality"] or 1,
                "rely_on_referential_integrity": bool(
                    rrow["RelyOnReferentialIntegrity"]),
                "join_on_date_behavior": rrow["JoinOnDateBehavior"] or 1,
                "security_filtering_behavior": rrow["SecurityFilteringBehavior"] or 1,
                "relationship_type": rrow["Type"] or 1,
            })

        # Get existing user hierarchies
        user_hierarchies = []
        for hrow in conn.execute(
            "SELECT h.Name, t.Name as TableName FROM Hierarchy h "
            "JOIN [Table] t ON h.TableID = t.ID "
            "WHERE t.ModelID = 1 ORDER BY h.ID"
        ):
            levels = []
            for lrow in conn.execute(
                # COALESCE, not ExplicitName: a calculated-table or auto-date
                # column carries its name in InferredName and leaves
                # ExplicitName NULL. Reading only ExplicitName produced levels
                # with no column, and the hierarchy was then dropped in
                # silence — it cost Agents_Performance.pbix both of its date
                # hierarchies on an edit that reported success.
                "SELECT l.Name, COALESCE(c.ExplicitName, c.InferredName) "
                "       AS ColName FROM Level l "
                "JOIN [Column] c ON l.ColumnID = c.ID "
                "JOIN Hierarchy h ON l.HierarchyID = h.ID "
                "JOIN [Table] t ON h.TableID = t.ID "
                "WHERE h.Name = ? AND t.Name = ? ORDER BY l.Ordinal",
                (hrow["Name"], hrow["TableName"]),
            ):
                levels.append({"name": lrow["Name"], "column": lrow["ColName"]})
            if levels and all(lv["column"] for lv in levels):
                user_hierarchies.append({
                    "table": hrow["TableName"],
                    "name": hrow["Name"],
                    "levels": levels,
                })
            else:
                # Losing a drill-down hierarchy is visible to anyone using the
                # report, so it is never silent.
                lost_metadata.append(
                    f"hierarchy '{hrow['TableName']}'[{hrow['Name']}] could not "
                    f"be reproduced (a level references a column the rebuild "
                    f"cannot name) — the drill-down is gone from the model")

        # Get existing RLS roles
        rls_roles = []
        for rrow in conn.execute(
            "SELECT r.Name, r.Description FROM Role r WHERE r.ModelID = 1"
        ):
            perms = conn.execute(
                "SELECT t.Name as TableName, tp.FilterExpression "
                "FROM TablePermission tp JOIN [Table] t ON tp.TableID = t.ID "
                "WHERE tp.RoleID = (SELECT ID FROM Role WHERE Name = ?)",
                (rrow["Name"],),
            ).fetchall()
            for p in perms:
                rls_roles.append({
                    "role_name": rrow["Name"],
                    "description": rrow["Description"] or "",
                    "table_name": p["TableName"],
                    "filter_expression": p["FilterExpression"],
                })

        # Capture everything the builder does not emit, while the pre-rebuild
        # IDs are still meaningful.
        carry_snapshot = _snapshot_carryable_metadata(conn)
        # …and the authoring properties on the rows it DOES emit, which it
        # writes back as defaults.
        property_snapshot = _snapshot_object_properties(conn)

        conn.close()
    finally:
        os.unlink(tmp.name)

    # Build new DataModel via builder
    builder = PBIXBuilder()
    # Carry the source model's db.xml so the rebuilt database keeps its own
    # declared CompatibilityLevel. The generator hardcodes 1550, which told
    # Analysis Services that a model authored at 1455 was a 1550 database and
    # made it refuse to load the file at all.
    try:
        from pbix_mcp.formats.abf_rebuild import list_abf_files, read_abf_file
        _dbx = [f for f in list_abf_files(abf)
                if f.get("FileName", "").endswith(".db.xml")]
        if _dbx:
            builder._source_db_xml = read_abf_file(abf, _dbx[0])
        # …and the model's own metadata.sqlitedb, so the rebuild keeps its
        # schema era instead of imposing the builder's fixed 63-table one.
        # 20 of the 24 corpus models are an older era than that; rebuilding
        # them from a blank schema invented tables their compatibility level
        # never had, and the service refused to load the result.
        builder._source_metadata = meta_bytes
        # …and its encryption key. Carrying the metadata without the key left
        # DataSource connection strings encrypted under a key the file no
        # longer contained, and the service refused the model with "Failed to
        # decrypt sensitive data".
        _ck = [f for f in list_abf_files(abf)
               if "CryptKey" in f.get("FileName", "")]
        if _ck:
            builder._source_cryptkey = read_abf_file(abf, _ck[0])
    except Exception:
        # Falling back to the generated db.xml is the pre-existing behaviour;
        # never fail an edit over this.
        pass

    # Add existing tables (with optional row updates), skip removed tables
    for tinfo in tables:
        tname = tinfo["name"]
        if tname in remove_tables:
            continue
        if tname in table_updates:
            upd = table_updates[tname]
            builder.add_table(tname, upd["columns"], rows=upd["rows"],
                              calc_columns=upd.get("calc_columns"))
        else:
            # Read existing row data from VertiPaq. If a column cannot be
            # decoded we must NOT fall back to rebuilding the table with no
            # rows — that would silently turn an unrelated metadata edit
            # (add_measure, add_relationship, …) into whole-table data loss.
            # Abort loudly instead; the DataModel on disk is only rewritten at
            # the very end of the rebuild, so the original file stays intact.
            try:
                td = read_table_from_abf(abf, tname, meta_bytes)
            except Exception as e:
                raise InvalidPBIXError(
                    f"Cannot preserve the rows of table '{tname}' during this "
                    f"edit: {e}. Aborting rather than silently rebuilding the "
                    f"table with no data (the file on disk is left unchanged)."
                ) from e
            existing_rows = [
                dict(zip(td["columns"], row_vals))
                for row_vals in td.get("rows", [])
            ]
            builder.add_table(tname, tinfo["columns"], rows=existing_rows)

    # Add new tables
    for et in extra_tables:
        builder.add_table(et["name"], et["columns"], rows=et.get("rows", []))

    # Add all measures (existing + new), skip measures on removed tables
    for m in measures:
        if m["table"] not in remove_tables:
            builder.add_measure(m["table"], m["name"], m["expression"],
                                m.get("description", ""),
                                format_string=m.get("format_string"),
                                data_category=m.get("data_category"))
    for m in extra_measures:
        builder.add_measure(m["table"], m["name"], m["expression"],
                            m.get("description", ""),
                            format_string=m.get("format_string"),
                            data_category=m.get("data_category"))

    # Add all relationships (existing + new), skip removed ones and those referencing removed tables
    remove_rel_set = {(r[0], r[1], r[2], r[3]) for r in remove_relationships}
    for r in rels:
        key = (r["from_table"], r["from_column"], r["to_table"], r["to_column"])
        if key in remove_rel_set:
            continue
        if r["from_table"] in remove_tables or r["to_table"] in remove_tables:
            continue
        # Preserve the existing relationship's full semantics verbatim; keep
        # Desktop's From/To (Many/One) orientation rather than re-deriving it.
        builder.add_relationship(
            r["from_table"], r["from_column"], r["to_table"], r["to_column"],
            is_active=r.get("is_active", True),
            cross_filter_behavior=r.get("cross_filter_behavior", 1),
            from_cardinality=r.get("from_cardinality", 2),
            to_cardinality=r.get("to_cardinality", 1),
            rely_on_referential_integrity=r.get(
                "rely_on_referential_integrity", False),
            join_on_date_behavior=r.get("join_on_date_behavior", 1),
            security_filtering_behavior=r.get("security_filtering_behavior", 1),
            relationship_type=r.get("relationship_type", 1),
            auto_orient=False,
        )
    for r in extra_relationships:
        # Newly-added relationships may carry explicit semantics (from
        # pbix_datamodel_add_relationship); fall back to the historical
        # active/single/many-to-one defaults when they don't. auto_orient stays
        # True: the builder only auto-detects Many/One when the cardinality was
        # left at the default many-to-one, so an explicit 1:*/1:1/*:* keeps the
        # caller's orientation while a plain add still gets the old auto-correct.
        builder.add_relationship(
            r["from_table"], r["from_column"], r["to_table"], r["to_column"],
            is_active=r.get("is_active", True),
            cross_filter_behavior=r.get("cross_filter_behavior", 1),
            from_cardinality=r.get("from_cardinality", 2),
            to_cardinality=r.get("to_cardinality", 1),
            rely_on_referential_integrity=r.get(
                "rely_on_referential_integrity", False),
            join_on_date_behavior=r.get("join_on_date_behavior", 1),
            security_filtering_behavior=r.get("security_filtering_behavior", 1),
            relationship_type=r.get("relationship_type", 1),
            auto_orient=True,
        )

    # Add all user hierarchies (existing, preserved across rebuild)
    for uh in user_hierarchies:
        if uh["table"] not in remove_tables:
            builder.add_user_hierarchy(uh["table"], uh["name"], uh["levels"])

    new_pbix = builder.build()

    # Extract new DataModel from builder output
    import io
    import zipfile
    new_z = zipfile.ZipFile(io.BytesIO(new_pbix))
    new_dm = new_z.read("DataModel")

    # Write new DataModel
    with open(dm_path, "wb") as f:
        f.write(new_dm)

    # Re-attach the model metadata the builder does not emit (perspectives,
    # Q&A, KPIs, dynamic format strings, shared M expressions, auto-date
    # variations, annotations). Snapshotted before the rebuild with every
    # foreign key expressed by name, because the rebuild renumbers every ID.
    if carry_snapshot:
        lost_metadata.extend(
            _restore_carryable_metadata(dm_path, carry_snapshot))

    # Restore hidden-ness, format strings, sort-by targets, summarize-by,
    # data categories and field order. The builder writes defaults for all of
    # them, which turned every rebuild-path edit into a quiet reformatting of
    # the model.
    _restore_object_properties(dm_path, property_snapshot)

    # Re-apply RLS roles (builder doesn't support them, so splice after rebuild)
    if rls_roles:
        def _restore_rls(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0

            roles_created = {}  # role_name -> role_id
            for rls in rls_roles:
                rname = rls["role_name"]
                if rname not in roles_created:
                    max_id += 1
                    roles_created[rname] = max_id
                    c.execute(
                        "INSERT INTO Role (ID, ModelID, Name, Description) VALUES (?, 1, ?, ?)",
                        (max_id, rname, rls.get("description") or None),
                    )

                role_id = roles_created[rname]
                trow = c.execute(
                    "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                    (rls["table_name"],),
                ).fetchone()
                if trow and rls.get("filter_expression"):
                    max_id += 1
                    c.execute(
                        "INSERT INTO TablePermission (ID, RoleID, TableID, FilterExpression) "
                        "VALUES (?, ?, ?, ?)",
                        (max_id, role_id, trow["ID"], rls["filter_expression"]),
                    )

            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        _modify_metadata_only(dm_path, _restore_rls)

    # Re-stamp preserved field parameters (builder rebuilt them as plain
    # static tables; restore the Desktop calculated-partition shape). A table
    # explicitly replaced via table_updates is deliberately demoted to a
    # plain data table — the caller overwrote its columns/rows.
    new_size = len(new_dm)
    if field_param_specs:
        specs = [s for s in field_param_specs
                 if s["table"] not in remove_tables
                 and s["table"] not in table_updates]
        if specs:
            _, new_size = _apply_field_parameter_metadata(dm_path, specs)

    # Clear DAX cache — rebuild changes data
    _dax_cache.clear()

    if lost_report is not None:
        lost_report.extend(lost_metadata)
    for msg in lost_metadata:
        _responses.add_pending_warning(msg)

    return len(dm_bytes), new_size


def _apply_table_source(alias: str, table_name: str,
                        source_json: str) -> str | None:
    """Apply an optional post-write source update for pbix_set_table_data.

    Returns None when no source was requested, a "\\n  Source ..." suffix on
    success, or an "ERROR:<message>" marker the caller converts to a tool
    error (the rows ARE written at that point — the message must say so)."""
    if not source_json:
        return None
    try:
        res = json.loads(pbix_update_data_source(alias, table_name,
                                                 source_json))
    except Exception as e:                          # noqa: BLE001
        return f"ERROR:{e}"
    if not res.get("success"):
        return "ERROR:" + str(res.get("error") or res.get("message"))
    return "\n  Source updated: " + str(res.get("message", "")).split("\n")[0]


@mcp.tool()
def pbix_set_table_data(alias: str, table_name: str, data_json: str,
                        source_json: str = "") -> str:
    """Write/replace actual row data in a table in the DataModel (VertiPaq).

    This encodes the data into VertiPaq column format (IDF + IDFMETA +
    dictionary + HIDX) and rebuilds the ABF with the new column files.
    The DataModel is then XPress9 recompressed.

    Args:
        alias: The alias of the open file
        table_name: Name of the table to write data to
        data_json: JSON object with 'columns' and 'rows':
            {
              "columns": [
                {"name": "Col1", "data_type": "String", "nullable": true},
                {"name": "Col2", "data_type": "Int64", "nullable": false},
                {"name": "Img", "data_type": "String", "data_category": "ImageUrl"}
              ],
              "rows": [
                {"Col1": "hello", "Col2": 42, "Img": "data:image/svg+xml;utf8,..."},
                {"Col1": "world", "Col2": 99, "Img": "data:image/svg+xml;utf8,..."}
              ]
            }
            Supported data_types: String, Int64, Float64, DateTime, Decimal, Boolean.
            The type name is case-insensitive, and the type key may be
            "data_type", "dataType", or "type" — but it must be present on any
            column you want typed: a numeric column with no recognized type
            becomes String (so aggregating measures over it return BLANK). An
            unrecognized type name is rejected with a clear error rather than
            silently defaulting to String.
            Optional per-column "data_category" sets Column.DataCategory —
            e.g. "ImageUrl" so table/matrix cells (and the Power BI service)
            render the value as an image, or "WebUrl" for clickable links.
            It survives later rebuild-based edits.
        source_json: Optional connection parameters applied to the table's
            partition AFTER the rows are written — same format as
            pbix_update_data_source (ledger issues-12: writing a data
            snapshot and pointing the partition at its live source is one
            operation, not two). Empty (default) leaves the partition
            M untouched.
    """
    try:
        info = _ensure_open(alias)
        data = json.loads(data_json)
        if not isinstance(data, dict):
            # A bare row list is the obvious guess and used to surface as
            # "'list' object has no attribute 'get'" — a Python error, not an
            # answer.
            got = "an array of rows" if isinstance(data, list) else type(data).__name__
            return ToolResponse.error(
                f"data_json must be an object with 'columns' and 'rows'; got "
                f"{got}. Expected: {{\"columns\": [{{\"name\": \"Col1\", "
                f"\"data_type\": \"String\"}}], \"rows\": [{{\"Col1\": "
                f"\"hello\"}}]}}. Supported data_types: String, Int64, Float64, "
                f"DateTime, Decimal, Boolean.", ABFRebuildError.code).to_text()
        columns = data.get("columns", [])
        rows = data.get("rows", [])
        if not columns or not rows:
            return ToolResponse.error("'columns' and 'rows' are required and must not be empty.", ABFRebuildError.code).to_text()
        # Normalize column type keys/values up front (accept dataType/type and
        # case-insensitive names) so a mistyped key can't silently ship a
        # numeric column as String — the failure mode that made bound measures
        # return BLANK and Desktop show "Error fetching data" (OpenBI #21). This
        # also turns a columns-as-strings payload into a clear message instead
        # of a raw "string indices must be integers" TypeError.
        from pbix_mcp.builder import normalize_column_defs
        try:
            columns = normalize_column_defs(columns, table_name)
        except (ValueError, TypeError) as e:
            return ToolResponse.error(str(e), ABFRebuildError.code).to_text()

        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Check if table exists — update existing or add new
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        with open(dm_path, "rb") as f:
            dm_check = f.read()
        abf_check = decompress_datamodel(dm_check)
        meta_check = read_metadata_sqlite(abf_check)
        tmp_check = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp_check.write(meta_check)
        tmp_check.close()
        try:
            conn_check = sqlite3.connect(tmp_check.name)
            exists = conn_check.execute(
                "SELECT 1 FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            conn_check.close()
        finally:
            os.unlink(tmp_check.name)

        if exists:
            old_size, new_size = _rebuild_preserving_calc(
                alias, info,
                table_updates={table_name: {"columns": columns, "rows": rows}},
            )
            action = "updated"
        else:
            old_size, new_size = _rebuild_preserving_calc(
                alias, info,
                extra_tables=[{"name": table_name, "columns": columns, "rows": rows}],
            )
            action = "created"

        info["modified"] = True
        src_note = _apply_table_source(alias, table_name, source_json)
        if src_note is not None and src_note.startswith("ERROR:"):
            return ToolResponse.error(
                f"Rows written, but source update failed: {src_note[6:]}",
                "SOURCE_UPDATE_FAILED").to_text()
        return ToolResponse.ok(
            f"Table '{table_name}' {action}: {len(rows)} rows, {len(columns)} columns\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
            + (src_note or "")
        ).to_text()
    except json.JSONDecodeError as e:
        return ToolResponse.error(f"Invalid JSON: {e}", ABFRebuildError.code).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_update_table_rows(alias: str, table_name: str, rows_json: str) -> str:
    """Update row data in an existing table, inferring column types from current schema.

    Reads the existing column definitions from the DataModel metadata,
    then encodes the new rows into VertiPaq format.

    Args:
        alias: The alias of the open file
        table_name: Name of the existing table
        rows_json: JSON array of row objects, e.g. [{"Col1": "val", "Col2": 42}, ...]
    """
    try:
        info = _ensure_open(alias)
        rows = json.loads(rows_json)
        if not rows:
            return ToolResponse.error("rows must not be empty.", ABFRebuildError.code).to_text()

        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Read column definitions from existing metadata
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()
        abf = decompress_datamodel(dm_bytes)
        meta_bytes = read_metadata_sqlite(abf)

        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta_bytes)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            _AMO_TO_TYPE = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                            10: "Decimal", 11: "Boolean"}
            col_rows = conn.execute(
                """SELECT c.ExplicitName, c.ExplicitDataType, c.DataCategory
                   FROM [Column] c
                   JOIN [Table] t ON c.TableID = t.ID
                   WHERE t.Name = ? AND c.Type = 1
                   ORDER BY c.ID""",
                (table_name,)
            ).fetchall()
            conn.close()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        if not col_rows:
            return ToolResponse.error(
                f"Table '{table_name}' not found or has no user columns.",
                "TABLE_NOT_FOUND"
            ).to_text()

        columns = [{"name": cr["ExplicitName"],
                     "data_type": _AMO_TO_TYPE.get(cr["ExplicitDataType"], "String"),
                     "data_category": cr["DataCategory"]}
                    for cr in col_rows]

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            table_updates={table_name: {"columns": columns, "rows": rows}},
        )
        info["modified"] = True
        col_names = [c["name"] for c in columns]
        return ToolResponse.ok(
            f"Table '{table_name}' updated: {len(rows)} rows, {len(columns)} columns\n"
            f"  Columns: {', '.join(col_names)}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except json.JSONDecodeError as e:
        return ToolResponse.error(f"Invalid JSON: {e}", ABFRebuildError.code).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_append_table_rows(alias: str, table_name: str, rows_json: str = "",
                           rows_path: str = "") -> str:
    """APPEND rows to an existing table (issue #46 — the batching load).

    Unlike pbix_set_table_data / pbix_update_table_rows, which REPLACE the
    table's rows, this reads the current rows from VertiPaq, extends them
    with the batch, and re-encodes — so a large load can be pushed in
    batches, each call freeing the previous batch on the caller's side.
    Column schema is inferred from the existing table (like
    pbix_update_table_rows); missing keys in a batch row store as NULL.

    Cost note: every call re-encodes the WHOLE table (VertiPaq needs the
    full column view), so appending N batches costs O(N * table size). For
    a single huge initial load, prefer the streaming source instead:
    pbix_create's per-table "rows_path" (an NDJSON file) encodes once.

    Args:
        alias: The alias of the open file
        table_name: Name of the existing table
        rows_json: JSON array of row objects to append
        rows_path: Path to an NDJSON file (one JSON row object per line) to
            append instead of rows_json — mutually exclusive with it.
    """
    try:
        info = _ensure_open(alias)
        if bool(rows_json) == bool(rows_path):
            return ToolResponse.error(
                "Pass exactly one of rows_json or rows_path.",
                "INVALID_INPUT").to_text()
        if rows_path:
            if not os.path.exists(rows_path):
                return ToolResponse.error(
                    f"rows_path not found: {rows_path}",
                    "INVALID_INPUT").to_text()
            new_rows = _read_ndjson_rows(rows_path)
        else:
            new_rows = json.loads(rows_json)
            if not isinstance(new_rows, list):
                return ToolResponse.error(
                    "rows_json must be a JSON array of row objects.",
                    "INVALID_INPUT").to_text()
        if not new_rows:
            return ToolResponse.error(
                "No rows to append.", "INVALID_INPUT").to_text()

        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()
        abf = decompress_datamodel(dm_bytes)
        meta_bytes = read_metadata_sqlite(abf)

        # Existing schema (data columns only — calculated columns are
        # recomputed by the rebuild's preservation pass).
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta_bytes)
        tmp.close()
        conn = None
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            _AMO_TO_TYPE = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                            10: "Decimal", 11: "Boolean"}
            col_rows = conn.execute(
                """SELECT c.ExplicitName, c.ExplicitDataType, c.DataCategory
                   FROM [Column] c
                   JOIN [Table] t ON c.TableID = t.ID
                   WHERE t.Name = ? AND c.Type = 1
                   ORDER BY c.ID""",
                (table_name,)).fetchall()
        finally:
            try:
                if conn is not None:
                    conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        if not col_rows:
            return ToolResponse.error(
                f"Table '{table_name}' not found or has no user columns.",
                "TABLE_NOT_FOUND").to_text()
        columns = [{"name": cr["ExplicitName"],
                    "data_type": _AMO_TO_TYPE.get(cr["ExplicitDataType"], "String"),
                    "data_category": cr["DataCategory"]}
                   for cr in col_rows]
        col_names = [c["name"] for c in columns]

        # Current rows from VertiPaq (data columns only, matching the schema).
        td = read_table_from_abf(abf, table_name, meta_bytes,
                                 include_calculated=False)
        existing_cols = td.get("columns") or []
        idx = {c: existing_cols.index(c) for c in col_names
               if c in existing_cols}
        combined = [
            {c: (row[idx[c]] if c in idx else None) for c in col_names}
            for row in (td.get("rows") or [])
        ]
        prev_count = len(combined)
        combined.extend(new_rows)

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            table_updates={table_name: {"columns": columns, "rows": combined}},
        )
        info["modified"] = True
        return ToolResponse.ok(
            f"Table '{table_name}': appended {len(new_rows)} rows "
            f"({prev_count:,} -> {len(combined):,})\n"
            f"  Columns: {', '.join(col_names)}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except json.JSONDecodeError as e:
        return ToolResponse.error(f"Invalid JSON: {e}", ABFRebuildError.code).to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "INVALID_INPUT").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_list_tables(alias: str) -> str:
    """List all tables in the data model with row/column counts.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader, format_statistics_table
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))
        stats = model.statistics
        return ToolResponse.ok(format_statistics_table(stats)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), DataModelCompressionError.code).to_text()


# ---- Section 7b: Lightweight metadata-only modification ----

def _query_metadata_rows(dm_path: str, sql: str, params: tuple = ()) -> list:
    """Run a read-only query against the DataModel's metadata SQLite.

    Returns plain tuples; the caller owns the formatting.
    """
    import tempfile

    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    with open(dm_path, "rb") as f:
        dm_bytes = f.read()
    meta_bytes = read_metadata_sqlite(decompress_datamodel(dm_bytes))
    if not meta_bytes:
        raise DataModelCompressionError(
            "Could not extract metadata.sqlitedb from ABF.")

    fd, tmp = tempfile.mkstemp(suffix=".sqlitedb",
                               dir=os.path.dirname(dm_path))
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(meta_bytes)
        conn = sqlite3.connect(tmp)
        try:
            return list(conn.execute(sql, params).fetchall())
        finally:
            conn.close()
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass


def _modify_metadata_only(
    dm_path: str, modifier_fn: Callable[[sqlite3.Connection], None]
) -> tuple[int, int]:
    """Lightweight metadata modification — no full DataModel rebuild.

    Only modifies metadata.sqlitedb inside the ABF. Does NOT regenerate
    VertiPaq binary data, H$ hierarchies, or R$ relationship indexes.

    Safe for: Partition.QueryDefinition, Partition.Mode changes.
    NOT safe for: adding/removing tables, columns, or relationships.

    Returns (old_dm_size, new_dm_size).
    """
    from pbix_mcp.formats.abf_rebuild import rebuild_abf_with_modified_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import (
        compress_datamodel,
        decompress_datamodel,
    )

    with open(dm_path, "rb") as f:
        dm_bytes = f.read()

    abf = decompress_datamodel(dm_bytes)
    new_abf = rebuild_abf_with_modified_sqlite(abf, modifier_fn)
    # original_dm enables chunk reuse when the encoder can't re-emit the
    # model's incompressible VertiPaq chunks (real-world files; the
    # "Compression failed or not effective" class).
    new_dm = compress_datamodel(new_abf, original_dm=dm_bytes)

    from pbix_mcp.formats.datamodel_roundtrip import _detect_format
    if _detect_format(new_dm) == "uncompressed":
        # The uncompressed fallback engaged — the edit is intact, but flag it.
        _responses.add_pending_warning(
            "DataModel stored UNCOMPRESSED: the model contains chunks the "
            "XPress9 encoder cannot re-emit and no reusable prefix covered "
            "the edit. The file stays valid (the PBIX ZIP layer still "
            "deflates it) but is larger on disk.")

    with open(dm_path, "wb") as f:
        f.write(new_dm)

    # Clear DAX cache — metadata changes may affect measure evaluation
    _dax_cache.clear()

    return len(dm_bytes), len(new_dm)


# ---- Section 8: DataModel WRITE tools (via XPress9 round-trip) ----

def _modify_metadata_sqlite(
    dm_path: str, modifier_fn: Callable[[sqlite3.Connection], None],
    info: dict | None = None,
) -> tuple:
    """Modify metadata via full DataModel rebuild.

    Applies modifier_fn to a temporary copy of the current metadata to
    determine the changes, then reads the modified measures/relationships
    and rebuilds the entire DataModel via the builder pipeline.

    This avoids ALL post-build ABF modification which causes
    NullReferenceException at RunModelSchemaValidation.

    Args:
        dm_path: Path to the DataModel file inside the work_dir
        modifier_fn: Function that receives a sqlite3.Connection and should
                     make changes + commit.
        info: Open file info dict (required for full rebuild)

    Returns:
        Tuple of (original_dm_bytes, new_dm_bytes, None)
    """
    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    with open(dm_path, "rb") as f:
        dm_bytes = f.read()

    if info is None:
        work_dir = os.path.dirname(dm_path)
        info = {"work_dir": work_dir, "path": dm_path}

    # Apply the modifier to a TEMPORARY copy of metadata to see what changed.
    # Then rebuild the entire DataModel with the modified metadata's
    # measures and relationships baked in.
    abf = decompress_datamodel(dm_bytes)
    meta_bytes = read_metadata_sqlite(abf)

    fd, tmp_path = tempfile.mkstemp(suffix=".sqlitedb")
    try:
        os.write(fd, meta_bytes)
        os.close(fd)
        fd = None

        conn = sqlite3.connect(tmp_path)
        conn.row_factory = sqlite3.Row
        try:
            # Apply the modifier
            modifier_fn(conn)
            conn.commit()

            # Read the MODIFIED measures (these will replace the builder's measures)
            _AMO_TO_TYPE = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                            10: "Decimal", 11: "Boolean"}

            modified_measures = []
            for mrow in conn.execute(
                "SELECT t.Name as tbl, m.Name, m.Expression, m.FormatString, "
                "m.Description "
                "FROM Measure m JOIN [Table] t ON m.TableID = t.ID"
            ):
                modified_measures.append({
                    "table": mrow["tbl"], "name": mrow["Name"],
                    "expression": mrow["Expression"],
                    "format_string": mrow["FormatString"] or "",
                    "description": mrow["Description"] or "",
                })

            # Read tables and relationships from modified metadata
            modified_tables = []
            for trow in conn.execute(
                "SELECT ID, Name FROM [Table] WHERE ModelID = 1 "
                "AND Name NOT LIKE 'H$%' AND Name NOT LIKE 'R$%' ORDER BY ID"
            ):
                cols = [{"name": c["ExplicitName"],
                         "data_type": _AMO_TO_TYPE.get(c["ExplicitDataType"], "String")}
                        for c in conn.execute(
                            "SELECT ExplicitName, ExplicitDataType FROM [Column] "
                            "WHERE TableID = ? AND Type = 1 ORDER BY ID", (trow["ID"],))]
                modified_tables.append({"name": trow["Name"], "columns": cols})

            modified_rels = []
            for rrow in conn.execute(
                "SELECT ft.Name as ft, "
                "COALESCE(fc.ExplicitName, fc.InferredName) as fc, "
                "tt.Name as tt, "
                "COALESCE(tc.ExplicitName, tc.InferredName) as tc "
                "FROM Relationship r "
                "JOIN [Table] ft ON r.FromTableID = ft.ID "
                "JOIN [Column] fc ON r.FromColumnID = fc.ID "
                "JOIN [Table] tt ON r.ToTableID = tt.ID "
                "JOIN [Column] tc ON r.ToColumnID = tc.ID"
            ):
                modified_rels.append({
                    "from_table": rrow["ft"], "from_column": rrow["fc"],
                    "to_table": rrow["tt"], "to_column": rrow["tc"],
                })
        finally:
            conn.close()
    finally:
        if fd is not None:
            os.close(fd)
        try:
            os.unlink(tmp_path)
        except OSError:
            pass

    # Now rebuild using the builder with the modified state.
    # We pass all measures/rels as the "current" state — _rebuild_datamodel
    # reads its own measures/rels from metadata, so we need to override.
    from pbix_mcp.builder import PBIXBuilder
    from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf

    builder = PBIXBuilder()
    for tinfo in modified_tables:
        tname = tinfo["name"]
        # See _rebuild_datamodel: never silently substitute rows=[] on a read
        # failure — that would destroy the table's data during an unrelated
        # metadata edit. Abort loudly and leave the file unchanged.
        try:
            td = read_table_from_abf(abf, tname, meta_bytes)
        except Exception as e:
            raise InvalidPBIXError(
                f"Cannot preserve the rows of table '{tname}' during this "
                f"edit: {e}. Aborting rather than silently rebuilding the "
                f"table with no data (the file on disk is left unchanged)."
            ) from e
        existing_rows = [dict(zip(td["columns"], rv))
                         for rv in td.get("rows", [])]
        builder.add_table(tname, tinfo["columns"], rows=existing_rows)

    for m in modified_measures:
        builder.add_measure(m["table"], m["name"], m["expression"],
                            m.get("description", ""),
                            format_string=m.get("format_string"),
                            data_category=m.get("data_category"))

    for r in modified_rels:
        builder.add_relationship(
            r["from_table"], r["from_column"], r["to_table"], r["to_column"]
        )

    new_pbix = builder.build()

    import io
    import zipfile
    new_z = zipfile.ZipFile(io.BytesIO(new_pbix))
    new_dm = new_z.read("DataModel")

    with open(dm_path, "wb") as f:
        f.write(new_dm)

    return dm_bytes, new_dm, None


@mcp.tool()
def pbix_datamodel_query_metadata(alias: str, sql_query: str) -> str:
    """Run a read-only SQL query on the DataModel's metadata SQLite.

    The metadata store is a **SQLite database** (the ABF's metadata.sqlitedb
    with tables like [Table], [Column], Measure, Relationship, Partition), NOT
    an Analysis Services $SYSTEM rowset — use SQLite syntax, and quote the
    reserved names: ``SELECT Name FROM [Table]``.

    Args:
        alias: The alias of the open file
        sql_query: SQLite query to run (e.g., "SELECT Name, Expression FROM Measure")
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found in this file.", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        abf = decompress_datamodel(dm_bytes)
        meta_bytes = read_metadata_sqlite(abf)

        if not meta_bytes:
            return ToolResponse.error("Could not extract metadata.sqlitedb from ABF.", DataModelCompressionError.code).to_text()

        # Write to temp file for sqlite3
        tmp = os.path.join(info["work_dir"], "_meta_query.tmp")
        with open(tmp, "wb") as f:
            f.write(meta_bytes)

        conn = None
        try:
            conn = sqlite3.connect(tmp)
            conn.row_factory = sqlite3.Row
            cursor = conn.execute(sql_query)
            rows = cursor.fetchall()
            columns = [d[0] for d in cursor.description] if cursor.description else []
        finally:
            # Close BEFORE removing, on every path: with the close on the
            # success path only, a bad query left the handle open and the
            # os.remove raised "[WinError 32] file in use" on Windows —
            # masking the caller's real SQL error (issue #24 r22#3).
            if conn is not None:
                try:
                    conn.close()
                except Exception:
                    pass
            try:
                os.remove(tmp)
            except OSError:
                pass

        if not rows:
            return ToolResponse.ok("Query returned no results.").to_text()

        # Format output
        lines = [" | ".join(columns)]
        lines.append("-" * len(lines[0]))
        for row in rows[:200]:
            lines.append(" | ".join(str(row[c]) for c in columns))
        result = "\n".join(lines)
        if len(rows) > 200:
            result += f"\n... ({len(rows)} total rows, showing first 200)"
        return ToolResponse.ok(result).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_modify_metadata(alias: str, sql_statement: str) -> str:
    """Execute a SQL DDL/DML statement on the DataModel's metadata SQLite.

    This allows direct manipulation of the metadata database (tables, measures,
    columns, relationships, etc.). The ABF is fully rebuilt.

    Args:
        alias: The alias of the open file
        sql_statement: SQL statement to execute (INSERT, UPDATE, DELETE, ALTER, etc.)
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        changes = [0]

        def _do_sql(conn: sqlite3.Connection):
            conn.execute(sql_statement)
            changes[0] = conn.total_changes
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_sql)
        info["modified"] = True
        return ToolResponse.ok(
            f"SQL executed successfully.\n"
            f"  Changes: {changes[0]}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_modify_measure(
    alias: str, measure_name: str, new_expression: str = "",
    new_format_string: str = "", new_data_category: str = "",
    new_data_type: str = ""
) -> str:
    """Modify a DAX measure's expression / format / data category / data type.

    Performs a metadata-only splice so expressions of any length are
    supported. Every parameter is optional — empty string means "leave
    unchanged", so a format-string or DataCategory change no longer forces
    the caller to read and re-send the current expression. At least one
    change must be provided. To CLEAR an existing DataCategory use
    pbix_datamodel_set_measure_category.

    Args:
        alias: The alias of the open file
        measure_name: Name of the measure to modify
        new_expression: New DAX expression (empty = leave unchanged)
        new_format_string: New format string (empty = leave unchanged)
        new_data_category: DataCategory to set — e.g. "ImageUrl" so
            table/matrix cells (and the Power BI service) render the measure's
            data-URI string as an image. Empty = leave unchanged; clear with
            pbix_datamodel_set_measure_category.
        new_data_type: Measure result type to set — a name
            (String/Int64/Double/Decimal/DateTime/Boolean) or AMO code. Empty =
            leave unchanged. Use this to fix a measure whose stored DataType
            lies about its result (e.g. a decimal measure stuck at Int64, which
            the Power BI service truncates to a whole number).
    """
    try:
        from pbix_mcp.builder import (
            find_reserved_var_names,
            normalize_measure_data_type,
        )
        resolved_new_dt = normalize_measure_data_type(new_data_type)

        # Reject reserved DAX/MDX names used as VARs (see pbix_datamodel_add_measure).
        if new_expression:
            reserved = find_reserved_var_names(new_expression)
            if reserved:
                names = ", ".join(f"'{n}'" for n in reserved)
                return ToolResponse.error(
                    f"new_expression uses reserved DAX/MDX name(s) as VAR "
                    f"variable(s): {names}. Analysis Services rejects them and "
                    f"the visual goes blank in the Power BI service. Rename the "
                    f"VAR(s), e.g. VAR {reserved[0]} -> VAR _{reserved[0]}.",
                    "RESERVED_VAR_NAME").to_text()

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()
        if not (new_expression or new_format_string or new_data_category
                or resolved_new_dt):
            return ToolResponse.error(
                "Nothing to change — provide new_expression, "
                "new_format_string, new_data_category, and/or new_data_type.",
                "NOTHING_TO_CHANGE").to_text()

        old_info = {}

        def _do_modify(conn: sqlite3.Connection):
            c = conn.cursor()
            c.execute("SELECT ID, Expression FROM Measure WHERE Name = ?", (measure_name,))
            row = c.fetchone()
            if not row:
                raise ValueError(f"Measure '{measure_name}' not found")
            old_info["id"] = row[0]
            old_info["expression"] = row[1]

            updates = []
            params: list = []
            if new_expression:
                updates.append("Expression = ?")
                params.append(new_expression)
            if new_format_string:
                updates.append("FormatString = ?")
                params.append(new_format_string)
            if new_data_category:
                updates.append("DataCategory = ?")
                params.append(new_data_category)
            if resolved_new_dt:
                updates.append("DataType = ?")
                params.append(resolved_new_dt)
            params.append(measure_name)

            c.execute(f"UPDATE Measure SET {', '.join(updates)} WHERE Name = ?", params)
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_modify)
        info["modified"] = True
        changed = []
        if new_expression:
            changed.append(f"  Old: {old_info.get('expression', '?')}\n"
                           f"  New: {new_expression}")
        if new_format_string:
            changed.append(f"  FormatString: {new_format_string}")
        if new_data_category:
            changed.append(f"  DataCategory: {new_data_category}")
        if resolved_new_dt:
            _dt_names = {2: "String", 6: "Int64", 8: "Double",
                         9: "DateTime", 10: "Decimal", 11: "Boolean"}
            changed.append(
                f"  DataType: {_dt_names.get(resolved_new_dt, resolved_new_dt)}")
        return ToolResponse.ok(
            f"Measure '{measure_name}' updated:\n"
            + "\n".join(changed) + "\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_set_measure_category(
    alias: str, measure_name: str, data_category: str = ""
) -> str:
    """Set or CLEAR a measure's DataCategory — no expression required.

    A metadata-only splice touching nothing but Measure.DataCategory.
    Empty ``data_category`` clears it (SQL NULL) — the only first-class
    clearing path (pbix_datamodel_modify_measure treats empty as "leave
    unchanged"). Common values: "ImageUrl" (cells render a data-URI string
    as an image), "WebUrl" (clickable link).

    Args:
        alias: The alias of the open file
        measure_name: Name of the measure
        data_category: The category to set, or empty to clear it
    """
    try:
        logger.info("pbix_datamodel_set_measure_category %r -> %r",
                    measure_name, data_category)
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        def _do_set(conn: sqlite3.Connection):
            c = conn.cursor()
            c.execute("SELECT ID FROM Measure WHERE Name = ?", (measure_name,))
            if not c.fetchone():
                raise ValueError(f"Measure '{measure_name}' not found")
            c.execute("UPDATE Measure SET DataCategory = ? WHERE Name = ?",
                      (data_category or None, measure_name))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_set)
        info["modified"] = True
        return ToolResponse.ok(
            f"Measure '{measure_name}' DataCategory "
            + (f"set to '{data_category}'." if data_category else "cleared.")
            + f"\n  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_add_measure(
    alias: str, table_name: str, measure_name: str, expression: str,
    format_string: str = "", description: str = "", data_category: str = "",
    data_type: str = ""
) -> str:
    """Create a new DAX measure in the specified table.

    The ABF is fully rebuilt so measures of any size are supported.

    Args:
        alias: The alias of the open file
        table_name: Table to add the measure to
        measure_name: Name of the new measure
        expression: DAX expression
        format_string: Optional format string
        description: Optional description
        data_category: Optional DataCategory — e.g. "ImageUrl" so table/matrix
            cells (and the Power BI service) render the measure's
            ``data:image/svg+xml;utf8,...`` string as an image, or "WebUrl"
            for clickable links. Default: none (current behavior).
        data_type: Optional measure result type — a name
            (String/Int64/Double/Decimal/DateTime/Boolean) or AMO code. When
            omitted, the type is inferred from the expression (text -> String,
            otherwise Double). Double is used instead of a hardcoded Int64 so
            decimal and percentage measures are not truncated in the service.
    """
    try:
        from pbix_mcp.builder import (
            find_reserved_var_names,
            infer_measure_data_type,
            normalize_measure_data_type,
        )
        resolved_dt = normalize_measure_data_type(data_type)
        if not resolved_dt:
            resolved_dt = infer_measure_data_type(expression)

        # A VAR whose name is a DAX function or reserved keyword compiles in our
        # lenient engine but makes Analysis Services fail the whole visual in the
        # Power BI service (MdxScript "Failed to resolve name 'SYNTAXERROR'").
        # Reject it here, matching what Power BI Desktop enforces.
        reserved = find_reserved_var_names(expression)
        if reserved:
            names = ", ".join(f"'{n}'" for n in reserved)
            return ToolResponse.error(
                f"Measure '{measure_name}' uses reserved DAX/MDX name(s) as VAR "
                f"variable(s): {names}. These are valid in the local engine but "
                f"Analysis Services rejects them, blanking the visual in the "
                f"Power BI service. Rename the VAR(s) (e.g. prefix with an "
                f"underscore or a word like 'v'): VAR {reserved[0]} -> "
                f"VAR _{reserved[0]}.",
                "RESERVED_VAR_NAME").to_text()

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        def _do_add(conn: sqlite3.Connection):
            c = conn.cursor()

            # Get table ID
            c.execute("SELECT ID FROM [Table] WHERE Name = ?", (table_name,))
            trow = c.fetchone()
            if not trow:
                raise ValueError(f"Table '{table_name}' not found")
            table_id = trow[0]

            # Check if measure already exists
            c.execute("SELECT ID FROM Measure WHERE Name = ?", (measure_name,))
            if c.fetchone():
                raise ValueError(f"Measure '{measure_name}' already exists")

            # A measure cannot share its name (case-insensitively) with a
            # column on the SAME table. Analysis Services rejects the resulting
            # calculation script — the model fails to process with "One or more
            # errors were encountered in the MDX script" when opened in the
            # Power BI service, even though our lenient local engine renders it
            # fine. Power BI Desktop's UI prevents this collision; enforce it
            # here so we never emit a file that loads locally but breaks online.
            c.execute(
                "SELECT ExplicitName, InferredName FROM [Column] "
                "WHERE TableID = ?",
                (table_id,),
            )
            mname_lower = measure_name.lower()
            for exp_name, inf_name in c.fetchall():
                col_disp = exp_name if exp_name is not None else inf_name
                if col_disp and col_disp.lower() == mname_lower:
                    raise ValueError(
                        f"Measure '{measure_name}' collides with column "
                        f"{table_name}[{col_disp}] on the same table "
                        f"(same name, case-insensitive). Analysis Services "
                        f"rejects this and the model will fail to load in the "
                        f"Power BI service. Rename the measure or the column."
                    )

            # Get next ID from MAXID (PBI's global ID counter).
            # MAXID is always >= the highest ID across all tables.
            # Using MAX(ID) per table misses IDs in system tables like
            # AttributeHierarchyStorage, SegmentMapStorage, etc.
            maxid_row = c.execute(
                "SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'"
            ).fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0
            new_id = max_id + 1

            # Use Windows FILETIME timestamp (matching builder format).
            # NAIVE UTC (utcnow() is deprecated): the delta's epoch is naive.
            import datetime
            now = datetime.datetime.now(
                datetime.timezone.utc).replace(tzinfo=None)
            epoch = datetime.datetime(1601, 1, 1)
            filetime = int((now - epoch).total_seconds() * 10_000_000)

            # Generate a LineageTag UUID
            import uuid
            lineage_tag = str(uuid.uuid4())

            c.execute(
                """INSERT INTO Measure (ID, TableID, Name, Description, DataType,
                    Expression, FormatString, IsHidden, State, ModifiedTime,
                    StructureModifiedTime, KPIID, IsSimpleMeasure, ErrorMessage,
                    DisplayFolder, DetailRowsDefinitionID, DataCategory,
                    FormatStringDefinitionID, LineageTag, SourceLineageTag)
                VALUES (?, ?, ?, ?, ?, ?, ?, 0, 1, ?, ?, 0, 0, NULL,
                    NULL, 0, ?, 0, ?, NULL)""",
                (new_id, table_id, measure_name, description or None,
                 resolved_dt, expression, format_string or None,
                 filetime, filetime, data_category or None, lineage_tag)
            )
            # Update MAXID so subsequent adds get a fresh ID
            c.execute(
                "UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'",
                (str(new_id),)
            )
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        _dt_names = {2: "String", 6: "Int64", 8: "Double",
                     9: "DateTime", 10: "Decimal", 11: "Boolean"}
        dt_label = _dt_names.get(resolved_dt, str(resolved_dt))
        dt_note = "inferred" if not normalize_measure_data_type(data_type) \
            else "explicit"
        return ToolResponse.ok(
            f"Measure '{measure_name}' added to table '{table_name}':\n"
            f"  Expression: {expression}\n"
            f"  DataType: {dt_label} ({dt_note})\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_remove_measure(alias: str, measure_name: str) -> str:
    """Delete a DAX measure from the DataModel.

    The ABF is fully rebuilt.

    Args:
        alias: The alias of the open file
        measure_name: Name of the measure to remove
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        old_info = {}

        def _do_remove(conn: sqlite3.Connection):
            c = conn.cursor()
            c.execute(
                "SELECT m.ID, m.Expression, t.Name FROM Measure m "
                "JOIN [Table] t ON m.TableID = t.ID "
                "WHERE m.Name = ?",
                (measure_name,)
            )
            row = c.fetchone()
            if not row:
                raise ValueError(f"Measure '{measure_name}' not found")
            old_info["id"] = row[0]
            old_info["expression"] = row[1]
            old_info["table"] = row[2]

            c.execute("DELETE FROM Measure WHERE Name = ?", (measure_name,))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_remove)
        info["modified"] = True
        return ToolResponse.ok(
            f"Measure '{measure_name}' removed from table '{old_info.get('table', '?')}':\n"
            f"  Old expression: {old_info.get('expression', '?')}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_add_relationship(
    alias: str,
    from_table: str,
    from_column: str,
    to_table: str,
    to_column: str,
    cardinality: str = "ManyToOne",
    cross_filter_direction: str = "single",
    is_active: bool = True,
) -> str:
    """Add a relationship between two tables. Rebuilds the DataModel.

    Creates a cross-table relationship with R$ index tables in VertiPaq
    (except many-to-many, which by design has no join index).

    Args:
        alias: The alias of the open file
        from_table: "From" table name — the many side for the default ManyToOne
        from_column: Join column in the from table
        to_table: "To" table name — the one side for the default ManyToOne
        to_column: Join column in the to table
        cardinality: One of "ManyToOne" (default), "OneToMany", "OneToOne",
            "ManyToMany". Accepts the glyph forms "*:1", "1:*", "1:1", "*:*".
            For OneToOne/ManyToMany both join columns should be unique / both
            non-unique respectively, matching Power BI Desktop's rules.
        cross_filter_direction: "single" (default, filters flow one way) or
            "both" (bidirectional). Accepts 1/2.
        is_active: False creates an inactive relationship (activate it from DAX
            with USERELATIONSHIP). Only one active relationship may exist per
            table pair+path, exactly as in Desktop.
    """
    _CARD = {
        "manytoone": (2, 1), "*:1": (2, 1), "m:1": (2, 1),
        "onetomany": (1, 2), "1:*": (1, 2), "1:m": (1, 2),
        "onetoone": (1, 1), "1:1": (1, 1),
        "manytomany": (2, 2), "*:*": (2, 2), "m:m": (2, 2),
    }
    key = str(cardinality).strip().lower().replace(" ", "").replace("-", "")
    if key not in _CARD:
        return ToolResponse.error(
            f"Invalid cardinality {cardinality!r}. Use one of: ManyToOne, "
            "OneToMany, OneToOne, ManyToMany.", "INVALID_ARGUMENT").to_text()
    from_card, to_card = _CARD[key]

    xf = str(cross_filter_direction).strip().lower()
    if xf in ("single", "onedirection", "one", "1"):
        cross_filter = 1
    elif xf in ("both", "bothdirections", "bidirectional", "2"):
        cross_filter = 2
    else:
        return ToolResponse.error(
            f"Invalid cross_filter_direction {cross_filter_direction!r}. "
            "Use 'single' or 'both'.", "INVALID_ARGUMENT").to_text()

    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            extra_relationships=[{
                "from_table": from_table, "from_column": from_column,
                "to_table": to_table, "to_column": to_column,
                "is_active": bool(is_active),
                "cross_filter_behavior": cross_filter,
                "from_cardinality": from_card,
                "to_cardinality": to_card,
            }],
        )
        info["modified"] = True
        _note = ""
        if (from_card, to_card) == (1, 1):
            # A 1:1 is stored as a genuine one-to-one: Power BI forces
            # cross-filter Both and the builder emits the reverse R$ index.
            _card_label, cross_filter = "1:1", 2
        else:
            _card_label = {(2, 1): "*:1", (1, 2): "1:*", (2, 2): "*:*"}[
                (from_card, to_card)]
        _dir_label = "both" if cross_filter == 2 else "single"
        _act_label = "active" if is_active else "inactive"
        return ToolResponse.ok(
            f"Relationship added: {from_table}.{from_column} → {to_table}.{to_column}\n"
            f"  cardinality={_card_label}, cross-filter={_dir_label}, {_act_label}"
            f"{_note}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_remove_relationship(
    alias: str,
    from_table: str,
    from_column: str,
    to_table: str,
    to_column: str,
) -> str:
    """Remove a relationship between two tables. Rebuilds the DataModel.

    Args:
        alias: The alias of the open file
        from_table: Fact table name (many side)
        from_column: Foreign key column in fact table
        to_table: Dimension table name (one side)
        to_column: Primary key column in dimension table
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            remove_relationships=[(from_table, from_column, to_table, to_column)],
        )
        info["modified"] = True
        return ToolResponse.ok(
            f"Relationship removed: {from_table}.{from_column} → {to_table}.{to_column}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_modify_relationship(
    alias: str,
    from_table: str,
    from_column: str,
    to_table: str,
    to_column: str,
    cardinality: str = "",
    cross_filter_direction: str = "",
    is_active: str = "",
) -> str:
    """Change an existing relationship in place (no remove + re-add).

    Every parameter is optional — empty means "leave unchanged"; at least one
    change must be given. Toggling ``is_active`` / ``cross_filter_direction`` is
    a metadata-only splice, so it also works on models the rebuild path refuses
    (those containing calculated tables/columns). Changing ``cardinality``
    re-runs the rebuild so the R$ join indexes are regenerated to match.

    Args:
        alias: The alias of the open file
        from_table: "From" table of the existing relationship
        from_column: Join column in the from table
        to_table: "To" table of the existing relationship
        to_column: Join column in the to table
        cardinality: New cardinality — "ManyToOne", "OneToMany", "OneToOne",
            "ManyToMany" (or "*:1", "1:*", "1:1", "*:*"). Empty = unchanged.
        cross_filter_direction: "single" or "both". Empty = unchanged.
        is_active: "true" or "false". Empty = unchanged.
    """
    try:
        _CARD = {
            "manytoone": (2, 1), "*:1": (2, 1), "m:1": (2, 1),
            "onetomany": (1, 2), "1:*": (1, 2), "1:m": (1, 2),
            "onetoone": (1, 1), "1:1": (1, 1),
            "manytomany": (2, 2), "*:*": (2, 2), "m:m": (2, 2),
        }
        new_card = None
        if cardinality:
            key = str(cardinality).strip().lower().replace(" ", "").replace("-", "")
            if key not in _CARD:
                return ToolResponse.error(
                    f"Invalid cardinality {cardinality!r}. Use one of: ManyToOne, "
                    "OneToMany, OneToOne, ManyToMany.", "INVALID_ARGUMENT").to_text()
            new_card = _CARD[key]

        new_xf = None
        if cross_filter_direction:
            xf = str(cross_filter_direction).strip().lower()
            if xf in ("single", "onedirection", "one", "1"):
                new_xf = 1
            elif xf in ("both", "bothdirections", "bidirectional", "2"):
                new_xf = 2
            else:
                return ToolResponse.error(
                    f"Invalid cross_filter_direction {cross_filter_direction!r}. "
                    "Use 'single' or 'both'.", "INVALID_ARGUMENT").to_text()

        new_active = None
        if str(is_active).strip():
            av = str(is_active).strip().lower()
            if av in ("true", "1", "yes"):
                new_active = True
            elif av in ("false", "0", "no"):
                new_active = False
            else:
                return ToolResponse.error(
                    f"Invalid is_active {is_active!r}. Use 'true' or 'false'.",
                    "INVALID_ARGUMENT").to_text()

        if new_card is None and new_xf is None and new_active is None:
            return ToolResponse.error(
                "Nothing to change — provide cardinality, "
                "cross_filter_direction, and/or is_active.",
                "NOTHING_TO_CHANGE").to_text()

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.", DataModelCompressionError.code).to_text()

        found: dict = {}

        def _locate(c: sqlite3.Cursor):
            """Find the relationship row id + current semantics."""
            row = c.execute(
                "SELECT r.ID, r.IsActive, r.CrossFilteringBehavior, "
                "r.FromCardinality, r.ToCardinality "
                "FROM Relationship r "
                "JOIN [Table] ft ON r.FromTableID = ft.ID "
                "JOIN [Column] fc ON r.FromColumnID = fc.ID "
                "JOIN [Table] tt ON r.ToTableID = tt.ID "
                "JOIN [Column] tc ON r.ToColumnID = tc.ID "
                "WHERE ft.Name = ? AND COALESCE(fc.ExplicitName, fc.InferredName) = ? "
                "AND tt.Name = ? "
                "AND COALESCE(tc.ExplicitName, tc.InferredName) = ?",
                (from_table, from_column, to_table, to_column)).fetchone()
            if row is None:
                # Desktop may store the pair in the opposite orientation.
                row = c.execute(
                    "SELECT r.ID, r.IsActive, r.CrossFilteringBehavior, "
                    "r.FromCardinality, r.ToCardinality "
                    "FROM Relationship r "
                    "JOIN [Table] ft ON r.FromTableID = ft.ID "
                    "JOIN [Column] fc ON r.FromColumnID = fc.ID "
                    "JOIN [Table] tt ON r.ToTableID = tt.ID "
                    "JOIN [Column] tc ON r.ToColumnID = tc.ID "
                    "WHERE ft.Name = ? "
                    "AND COALESCE(fc.ExplicitName, fc.InferredName) = ? "
                    "AND tt.Name = ? "
                    "AND COALESCE(tc.ExplicitName, tc.InferredName) = ?",
                    (to_table, to_column, from_table, from_column)).fetchone()
                if row is not None:
                    found["swapped"] = True
            return row

        def _do_update(conn: sqlite3.Connection):
            c = conn.cursor()
            row = _locate(c)
            if row is None:
                raise ValueError(
                    f"Relationship {from_table}.{from_column} → "
                    f"{to_table}.{to_column} not found")
            rid, cur_active, cur_xf, cur_fc, cur_tc = row
            found.update({"id": rid, "active": cur_active, "xf": cur_xf,
                          "card": (cur_fc, cur_tc)})
            sets, params = [], []
            if new_active is not None:
                sets.append("IsActive = ?")
                params.append(1 if new_active else 0)
            if new_xf is not None:
                sets.append("CrossFilteringBehavior = ?")
                params.append(new_xf)
            if new_card is not None:
                fcard, tcard = new_card
                if found.get("swapped"):
                    fcard, tcard = tcard, fcard
                sets.append("FromCardinality = ?")
                params.append(fcard)
                sets.append("ToCardinality = ?")
                params.append(tcard)
            params.append(rid)
            c.execute(
                f"UPDATE Relationship SET {', '.join(sets)} WHERE ID = ?",
                params)
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_update)

        # A cardinality change alters which side carries the join index, so
        # regenerate the model's R$ index tables. is_active / cross-filter are
        # pure metadata and need no rebuild (and stay usable on models the
        # rebuild path refuses).
        rebuilt = False
        if new_card is not None and tuple(found["card"]) != tuple(
                sorted(new_card) if found.get("swapped") else new_card):
            try:
                old_size, new_size = _rebuild_datamodel(info)
                rebuilt = True
            except PBIXMCPError:
                # Model can't be rebuilt (calc tables/columns). The metadata
                # change is applied; say so rather than failing the edit.
                rebuilt = False

        info["modified"] = True
        _dax_cache.pop(alias, None)
        changed = []
        if new_active is not None:
            changed.append(f"  IsActive: {bool(found['active'])} → {new_active}")
        if new_xf is not None:
            changed.append(
                f"  CrossFilter: {'both' if found['xf'] == 2 else 'single'} → "
                f"{'both' if new_xf == 2 else 'single'}")
        if new_card is not None:
            changed.append(
                f"  Cardinality: {found['card']} → {new_card}"
                + ("" if rebuilt else " (metadata only — join indexes NOT "
                   "regenerated; this model can't be rebuilt)"))
        return ToolResponse.ok(
            f"Relationship {from_table}.{from_column} → {to_table}.{to_column} "
            f"updated:\n" + "\n".join(changed) +
            f"\n  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "RELATIONSHIP_NOT_FOUND").to_text()
    except Exception as e:
        return ToolResponse.error(
            f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_remove_table(alias: str, table_name: str) -> str:
    """Remove a table and its measures/relationships from the DataModel.

    Rebuilds the DataModel without the specified table. All measures hosted
    on the table and all relationships referencing it are also removed.

    Args:
        alias: The alias of the open file
        table_name: Name of the table to remove
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            remove_tables={table_name},
        )
        info["modified"] = True
        return ToolResponse.ok(
            f"Table '{table_name}' removed (with its measures and relationships)\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_add_field_parameter(
    alias: str, parameter_name: str, fields_json: str
) -> str:
    """Create a field parameter — a slicer-driven column/measure switcher.

    Authors the COMPLETE Desktop shape (diffed against Desktop-authored ground
    truth): a calculated-table partition holding the
    ``{("Display", NAMEOF('Table'[Field]), n), ...}`` tuple set with full
    static VertiPaq storage, the ``ParameterMetadata`` ExtendedProperty on the
    hidden Fields column, the display column sorted by the hidden Order
    column, and the display→Fields group-by wiring — so Power BI Desktop and
    the service treat it as a REAL field parameter (field-swapping in
    visuals), not just a lookup table. Field parameters survive later
    rebuild-based edits (they are recognized and re-stamped).

    Args:
        alias: The alias of the open file
        parameter_name: Name for the field parameter table (e.g. "Metric Selector")
        fields_json: JSON array of fields to include, e.g.
            '[{"display": "Revenue", "ref": "Sales[Revenue]"},
              {"display": "Profit",  "ref": "'Sales'[Profit]"},
              {"display": "Units",   "ref": "Sales[Units]"}]'
            Each entry has "display" (label shown in slicer; must not contain
            double quotes) and "ref" (Table[Field] or 'Table'[Field], where
            Field is a column or measure that must exist in the model).
    """
    try:
        fields = json.loads(fields_json)
        if not fields or not isinstance(fields, list):
            raise ValueError("fields_json must be a non-empty JSON array")

        for f in fields:
            if "display" not in f or "ref" not in f:
                raise ValueError("Each field must have 'display' and 'ref' keys")
            if '"' in str(f["display"]):
                raise ValueError(
                    f"display {f['display']!r} must not contain double quotes")

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Normalize refs to the DAX-quoted 'Table'[Field] form (what NAMEOF()
        # evaluates to — verified against Desktop-authored files) and validate
        # every target against the model. Silently authoring a parameter over
        # a typo'd field would be another silent-wrong.
        norm_fields = []
        for f in fields:
            table, name, canonical = _normalize_field_ref(str(f["ref"]))
            norm_fields.append({"display": str(f["display"]),
                                "ref": canonical, "table": table, "name": name})

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        with open(dm_path, "rb") as fh:
            meta_bytes = read_metadata_sqlite(decompress_datamodel(fh.read()))
        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        os.write(fd, meta_bytes)
        os.close(fd)
        try:
            conn = sqlite3.connect(tmp_path)
            conn.row_factory = sqlite3.Row
            if conn.execute("SELECT 1 FROM [Table] WHERE Name = ? AND ModelID = 1",
                            (parameter_name,)).fetchone():
                raise ValueError(f"Table '{parameter_name}' already exists")
            missing = []
            for nf in norm_fields:
                trow = conn.execute(
                    "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                    (nf["table"],)).fetchone()
                if not trow:
                    missing.append(f"{nf['ref']} (no table '{nf['table']}')")
                    continue
                hit = conn.execute(
                    "SELECT 1 FROM [Column] WHERE TableID = ? AND ExplicitName = ? "
                    "UNION SELECT 1 FROM Measure WHERE TableID = ? AND Name = ?",
                    (trow["ID"], nf["name"], trow["ID"], nf["name"])).fetchone()
                if not hit:
                    missing.append(nf["ref"])
            conn.close()
            if missing:
                raise ValueError(
                    "Field ref(s) not found in the model: " + ", ".join(missing))
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass

        # Static row data — the same constants NAMEOF() evaluates to.
        rows = []
        for i, nf in enumerate(norm_fields):
            rows.append({
                parameter_name: nf["display"],
                f"{parameter_name} Fields": nf["ref"],
                f"{parameter_name} Order": i,
            })

        # 1) Create the table via _rebuild_datamodel (full VertiPaq storage).
        extra_table = {
            "name": parameter_name,
            "columns": [
                {"name": parameter_name, "data_type": "String"},
                {"name": f"{parameter_name} Fields", "data_type": "String"},
                {"name": f"{parameter_name} Order", "data_type": "Int64"},
            ],
            "rows": rows,
        }
        old_size, _ = _rebuild_preserving_calc(alias, info, extra_tables=[extra_table])

        # 2) Stamp the Desktop field-parameter metadata shape on top.
        qd = _field_parameter_query_definition(norm_fields)
        _, new_size = _apply_field_parameter_metadata(
            dm_path, [{"table": parameter_name, "query_definition": qd}])
        info["modified"] = True

        field_list = ", ".join(f["display"] for f in fields)
        return ToolResponse.ok(
            f"Field parameter '{parameter_name}' created with {len(fields)} fields: {field_list}\n"
            f"  Calculated partition: {{(\"…\", NAMEOF(…), n), …}} + ParameterMetadata\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes\n"
            f"Use as a slicer to let users switch between these fields in visuals."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_datamodel_add_calculation_group(
    alias: str, group_name: str, items_json: str, precedence: int = 0
) -> str:
    """Create a calculation group — dynamic measure modifiers (YTD, QTD, PY, etc.).

    Calculation groups apply DAX transformations to any measure used in a visual.
    For example, a "Time Intelligence" group with items "Current", "YTD", "PY"
    lets users switch between time calculations via a slicer.

    Args:
        alias: The alias of the open file
        group_name: Name for the calculation group table (e.g. "Time Intelligence")
        items_json: JSON array of calculation items, e.g.
            '[{"name": "Current", "expression": "SELECTEDMEASURE()"},
              {"name": "YTD", "expression": "CALCULATE(SELECTEDMEASURE(), DATESYTD(''Date''[Date]))"},
              {"name": "PY",  "expression": "CALCULATE(SELECTEDMEASURE(), SAMEPERIODLASTYEAR(''Date''[Date]))"}]'
            Each item has "name" (display label) and "expression" (DAX using SELECTEDMEASURE()).
        precedence: Evaluation order when multiple calc groups exist (default 0)
    """
    try:
        items = json.loads(items_json)
        if not items or not isinstance(items, list):
            raise ValueError("items_json must be a non-empty JSON array")
        for item in items:
            if "name" not in item or "expression" not in item:
                raise ValueError("Each item must have 'name' and 'expression' keys")

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Build row data for the calculation group table
        # Calc groups have 2 columns: Name (text) and Ordinal (int)
        rows = []
        for i, item in enumerate(items):
            rows.append({
                "Name": item["name"],
                "Ordinal": i,
            })

        extra_table = {
            "name": group_name,
            "columns": [
                {"name": "Name", "data_type": "String"},
                {"name": "Ordinal", "data_type": "Int64"},
            ],
            "rows": rows,
        }

        # Create table via _rebuild_datamodel (full VertiPaq storage)
        old_size, new_size = _rebuild_preserving_calc(
            alias, info,
            extra_tables=[extra_table],
        )

        # Now add CalculationGroup + CalculationItem metadata via splice
        # (these are metadata-only — no VertiPaq impact)
        def _do_calc_group(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0

            # Find the table we just created
            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (group_name,)
            ).fetchone()
            if not trow:
                raise PBIXMCPError(f"Table '{group_name}' not found after rebuild", "INTERNAL_ERROR")
            table_id = trow["ID"]

            # Create CalculationGroup
            max_id += 1
            cg_id = max_id
            c.execute(
                "INSERT INTO CalculationGroup (ID, TableID, Precedence, ModifiedTime) "
                "VALUES (?, ?, ?, ?)",
                (cg_id, table_id, precedence, int(datetime.now().timestamp() * 1e7)),
            )

            # Link table to calculation group
            c.execute(
                "UPDATE [Table] SET CalculationGroupID = ? WHERE ID = ?",
                (cg_id, table_id),
            )

            # Calculation group tables require Partition Type=7 (CalculationGroup source)
            c.execute(
                "UPDATE Partition SET Type = 7, QueryDefinition = NULL, DataView = 3 WHERE TableID = ?",
                (table_id,),
            )

            # Create CalculationItems
            for i, item in enumerate(items):
                max_id += 1
                c.execute(
                    "INSERT INTO CalculationItem (ID, CalculationGroupID, Name, "
                    "Expression, Ordinal, ModifiedTime) "
                    "VALUES (?, ?, ?, ?, ?, ?)",
                    (max_id, cg_id, item["name"], item["expression"], i,
                     int(datetime.now().timestamp() * 1e7)),
                )

            # Calculation groups require DiscourageImplicitMeasures=1 on the Model
            c.execute("UPDATE Model SET DiscourageImplicitMeasures = 1 WHERE ID = 1")

            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        dm_path = os.path.join(info["work_dir"], "DataModel")
        _modify_metadata_only(dm_path, _do_calc_group)
        info["modified"] = True

        item_list = ", ".join(item["name"] for item in items)
        return ToolResponse.ok(
            f"Calculation group '{group_name}' created with {len(items)} items: {item_list}\n"
            f"  Precedence: {precedence}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes\n"
            f"Add to a slicer — measures in visuals will be modified by the selected item."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_modify_column(
    alias: str, table_name: str, column_name: str,
    property_name: str, new_value: str
) -> str:
    """Modify a column property in the DataModel metadata.

    Supports string, integer, and float columns. The ABF is fully rebuilt.

    Args:
        alias: The alias of the open file
        table_name: Name of the table containing the column
        column_name: Name of the column to modify
        property_name: Property to change (e.g., 'FormatString', 'IsHidden', 'Description')
        new_value: New value for the property
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        def _do_modify(conn: sqlite3.Connection):
            c = conn.cursor()
            c.execute(
                "SELECT c.ID FROM [Column] c "
                "JOIN [Table] t ON c.TableID = t.ID "
                "WHERE t.Name = ? AND c.ExplicitName = ?",
                (table_name, column_name)
            )
            row = c.fetchone()
            if not row:
                raise ValueError(
                    f"Column '{column_name}' not found in table '{table_name}'"
                )

            # Try numeric conversion
            try:
                val = int(new_value)
            except ValueError:
                try:
                    val = float(new_value)
                except ValueError:
                    val = new_value

            c.execute(
                f"UPDATE [Column] SET [{property_name}] = ? "
                f"WHERE ID = ?",
                (val, row[0])
            )
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_modify)
        info["modified"] = True
        return ToolResponse.ok(
            f"Column '{table_name}'.'{column_name}' updated:\n"
            f"  {property_name} = {new_value}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_set_sort_by_column(
    alias: str, table_name: str, column_name: str, sort_by_column: str = ""
) -> str:
    """Sort one column by the values of another (e.g. "Month Name" by "Month
    Number"), or clear an existing sort-by.

    The model stores this as a column ID, so it cannot be set by name through
    pbix_datamodel_modify_column; this resolves the name for you and rejects
    the combinations the engine refuses.

    Args:
        alias: The alias of the open file
        table_name: Table holding both columns
        column_name: Column whose display order changes
        sort_by_column: Column supplying the order. Empty string clears the
                        sort-by so the column sorts by its own values.
    """
    logger.info(
        f"pbix_set_sort_by_column: {alias} {table_name}[{column_name}] "
        f"by {sort_by_column!r}")
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.",
                DataModelCompressionError.code).to_text()

        if sort_by_column and sort_by_column == column_name:
            raise ValueError(
                f"A column cannot sort by itself ('{column_name}')")

        def _do_modify(conn: sqlite3.Connection):
            c = conn.cursor()

            def _column_id(name: str):
                c.execute(
                    "SELECT c.ID FROM [Column] c "
                    "JOIN [Table] t ON c.TableID = t.ID "
                    "WHERE t.Name = ? "
                    "  AND COALESCE(c.ExplicitName, c.InferredName) = ?",
                    (table_name, name))
                row = c.fetchone()
                if not row:
                    raise ValueError(
                        f"Column '{name}' not found in table '{table_name}'")
                return row[0]

            target = _column_id(column_name)
            # Desktop stores "no sort-by" as 0, not NULL — every column in a
            # Desktop-authored model has a non-NULL SortByColumnID.
            sort_id = _column_id(sort_by_column) if sort_by_column else 0

            # Power BI refuses a sort-by cycle (A sorts by B while B sorts by
            # A) — it would have no defined order.
            if sort_id:
                c.execute("SELECT SortByColumnID FROM [Column] WHERE ID = ?",
                          (sort_id,))
                row = c.fetchone()
                if row and row[0] == target:
                    raise ValueError(
                        f"'{sort_by_column}' already sorts by "
                        f"'{column_name}' — this would create a cycle")

            c.execute("UPDATE [Column] SET SortByColumnID = ? WHERE ID = ?",
                      (sort_id, target))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_modify)
        info["modified"] = True
        what = (f"now sorts by '{sort_by_column}'" if sort_by_column
                else "sort-by cleared")
        return ToolResponse.ok(
            f"Column '{table_name}'.'{column_name}' {what}.\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e)).to_text()


@mcp.tool()
def pbix_get_sort_by_columns(alias: str) -> str:
    """List every column that is sorted by another column.

    Args:
        alias: The alias of the open file
    """
    logger.info(f"pbix_get_sort_by_columns: {alias}")
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.",
                DataModelCompressionError.code).to_text()

        rows = _query_metadata_rows(
            dm_path,
            # Calculated and auto-date columns carry no ExplicitName — their
            # name lives in InferredName.
            "SELECT t.Name, "
            "       COALESCE(c.ExplicitName, c.InferredName), "
            "       COALESCE(s.ExplicitName, s.InferredName) "
            "FROM [Column] c "
            "JOIN [Table] t ON c.TableID = t.ID "
            "JOIN [Column] s ON c.SortByColumnID = s.ID "
            # 0 (not NULL) is how "no sort-by" is stored.
            "WHERE COALESCE(c.SortByColumnID, 0) <> 0 "
            "ORDER BY t.Name, 2")

        if not rows:
            return ToolResponse.ok(
                "No columns use a sort-by column.").to_text()
        lines = [f"  {t}[{col}] sorts by [{by}]" for t, col, by in rows]
        return ToolResponse.ok(
            f"{len(rows)} sort-by column(s):\n" + "\n".join(lines),
            data={"sort_by": [{"table": t, "column": c, "sort_by": s}
                              for t, c, s in rows]}).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e)).to_text()


_CALC_TYPE_NAME_TO_AMO = {
    "String": 2, "Int64": 6, "Double": 8, "Decimal": 10,
    "DateTime": 9, "Boolean": 11,
}
_CALC_COL_REF_RE = re.compile(r"(?:'([^']+)'|([A-Za-z_][A-Za-z0-9_]*))\s*\[\s*([^\]]+?)\s*\]")


def _infer_calc_type_name(values: list) -> str:
    """Best-effort AMO type NAME for a materialized calc column's values."""
    vals = [v for v in values if v is not None]
    if not vals:
        return "String"
    if all(isinstance(v, bool) for v in vals):
        return "Boolean"
    if all(isinstance(v, int) and not isinstance(v, bool) for v in vals):
        return "Int64"
    if all(isinstance(v, (int, float)) and not isinstance(v, bool) for v in vals):
        return "Double"
    return "String"


def _calc_col_refs(expression: str, table_name: str) -> set:
    """Column names referenced as table_name[col] in an expression."""
    out = set()
    for quoted, bare, col in _CALC_COL_REF_RE.findall(expression or ""):
        tbl = quoted or bare
        if tbl.lower() == table_name.lower():
            out.add(col)
    return out


_BARE_COL_RE = re.compile(r"(?<![\w\]'\"])\[([^\]\[]+)\]")


def _qualify_bare_column_refs(expression: str, table_name: str,
                              known_columns: list) -> str:
    """Rewrite an unqualified ``[Column]`` to ``'Table'[Column]``.

    Inside a calculated column, `[Date]` means "this table's Date" — it is the
    idiomatic way Power BI writes them, and it is what Desktop generates for
    auto date/time tables (`YEAR([Date])`). The evaluator only resolved the
    qualified forms, so every such column materialized as blank and the edit
    was refused. Only names that are actually columns of this table are
    rewritten, so measure references are left alone, and text inside string
    literals is never touched.
    """
    known = set(known_columns)
    if not known or not expression:
        return expression

    def _rewrite(segment: str) -> str:
        return _BARE_COL_RE.sub(
            lambda m: (f"'{table_name}'[{m.group(1)}]"
                       if m.group(1) in known else m.group(0)),
            segment)

    # Split on double-quoted DAX string literals and leave those untouched, so
    # a "[Total]" inside a label is never rewritten. Single quotes delimit
    # table names and are handled by the pattern's lookbehind, which also
    # keeps already-qualified `T[Col]` / `'T'[Col]` refs as they are.
    parts = re.split(r'("(?:[^"]|"")*")', expression)
    return "".join(p if p.startswith('"') else _rewrite(p) for p in parts)


def _materialize_table_calc_columns(
    table_name: str,
    data_columns: list[dict],
    data_rows: list[dict],
    calc_specs: list[dict],
    relationships: list,
    lookup_provider=None,
) -> tuple[list[dict], list[dict], list[dict]]:
    """Evaluate a table's calculated columns in dependency order.

    ``calc_specs`` = [{"column", "expression", "data_type"(optional)}]. Returns
    (columns_out, rows_out, restamp_specs) where columns_out/rows_out are the
    builder's add_table inputs (data + calc columns as plain data) and
    restamp_specs feed _apply_calculated_column_metadata. Raises ValueError on
    an unreliable expression, an unresolvable/circular dependency, or an eval
    failure — the caller aborts rather than materialize wrong values.
    """
    from pbix_mcp.dax.calc_tables import (
        calc_column_unsupported_reason,
        evaluate_row_context_column,
        expand_variation_accessors,
        lookupvalue_table_names,
        related_table_names,
    )

    # LOOKUPVALUE reads another table, so that table's rows have to be here.
    # Loaded LAZILY and cached: only tables a LOOKUPVALUE actually names get
    # decoded, so an ordinary same-table calc column still costs one read.
    lookup_data: dict[str, dict] = {}
    if lookup_provider is not None:
        wanted: set[str] = set()
        for spec in calc_specs:
            expr = spec.get("expression") or ""
            wanted |= lookupvalue_table_names(expr)
            wanted |= related_table_names(expr)
        for tn in wanted:
            if tn.lower() == table_name.lower():
                continue
            try:
                got = lookup_provider(tn)
            except Exception:  # noqa: BLE001 - a miss just means "refuse later"
                got = None
            if got:
                lookup_data[tn] = got

    col_names = [c["name"] for c in data_columns]
    type_by_name = {c["name"]: c.get("data_type", "String") for c in data_columns}
    rows = [[r.get(cn) for cn in col_names] for r in data_rows]

    pending = list(calc_specs)
    restamp: list[dict] = []
    progressed = True
    while pending and progressed:
        progressed = False
        still = []
        for spec in pending:
            refs = _calc_col_refs(spec["expression"], table_name)
            if not refs.issubset(set(col_names)):
                still.append(spec)
                continue
            snapshot = {table_name: {"columns": list(col_names),
                                     "rows": [list(r) for r in rows]}}
            snapshot.update(lookup_data)
            known_tables = {t: list(v.get("columns") or [])
                            for t, v in snapshot.items()}
            # Expand `X.[Date]` BEFORE qualifying bare references. When the
            # table ALSO has a real column named Date, the qualifier reads the
            # accessor's `[Date]` as a bare same-table reference and rewrites it
            # to `'T'[EstimatedCloseDate].'T'[Date]`, which no longer matches
            # the accessor -- expansion silently stops firing and the column is
            # refused as unresolvable. Expanding first is immune to the
            # collision. (Tables without a Date column qualify fine either way.)
            qualified = _qualify_bare_column_refs(
                expand_variation_accessors(spec["expression"]),
                table_name, col_names)
            # Re-run the gate now that the column list is known. Only here can
            # it tell MIN('T'[Year]) from MIN('T'[Yeer]) -- the engine answers 0
            # for the misspelling rather than failing, so without this the
            # column would materialize as zeros and report success.
            why = calc_column_unsupported_reason(
                qualified, table_name, col_names, known_tables, relationships)
            if why:
                raise ValueError(f"'{table_name}'[{spec['column']}]: {why}")
            vals, err = evaluate_row_context_column(
                col_names, rows, qualified,
                table_name, snapshot, relationships)
            if err:
                raise ValueError(
                    f"'{table_name}'[{spec['column']}]: {err}")
            dt = spec.get("data_type") or _infer_calc_type_name(vals)
            col_names.append(spec["column"])
            for i, r in enumerate(rows):
                r.append(vals[i])
            type_by_name[spec["column"]] = dt
            restamp.append({
                "table": table_name, "column": spec["column"],
                "expression": spec["expression"],
                "amo_type": _CALC_TYPE_NAME_TO_AMO.get(dt, 2),
            })
            progressed = True
        pending = still
    if pending:
        raise ValueError(
            f"calculated column(s) {[s['column'] for s in pending]} on "
            f"'{table_name}' reference columns that don't exist (or form a "
            f"dependency cycle)")

    columns_out = [{"name": cn, "data_type": type_by_name.get(cn, "String")}
                   for cn in col_names]
    rows_out = [dict(zip(col_names, r)) for r in rows]
    return columns_out, rows_out, restamp


@mcp.tool()
def pbix_datamodel_add_calculated_column(
    alias: str, table_name: str, column_name: str, dax: str, data_type: str = ""
) -> str:
    """Add a DAX calculated column to a table and materialize its values.

    Writes a Desktop-shape calculated column (Column.Type=2 + the DAX
    expression) AND stores the evaluated values in VertiPaq, so the file opens
    with correct data and the Power BI service recomputes it on Refresh.

    SCOPE: only row-context expressions over the table's OWN columns are
    supported — e.g. ``Margin = fct[Sales] - fct[Cost]``,
    ``IF(t[Qty] > 0, "Yes", "No")``, ``ROUND(t[X] * 0.1, 2)``.

    SUPPORTED (issue #25 r23#6 — the exact surface, so you don't discover it
    by probing): scalar row-context functions over THIS table's own columns —
    math/rounding, text incl. ``&`` concatenation, date/time (incl. DATEVALUE /
    DATEDIFF over columns), IF/SWITCH and the logical family — plus VAR/RETURN,
    and — only when NO filter-context function appears anywhere in the
    expression — a plain same-table aggregate ``SUM/AVERAGE/MIN/MAX(own[Col])``
    (deterministic per-row in a calc column).

    REFUSED with a clear reason (never stored wrong): ANY cross-table
    reference (incl. RELATED and LOOKUPVALUE), CALCULATE / CALCULATETABLE,
    X-iterators (SUMX/…/CONCATENATEX), RANKX/TOPN, COUNTROWS/COUNT/
    DISTINCTCOUNT, ALL/ALLSELECTED/FILTER/VALUES/KEEPFILTERS/TREATAS and the
    other filter/table functions, EARLIER, PATH/PATHITEM — Power BI computes
    these server-side across rows, which cannot be reproduced per-row here.
    Author those in Desktop, or express them as a MEASURE instead. (Existing
    Desktop-authored calc columns using RELATED/LOOKUPVALUE/CALCULATE are
    still PRESERVED across rebuild-path edits — this scope applies to
    authoring NEW columns.)

    Args:
        alias: The alias of the open file
        table_name: Table to add the column to
        column_name: Name of the new calculated column
        dax: The DAX expression (referencing this table's columns)
        data_type: Optional result type (String/Int64/Double/Decimal/
            DateTime/Boolean). Omitted = inferred from the evaluated values.
    """
    try:
        from pbix_mcp.dax.calc_tables import calc_column_unsupported_reason
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        if data_type and data_type not in _CALC_TYPE_NAME_TO_AMO:
            return ToolResponse.error(
                f"Invalid data_type '{data_type}'. Expected one of "
                f"{sorted(_CALC_TYPE_NAME_TO_AMO)}.", "INVALID_DATA_TYPE").to_text()

        reason = calc_column_unsupported_reason(dax, table_name)
        if reason:
            return ToolResponse.error(
                f"Calculated column '{table_name}'[{column_name}] cannot be "
                f"materialized: it {reason} Author it in Power BI Desktop, or "
                f"express it as a measure instead.", "UNSUPPORTED_CALC").to_text()

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta = read_metadata_sqlite(abf)
        _AMO_TO_NAME = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                        10: "Decimal", 11: "Boolean"}

        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            trow = conn.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (table_name,)).fetchone()
            if not trow:
                return ToolResponse.error(
                    f"Table '{table_name}' not found.", "TABLE_NOT_FOUND").to_text()
            tid = trow["ID"]
            # Name must be free on this table (no data/calc column, no measure).
            dup_col = conn.execute(
                "SELECT 1 FROM [Column] WHERE TableID = ? AND "
                "lower(ExplicitName) = lower(?)", (tid, column_name)).fetchone()
            if dup_col:
                return ToolResponse.error(
                    f"Column '{column_name}' already exists on '{table_name}'.",
                    "COLUMN_EXISTS").to_text()
            dup_meas = conn.execute(
                "SELECT 1 FROM Measure WHERE TableID = ? AND "
                "lower(Name) = lower(?)", (tid, column_name)).fetchone()
            if dup_meas:
                return ToolResponse.error(
                    f"A measure named '{column_name}' already exists on "
                    f"'{table_name}' — a column and measure sharing a name on "
                    f"one table breaks the Power BI service. Pick another name.",
                    "NAME_COLLIDES_MEASURE").to_text()

            # Re-materialize everything the rebuild would otherwise lose
            # (existing calc columns AND calc tables), plus the new column.
            rels_ctx = _get_dax_context(alias).get("relationships", [])
            table_updates, restamp_specs, table_restamp = (
                _plan_calc_preservation(
                    conn, abf, meta, rels_ctx,
                    extra_columns={table_name: [
                        {"column": column_name, "expression": dax,
                         "data_type": data_type or None}]}))
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        old_size, _ = _rebuild_datamodel(
            info, table_updates=table_updates, calc_authoring=True,
            restamp_calc_tables={s["table"] for s in table_restamp})
        # TABLE metadata first, then COLUMN metadata — the table stamp writes
        # flags the column stamp then refines, and doing it the other way round
        # is what silently demoted calculated columns to plain data elsewhere.
        if table_restamp:
            _, new_size = _apply_calculated_table_metadata(
                dm_path, table_restamp)
        _, new_size = _apply_calculated_column_metadata(dm_path, restamp_specs)
        info["modified"] = True
        global _dax_cache
        _dax_cache.pop(alias, None)

        new_spec = next(s for s in restamp_specs
                        if s["table"] == table_name and s["column"] == column_name)
        dt_name = {v: k for k, v in _CALC_TYPE_NAME_TO_AMO.items()}.get(
            new_spec["amo_type"], "String")
        return ToolResponse.ok(
            f"Calculated column '{table_name}'[{column_name}] added:\n"
            f"  Expression: {dax}\n"
            f"  DataType: {dt_name} "
            f"({'explicit' if data_type else 'inferred'}), values materialized\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "UNSUPPORTED_EXISTING_CALC").to_text()
    except Exception as e:
        return ToolResponse.error(
            f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


def _calc_column_dependents(conn, tid: int, table_name: str,
                            column_name: str) -> list[str]:
    """Calculated columns on the same table whose DAX reads ``column_name``.

    Matches the three forms Power BI writes: bare ``[Col]``, ``Table[Col]`` and
    ``'Table'[Col]``. Case-insensitive, because DAX identifiers are.
    """
    pat = re.compile(
        r"(?:'" + re.escape(table_name) + r"'|" + re.escape(table_name) + r")?"
        r"\s*\[" + re.escape(column_name) + r"\]", re.IGNORECASE)
    out = []
    for r in conn.execute(
        "SELECT ExplicitName, Expression FROM [Column] WHERE TableID = ? "
        "AND Type = 2 AND ExplicitName IS NOT NULL", (tid,)
    ):
        if r[0].lower() == column_name.lower():
            continue
        if r[1] and pat.search(r[1]):
            out.append(r[0])
    return out


@mcp.tool()
def pbix_datamodel_remove_calculated_column(
    alias: str, table_name: str, column_name: str
) -> str:
    """Remove a DAX calculated column, undoing pbix_datamodel_add_calculated_column.

    Drops both halves of the column: the Type=2 metadata carrying the DAX, and
    the materialized values in VertiPaq. Every other calculated column and
    calculated table in the model is re-evaluated and re-stamped, so the result
    is the model as it would have been had the column never been added.

    Only CALCULATED columns (Column.Type=2) can be removed this way. A plain
    data column is part of the table's source data — use pbix_set_table_data to
    rewrite the table without it.

    Refuses when another calculated column on the same table reads this one,
    naming the dependents, rather than leaving them evaluating against a column
    that no longer exists.

    Args:
        alias: The alias of the open file
        table_name: Table owning the column
        column_name: The calculated column to remove
    """
    try:
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.", DataModelCompressionError.code).to_text()

        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta = read_metadata_sqlite(abf)

        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            trow = conn.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (table_name,)).fetchone()
            if not trow:
                return ToolResponse.error(
                    f"Table '{table_name}' not found.",
                    "TABLE_NOT_FOUND").to_text()
            tid = trow["ID"]
            crow = conn.execute(
                "SELECT ExplicitName, Type, Expression FROM [Column] "
                "WHERE TableID = ? AND lower(ExplicitName) = lower(?)",
                (tid, column_name)).fetchone()
            if not crow:
                return ToolResponse.error(
                    f"Column '{column_name}' not found on '{table_name}'.",
                    "COLUMN_NOT_FOUND").to_text()
            if crow["Type"] != 2:
                kind = {1: "a data column", 3: "the internal row-number column",
                        4: "a calculated TABLE's data column"}.get(
                            crow["Type"], f"Type={crow['Type']}")
                return ToolResponse.error(
                    f"'{table_name}'[{crow['ExplicitName']}] is {kind}, not a "
                    f"calculated column, so removing it here would change the "
                    f"table's source data rather than undo an authored column. "
                    f"Use pbix_set_table_data to rewrite the table without it.",
                    "NOT_A_CALCULATED_COLUMN").to_text()

            actual = crow["ExplicitName"]
            dependents = _calc_column_dependents(conn, tid, table_name, actual)
            if dependents:
                return ToolResponse.error(
                    f"'{table_name}'[{actual}] is read by calculated column(s) "
                    f"{dependents} on the same table. Removing it would leave "
                    f"them evaluating against a column that no longer exists, "
                    f"so it was refused — remove those first.",
                    "CALC_COLUMN_HAS_DEPENDENTS").to_text()

            rels_ctx = _get_dax_context(alias).get("relationships", [])
            table_updates, restamp_specs, table_restamp = (
                _plan_calc_preservation(
                    conn, abf, meta, rels_ctx,
                    drop_columns={table_name: {actual}}))
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        old_size, new_size = _rebuild_datamodel(
            info, table_updates=table_updates, calc_authoring=True,
            restamp_calc_tables={s["table"] for s in table_restamp})
        if table_restamp:
            _, new_size = _apply_calculated_table_metadata(
                dm_path, table_restamp)
        if restamp_specs:
            _, new_size = _apply_calculated_column_metadata(
                dm_path, restamp_specs)
        info["modified"] = True
        global _dax_cache
        _dax_cache.pop(alias, None)

        kept = len(restamp_specs)
        return ToolResponse.ok(
            f"Calculated column '{table_name}'[{actual}] removed:\n"
            f"  Expression was: {crow['Expression']}\n"
            f"  Metadata and materialized values both dropped\n"
            f"  {kept} other calculated column(s) re-evaluated and preserved\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "UNSUPPORTED_EXISTING_CALC").to_text()
    except Exception as e:
        return ToolResponse.error(
            f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


def _plan_calc_preservation(conn, abf, meta, relationships, extra_columns=None,
                            base_data=None, skip_tables=None,
                            drop_columns=None):
    """Plan the re-materialization of everything a rebuild would otherwise lose.

    A from-scratch rebuild drops Type=2 calculated columns (they aren't read
    back from VertiPaq) and demotes calculated tables to plain data tables.
    Returns ``(table_updates, column_restamp, table_restamp)`` so the caller can
    rebuild and then re-stamp both shapes. Raises ValueError when an existing
    calculated column can't be faithfully reproduced — refuse, never corrupt.

    ``extra_columns`` = {table: [{"column", "expression", "data_type"}]} adds
    NEW calculated columns to the same plan, so authoring a calculated column
    and a calculated table compose in either order.

    ``base_data`` = {table: {"columns", "rows"}} supplies the plain data a
    table's calculated columns should be computed FROM, instead of what is
    currently in VertiPaq. A caller replacing a table's rows passes its new
    rows here so the calc columns are recomputed against them rather than
    silently carrying values derived from the old data.

    ``skip_tables`` names tables the caller is deleting, so the plan does not
    resurrect them.

    ``drop_columns`` = {table: {column, ...}} names calculated columns the
    caller is deleting. They are left out of the re-materialization, which is
    what actually removes them: a calculated column exists because the plan
    re-creates it, so omitting it from the plan is the removal.
    """
    from pbix_mcp.dax.calc_tables import calc_column_unsupported_reason
    from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf

    _AMO_TO_NAME = {2: "String", 6: "Int64", 8: "Double", 9: "DateTime",
                    10: "Decimal", 11: "Boolean"}
    table_updates: dict[str, dict] = {}
    table_restamp: list[dict] = []
    base_data = base_data or {}
    skip = {t.lower() for t in (skip_tables or set())}

    # --- existing calculated TABLES (excluding field parameters, which the
    # rebuild already detects and re-stamps on its own) ---
    for trow in conn.execute(
        "SELECT t.ID, t.Name FROM [Table] t JOIN [Partition] p "
        "ON p.TableID = t.ID WHERE t.ModelID = 1 AND p.Type = 2"
    ).fetchall():
        tid, tname = trow["ID"], trow["Name"]
        if _detect_field_parameter_shape(conn, tid) is not None:
            continue
        if tname.lower() in skip:
            continue
        qd = conn.execute(
            "SELECT QueryDefinition FROM [Partition] WHERE TableID = ?",
            (tid,)).fetchone()[0]

        # A table can be BOTH a calculated table and the owner of calculated
        # columns — every Power BI auto date/time table is: a Date column from
        # the partition's CALENDAR expression plus Year/MonthNo/Month/
        # QuarterNo/Quarter/Day computed on top. That needs two metadata stamps
        # on one table, and the calc-column stamp must not be overwritten by
        # the calc-table stamp, so the names are handed to the table stamper to
        # leave alone.
        owned_calc_columns = [
            r[0] for r in conn.execute(
                "SELECT ExplicitName FROM [Column] WHERE TableID = ? "
                "AND Type = 2 AND ExplicitName NOT LIKE 'RowNumber%' "
                "ORDER BY ID", (tid,))
            if r[0]]

        td = read_table_from_abf(abf, tname, meta)
        cols = td.get("columns") or []
        rows = td.get("rows") or []
        generated = False
        if not cols:
            # No materialized data in VertiPaq — Power BI generates this table
            # at load. Its generating DAX is declared right here in the
            # partition, so evaluate THAT rather than refusing outright.
            generated = True
            cols, rows, why = _evaluate_calc_table_partition(qd, abf, meta)
            if not cols:
                hint = ""
                if _is_auto_date_table(tname):
                    hint = (" This is one of Power BI's auto date/time tables. "
                            "Turning off File > Options > Data Load > Auto "
                            "date/time in Desktop and re-saving removes them; "
                            "the surgical tools (add_measure / modify_measure "
                            "/ remove_measure / modify_column / "
                            "set_sort_by_column) never rebuild and always work.")
                raise ValueError(
                    f"Calculated table '{tname}' has no materialized data and "
                    f"its expression could not be reproduced ({why}), so this "
                    f"edit would write a table that reopens empty. Refused."
                    f"{hint}")
        # Prefer the type the model already declares for the column. Inferring
        # from the regenerated values retyped every auto-date table's `Date`
        # column from DateTime to String, because the CALENDAR expression hands
        # dates back as ISO strings — the table looked intact while its date
        # semantics were gone.
        declared = {}
        for cr in conn.execute(
            "SELECT COALESCE(ExplicitName, InferredName) AS nm, "
            "       ExplicitDataType AS edt, InferredDataType AS idt "
            "FROM [Column] WHERE TableID = ?", (tid,)
        ):
            amo = cr["edt"] if cr["edt"] in _AMO_TO_NAME else cr["idt"]
            if amo in _AMO_TO_NAME:
                declared[cr["nm"]] = _AMO_TO_NAME[amo]
        col_defs = [{"name": c,
                     "data_type": declared.get(
                         c, _infer_calc_type_name([r[i] for r in rows]))}
                    for i, c in enumerate(cols)]
        table_updates[tname] = {
            "columns": col_defs,
            "rows": [dict(zip(cols, r)) for r in rows],
        }
        if generated:
            # These rows came from the expression, not from VertiPaq, so the
            # table's own calculated columns (an auto-date table's Year/Month/
            # Quarter/Day) have nothing to compute against unless we hand them
            # the rows we just generated.
            base_data.setdefault(tname, table_updates[tname])
        # SourceColumn as the model already spells it. A calc table built on
        # ANOTHER table qualifies its columns — Power BI's per-table auto-date
        # tables all read 'DateAutoTemplate[Year]' because they are copies of
        # that template. Synthesising a bare '[Year]' instead loses the
        # qualifier, the engine cannot resolve the column, and Desktop refuses
        # the whole model with "Relationship '<guid>' points to deleted column
        # 'Date' in table 'Date'".
        source_columns = {
            r["nm"]: r["sc"] for r in conn.execute(
                "SELECT COALESCE(ExplicitName, InferredName) AS nm, "
                "       SourceColumn AS sc FROM [Column] WHERE TableID = ?",
                (tid,))
            if r["nm"] and r["sc"]}
        table_restamp.append({"table": tname, "expression": qd,
                              "calc_columns": owned_calc_columns,
                              "source_columns": source_columns})

    # Tables that are BOTH a calculated table and a calc-column owner: their
    # calculated columns are stamped separately and must carry the calc-table
    # system flag, which an ordinary calculated column does not.
    calc_table_owners = {s["table"] for s in table_restamp
                         if s.get("calc_columns")}

    # --- existing calculated COLUMNS: re-evaluate from their DAX ---
    existing = conn.execute(
        "SELECT t.Name AS tbl, c.ExplicitName AS col, c.Expression AS expr "
        "FROM [Column] c JOIN [Table] t ON c.TableID = t.ID "
        "WHERE c.Type = 2 AND t.ModelID = 1 "
        "AND c.ExplicitName NOT LIKE 'RowNumber%'").fetchall()
    dropped = {t.lower(): {c.lower() for c in cols}
               for t, cols in (drop_columns or {}).items()}
    # Table -> column names for the WHOLE model, from metadata only (no data
    # decode). This is what lets the gate accept a LOOKUPVALUE: it can confirm
    # the table and columns it names really exist before anything is read.
    model_tables: dict[str, list[str]] = {}
    for r in conn.execute(
        "SELECT t.Name AS tbl, "
        "       COALESCE(c.ExplicitName, c.InferredName) AS col "
        "FROM [Column] c JOIN [Table] t ON c.TableID = t.ID "
        "WHERE t.ModelID = 1 "
        "AND COALESCE(c.ExplicitName, c.InferredName) NOT LIKE 'RowNumber%'"):
        model_tables.setdefault(r["tbl"], []).append(r["col"])
    calc_by_table: dict[str, list[dict]] = {}
    for r in existing:
        if r["tbl"].lower() in skip:
            continue
        if (r["col"] or "").lower() in dropped.get(r["tbl"].lower(), ()):
            continue
        bad = calc_column_unsupported_reason(
            r["expr"] or "", r["tbl"], model_tables.get(r["tbl"]),
            model_tables, relationships)
        if bad:
            raise ValueError(
                f"The model has a calculated column '{r['tbl']}'[{r['col']}] "
                f"this engine can't reproduce ({bad.split('.')[0]}). This edit "
                f"would rebuild it and risk corrupting its values, so it was "
                f"refused.")
        calc_by_table.setdefault(r["tbl"], []).append(
            {"column": r["col"], "expression": r["expr"]})
    # A table losing its LAST calculated column still has to be re-supplied, or
    # the rebuild guard sees a table that carries Type=2 columns and is not in
    # table_updates, and refuses the very edit that removes them.
    for tname in (drop_columns or {}):
        if tname.lower() not in skip:
            calc_by_table.setdefault(tname, [])
    for tname, specs in (extra_columns or {}).items():
        if tname in table_updates and any(
                s["table"] == tname for s in table_restamp):
            raise ValueError(
                f"'{tname}' is a calculated table — its columns come from its "
                f"own DAX expression, so a calculated column cannot be added "
                f"to it.")
        calc_by_table.setdefault(tname, []).extend(specs)

    column_restamp: list[dict] = []
    for tname, specs in calc_by_table.items():
        trow = conn.execute(
            "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
            (tname,)).fetchone()
        data_cols = []
        # Type 1 = data, 3 = RowNumber, 4 = CalculatedTableColumn. Type 4 is
        # the DATA column of a calculated table (an auto-date table's `Date`),
        # and omitting it left the materializer with no schema, so every
        # calculated column on such a table evaluated to blank. Those columns
        # also carry their name in InferredName rather than ExplicitName.
        for cr in conn.execute(
            "SELECT COALESCE(ExplicitName, InferredName) AS nm, "
            "       ExplicitDataType, InferredDataType "
            "FROM [Column] WHERE TableID = ? AND Type IN (1, 3, 4) "
            "AND COALESCE(ExplicitName, InferredName) NOT LIKE 'RowNumber%' "
            "ORDER BY ID",
                (trow["ID"],)):
            edt = cr["ExplicitDataType"]
            amo = edt if edt in _AMO_TO_NAME else cr["InferredDataType"]
            data_cols.append({"name": cr["nm"],
                              "data_type": _AMO_TO_NAME.get(amo, "String")})
        override = base_data.get(tname)
        if override:
            # The caller is replacing this table's rows; the calc columns must
            # be computed from the NEW data, not from what VertiPaq still holds.
            data_cols = list(override["columns"])
            data_rows = list(override["rows"])
        else:
            td = read_table_from_abf(abf, tname, meta)
            data_rows = [dict(zip(td["columns"], rv))
                         for rv in td.get("rows", [])]
        def _lookup_table(want: str, _tname=tname):
            """Rows of another table, for a LOOKUPVALUE in _tname's calc columns.

            Prefers the caller's REPLACEMENT rows when it is also rewriting that
            table -- looking the value up in the VertiPaq copy the caller is
            about to overwrite would resolve against stale data.
            """
            if want.lower() == _tname.lower():
                return None
            for cand, ov in base_data.items():
                if cand.lower() == want.lower():
                    return {"columns": [c["name"] if isinstance(c, dict) else c
                                        for c in ov["columns"]],
                            "rows": [list(r) for r in ov["rows"]]}
            hit = conn.execute(
                "SELECT Name FROM [Table] WHERE ModelID = 1 "
                "AND lower(Name) = lower(?)", (want,)).fetchone()
            if not hit:
                return None
            td2 = read_table_from_abf(abf, hit["Name"], meta,
                                      include_calculated=True)
            return {"columns": list(td2.get("columns") or []),
                    "rows": [list(r) for r in td2.get("rows") or []]}

        cols_out, rows_out, restamp = _materialize_table_calc_columns(
            tname, data_cols, data_rows, specs, relationships,
            lookup_provider=_lookup_table)
        # The calc columns' values go into VertiPaq, but must NOT be embedded in
        # the partition's Enter-data M — Power BI recomputes them from DAX.
        table_updates[tname] = {"columns": cols_out, "rows": rows_out,
                                "calc_columns": [r["column"] for r in restamp]}
        if tname in calc_table_owners:
            for spec in restamp:
                spec["system_flags"] = 2
        column_restamp.extend(restamp)

    return table_updates, column_restamp, table_restamp


_DAX_TABLE_REF_RE = re.compile(r"'([^']+)'\s*\[|(?<![\w'])([A-Za-z_]\w*)\s*\[")


def _evaluate_calc_table_partition(query_definition: str, abf, meta) -> tuple:
    """Evaluate a calculated table's own declared DAX to regenerate its rows.

    Returns ``(columns, rows, reason)``; ``columns`` is empty when it could not
    be reproduced and ``reason`` says why.

    Only the tables the expression actually references are read out of the ABF.
    Building the full DAX context here would decode the whole model's VertiPaq
    data — minutes of work, and the source of a hang this project already had.
    """
    from pbix_mcp.dax.calc_tables import evaluate_calc_table_expression
    from pbix_mcp.formats.vertipaq_decoder import read_table_from_abf

    expr = (query_definition or "").strip()
    if not expr:
        return [], [], "the partition declares no expression"

    referenced = {m.group(1) or m.group(2)
                  for m in _DAX_TABLE_REF_RE.finditer(expr)}
    tables: dict = {}
    for name in referenced:
        try:
            td = read_table_from_abf(abf, name, meta)
        except Exception:
            continue
        if td and td.get("columns"):
            tables[name] = {"columns": td["columns"],
                            "rows": td.get("rows") or []}

    try:
        result, err = evaluate_calc_table_expression(expr, tables, {}, [])
    except Exception as exc:
        return [], [], f"{type(exc).__name__}: {exc}"
    if err or not result:
        return [], [], (err or "no result")
    cols = result.get("columns") or []
    rows = result.get("rows") or []
    if not cols or not rows:
        return [], [], "expression produced no rows"
    return cols, rows, ""


def _is_auto_date_table(name: str) -> bool:
    """Power BI's generated auto date/time tables.

    `DateTableTemplate_<guid>` is the hidden template; `LocalDateTable_<guid>`
    is the per-date-column instance. Both are machine-generated and rebuilt at
    refresh, so their rows cannot be reproduced from the file alone.
    """
    return bool(re.match(r"^(LocalDateTable|DateTableTemplate)_", name or ""))


def _relationships_from_metadata(conn) -> list:
    """Relationships in the DAX engine's shape, straight from metadata SQL.

    `_get_dax_context` would give the same list, but it materializes EVERY
    table's rows to do it — decoding the whole model's VertiPaq data just to
    learn which columns join. On a large model that is minutes of work for a
    list of a dozen dicts, and on one real report (Microsoft's Competitive
    Marketing sample) a column segment decodes so slowly that the edit never
    returns at all.
    """
    rows = conn.execute(
        "SELECT ft.Name AS ft, "
        "       COALESCE(fc.ExplicitName, fc.InferredName) AS fc, "
        "       tt.Name AS tt, "
        "       COALESCE(tc.ExplicitName, tc.InferredName) AS tc, "
        "       r.IsActive, r.CrossFilteringBehavior "
        "FROM Relationship r "
        "JOIN [Table] ft ON r.FromTableID = ft.ID "
        "JOIN [Column] fc ON r.FromColumnID = fc.ID "
        "JOIN [Table] tt ON r.ToTableID = tt.ID "
        "JOIN [Column] tc ON r.ToColumnID = tc.ID").fetchall()
    return [{
        "FromTable": r["ft"] or "", "FromColumn": r["fc"] or "",
        "ToTable": r["tt"] or "", "ToColumn": r["tc"] or "",
        "IsActive": 1 if r["IsActive"] is None else r["IsActive"],
        "CrossFilteringBehavior": r["CrossFilteringBehavior"] or 1,
    } for r in rows]


def _rebuild_preserving_calc(alias: str, info: dict, **rebuild_kwargs):
    """Rebuild the DataModel with calculated tables/columns carried through.

    A from-scratch rebuild reconstructs the model from data, which drops Type=2
    calculated columns and demotes calculated tables to plain data — so every
    rebuild-path edit used to be REFUSED outright on a model containing either.
    That is three of the four reports in the public corpus, which made adding a
    table, adding a relationship or removing a table impossible on most real
    files.

    This plans the re-materialization first (the same plan the calc-authoring
    tools use), runs the rebuild with it, then re-stamps the calc metadata, so
    the edit lands with the calc objects intact. When the plan cannot reproduce
    an existing calculated column or table, `_plan_calc_preservation` raises and
    the edit is still refused — never corrupted.

    Returns ``(old_dm_size, new_dm_size)``.
    """
    import tempfile

    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    dm_path = os.path.join(info["work_dir"], "DataModel")
    caller_updates = dict(rebuild_kwargs.pop("table_updates", None) or {})
    remove_tables = rebuild_kwargs.get("remove_tables") or set()

    with open(dm_path, "rb") as f:
        abf = decompress_datamodel(f.read())
    meta = read_metadata_sqlite(abf)

    tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
    tmp.write(meta)
    tmp.close()
    try:
        conn = sqlite3.connect(tmp.name)
        conn.row_factory = sqlite3.Row
        try:
            planned, col_restamp, table_restamp = _plan_calc_preservation(
                conn, abf, meta, _relationships_from_metadata(conn),
                base_data=caller_updates, skip_tables=remove_tables)
        finally:
            conn.close()
    finally:
        os.unlink(tmp.name)

    if not planned and not col_restamp and not table_restamp:
        # No calc objects to preserve — the plain rebuild is exact.
        return _rebuild_datamodel(info, table_updates=caller_updates,
                                  **rebuild_kwargs)

    # The plan already folded in the caller's rows for any table it touched
    # (via base_data); the caller's other tables pass straight through.
    merged = dict(caller_updates)
    merged.update(planned)

    old_size, new_size = _rebuild_datamodel(
        info, table_updates=merged, calc_authoring=True,
        restamp_calc_tables={s["table"] for s in table_restamp},
        **rebuild_kwargs)
    # TABLE metadata first, then COLUMN metadata. An auto-date table is both a
    # calculated table AND the owner of calculated columns; stamping the table
    # shape second re-wrote its columns as calc-table data columns and silently
    # demoted the Type=2 stamps applied moments earlier — the edit "succeeded"
    # while dropping six calculated columns.
    if table_restamp:
        _, new_size = _apply_calculated_table_metadata(dm_path, table_restamp)
    if col_restamp:
        _, new_size = _apply_calculated_column_metadata(dm_path, col_restamp)

    global _dax_cache
    _dax_cache.pop(alias, None)
    return old_size, new_size


@mcp.tool()
def pbix_datamodel_add_calculated_table(
    alias: str, table_name: str, dax: str
) -> str:
    """Add a DAX calculated table and materialize its rows.

    Evaluates the table expression, stores the resulting rows in VertiPaq, and
    stamps Desktop's calculated-table metadata (partition Type=2 carrying the
    DAX, data columns Type=4) so the file opens with data and Power BI
    recomputes the table on Refresh.

    SCOPE: the expression must be one this engine reproduces faithfully — e.g.
    ``DATATABLE(...)``, ``GENERATESERIES(1, 12, 1)``, ``DISTINCT(Sales[Cat])``,
    ``VALUES(Sales[Product])``, ``FILTER(Sales, Sales[Amount] > 100)``,
    ``TOPN(10, Sales, Sales[Amount])``, ``ADDCOLUMNS(...)``, or a bare table
    reference. Shapes whose result this engine cannot reproduce exactly
    (SUMMARIZE / SUMMARIZECOLUMNS / SELECTCOLUMNS / GROUPBY, or anything using
    an unsupported function) are REFUSED with a reason rather than persisted
    with silently-wrong rows.

    Args:
        alias: The alias of the open file
        table_name: Name for the new calculated table
        dax: The table-valued DAX expression
    """
    try:
        from pbix_mcp.dax.calc_tables import (
            calc_table_unsupported_reason,
            evaluate_calc_table_expression,
        )
        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        reason = calc_table_unsupported_reason(dax)
        if reason:
            return ToolResponse.error(
                f"Calculated table '{table_name}' cannot be materialized: it "
                f"{reason}.", "UNSUPPORTED_CALC_TABLE").to_text()

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error(
                "No DataModel found.", DataModelCompressionError.code).to_text()

        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta = read_metadata_sqlite(abf)

        ctx = _get_dax_context(alias)
        result, err = evaluate_calc_table_expression(
            dax, ctx["tables"], ctx.get("measure_defs"),
            ctx.get("relationships"))
        if err:
            return ToolResponse.error(
                f"Calculated table '{table_name}' cannot be materialized: "
                f"{err}.", "CALC_TABLE_EVAL_FAILED").to_text()

        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(meta)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row
            if conn.execute(
                "SELECT 1 FROM [Table] WHERE lower(Name) = lower(?) "
                "AND ModelID = 1", (table_name,)).fetchone():
                return ToolResponse.error(
                    f"Table '{table_name}' already exists.",
                    "TABLE_EXISTS").to_text()
            table_updates, col_restamp, table_restamp = _plan_calc_preservation(
                conn, abf, meta, ctx.get("relationships") or [])
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        cols, rows = result["columns"], result["rows"]
        new_cols = [{"name": c,
                     "data_type": _infer_calc_type_name([r[i] for r in rows])}
                    for i, c in enumerate(cols)]
        new_table = {"name": table_name, "columns": new_cols,
                     "rows": [dict(zip(cols, r)) for r in rows]}

        old_size, _ = _rebuild_datamodel(
            info, table_updates=table_updates, extra_tables=[new_table],
            calc_authoring=True,
            restamp_calc_tables={s["table"] for s in table_restamp})
        if col_restamp:
            _apply_calculated_column_metadata(dm_path, col_restamp)
        _, new_size = _apply_calculated_table_metadata(
            dm_path, table_restamp + [{"table": table_name, "expression": dax}])
        info["modified"] = True
        global _dax_cache
        _dax_cache.pop(alias, None)

        return ToolResponse.ok(
            f"Calculated table '{table_name}' added:\n"
            f"  Expression: {dax}\n"
            f"  Columns: {', '.join(c['name'] for c in new_cols)}\n"
            f"  Rows materialized: {len(rows):,}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except ValueError as e:
        return ToolResponse.error(str(e), "UNSUPPORTED_EXISTING_CALC").to_text()
    except Exception as e:
        return ToolResponse.error(
            f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_datamodel_decompress(alias: str) -> str:
    """Decompress the DataModel from a PBIX into raw ABF format.

    This decompresses the XPress9-compressed DataModel and saves the
    raw ABF file for inspection. The ABF contains the full VertiPaq
    storage engine data.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import list_abf_files
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        logger.info("Decompressing DataModel (%d bytes) for '%s'", len(dm_bytes), alias)
        abf = decompress_datamodel(dm_bytes)
        logger.debug("Decompressed to %d bytes ABF", len(abf))
        abf_path = dm_path + ".abf"
        with open(abf_path, "wb") as f:
            f.write(abf)

        file_log = list_abf_files(abf)
        summary = [f"Decompressed DataModel: {len(dm_bytes):,} → {len(abf):,} bytes"]
        summary.append(f"ABF saved to: {abf_path}")
        summary.append(f"\nABF contains {len(file_log)} files:")
        for entry in file_log:
            summary.append(f"  {entry['Path']} ({entry['Size']:,} bytes)")
        return ToolResponse.ok("\n".join(summary)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_datamodel_recompress(alias: str, abf_path: str = "") -> str:
    """Recompress a modified ABF file back into the DataModel.

    After using pbix_datamodel_decompress to get the ABF, you can
    decompress and modify the ABF (or any of its internal files), call
    this to XPress9-compress it back into the DataModel. The next
    pbix_save will include the updated DataModel.

    Workflow:
      1. pbix_datamodel_decompress(alias)  ->  saves .abf
      2. Modify the .abf (directly, or via modify_measure / modify_metadata)
      3. pbix_datamodel_recompress(alias)   ->  compresses .abf back into DataModel

    Args:
        alias: The alias of the open file
        abf_path: Path to the ABF file to compress. Default: the .abf
                  next to the DataModel.
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        from pbix_mcp.formats.datamodel_roundtrip import compress_datamodel

        if not abf_path:
            abf_path = dm_path + ".abf"

        if not os.path.exists(abf_path):
            return ToolResponse.error(
                f"ABF file not found at {abf_path}. Run pbix_datamodel_decompress first.",
                ABFRebuildError.code
            ).to_text()

        with open(abf_path, "rb") as f:
            abf_bytes = f.read()

        logger.info("Recompressing ABF (%d bytes) for '%s'", len(abf_bytes), alias)

        # Validate ABF starts with BOM
        if not abf_bytes[:2] == b"\xff\xfe":
            return ToolResponse.error(
                f"File does not look like a valid ABF (expected \\xff\\xfe BOM, got {abf_bytes[:2].hex()}).",
                ABFRebuildError.code
            ).to_text()

        # Read original DataModel size for comparison (and for chunk reuse
        # when the encoder can't re-emit incompressible chunks).
        orig_size = os.path.getsize(dm_path) if os.path.exists(dm_path) else 0
        orig_dm = None
        if orig_size:
            with open(dm_path, "rb") as f:
                orig_dm = f.read()

        new_dm = compress_datamodel(abf_bytes, original_dm=orig_dm)

        with open(dm_path, "wb") as f:
            f.write(new_dm)

        info["modified"] = True
        return ToolResponse.ok(
            f"Recompressed ABF -> DataModel:\n"
            f"  ABF size:          {len(abf_bytes):>12,} bytes\n"
            f"  Old DataModel:     {orig_size:>12,} bytes\n"
            f"  New DataModel:     {len(new_dm):>12,} bytes\n"
            f"  XPress9 blocks:    {(len(abf_bytes) + 2097151) // 2097152}\n"
            f"  Saved to: {dm_path}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", ABFRebuildError.code).to_text()


@mcp.tool()
def pbix_datamodel_replace_file(alias: str, internal_path: str, new_content_path: str) -> str:
    """Replace a specific file inside the ABF (decompressed DataModel).

    This lets you swap out any internal ABF file — for example, replace
    metadata.sqlitedb with a modified version.

    Files can be ANY size — the ABF is fully rebuilt with updated offsets
    and headers.

    Args:
        alias: The alias of the open file
        internal_path: Partial path to match inside the ABF (e.g. 'metadata.sqlitedb')
        new_content_path: Path to the replacement file on disk
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        if not os.path.exists(new_content_path):
            return ToolResponse.error(f"Replacement file not found: {new_content_path}", ABFRebuildError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import (
            find_abf_file,
            list_abf_files,
            rebuild_abf_with_replacement,
        )
        from pbix_mcp.formats.datamodel_roundtrip import compress_datamodel, decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        with open(new_content_path, "rb") as f:
            new_content = f.read()

        abf = decompress_datamodel(dm_bytes)
        file_log = list_abf_files(abf)
        entry = find_abf_file(file_log, internal_path)
        if not entry:
            return ToolResponse.error(f"No file matching '{internal_path}' in ABF.", ABFRebuildError.code).to_text()

        fname = entry["Path"]
        new_abf = rebuild_abf_with_replacement(abf, {internal_path: new_content})
        new_dm = compress_datamodel(new_abf, original_dm=dm_bytes)

        with open(dm_path, "wb") as f:
            f.write(new_dm)

        info["modified"] = True
        return ToolResponse.ok(
            f"Replaced '{fname}' in ABF (full rebuild):\n"
            f"  Old file size: {entry['Size']:,} bytes\n"
            f"  New file size: {len(new_content):,} bytes\n"
            f"  ABF: {len(abf):,} -> {len(new_abf):,} bytes\n"
            f"  DataModel recompressed: {len(dm_bytes):,} -> {len(new_dm):,} bytes"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", ABFRebuildError.code).to_text()


@mcp.tool()
def pbix_datamodel_extract_file(alias: str, internal_path: str, output_path: str = "") -> str:
    """Extract a specific file from inside the ABF (decompressed DataModel).

    Args:
        alias: The alias of the open file
        internal_path: Partial path to match inside the ABF (e.g. 'metadata.sqlitedb')
        output_path: Where to save the extracted file. Default: next to DataModel.
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import find_abf_file, list_abf_files, read_abf_file
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        abf = decompress_datamodel(dm_bytes)
        file_log = list_abf_files(abf)
        entry = find_abf_file(file_log, internal_path)
        if not entry:
            return ToolResponse.error(f"No file matching '{internal_path}' in ABF.", DataModelCompressionError.code).to_text()

        content = read_abf_file(abf, entry)

        if not output_path:
            fname = os.path.basename(entry["Path"])
            output_path = os.path.join(info["work_dir"], fname)

        with open(output_path, "wb") as f:
            f.write(content)

        return ToolResponse.ok(
            f"Extracted '{entry['Path']}' ({len(content):,} bytes)\n"
            f"  ABF path: {entry['Path']}\n"
            f"  Saved to: {output_path}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", DataModelCompressionError.code).to_text()


@mcp.tool()
def pbix_datamodel_list_abf_files(alias: str) -> str:
    """List all files inside the ABF (decompressed DataModel).

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import list_abf_files
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        abf = decompress_datamodel(dm_bytes)
        files = list_abf_files(abf)

        lines = [f"ABF contains {len(files)} files ({len(abf):,} bytes decompressed):\n"]
        for entry in files:
            lines.append(f"  {entry['Path']} ({entry['Size']:,} bytes)")
        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", DataModelCompressionError.code).to_text()


# ---- Section 9: DAX Evaluation Engine ----

# Cache for DAX context data per alias (tables + measures + relationships)
_dax_cache: dict = {}

def _get_dax_context(alias: str) -> dict:
    """Build or retrieve cached DAX context (tables, measures, relationships)."""
    if alias in _dax_cache:
        return _dax_cache[alias]

    info = _ensure_open(alias)
    from pbix_mcp.formats.model_reader import ModelReader
    model = ModelReader(info["path"], work_dir=info.get("work_dir"))

    # Load measures
    measures_list = model.dax_measures
    measure_defs = {}
    measure_tables = {}
    measure_types = {}
    for m in measures_list:
        measure_defs[m.get('Name', '')] = m.get('Expression', '')
        # A measure's HOME TABLE is what DAX uses to disambiguate an
        # unqualified [Column] inside it. Dropping it here made
        # MS_Revenue_Opportunities' `Revenue = SUM([ProductRevenue])` blank:
        # three tables own a ProductRevenue column, so the reference resolved
        # to nothing and all six measures on that table read BLANK.
        if m.get('TableName'):
            measure_tables[m.get('Name', '')] = m.get('TableName')
        # Declared Measure.DataType (AMO code), so a declared Int64 can flow
        # through to the evaluate tools' result typing (issue #40).
        if m.get('DataType') is not None:
            measure_types[m.get('Name', '')] = m.get('DataType')

    # Load relationships
    rels_list = model.relationships
    relationships = []
    for r in rels_list:
        relationships.append({
            'FromTable': r.get('FromTableName', ''),
            'FromColumn': r.get('FromColumnName', ''),
            'ToTable': r.get('ToTableName', ''),
            'ToColumn': r.get('ToColumnName', ''),
            'IsActive': r.get('IsActive', 1),
            # Carry cross-filter direction so the engine honors bidirectional
            # (CrossFilteringBehavior=2) relationships. 1 = single (default).
            'CrossFilteringBehavior': r.get('CrossFilteringBehavior', 1),
        })

    # Load all user-facing tables
    schema_list = model.schema
    table_names = sorted(set(r['TableName'] for r in schema_list))
    tables = {}
    for tname in table_names:
        if tname.startswith('H$') or tname.startswith('R$'):
            continue
        try:
            td = model.get_table(tname)
            if td and td.get('columns') and td.get('rows'):
                tables[tname] = {
                    'columns': td['columns'],
                    'rows': td['rows'],
                }
        except Exception:
            continue

    # --- Load calculated tables from ABF metadata ---
    # Uses calc_tables.py as the single source of truth for evaluating
    # DATATABLE, GENERATESERIES, CALENDAR, and other calculated table expressions
    # that exist only as DAX in metadata, not in VertiPaq column stores.
    try:
        from pbix_mcp.dax.calc_tables import load_calculated_tables
        tables = load_calculated_tables(info["path"], tables, relationships)
    except Exception:
        pass  # If calculated table loading fails, continue without them

    # Performance warning for large tables
    _LARGE_TABLE_THRESHOLD = 100_000
    for tname, tdata in tables.items():
        row_count = len(tdata.get('rows', []))
        if row_count > _LARGE_TABLE_THRESHOLD:
            logger.warning("Table '%s' has %d rows — DAX evaluation may be slow", tname, row_count)

    # Detect date table — try multiple heuristics
    date_table = None
    date_column = None
    # Pass 1: table name contains 'date' AND has a 'Date' column
    for tname, tdata in tables.items():
        if 'date' in tname.lower():
            if 'Date' in tdata['columns']:
                date_table = tname
                date_column = 'Date'
                break
    # Pass 2: table name starts with common date-table prefixes (dimDate, DimDate, DateTable, Calendar)
    if not date_table:
        for tname, tdata in tables.items():
            tlow = tname.lower().replace(' ', '').replace('-', '').replace('_', '')
            if tlow in ('dimdate', 'datetable', 'calendar', 'datekey', 'dates'):
                for cname in tdata['columns']:
                    if cname.lower() == 'date':
                        date_table = tname
                        date_column = cname
                        break
                if date_table:
                    break
    # Pass 3: any table with a 'Date' column that also has Year/Month columns (likely a date dimension)
    if not date_table:
        for tname, tdata in tables.items():
            cols_lower = [c.lower() for c in tdata['columns']]
            if 'date' in cols_lower and ('year' in cols_lower or 'month' in cols_lower):
                date_col_idx = cols_lower.index('date')
                date_table = tname
                date_column = tdata['columns'][date_col_idx]
                break

    # --- Load default slicer filters from report layout ---
    # These are the slicer values that Power BI applies when you first open
    # the report (before any user interaction). Without them, measures using
    # SELECTEDVALUE on parameter tables return BLANK.
    default_filters = {}
    try:
        default_filters = _get_all_default_filters(info["work_dir"])
    except Exception:
        pass

    ctx = {
        'tables': tables,
        'measure_defs': measure_defs,
        'measure_tables': measure_tables,
        # name -> declared Measure.DataType AMO code (issue #40)
        'measure_types': measure_types,
        # The MODEL's own column list. `tables` above is a materialized subset,
        # so it cannot be used to decide that a reference is unresolvable.
        'model_columns': model.all_column_names,
        'date_table': date_table,
        'date_column': date_column,
        'relationships': relationships,
        'default_filters': default_filters,
        'work_dir': info["work_dir"],
    }
    _dax_cache[alias] = ctx
    return ctx


# Measure references accepted by the evaluate tools: bare (Pipeline Value),
# bracketed ([Pipeline Value]), and table-qualified ('SalesPipeline'[Pipeline
# Value] / SalesPipeline[Pipeline Value]). measure_defs is keyed by BARE names,
# and the engine treats an unknown measure as BLANK by design — so without
# normalization the DAX-style forms silently evaluated to (null) for every row.
# The quoted-table alternative accepts DAX's '' escape ('O''Brien Sales'[M]).
_MEASURE_NAME_RE = re.compile(
    r"^\s*(?:'(?:[^']|'')+'|[^'\[\]]+?)?\s*\[\s*([^\[\]]+?)\s*\]\s*$"
)


def _split_measure_list(measures: str) -> list[str]:
    """Split the comma-separated ``measures`` argument, ignoring commas inside
    [brackets] or 'quoted table names' so "[A, B],[C]" yields two names.

    A quote only OPENS at the start of a token (the 'Table'[Measure] form) —
    an apostrophe inside a bare name ("Tom's Margin, Sales") is plain text and
    must not swallow the following comma. Inside a quoted table name, DAX's
    doubled-quote escape ('') stays part of the name.
    """
    parts: list[str] = []
    buf: list[str] = []
    depth = 0
    in_quote = False
    i = 0
    while i < len(measures):
        ch = measures[i]
        if ch == "'" and depth == 0:
            if in_quote:
                if i + 1 < len(measures) and measures[i + 1] == "'":
                    buf.append("''")           # escaped quote, stay in quote
                    i += 2
                    continue
                in_quote = False
            elif not "".join(buf).strip():     # token start -> table qualifier
                in_quote = True
        elif not in_quote:
            if ch == "[":
                depth += 1
            elif ch == "]" and depth:
                depth -= 1
        if ch == "," and depth == 0 and not in_quote:
            parts.append("".join(buf))
            buf = []
        else:
            buf.append(ch)
        i += 1
    parts.append("".join(buf))
    return [p.strip() for p in parts if p.strip()]


def _parse_measure_names(measures: str, measure_defs: dict) -> list[str]:
    """Split + normalize measure references to bare names, validating every
    name against the model.

    An exact model match always wins (a real measure named "Cost [USD]" is
    never mis-parsed as table "Cost" + measure "USD"); otherwise DAX-style
    forms are unwrapped to the bare name, with a case-insensitive fallback to
    the model's canonical casing (Power BI names are case-insensitive).

    Raises DAXMeasureNotFoundError (with close-match hints) for unknown names,
    so a typo is distinguishable from a measure that genuinely returns BLANK.
    """
    lower_map = {k.lower(): k for k in measure_defs}

    # A measure NAME may itself contain a comma -- Power BI allows it, and
    # Agents_Performance really does have "KPI, Avg MTD Sales towards Target".
    # Splitting first made every such measure unevaluable: the pieces ('KPI',
    # 'Avg MTD Sales towards Target') are not measures, so the call failed with
    # MEASURE_NOT_FOUND even though the measure exists. Applying this function's
    # own rule -- an exact model match always wins -- to the WHOLE string before
    # splitting fixes it without affecting genuine comma-separated lists, which
    # cannot match a single measure name.
    whole = (measures or "").strip()
    if whole in measure_defs:
        return [whole]
    if whole.lower() in lower_map:
        return [lower_map[whole.lower()]]

    names = []
    unknown = []
    for raw in _split_measure_list(measures):
        if raw in measure_defs:
            names.append(raw)
            continue
        m = _MEASURE_NAME_RE.match(raw)
        n = m.group(1).strip() if m else raw
        if n not in measure_defs:
            n = lower_map.get(n.lower(), n)
        names.append(n)
        if n not in measure_defs:
            unknown.append(n)

    if unknown:
        hints = []
        for n in unknown:
            close = difflib.get_close_matches(
                n, list(measure_defs.keys()), n=3, cutoff=0.6)
            hints.append(f"'{n}'"
                         + (f" (did you mean: {', '.join(close)}?)" if close else ""))
        raise DAXMeasureNotFoundError(
            "Measure(s) not found in the model: " + "; ".join(hints)
            + ". Use pbix_get_model_measures to list available measures.")
    return names


def _resolve_topn_filters(ctx: dict, filters: dict) -> dict:
    """Materialize ``{"top_n": ...}`` filter specs into concrete In-sets.

    Ledger issues-14 (half B; predicates were half A): a filter value of
    ``{"top_n": {"n": 5, "by": "<measure or Table.Column>",
    "direction": "desc"}}`` ranks the key's distinct values by the aggregate
    -- evaluated under the OTHER filters -- keeps the top n, and replaces the
    spec with the plain value list, so the engine then applies ordinary
    set-membership. ``"asc"`` ranks smallest-first (bottom-N); blanks sort
    last either way; ties keep model order (stable sort). This is exactly
    the materialization OpenBI performed client-side, moved server-side."""
    if not any(isinstance(v, dict) and "top_n" in v for v in filters.values()):
        return filters
    from pbix_mcp.dax import engine as dax_engine
    out = dict(filters)
    for key, spec in filters.items():
        if not (isinstance(spec, dict) and "top_n" in spec):
            continue
        tn = spec.get("top_n") or {}
        n = int(tn.get("n", 0))
        by = str(tn.get("by", "")).strip()
        direction = str(tn.get("direction", "desc")).lower()
        if n <= 0 or not by:
            raise ValueError(
                f"top_n spec for '{key}' needs n >= 1 and a 'by' measure "
                "or Table.Column")
        t, _, c = key.partition(".")
        tbl = (ctx.get("tables") or {}).get(t)
        if not tbl:
            raise ValueError(f"top_n filter key '{key}': table '{t}' not found")
        cols = tbl["columns"]
        if c not in cols:
            raise ValueError(
                f"top_n filter key '{key}': column '{c}' not in '{t}'")
        ci = cols.index(c)
        distinct = list(dict.fromkeys(r[ci] for r in tbl["rows"]))
        measures = dict(ctx["measure_defs"])
        by_name = by
        if by not in measures:
            low = by.lower()
            match = next((m for m in measures if m.lower() == low), None)
            if match is not None:
                by_name = match
            elif "." in by:
                bt, _, bc = by.partition(".")
                by_name = "__topn_by__"
                measures[by_name] = f"SUM('{bt}'[{bc}])"
            else:
                raise ValueError(
                    f"top_n 'by' = '{by}' is neither a measure nor a "
                    "Table.Column reference")
        # ranking context: every OTHER plain filter applies; other top_n
        # keys are excluded rather than half-resolved
        base_fc = {k: v for k, v in out.items()
                   if k != key and not (isinstance(v, dict) and "top_n" in v)}
        ranked = []
        for v in distinct:
            fc = dict(base_fc)
            fc[key] = [v]
            res = dax_engine.evaluate_measures_smart(
                [by_name], ctx["tables"], measures, fc,
                ctx.get("date_table"), ctx.get("date_column"),
                ctx.get("relationships"), simulate_row_context=False,
                measure_tables=ctx.get("measure_tables"),
                model_columns=ctx.get("model_columns"))
            val = res.get(by_name)
            if not isinstance(val, (int, float)) or isinstance(val, bool):
                val = None
            ranked.append((v, val))
        if direction == "asc":
            ranked.sort(key=lambda pr: (pr[1] is None,
                                        pr[1] if pr[1] is not None else 0.0))
        else:
            ranked.sort(key=lambda pr: (pr[1] is not None,
                                        pr[1] if pr[1] is not None else 0.0),
                        reverse=True)
        out[key] = [v for v, _ in ranked[:n]]
    return out


def _resolve_default_filters(ctx: dict, page_index: int) -> dict | None:
    """Default slicer filters for an evaluation.

    ``page_index=-1`` merges every page's slicer defaults (the historical
    behavior; mtime-cached). ``page_index>=0`` uses ONLY that page's slicers —
    the Power BI service scopes a slicer's default selection to the slicer's
    own page, so page-scoped application is what matches a visual on that
    page. An out-of-range page_index raises LayoutParseError (a typo'd index
    must not silently mean "raw model" — indistinguishable from a valid
    slicer-less page). On PARSE errors, fail CLOSED (no defaults) rather than
    reuse a possibly stale snapshot."""
    if page_index >= 0:
        work_dir = ctx.get("work_dir")
        # _get_layout falls back to the PBIR reader itself.
        layout = _get_layout(work_dir) if work_dir else None
        n_pages = len((layout or {}).get("sections", []))
        if page_index >= n_pages:
            raise LayoutParseError(
                f"Page index {page_index} out of range ({n_pages} page(s))")
        try:
            return _extract_default_filters_dict(
                work_dir, page_index, layout=layout) or None
        except Exception:
            return None
    try:
        return _get_default_filters_current(ctx) or None
    except Exception:
        return None


@mcp.tool()
def pbix_evaluate_dax(
    alias: str,
    measures: str,
    filter_context: str = "",
    apply_default_filters: bool = True,
    page_index: int = -1,
) -> str:
    """Evaluate one or more DAX measures against the data model.

    Uses the built-in DAX engine to compute measure values, supporting:
    SUM, AVERAGE, DIVIDE, IF, CALCULATE, DATEADD, REMOVEFILTERS, ALL,
    MAXX, SUMX, VAR/RETURN, and 25+ other DAX functions.

    Supports relationship-based filter propagation (star-schema joins).

    Args:
        alias: The alias of the open file
        measures: Comma-separated measure names to evaluate, e.g. "Sales,Profit Margin,Sales LY"
        filter_context: Optional JSON filter context, e.g. '{"dim-Date.Year": [2015]}'.
            A non-empty filter_context always wins (defaults are not merged in).
        apply_default_filters: When filter_context is empty: True (default)
            auto-applies the report's persisted default slicer selections;
            False evaluates against the RAW, truly unfiltered model.
            NOTE: pbix_evaluate_dax_per_dimension defaults this to False —
            pass the flag explicitly when you need identical behavior across
            both tools.
        page_index: Scope for the auto-applied defaults: -1 (default) merges
            every page's slicer defaults; >= 0 applies ONLY that page's
            slicers — the service scopes a slicer's default selection to its
            own page, so pass the page a visual lives on to reproduce the
            number that visual shows in the service.
    """
    try:
        info = _ensure_open(alias)
        if info.get("is_directquery"):
            return ToolResponse.error(
                "This file uses DirectQuery — DAX evaluation requires local data. "
                "Use layout, measure, and metadata tools instead.",
                UnsupportedFormatError.code,
            ).to_text()

        from pbix_mcp.dax import engine as dax_engine

        ctx = _get_dax_context(alias)
        measure_names = _parse_measure_names(measures, ctx['measure_defs'])

        parsed_fc = FilterContext.from_json_str(filter_context)
        if parsed_fc.filters:
            parsed_fc.filters = _resolve_topn_filters(ctx, parsed_fc.filters)
        fc: dict | None
        if parsed_fc.filters:
            fc = parsed_fc.filters
        elif apply_default_filters:
            # Auto-apply default slicer filters from the report layout, re-derived
            # from the current layout (not the open-time snapshot) so a slicer/
            # filter edit is honored on the next evaluate (OpenBI #7).
            fc = _resolve_default_filters(ctx, page_index)
        else:
            fc = None

        # Reset unsupported + error trackers before evaluation
        dax_engine._engine.unsupported_functions.clear()
        dax_engine._engine.eval_errors.clear()
        dax_engine._engine.timed_out.clear()
        logger.info("Evaluating %d measures for '%s'", len(measure_names), alias)

        # simulate_row_context=False: report what the MODEL returns, exactly as
        # Desktop's own engine does. The fallback that guesses a parameter-table
        # selection when a measure is BLANK produced numbers Desktop never shows
        # at the grand total (see evaluate_measures_smart).
        results = dax_engine.evaluate_measures_smart(
            measure_names, ctx['tables'], ctx['measure_defs'],
            fc, ctx['date_table'], ctx['date_column'],
            ctx.get('relationships'), simulate_row_context=False,
            measure_tables=ctx.get('measure_tables'),
            model_columns=ctx.get('model_columns')
        )

        # Build structured response with DAXResult objects
        unsupported = set(dax_engine._engine.unsupported_functions)
        timed_out = set(dax_engine._engine.timed_out)
        dax_results = []
        for name, val in results.items():
            if val is not None:
                # A declared Measure.DataType of Int64 (6) coerces an integral
                # float result to int, exactly as Analysis Services casts to
                # the declared type — so DISTINCTCOUNT-style measures report
                # data_type "Int64", not "Double" (issue #40). Only the
                # value-consistent direction is trusted: Measure.DataType is
                # unreliable the other way (currency measures stored as Int64
                # with non-integral values keep their real Double typing).
                if (ctx.get('measure_types', {}).get(name) == 6
                        and isinstance(val, float) and val.is_integer()):
                    val = int(val)
                dax_results.append(DAXResult(name=name, value=val, status="ok"))
            elif name in timed_out:
                # NOT a blank: the evaluation was abandoned on the wall-clock
                # budget. Reporting it as blank made "no value" and "we ran out
                # of time" look identical in the tool output.
                dax_results.append(DAXResult(
                    name=name, value=None, status="error",
                    error_message=(
                        "evaluation budget exceeded "
                        f"({dax_engine._engine._max_eval_seconds:.0f}s); raise "
                        "PBIX_DAX_MAX_SECONDS to allow longer"),
                ))
            elif name in dax_engine._engine.eval_errors:
                # NOT a blank: evaluation RAISED (unresolvable reference,
                # circular definition, ...). Reporting it as blank made a
                # broken measure indistinguishable from a legitimately empty
                # one (ledger issues-7).
                dax_results.append(DAXResult(
                    name=name, value=None, status="error",
                    error_message=dax_engine._engine.eval_errors[name],
                ))
            elif unsupported:
                # Value is None and unsupported functions were hit — mark as unsupported
                dax_results.append(DAXResult(
                    name=name, value=None, status="unsupported",
                    error_message=f"Uses unsupported function(s): {', '.join(sorted(unsupported))}",
                ))
            else:
                dax_results.append(DAXResult(name=name, value=None, status="blank"))

        warnings = []
        if unsupported:
            warnings.append(f"{len(unsupported)} unsupported DAX function(s): {', '.join(sorted(unsupported))}")

        response = DAXEvalResponse(
            success=True,
            results=dax_results,
            warnings=warnings,
        )
        logger.debug("DAX eval complete: %d ok, %d blank",
                      sum(1 for r in dax_results if r.status == "ok"),
                      sum(1 for r in dax_results if r.status == "blank"))
        return response.to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_evaluate_dax_per_dimension(
    alias: str,
    measures: str,
    dimension: str,
    filter_context: str = "",
    max_values: int = 20,
    apply_default_filters: bool = False,
    page_index: int = -1,
) -> str:
    """Evaluate DAX measures for each value of a dimension (e.g., Sales per State).

    Iterates over unique values of a dimension column and evaluates measures
    with that dimension value as a filter. Uses relationship-based propagation.

    Args:
        alias: The alias of the open file
        measures: Comma-separated measure names, e.g. "Sales,Sales LY,Sales change"
        dimension: Table.Column to iterate over, e.g. "dim-Geo.State"
        filter_context: Optional JSON base filter, e.g. '{"dim-Date.Year": [2015]}'.
            A non-empty filter_context always wins (defaults are not merged in).
        max_values: Maximum dimension values to evaluate (default 20)
        apply_default_filters: When filter_context is empty: False (default)
            iterates against the raw model — the historical behavior, which
            matches a visual on a page without restrictive slicers; True
            auto-applies the report's persisted default slicer selections
            (same machinery as pbix_evaluate_dax).
            NOTE: pbix_evaluate_dax defaults this to True — pass the flag
            explicitly when you need identical behavior across both tools.
        page_index: Scope for the auto-applied defaults: -1 (default) merges
            every page's slicer defaults; >= 0 applies ONLY that page's
            slicers (service semantics — a slicer's default selection scopes
            to its own page).
    """
    try:
        from pbix_mcp.dax import engine as dax_engine

        ctx = _get_dax_context(alias)
        measure_names = _parse_measure_names(measures, ctx['measure_defs'])
        parsed_fc = FilterContext.from_json_str(filter_context)
        if parsed_fc.filters:
            parsed_fc.filters = _resolve_topn_filters(ctx, parsed_fc.filters)
        base_fc = parsed_fc.filters
        if not base_fc and apply_default_filters:
            base_fc = _resolve_default_filters(ctx, page_index) or {}

        try:
            dim_ref = DimensionRef.parse(dimension)
        except ValueError as e:
            # DimensionParseError carries .message/.code; a plain ValueError
            # does not — never let the handler itself raise AttributeError.
            return ToolResponse.error(
                getattr(e, "message", None) or str(e),
                getattr(e, "code", None) or "INVALID_INPUT",
            ).to_text()
        dim_table, dim_col = dim_ref.table, dim_ref.column

        # Get unique dimension values
        tbl = ctx['tables'].get(dim_table)
        if not tbl:
            return ToolResponse.error(f"Table '{dim_table}' not found", PBIXMCPError.code).to_text()
        col_idx = next((i for i, c in enumerate(tbl['columns']) if c == dim_col), -1)
        if col_idx < 0:
            return ToolResponse.error(f"Column '{dim_col}' not found in '{dim_table}'", PBIXMCPError.code).to_text()

        unique_vals = list(set(row[col_idx] for row in tbl['rows'] if row[col_idx] is not None))
        unique_vals.sort(key=lambda x: str(x))

        lines = [f"DAX per {dimension} ({len(unique_vals)} values, showing {min(len(unique_vals), max_values)}):\n"]

        # Header
        header = f"{'Value':<25s}"
        for m in measure_names:
            header += f"  {m:>15s}"
        lines.append(header)
        lines.append("-" * len(header))

        capped_vals = unique_vals[:max_values]

        # Fast path: group the fact rows by the propagated join key ONCE and
        # aggregate per bucket, instead of re-filtering the whole fact table for
        # every dimension value. Only covers simple aggregations (SUM/COUNT/…);
        # measures it cannot safely bucket are absent from `fast` and are
        # evaluated the exact per-value way below, so results never change.
        fast = dax_engine.evaluate_per_dimension(
            measure_names, ctx['tables'], ctx['measure_defs'], base_fc,
            dimension, dim_table, dim_col, capped_vals,
            ctx['date_table'], ctx['date_column'], ctx.get('relationships')
        )
        fallback_measures = [m for m in measure_names if m not in fast]

        for val in capped_vals:
            if fallback_measures:
                fc = dict(base_fc)
                fc[dimension] = [val]
                fb = dax_engine.evaluate_measures_batch(
                    fallback_measures, ctx['tables'], ctx['measure_defs'],
                    fc, ctx['date_table'], ctx['date_column'],
                    ctx.get('relationships'), group_keys={dimension},
                    selected_filters=base_fc
                )
            else:
                fb = {}
            results = {m: (fast[m].get(val) if m in fast else fb.get(m))
                       for m in measure_names}

            row_str = f"{str(val):<25s}"
            for m in measure_names:
                v = results.get(m)
                if isinstance(v, float):
                    if abs(v) < 2 and abs(v) > 0.001:
                        row_str += f"  {v:>14.1%}"
                    else:
                        row_str += f"  {v:>15,.2f}"
                elif isinstance(v, int):
                    row_str += f"  {v:>15,}"
                elif v is None:
                    row_str += f"  {'(null)':>15s}"
                else:
                    row_str += f"  {str(v):>15s}"
            lines.append(row_str)

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_evaluate_dax_grouped(
    alias: str,
    measures: str,
    group_by: str,
    filter_context: str = "",
    max_groups: int = 3500,
    apply_default_filters: bool = False,
    page_index: int = -1,
    top_n: int = 0,
    order_by: str = "",
    order: str = "desc",
) -> str:
    """Evaluate measures for every group key in ONE call, returning STRUCTURED rows.

    The GROUP-BY entry point for chart-shaped work: instead of a separate
    evaluation per category value (O(distinct values) round-trips), the fact
    rows are bucketed by the propagated join key once and each bucket is
    aggregated — so a chart bound to thousands of categories costs a single
    call. Results come back as machine-readable data (one object per group),
    not a formatted table, so a client can bind them directly.

    Measures the fast path can't bucket (anything that isn't a simple
    aggregation, or an ambiguous join) are still evaluated exactly, per group,
    so values always match ``pbix_evaluate_dax``.

    Args:
        alias: The alias of the open file
        measures: Comma-separated measure names, e.g. "Sales,Sales LY"
        group_by: Grouping column "Table.Column". Comma-separate for a
            composite key, e.g. "dim-Geo.Country,dim-Geo.State" (a composite
            key evaluates per group combination — this is also the MATRIX
            recipe: "RowDim.Col,ColDim.Col" returns one structured row per
            (row, column) cell, which a client pivots into the grid; a
            series chart is the same with the series column as the second
            key — the single-column form is
            the one with the single-pass fast path).
        filter_context: Optional JSON base filter, e.g. '{"dim-Date.Year": [2015]}'
        max_groups: Cap on groups returned (default 3500 — Power BI's own
            data-reduction window). Groups beyond it are dropped and reported
            via ``truncated``/``group_count``.
        apply_default_filters: When filter_context is empty, apply the report's
            persisted default slicer selections (default False = raw model).
            NOTE: pbix_evaluate_dax defaults this to True; the per-dimension and
            grouped tools default to False. That is deliberate — a sweep over
            many group keys is normally asked against the RAW model — but it
            means the same measure can differ between the two tools unless you
            pass the flag explicitly. Do so whenever you compare their output.
        page_index: Scope for those defaults: -1 merges every page, >= 0 scopes
            to that page (service semantics).
        top_n: Keep only the highest/lowest N groups after evaluation — a live
            Top-N ranking (0 = keep all). Every group is still evaluated, so the
            ranking reflects real measure values.
        order_by: Measure to rank by when top_n is set (defaults to the first
            measure). Also sorts the returned rows.
        order: "desc" (default) or "asc".
    """
    try:
        from pbix_mcp.dax import engine as dax_engine

        ctx = _get_dax_context(alias)
        rels: list = ctx.get('relationships') or []
        measure_names = _parse_measure_names(measures, ctx['measure_defs'])
        parsed_fc = FilterContext.from_json_str(filter_context)
        if parsed_fc.filters:
            parsed_fc.filters = _resolve_topn_filters(ctx, parsed_fc.filters)
        base_fc = parsed_fc.filters
        if not base_fc and apply_default_filters:
            base_fc = _resolve_default_filters(ctx, page_index) or {}

        refs = [g.strip() for g in (group_by or "").split(",") if g.strip()]
        if not refs:
            return ToolResponse.error(
                "group_by is required, e.g. 'dim-Geo.State'.",
                "INVALID_INPUT").to_text()
        keys = []
        for ref in refs:
            try:
                d = DimensionRef.parse(ref)
            except ValueError as e:
                return ToolResponse.error(
                    getattr(e, "message", None) or str(e),
                    getattr(e, "code", None) or "INVALID_INPUT").to_text()
            tbl = ctx['tables'].get(d.table)
            if not tbl:
                return ToolResponse.error(
                    f"Table '{d.table}' not found", "TABLE_NOT_FOUND").to_text()
            if d.column not in tbl['columns']:
                return ToolResponse.error(
                    f"Column '{d.column}' not found in '{d.table}'",
                    "COLUMN_NOT_FOUND").to_text()
            keys.append((f"{d.table}.{d.column}", d.table, d.column,
                         tbl['columns'].index(d.column), tbl))

        # Distinct group keys, in a stable order.
        if len(keys) == 1:
            _ref, _t, _c, idx, tbl = keys[0]
            uniq = list(dict.fromkeys(
                r[idx] for r in tbl['rows'] if r[idx] is not None))
        else:
            if len({k[1] for k in keys}) > 1:
                return ToolResponse.error(
                    "A composite group_by must use columns from ONE table.",
                    "INVALID_INPUT").to_text()
            tbl = keys[0][4]
            idxs = [k[3] for k in keys]
            uniq = list(dict.fromkeys(
                tuple(r[i] for i in idxs) for r in tbl['rows']
                if all(r[i] is not None for i in idxs)))
        uniq.sort(key=lambda v: tuple(str(x) for x in v)
                  if isinstance(v, tuple) else str(v))
        total = len(uniq)
        capped = uniq[:max_groups]

        results: list[dict] = []
        if len(keys) == 1:
            ref, dim_table, dim_col = keys[0][0], keys[0][1], keys[0][2]
            fast = dax_engine.evaluate_per_dimension(
                measure_names, ctx['tables'], ctx['measure_defs'], base_fc,
                ref, dim_table, dim_col, capped,
                ctx['date_table'], ctx['date_column'], rels)
            slow = [m for m in measure_names if m not in fast]
            for val in capped:
                vals = {}
                if slow:
                    fc = dict(base_fc)
                    fc[ref] = [val]
                    vals = dax_engine.evaluate_measures_batch(
                        slow, ctx['tables'], ctx['measure_defs'], fc,
                        ctx['date_table'], ctx['date_column'], rels,
                        group_keys={ref}, selected_filters=base_fc)
                results.append({
                    "key": {dim_col: val},
                    "values": {m: (fast[m].get(val) if m in fast else
                                   vals.get(m)) for m in measure_names},
                })
        else:
            for combo in capped:
                fc = dict(base_fc)
                for (ref, _t, _c, _i, _tb), v in zip(keys, combo):
                    fc[ref] = [v]
                vals = dax_engine.evaluate_measures_batch(
                    measure_names, ctx['tables'], ctx['measure_defs'], fc,
                    ctx['date_table'], ctx['date_column'], rels,
                    group_keys={k[0] for k in keys}, selected_filters=base_fc)
                results.append({
                    "key": {k[2]: v for k, v in zip(keys, combo)},
                    "values": {m: vals.get(m) for m in measure_names},
                })

        # Live Top-N / ordering over the evaluated groups.
        ranked_by = ""
        if top_n or order_by:
            ranked_by = order_by or (measure_names[0] if measure_names else "")
            if ranked_by and ranked_by not in measure_names:
                return ToolResponse.error(
                    f"order_by '{ranked_by}' is not one of the evaluated "
                    f"measures: {measure_names}", "INVALID_INPUT").to_text()
            if ranked_by:
                # Groups with no value (BLANK) always sink to the bottom, in
                # BOTH directions — otherwise a null would outrank a real
                # number in a Top-N.
                def _has_val(g):
                    return isinstance(g["values"].get(ranked_by), (int, float))
                scored = [g for g in results if _has_val(g)]
                blanks = [g for g in results if not _has_val(g)]
                scored.sort(key=lambda g: g["values"][ranked_by],
                            reverse=str(order).lower() != "asc")
                results = scored + blanks
            if top_n and top_n > 0:
                results = results[:top_n]

        return ToolResponse.ok(
            f"Evaluated {len(measure_names)} measure(s) across "
            f"{len(results):,} of {total:,} group(s)."
            + (f" Top {top_n} by '{ranked_by}' ({order})." if top_n else ""),
            data={
                "group_by": [k[0] for k in keys],
                "measures": measure_names,
                "group_count": total,
                "returned": len(results),
                "truncated": total > len(results),
                "order_by": ranked_by or None,
                "order": order if ranked_by else None,
                "groups": results,
            },
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(
            f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


def _pbir_alias_factory():
    """Allocate short From-aliases per entity, the way a classic query does."""
    aliases: dict = {}

    def alias_for(entity: str) -> str:
        if entity not in aliases:
            base = (entity[:1] or "t").lower()
            cand, n = base, 1
            while cand in aliases.values():
                n += 1
                cand = f"{base}{n}"
            aliases[entity] = cand
        return str(aliases[entity])

    return aliases, alias_for


def _pbir_rewrite_sourcerefs(node, alias_for):
    """Rewrite PBIR ``SourceRef.Entity`` into the classic ``SourceRef.Source``
    alias form, preserving everything else (works for Column, Measure and
    Aggregation shapes alike)."""
    if isinstance(node, dict):
        out = {}
        for k, v in node.items():
            if k == "SourceRef" and isinstance(v, dict) and "Entity" in v:
                out[k] = {"Source": alias_for(v["Entity"])}
            else:
                out[k] = _pbir_rewrite_sourcerefs(v, alias_for)
        return out
    if isinstance(node, list):
        return [_pbir_rewrite_sourcerefs(x, alias_for) for x in node]
    return node


def _pbir_visual_to_container(vdata: dict) -> dict:
    """Convert one PBIR ``visual.json`` into a legacy visualContainer.

    Carries across everything a consumer needs to actually use the visual: its
    NAME (so it can be addressed at all), geometry incl. z/tabOrder, the field
    bindings as both ``projections`` and a synthesized ``prototypeQuery`` (so
    column-vs-measure is recoverable), hidden state, and its filters.
    """
    visual_obj = dict(vdata.get("visual") or {})
    query = visual_obj.pop("query", None) or {}

    # Container-level formatting (title, background, border, shadow, header)
    # lives at `visual.visualContainerObjects` in PBIR — verified against 70/70
    # visuals in the service-authored corpus, and required there by
    # visualConfiguration/2.3.0 (visualContainer sets additionalProperties:
    # false and does NOT allow it at the top level). The classic spelling is
    # singleVisual.vcObjects. Without this rename every container format read
    # back empty, so recolor skipped container colours and format_visual's
    # title vanished on save.
    if "visualContainerObjects" in visual_obj:
        visual_obj["vcObjects"] = visual_obj.pop("visualContainerObjects")
    # Carry the PBIR sort so the writer can tell "unchanged" from "cleared".
    if query.get("sortDefinition") is not None:
        visual_obj["__pbir_sort__"] = query["sortDefinition"]

    aliases, alias_for = _pbir_alias_factory()
    projections: dict = {}
    selects: list = []
    seen_select: set = set()
    for role, rv in (query.get("queryState") or {}).items():
        entries = []
        for proj in (rv.get("projections") or []):
            qref = proj.get("queryRef")
            entry = {"queryRef": qref}
            if "active" in proj:
                entry["active"] = proj["active"]
            if proj.get("nativeQueryRef"):
                entry["nativeQueryRef"] = proj["nativeQueryRef"]
            entries.append(entry)
            field = proj.get("field") or {}
            if field and qref not in seen_select:
                sel = _pbir_rewrite_sourcerefs(field, alias_for)
                if isinstance(sel, dict):
                    sel = dict(sel)
                    sel["Name"] = qref
                    selects.append(sel)
                    seen_select.add(qref)
        projections[role] = entries

    if projections:
        visual_obj["projections"] = projections
    if selects:
        visual_obj["prototypeQuery"] = {
            "Version": 2,
            "From": [{"Name": a, "Entity": e, "Type": 0}
                     for e, a in aliases.items()],
            "Select": selects,
        }

    pos = vdata.get("position") or {}
    config: dict = {"name": vdata.get("name", ""), "singleVisual": visual_obj}
    if pos:
        config["layouts"] = [{"id": 0, "position": dict(pos)}]

    container: dict = {
        # Folder id this visual came from. The PBIR writer re-reads that
        # original file and patches only what changed, so fields this
        # converter doesn't model survive an edit untouched.
        "__pbir_visual__": vdata.get("name", ""),
        "config": json.dumps(config),
        "x": pos.get("x", 0),
        "y": pos.get("y", 0),
        "z": pos.get("z", 0),
        "width": pos.get("width", 0),
        "height": pos.get("height", 0),
        "tabOrder": pos.get("tabOrder", 0),
    }
    if vdata.get("isHidden"):
        container["isHidden"] = True
    vfilters = (vdata.get("filterConfig") or {}).get("filters")
    if vfilters:
        container["filters"] = json.dumps(vfilters)
    return container


def _get_layout_pbir(work_dir: str) -> dict | None:
    """Read a PBIR report as a legacy-compatible layout structure.

    PBIR (what the Power BI service produces) stores the report as a tree:
    ``Report/definition/pages/pages.json`` for the page order plus a
    ``page.json`` and one ``visuals/<id>/visual.json`` per visual. This
    converts that tree into the classic ``{sections: [{visualContainers: []}]}``
    document every other reader in pbix-mcp already understands.

    The result is marked ``__pbir__`` so ``_set_layout`` can refuse to write a
    synthesized layout back over a PBIR file.
    """
    pages_json = os.path.join(work_dir, "Report", "definition", "pages", "pages.json")
    if not os.path.exists(pages_json):
        return None

    try:
        with open(pages_json, "r", encoding="utf-8") as f:
            pages_meta = json.load(f)
    except Exception:
        return None

    pages_dir = os.path.dirname(pages_json)
    sections = []

    page_dirs = []
    if isinstance(pages_meta, list):
        for pm in pages_meta:
            pid = pm.get("id") or pm.get("name", "")
            if pid:
                page_dirs.append(pid)
    elif isinstance(pages_meta, dict):
        # Real PBIR pages.json is an object whose "pageOrder" array IS the
        # page order (Microsoft's pagesMetadata schema). Honor it — page
        # indices are load-bearing (page-scoped default filters) — and fall
        # back to a SORTED directory listing so the order is at least
        # deterministic when pageOrder is absent (os.listdir order is
        # filesystem-dependent).
        try:
            on_disk = [pid for pid in sorted(os.listdir(pages_dir))
                       if os.path.isdir(os.path.join(pages_dir, pid))]
        except OSError:
            on_disk = []
        page_order = pages_meta.get("pageOrder")
        if isinstance(page_order, list) and page_order:
            ordered = [pid for pid in page_order if pid in on_disk]
            page_dirs = ordered + [pid for pid in on_disk if pid not in ordered]
        else:
            page_dirs = on_disk

    active_page = pages_meta.get("activePageName") if isinstance(pages_meta, dict) else None

    for pid in page_dirs:
        page_dir = os.path.join(pages_dir, pid)
        pdata = {}
        page_json = os.path.join(page_dir, "page.json")
        if os.path.exists(page_json):
            try:
                with open(page_json, "r", encoding="utf-8") as f:
                    pdata = json.load(f)
            except Exception:
                pdata = {}

        containers = []
        visuals_dir = os.path.join(page_dir, "visuals")
        if os.path.isdir(visuals_dir):
            for vid in sorted(os.listdir(visuals_dir)):
                visual_json = os.path.join(visuals_dir, vid, "visual.json")
                if not os.path.exists(visual_json):
                    continue
                try:
                    with open(visual_json, "r", encoding="utf-8") as f:
                        vdata = json.load(f)
                    container = _pbir_visual_to_container(vdata)
                    mobile_json = os.path.join(visuals_dir, vid, "mobile.json")
                    if os.path.exists(mobile_json):
                        try:
                            with open(mobile_json, "r", encoding="utf-8") as f:
                                container["mobile"] = json.load(f)
                        except Exception:
                            pass
                    containers.append(container)
                except Exception:
                    continue

        section: dict = {
            "__pbir_page__": pid,
            "name": pdata.get("name", pid),
            "displayName": pdata.get("displayName", pid),
            "visualContainers": containers,
        }
        for key in ("width", "height", "displayOption", "type"):
            if key in pdata:
                section[key] = pdata[key]
        # Present the classic int form so callers see one type regardless of
        # which format the report is stored in (_pbir_patch_page inverts it).
        if isinstance(section.get("displayOption"), str):
            section["displayOption"] = _PBIR_DISPLAY_OPTION_INV.get(
                section["displayOption"], section["displayOption"])
        # PBIR puts page visibility on the page as a string enum; classic keeps
        # it in the section `config` JSON as an int. Normalize to the classic
        # form so one code path drives both formats.
        if "visibility" in pdata:
            section["config"] = json.dumps(
                {"visibility":
                 1 if pdata["visibility"] == "HiddenInViewMode" else 0},
                ensure_ascii=False)
        pfilters = (pdata.get("filterConfig") or {}).get("filters")
        if pfilters:
            section["filters"] = json.dumps(pfilters)
        if pdata.get("objects"):
            section["objects"] = pdata["objects"]
        if active_page and pdata.get("name", pid) == active_page:
            section["isActive"] = True
        sections.append(section)

    if not sections:
        return None
    layout = {"sections": sections, "__pbir__": True}

    # Report-level state lives in Report/definition/report.json on PBIR, but
    # every classic caller reads and mutates it on the LAYOUT. Surfacing it
    # here (and writing it back in _set_layout_pbir) is what makes image,
    # theme and custom-visual registration work on a PBIR report at all —
    # without it those tools mutated a throwaway dict and reported success.
    config: dict = {}
    bookmarks = _pbir_read_bookmarks(work_dir)
    if bookmarks:
        config["bookmarks"] = bookmarks

    rcfg = {}
    rpath = _report_config_path(work_dir)
    if os.path.exists(rpath):
        try:
            with open(rpath, "r", encoding="utf-8-sig") as f:
                rcfg = json.load(f)
        except Exception:
            rcfg = {}
    if rcfg.get("resourcePackages"):
        layout["resourcePackages"] = _pbir_packages_to_classic(
            rcfg["resourcePackages"])
    if rcfg.get("publicCustomVisuals"):
        layout["publicCustomVisuals"] = list(rcfg["publicCustomVisuals"])
    if rcfg.get("themeCollection"):
        config["themeCollection"] = copy.deepcopy(rcfg["themeCollection"])
    # Report-level filters: PBIR report.json.filterConfig.filters <-> the
    # classic layout's top-level `filters` JSON string.
    rfilters = (rcfg.get("filterConfig") or {}).get("filters")
    if rfilters:
        layout["filters"] = json.dumps(rfilters, ensure_ascii=False)
    if rcfg.get("settings"):
        config["settings"] = copy.deepcopy(rcfg["settings"])
    if rcfg.get("objects"):
        config["objects"] = copy.deepcopy(rcfg["objects"])

    if config:
        layout["config"] = json.dumps(config, ensure_ascii=False)
    return layout


# PBIR keeps each bookmark in its own file under `definition/bookmarks/`, with
# a sibling `bookmarks.json` holding the order and any groups. See
# https://learn.microsoft.com/power-bi/developer/projects/projects-report
_PBIR_BOOKMARK_SCHEMA = (
    "https://developer.microsoft.com/json-schemas/fabric/item/report/"
    "definition/bookmark/1.0.0/schema.json")
_PBIR_BOOKMARKS_META_SCHEMA = (
    "https://developer.microsoft.com/json-schemas/fabric/item/report/"
    "definition/bookmarksMetadata/1.0.0/schema.json")


def _pbir_bookmarks_dir(work_dir: str) -> str:
    return os.path.join(work_dir, "Report", "definition", "bookmarks")


def _pbir_read_bookmarks(work_dir: str) -> list:
    """Load `definition/bookmarks/` into the classic layout-config shape.

    Groups become a classic bookmark whose ``children`` array holds the nested
    bookmarks, which is how Report/Layout represents a bookmark group.
    """
    bdir = _pbir_bookmarks_dir(work_dir)
    if not os.path.isdir(bdir):
        return []

    by_name: dict = {}
    for fn in sorted(os.listdir(bdir)):
        if not fn.endswith(".bookmark.json"):
            continue
        try:
            with open(os.path.join(bdir, fn), "r", encoding="utf-8") as f:
                doc = json.load(f)
        except Exception:
            continue
        # Remember the declared $schema; the service stamps versions newer than
        # the one we default to, and overwriting it rewrites files the caller
        # never touched.
        schema = doc.pop("$schema", None)
        name = doc.get("name") or fn[: -len(".bookmark.json")]
        doc["name"] = name
        # Remember the file stem so a later write lands on the same file even
        # when it was renamed away from the default convention.
        doc["__pbir_file__"] = fn[: -len(".bookmark.json")]
        if schema:
            doc["__pbir_schema__"] = schema
        by_name[name] = doc

    meta_path = os.path.join(bdir, "bookmarks.json")
    order: list = []
    if os.path.exists(meta_path):
        try:
            with open(meta_path, "r", encoding="utf-8") as f:
                order = (json.load(f) or {}).get("items") or []
        except Exception:
            order = []

    out: list = []
    seen: set = set()
    for item in order:
        if not isinstance(item, dict):
            continue
        if "children" in item:
            kids = [by_name[c] for c in item.get("children", [])
                    if c in by_name]
            seen.update(k["name"] for k in kids)
            out.append({"name": item.get("name", ""),
                        "displayName": item.get("displayName", ""),
                        "children": kids})
        else:
            bm = by_name.get(item.get("name", ""))
            if bm is not None:
                seen.add(bm["name"])
                out.append(bm)
    # Bookmarks present on disk but missing from bookmarks.json still belong to
    # the report — appending them keeps a read/write cycle from dropping them.
    out.extend(bm for name, bm in by_name.items() if name not in seen)
    return out


def _pbir_write_bookmarks(work_dir: str, bookmarks: list) -> None:
    """Persist classic-shaped bookmarks back into `definition/bookmarks/`."""
    import uuid

    bdir = _pbir_bookmarks_dir(work_dir)
    if not bookmarks:
        # Nothing to write: drop a now-empty folder so the report matches.
        if os.path.isdir(bdir):
            shutil.rmtree(bdir, ignore_errors=True)
        return

    os.makedirs(bdir, exist_ok=True)
    items: list = []
    written: set = set()

    def _write_one(bm: dict) -> str:
        name = bm.get("name") or "Bookmark" + uuid.uuid4().hex[:20]
        stem = bm.get("__pbir_file__") or name
        doc = {k: v for k, v in bm.items()
               if k not in ("__pbir_file__", "__pbir_schema__", "children")}
        doc["name"] = name
        doc.setdefault("displayName", name)
        doc.setdefault("explorationState", {})
        # Keep whatever version the file declared; only a NEW bookmark gets our
        # default stamped on it.
        doc["$schema"] = bm.get("__pbir_schema__") or _PBIR_BOOKMARK_SCHEMA
        _pbir_write_json(os.path.join(bdir, f"{stem}.bookmark.json"), doc)
        written.add(f"{stem}.bookmark.json")
        return name

    for bm in bookmarks:
        if not isinstance(bm, dict):
            continue
        kids = bm.get("children")
        if isinstance(kids, list) and kids:
            child_names = [_write_one(k) for k in kids
                           if isinstance(k, dict)]
            items.append({"name": bm.get("name") or
                          "Group" + uuid.uuid4().hex[:16],
                          "displayName": bm.get("displayName", ""),
                          "children": child_names})
        else:
            items.append({"name": _write_one(bm)})

    _pbir_write_json(os.path.join(bdir, "bookmarks.json"),
                     {"$schema": _PBIR_BOOKMARKS_META_SCHEMA, "items": items})

    # Remove bookmark files the caller deleted.
    for fn in os.listdir(bdir):
        if fn.endswith(".bookmark.json") and fn not in written:
            try:
                os.remove(os.path.join(bdir, fn))
            except OSError:
                pass


def _pbir_entityize(node, alias2entity: dict):
    """Inverse of _pbir_rewrite_sourcerefs: classic ``SourceRef.Source`` alias
    back to the PBIR ``SourceRef.Entity`` form."""
    if isinstance(node, dict):
        out = {}
        for k, v in node.items():
            if k == "SourceRef" and isinstance(v, dict) and "Source" in v:
                alias = v["Source"]
                out[k] = {"Entity": alias2entity.get(alias, alias)}
            else:
                out[k] = _pbir_entityize(v, alias2entity)
        return out
    if isinstance(node, list):
        return [_pbir_entityize(x, alias2entity) for x in node]
    return node


def _pbir_query_state_from_single_visual(sv: dict) -> dict:
    """Rebuild PBIR ``query.queryState`` from legacy projections + prototypeQuery."""
    projections = sv.get("projections") or {}
    proto = sv.get("prototypeQuery") or {}
    alias2entity = {f.get("Name"): f.get("Entity")
                    for f in (proto.get("From") or [])}
    select_by_name = {s.get("Name"): s for s in (proto.get("Select") or [])}
    state: dict = {}
    for role, entries in projections.items():
        plist = []
        for e in entries or []:
            qref = e.get("queryRef")
            item: dict = {}
            sel = select_by_name.get(qref)
            if sel:
                field = {k: v for k, v in sel.items() if k != "Name"}
                item["field"] = _pbir_entityize(field, alias2entity)
            if qref is not None:
                item["queryRef"] = qref
            if e.get("nativeQueryRef"):
                item["nativeQueryRef"] = e["nativeQueryRef"]
            if "active" in e:
                item["active"] = e["active"]
            plist.append(item)
        state[role] = {"projections": plist}
    return state


# Fields the PBIR schemas type as string enums, where classic Report/Layout
# uses an int. Writing the int produces a file the Power BI service IMPORTS
# without complaint and then refuses to open ("Unable to load report") — a
# schema violation is a documented BLOCKING error. Checking here turns that
# into a loud failure at save time instead.
_PBIR_ENUM_FIELDS = {
    "page.json": {
        "displayOption": set(_PBIR_DISPLAY_OPTION.values()),
        "visibility": {"AlwaysVisible", "HiddenInViewMode"},
        "type": {"Drillthrough", "Tooltip"},
        "howCreated": {"Default", "Copilot"},
    },
}
_PBIR_REQUIRED = {
    "page.json": ("name", "displayName", "displayOption"),
    "visual.json": ("name",),
}


def _pbir_selfcheck(path: str, data: dict) -> None:
    """Structural check of one PBIR document before it is written.

    Deliberately offline and narrow: it enforces the handful of rules this
    converter can plausibly get wrong (an enum written as the classic int, a
    required field dropped), not the full published schema. For the full check
    against Microsoft's own schemas see scripts/validate_pbir_schemas.py.
    """
    kind = os.path.basename(path)
    for field in _PBIR_REQUIRED.get(kind, ()):
        if data.get(field) in (None, ""):
            raise LayoutParseError(
                f"Refusing to write {kind}: required field '{field}' is "
                f"missing. Power BI would import this report and then fail to "
                f"open it.")
    for field, allowed in _PBIR_ENUM_FIELDS.get(kind, {}).items():
        if field not in data:
            continue
        value = data[field]
        if isinstance(value, str) and value in allowed:
            continue
        hint = ""
        if isinstance(value, bool) or isinstance(value, int):
            hint = (" This looks like the classic Report/Layout integer form; "
                    "PBIR stores the enum NAME.")
        raise LayoutParseError(
            f"Refusing to write {kind}: '{field}' is {value!r}, which is not "
            f"one of {sorted(allowed)}.{hint} Power BI would import this "
            f"report and then fail to open it.")


def _pbir_write_json(path: str, data: dict) -> None:
    _pbir_selfcheck(path, data)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


# Keys the reader SYNTHESIZES onto singleVisual or relocates from elsewhere in
# visual.json. They must never be written back into the PBIR `visual` block, or
# the file gains classic-only fields that Power BI does not model.
_PBIR_SYNTHETIC_SV_KEYS = frozenset(
    {"projections", "prototypeQuery", "vcObjects", "__pbir_sort__"})


def _pbir_sort_definition(sv: dict) -> dict | None:
    """Translate a classic ``prototypeQuery.OrderBy`` into PBIR sortDefinition.

    Returns None when the caller authored no sort, so the caller can fall back
    to whatever the original file had rather than clearing it.
    """
    proto = sv.get("prototypeQuery") or {}
    order_by = proto.get("OrderBy")
    if order_by is None:
        return None

    alias2entity = {f.get("Name"): f.get("Entity")
                    for f in (proto.get("From") or [])}

    def _entityize(node: dict) -> None:
        ref = (node.get("Expression") or {}).get("SourceRef") or {}
        if "Source" in ref:
            node["Expression"]["SourceRef"] = {
                "Entity": alias2entity.get(ref["Source"], ref["Source"])}

    entries = []
    for ob in order_by:
        expr = copy.deepcopy(ob.get("Expression") or {})
        inner = expr.get("Measure") or expr.get("Column")
        if inner is None and "Aggregation" in expr:
            agg_inner = expr["Aggregation"].get("Expression") or {}
            inner = agg_inner.get("Column") or agg_inner.get("Measure")
        if inner is None:
            # Unknown shape (e.g. HierarchyLevel) — skip rather than leak an
            # alias-based SourceRef into the PBIR.
            continue
        _entityize(inner)
        entries.append({
            "field": expr,
            "direction": ("Ascending" if ob.get("Direction") == 1
                          else "Descending"),
        })
    if entries:
        return {"sort": entries, "isDefaultSort": False}
    return {"sort": [], "isDefaultSort": True}


def _pbir_patch_visual(orig: dict, container: dict) -> dict:
    """Apply a legacy visualContainer onto its ORIGINAL PBIR visual.json.

    Only fields that actually differ from what the reader would have produced
    for ``orig`` are written, so anything this converter doesn't model (custom
    visual settings, sortDefinition, howCreated, …) survives untouched.
    """
    out = copy.deepcopy(orig) if orig else {}
    cfg = _parse_visual_config(container)
    sv = cfg.get("singleVisual", {}) or {}
    baseline_cfg = _parse_visual_config(
        _pbir_visual_to_container(orig)) if orig else {}
    base_sv = baseline_cfg.get("singleVisual", {}) or {}

    name = cfg.get("name") or container.get("__pbir_visual__") or out.get("name")
    if name:
        out["name"] = name

    # Only persist geometry the caller actually CHANGED. The reader defaults
    # missing keys (e.g. tabOrder -> 0), and writing those back would invent
    # fields the original file never had.
    base_container = _pbir_visual_to_container(orig) if orig else {}
    pos = dict(out.get("position") or {})
    for key in ("x", "y", "z", "width", "height", "tabOrder"):
        if key not in container or container[key] is None:
            continue
        if orig and container[key] == base_container.get(key) and key not in pos:
            continue  # unchanged AND absent from the original — leave it out
        pos[key] = container[key]
    if pos:
        out["position"] = pos

    visual = dict(out.get("visual") or {})
    # Everything the caller left on singleVisual belongs in the PBIR `visual`
    # block, EXCEPT the keys this converter synthesizes or relocates. A closed
    # whitelist here silently dropped columnProperties, expansionStates,
    # activeProjections, showAllRoles, display and howCreated — and, for a
    # NEWLY created visual (no original to deep-copy), everything else too.
    for key, value in sv.items():
        if key in _PBIR_SYNTHETIC_SV_KEYS:
            continue
        visual[key] = value
    for key in list(visual):
        if (key not in sv and key in base_sv
                and key not in _PBIR_SYNTHETIC_SV_KEYS):
            visual.pop(key, None)

    # Container formatting goes back under `visual`, using the PBIR spelling.
    if "vcObjects" in sv:
        visual["visualContainerObjects"] = sv["vcObjects"]
    elif "vcObjects" in base_sv:
        visual.pop("visualContainerObjects", None)

    # Rewrite the query ONLY when the bindings actually changed — otherwise the
    # original query block (sortDefinition and friends) is preserved verbatim.
    if (sv.get("projections") != base_sv.get("projections")
            or sv.get("prototypeQuery") != base_sv.get("prototypeQuery")):
        state = _pbir_query_state_from_single_visual(sv)
        if state:
            query = dict(visual.get("query") or {})
            query["queryState"] = state
            # A rewritten query would otherwise lose the sort entirely, which
            # is how an added or duplicated visual came out unsorted.
            sort = _pbir_sort_definition(sv)
            if sort is not None:
                query["sortDefinition"] = sort
            elif sv.get("__pbir_sort__") is not None:
                query["sortDefinition"] = sv["__pbir_sort__"]
            visual["query"] = query
        else:
            visual.pop("query", None)
    elif sv.get("__pbir_sort__") is not None and "query" in visual:
        visual["query"].setdefault("sortDefinition", sv["__pbir_sort__"])
    # Some containers (groups/shapes) legitimately have NO visual block —
    # don't invent an empty one.
    if visual or "visual" in out:
        out["visual"] = visual

    if container.get("isHidden"):
        out["isHidden"] = True
    else:
        out.pop("isHidden", None)

    filters = container.get("filters")
    if filters:
        parsed = json.loads(filters) if isinstance(filters, str) else filters
        if parsed:
            fc = dict(out.get("filterConfig") or {})
            fc["filters"] = parsed
            out["filterConfig"] = fc
        else:
            out.pop("filterConfig", None)
    elif "filterConfig" in out and "filters" not in container:
        pass  # untouched by the caller — keep whatever the file had
    return out


def _pbir_patch_page(orig: dict, section: dict, pid: str) -> dict:
    """Apply a legacy section onto its ORIGINAL PBIR page.json."""
    out = copy.deepcopy(orig) if orig else {}
    out["name"] = section.get("name") or out.get("name") or pid
    if section.get("displayName"):
        out["displayName"] = section["displayName"]
    for key in ("width", "height", "displayOption", "type"):
        if key in section:
            out[key] = section[key]
    # PBIR requires the enum name; classic-shaped callers (and pbix_add_page)
    # supply the int. Unknown ints fall back to the format default rather than
    # writing a value the schema rejects.
    dopt = out.get("displayOption")
    if isinstance(dopt, bool) or isinstance(dopt, int):
        out["displayOption"] = _PBIR_DISPLAY_OPTION.get(int(dopt), "FitToPage")
    elif dopt is None:
        out["displayOption"] = "FitToPage"
    # Inverse of the reader: classic `config.visibility` int -> PBIR enum name.
    raw_cfg = section.get("config")
    if raw_cfg is not None:
        try:
            cfg = (json.loads(raw_cfg) if isinstance(raw_cfg, str)
                   else raw_cfg) or {}
        except (json.JSONDecodeError, TypeError):
            cfg = {}
        if isinstance(cfg, dict) and "visibility" in cfg:
            out["visibility"] = ("HiddenInViewMode" if cfg["visibility"]
                                 else "AlwaysVisible")
    if section.get("objects"):
        out["objects"] = section["objects"]
    filters = section.get("filters")
    if filters:
        parsed = json.loads(filters) if isinstance(filters, str) else filters
        if parsed:
            fc = dict(out.get("filterConfig") or {})
            fc["filters"] = parsed
            out["filterConfig"] = fc
    if not orig:
        # Only a brand-new page gets a $schema stamped; patching an existing
        # file must not invent fields it didn't have.
        out.setdefault(
            "$schema",
            "https://developer.microsoft.com/json-schemas/fabric/item/report/"
            "definition/page/2.1.0/schema.json")
    return out


def _set_layout_pbir(work_dir: str, layout: dict) -> None:
    """Write a legacy-shaped layout back into the PBIR Report/definition tree.

    Rather than regenerating the tree from the (lossy) legacy view, each page
    and visual is patched onto the ORIGINAL file it was read from, so fields
    this converter doesn't model survive an edit. Pages/visuals the caller
    added are created; ones it removed are deleted; pages.json is rewritten to
    match the new order.
    """
    import shutil

    pages_dir = os.path.join(work_dir, "Report", "definition", "pages")
    pages_json = os.path.join(pages_dir, "pages.json")
    meta = {}
    had_meta = os.path.exists(pages_json)
    if had_meta:
        try:
            with open(pages_json, "r", encoding="utf-8") as f:
                meta = json.load(f)
        except Exception:
            meta = {}
    if not isinstance(meta, dict):
        meta = {}

    sections = layout.get("sections", []) or []
    keep_pages = []
    for idx, section in enumerate(sections):
        pid = section.get("__pbir_page__") or _sanitize_pbir_name(
            section.get("name") or section.get("displayName") or f"page{idx}")
        keep_pages.append(pid)
        page_dir = os.path.join(pages_dir, pid)
        page_file = os.path.join(page_dir, "page.json")
        orig_page = {}
        if os.path.exists(page_file):
            try:
                with open(page_file, "r", encoding="utf-8") as f:
                    orig_page = json.load(f)
            except Exception:
                orig_page = {}
        _pbir_write_json(page_file, _pbir_patch_page(orig_page, section, pid))

        visuals_dir = os.path.join(page_dir, "visuals")
        keep_visuals = []
        for vidx, container in enumerate(section.get("visualContainers", []) or []):
            cfg = _parse_visual_config(container)
            vid = (container.get("__pbir_visual__") or cfg.get("name")
                   or f"visual{vidx}")
            vid = _sanitize_pbir_name(str(vid))
            keep_visuals.append(vid)
            vfile = os.path.join(visuals_dir, vid, "visual.json")
            orig_v = {}
            if os.path.exists(vfile):
                try:
                    with open(vfile, "r", encoding="utf-8") as f:
                        orig_v = json.load(f)
                except Exception:
                    orig_v = {}
            patched = _pbir_patch_visual(orig_v, container)
            if not orig_v:
                patched.setdefault(
                    "$schema",
                    "https://developer.microsoft.com/json-schemas/fabric/item/"
                    "report/definition/visualContainer/2.11.0/schema.json")
            _pbir_write_json(vfile, patched)
            mobile = container.get("mobile")
            if mobile:
                _pbir_write_json(
                    os.path.join(visuals_dir, vid, "mobile.json"), mobile)

        # drop visuals the caller removed
        if os.path.isdir(visuals_dir):
            for existing in os.listdir(visuals_dir):
                if existing not in keep_visuals and os.path.isdir(
                        os.path.join(visuals_dir, existing)):
                    shutil.rmtree(os.path.join(visuals_dir, existing),
                                  ignore_errors=True)

    # drop pages the caller removed
    if os.path.isdir(pages_dir):
        for existing in os.listdir(pages_dir):
            full = os.path.join(pages_dir, existing)
            if os.path.isdir(full) and existing not in keep_pages:
                shutil.rmtree(full, ignore_errors=True)

    meta["pageOrder"] = keep_pages
    active = meta.get("activePageName")
    explicit_active = next(
        (s.get("name") for s in sections if s.get("isActive")), None)
    if explicit_active:
        meta["activePageName"] = explicit_active
    elif active not in keep_pages and keep_pages:
        meta["activePageName"] = keep_pages[0]
    if not had_meta:
        meta.setdefault(
            "$schema",
            "https://developer.microsoft.com/json-schemas/fabric/item/report/"
            "definition/pagesMetadata/1.1.0/schema.json")
    _pbir_write_json(pages_json, meta)

    # Bookmarks live outside the pages tree, in the layout-level `config`
    # string that classic callers edit. Only touch the folder when the caller
    # actually supplied a config, so a page-only edit can't delete bookmarks.
    raw_config = layout.get("config")
    cfg: dict = {}
    if raw_config is not None:
        try:
            cfg = (json.loads(raw_config) if isinstance(raw_config, str)
                   else raw_config) or {}
        except (json.JSONDecodeError, TypeError):
            cfg = {}
        if isinstance(cfg, dict) and "bookmarks" in cfg:
            _pbir_write_bookmarks(work_dir, cfg.get("bookmarks") or [])

    # Report-level state back into report.json, in PBIR shape. Only keys the
    # caller actually supplied are touched, so a page-only edit cannot wipe the
    # report's resources or theme.
    rpath = _report_config_path(work_dir)
    rcfg: dict = {}
    if os.path.exists(rpath):
        try:
            with open(rpath, "r", encoding="utf-8-sig") as f:
                rcfg = json.load(f)
        except Exception:
            rcfg = {}
    before = json.dumps(rcfg, sort_keys=True)

    if "resourcePackages" in layout:
        rcfg["resourcePackages"] = _classic_packages_to_pbir(
            layout["resourcePackages"], rcfg.get("resourcePackages"))
    if "publicCustomVisuals" in layout:
        pcv = list(layout["publicCustomVisuals"] or [])
        if pcv:
            rcfg["publicCustomVisuals"] = pcv
        else:
            rcfg.pop("publicCustomVisuals", None)
    if isinstance(cfg, dict) and "themeCollection" in cfg:
        rcfg["themeCollection"] = _classic_theme_to_pbir(
            cfg["themeCollection"], rcfg.get("themeCollection"))
    if "filters" in layout:
        raw = layout["filters"]
        try:
            parsed = json.loads(raw) if isinstance(raw, str) else raw
        except (json.JSONDecodeError, TypeError):
            parsed = None
        if parsed:
            fc = dict(rcfg.get("filterConfig") or {})
            fc["filters"] = parsed
            rcfg["filterConfig"] = fc
        else:
            rcfg.pop("filterConfig", None)
    if isinstance(cfg, dict) and "settings" in cfg:
        rcfg["settings"] = cfg["settings"]
    if isinstance(cfg, dict) and "objects" in cfg:
        rcfg["objects"] = cfg["objects"]

    if json.dumps(rcfg, sort_keys=True) != before:
        _pbir_write_json(rpath, rcfg)



def _extract_default_filters_dict(work_dir: str, page_index: int = 0, layout: dict | None = None) -> dict:
    """Internal: extract default slicer filters as a dict for programmatic use.

    Handles both In-type (value list) and Comparison-type (equality/range) filters.
    Returns { 'Entity.Property': [values] } suitable for use as filter_context.

    ``layout`` may be a pre-parsed layout dict to avoid re-reading/re-parsing the
    (potentially multi-MB) report layout once per page.
    """
    if layout is None:
        layout = _get_layout(work_dir)   # handles PBIR too
    if not layout:
        return {}

    sections = layout.get("sections", [])
    if page_index < 0 or page_index >= len(sections):
        return {}

    page = sections[page_index]
    containers = page.get("visualContainers", [])
    filters = {}

    import re as _re

    def _parse_literal(lit):
        """Parse a literal value from filter JSON."""
        if lit is None:
            return None
        s = str(lit)
        # Numeric literals (possibly suffixed with D/L for double/long)
        num_match = _re.match(r'^(-?\d+(?:\.\d+)?)[DL]?$', s, _re.IGNORECASE)
        if num_match:
            return float(num_match.group(1)) if '.' in num_match.group(1) else int(num_match.group(1))
        # Datetime literals: datetime'2024-01-01T00:00:00'
        dt_match = _re.match(r"^datetime'([^']+)'$", s, _re.IGNORECASE)
        if dt_match:
            return dt_match.group(1)  # Return the ISO datetime string
        # Power BI escapes single quotes as '' in filter JSON —
        # normalize to single quotes to match actual data values
        s = s.replace("''", "'")
        if s.startswith("'") and s.endswith("'"):
            s = s[1:-1]
        return s

    def _resolve_column(col_expr, from_entries):
        """Resolve Entity.Property from a column expression and From entries."""
        source = col_expr.get("Expression", {}).get("SourceRef", {}).get("Source")
        prop = col_expr.get("Property")
        from_entry = next((f for f in from_entries if f.get("Name") == source), {})
        entity = from_entry.get("Entity")
        if entity and prop:
            return f"{entity}.{prop}"
        return None

    for vc in containers:
        config = _parse_visual_config(vc)
        sv = config.get("singleVisual", {})

        # Check for filter in objects.general
        general_arr = sv.get("objects", {}).get("general", [])
        for gen in general_arr:
            filter_obj = gen.get("properties", {}).get("filter", {}).get("filter", {})
            if not filter_obj or not filter_obj.get("Where"):
                continue

            from_entries = filter_obj.get("From", [])

            for where in filter_obj["Where"]:
                cond = where.get("Condition", {})

                # --- In-type: value list filters ---
                if "In" in cond:
                    expr = cond["In"].get("Expressions", [{}])[0]
                    values = cond["In"].get("Values", [])
                    col_expr = expr.get("Column", {})
                    key = _resolve_column(col_expr, from_entries)

                    if key and values:
                        vals = []
                        for v in values:
                            lit = v[0].get("Literal", {}).get("Value") if v else None
                            parsed = _parse_literal(lit)
                            if parsed is not None:
                                vals.append(parsed)
                        if vals:
                            filters[key] = vals

                # --- Comparison-type: equality / range filters ---
                if "Comparison" in cond:
                    comp = cond["Comparison"]
                    kind = comp.get("ComparisonKind", 0)  # 0=Equal, 1=GT, 2=GTE, 3=LT, 4=LTE
                    left = comp.get("Left", {})
                    right = comp.get("Right", {})

                    # Left side should be a column reference
                    col_expr = left.get("Column", {})
                    key = _resolve_column(col_expr, from_entries)

                    # Right side should be a literal value
                    lit = right.get("Literal", {}).get("Value")
                    parsed = _parse_literal(lit)

                    if key and parsed is not None:
                        if kind == 0:
                            # Equality: single value filter
                            filters[key] = [parsed]
                        else:
                            # Range filter (GT/GTE/LT/LTE) — store as single value
                            # for SELECTEDVALUE to work on numeric slicers
                            filters[key] = [parsed]

    return filters


def _get_all_default_filters(work_dir: str) -> dict:
    """Get default filters merged across all pages."""
    layout = _get_layout(work_dir)   # handles PBIR too
    if not layout:
        return {}

    all_filters = {}
    sections = layout.get("sections", [])
    for i in range(len(sections)):
        # Pass the already-parsed layout so each page doesn't re-open/re-parse it.
        page_filters = _extract_default_filters_dict(work_dir, i, layout=layout)
        # Merge — later pages don't overwrite earlier ones
        for k, v in page_filters.items():
            if k not in all_filters:
                all_filters[k] = v
    return all_filters


def _layout_stamp(work_dir: str | None):
    """Return a cheap change-stamp for the report layout files: a tuple of
    (mtime, size) per candidate layout file, or None if none can be stat'd.
    Used to detect layout/slicer edits without a full re-parse. Including size
    alongside mtime guards against two same-size edits landing in one mtime tick.
    """
    if not work_dir:
        return None
    parts = []
    for rel in (
        os.path.join("Report", "Layout"),
        os.path.join("Report", "definition.pbir"),
        os.path.join("definition", "report.json"),
    ):
        try:
            st = os.stat(os.path.join(work_dir, rel))
        except OSError:
            continue
        parts.append((rel, st.st_mtime_ns, st.st_size))
    if not parts:
        return None
    # See _layout_writes: mtime+size cannot distinguish two same-length edits
    # inside one timestamp tick, which is precisely what changing a slicer's
    # selected value looks like.
    return (_layout_writes.get(work_dir, 0), tuple(parts))


def _get_default_filters_current(ctx: dict) -> dict:
    """Return the report's current default slicer filters, re-deriving from the
    layout only when it has changed on disk since the last derivation.

    The DAX context caches an open-time snapshot, but slicer/filter edits
    (pbix_set_filters, set_layout_raw, …) rewrite the layout without rebuilding
    the context, which would otherwise leak a stale selection (OpenBI #7). We key
    a cached derivation on a cheap layout change-stamp so the steady state is a
    single stat() instead of a full multi-MB layout re-parse on every evaluation.
    """
    work_dir = ctx.get("work_dir")
    stamp = _layout_stamp(work_dir)
    if (
        stamp is not None
        and ctx.get("_default_filters_stamp") == stamp
        and "default_filters" in ctx
    ):
        return ctx["default_filters"] or {}
    df = _get_all_default_filters(work_dir) if work_dir else {}
    ctx["default_filters"] = df
    ctx["_default_filters_stamp"] = stamp
    return df


@mcp.tool()
def pbix_get_default_filters(alias: str, page_index: int = 0) -> str:
    """Extract default slicer filter selections from a report page.

    Reads the filter config from slicer visuals (advancedSlicerVisual, slicer)
    to determine what the dashboard's default filtered state is.
    Supports both In-type (value list) and Comparison-type (equality/range) filters.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index (default 0)
    """
    try:
        info = _ensure_open(alias)
        filters = _extract_default_filters_dict(info["work_dir"], page_index)

        if not filters:
            return ToolResponse.ok(
                "No default slicer filters found on this page.",
                data={"filters": {}},
            ).to_text()

        lines = ["Default slicer filters:\n"]
        for key, vals in filters.items():
            lines.append(f"  {key}: {vals}")
        lines.append("\nUse as filter_context in pbix_evaluate_dax:")
        lines.append(f"  {json.dumps(filters)}")
        # Wrap in the standard envelope like every other tool, and expose the
        # parsed filters as structured data so programmatic consumers do not
        # have to scrape the rendered message.
        return ToolResponse.ok("\n".join(lines), data={"filters": filters}).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(f"{str(e)}\n{traceback.format_exc()}")


@mcp.tool()
def pbix_get_visual_positions(alias: str, page_index: int = 0) -> str:
    """Get all visual positions with parent group offset resolution.

    For visuals inside groups, the raw x/y coordinates are relative to the
    parent group. This tool resolves them to absolute page coordinates.

    Args:
        alias: The alias of the open file
        page_index: Zero-based page index
    """
    try:
        info = _ensure_open(alias)
        layout = _get_layout(info["work_dir"])
        if not layout:
            raise LayoutParseError("No layout found")

        sections = layout.get("sections", [])
        if page_index < 0 or page_index >= len(sections):
            raise LayoutParseError(f"Page index {page_index} out of range")

        page = sections[page_index]
        containers = page.get("visualContainers", [])

        # Pass 1: build group positions map
        group_positions = {}
        for vc in containers:
            config = _parse_visual_config(vc)
            name = config.get("name", "")
            if config.get("singleVisualGroup"):
                group_positions[name] = {"x": vc.get("x", 0), "y": vc.get("y", 0)}

        # Pass 2: resolve absolute positions
        lines = [f"Visual positions (absolute, {len(containers)} visuals):\n"]
        for i, vc in enumerate(containers):
            config = _parse_visual_config(vc)
            vtype = _get_visual_type(config)
            x = vc.get("x", 0)
            y = vc.get("y", 0)
            w = vc.get("width", 0)
            h = vc.get("height", 0)

            parent_group = config.get("parentGroupName")
            if parent_group and parent_group in group_positions:
                x += group_positions[parent_group]["x"]
                y += group_positions[parent_group]["y"]
                lines.append(f"  [{i}] {vtype:<30s} at ({x:.0f},{y:.0f}) {w:.0f}x{h:.0f}  [child of group]")
            else:
                lines.append(f"  [{i}] {vtype:<30s} at ({x:.0f},{y:.0f}) {w:.0f}x{h:.0f}")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise LayoutParseError(f"{str(e)}\n{traceback.format_exc()}")


@mcp.tool()
def pbix_clear_dax_cache(alias: str = "") -> str:
    """Clear the DAX engine data cache.

    Call this after modifying measures or table data to force fresh evaluation.

    Args:
        alias: Clear cache for specific alias, or all if empty
    """
    global _dax_cache
    if alias:
        _dax_cache.pop(alias, None)
        return ToolResponse.ok(f"DAX cache cleared for '{alias}'").to_text()
    else:
        _dax_cache.clear()
        return ToolResponse.ok("DAX cache cleared for all files").to_text()


# ---- Section 10: Calculated Columns ----

@mcp.tool()
def pbix_evaluate_calculated_columns(alias: str) -> str:
    """Evaluate all calculated columns in the data model.

    Finds columns with DAX expressions in the metadata, evaluates them
    per-row against actual table data, and adds the results to the
    cached data context. This is useful when calculated columns were
    defined but their values aren't materialized in VertiPaq.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)

        # Force re-build of DAX context with calculated columns
        global _dax_cache
        _dax_cache.pop(alias, None)
        ctx = _get_dax_context(alias)

        # Check if any calculated columns were evaluated
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm = f.read()
        abf = decompress_datamodel(dm)
        db_bytes = read_metadata_sqlite(abf)

        import sqlite3
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(db_bytes)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            calc_cols = conn.execute("""
                SELECT c.ExplicitName, c.Expression, t.Name
                FROM [Column] c JOIN [Table] t ON c.TableID = t.ID
                WHERE c.Expression IS NOT NULL AND c.Expression != ''
                  AND c.ExplicitName IS NOT NULL
                  AND c.ExplicitName NOT LIKE 'RowNumber%'
                  AND t.ModelID = 1
            """).fetchall()
            conn.close()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        if not calc_cols:
            return ToolResponse.ok("No calculated columns found in the data model.").to_text()

        lines = [f"Calculated columns ({len(calc_cols)}):\n"]
        for cc in calc_cols:
            tname = cc[2]
            cname = cc[0]
            expr = cc[1][:60].strip().replace('\n', ' ')
            tbl = ctx['tables'].get(tname)
            if tbl and cname in tbl['columns']:
                lines.append(f"  ✅ {tname}[{cname}] = {expr}...")
            else:
                lines.append(f"  ⚠ {tname}[{cname}] = {expr}... (not evaluated)")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 11: RLS (Row-Level Security) ----

@mcp.tool()
def pbix_get_rls_roles(alias: str) -> str:
    """Get all Row-Level Security roles and their table filter expressions.

    Returns role definitions and the DAX filter expressions that define
    what data each role can see.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm = f.read()
        abf = decompress_datamodel(dm)
        db_bytes = read_metadata_sqlite(abf)

        import sqlite3
        fd, tmp_path = tempfile.mkstemp(suffix=".db")
        os.write(fd, db_bytes)
        os.close(fd)
        conn = None
        try:
            conn = sqlite3.connect(tmp_path)
            conn.row_factory = sqlite3.Row

            roles = conn.execute("SELECT * FROM [Role]").fetchall()
            if not roles:
                conn.close()
                conn = None
                return ToolResponse.ok("No RLS roles defined in this file.").to_text()

            lines = [f"RLS Roles ({len(roles)}):\n"]
            for role in roles:
                role_id = role["ID"]
                role_name = role["Name"] if "Name" in role.keys() else f"Role {role_id}"
                lines.append(f"  Role: {role_name} (ID={role_id})")

                perms = conn.execute(
                    "SELECT * FROM [TablePermission] WHERE RoleID = ?",
                    (role_id,)
                ).fetchall()
                for perm in perms:
                    table_id = perm["TableID"]
                    perm_keys = perm.keys()
                    filter_expr = perm["FilterExpression"] if "FilterExpression" in perm_keys else (perm["QueryExpression"] if "QueryExpression" in perm_keys else "")
                    trow = conn.execute(
                        "SELECT Name FROM [Table] WHERE ID = ?", (table_id,)
                    ).fetchone()
                    tname = trow["Name"] if trow else f"Table {table_id}"
                    lines.append(f"    {tname}: {filter_expr}")

                members = conn.execute(
                    "SELECT * FROM [RoleMembership] WHERE RoleID = ?",
                    (role_id,)
                ).fetchall()
                if members:
                    lines.append(f"    Members: {len(members)}")

            conn.close()
            conn = None
            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            if conn:
                conn.close()
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_set_rls_role(
    alias: str,
    role_name: str,
    table_name: str,
    filter_expression: str,
    description: str = "",
) -> str:
    """Create or update a Row-Level Security role with a DAX filter expression.

    The filter expression is a DAX boolean expression that determines which
    rows are visible to the role. For example:
      'dim-Geo'[Country] = "USA"
      'Sales'[Amount] > 1000

    Args:
        alias: The alias of the open file
        role_name: Name of the RLS role (e.g., "US Sales Only")
        table_name: Table to apply the filter to
        filter_expression: DAX filter expression (e.g., 'Sales'[Region] = "West")
        description: Optional role description
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found", DataModelCompressionError.code).to_text()

        def _do_set(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            role = conn.execute(
                "SELECT ID FROM [Role] WHERE Name = ?", (role_name,)
            ).fetchone()
            c = conn.cursor()
            # Get MAXID for safe ID allocation
            maxid_row = c.execute(
                "SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'"
            ).fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0

            if role:
                role_id = role["ID"]
            else:
                max_id += 1
                role_id = max_id
                c.execute(
                    "INSERT INTO [Role] (ID, ModelID, Name, Description) VALUES (?, 1, ?, ?)",
                    (role_id, role_name, description),
                )

            table_row = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (table_name,)
            ).fetchone()
            if not table_row:
                raise ValueError(f"Table '{table_name}' not found")
            table_id = table_row["ID"]

            existing = c.execute(
                "SELECT ID FROM [TablePermission] WHERE RoleID = ? AND TableID = ?",
                (role_id, table_id)
            ).fetchone()
            if existing:
                c.execute(
                    "UPDATE [TablePermission] SET FilterExpression = ? WHERE ID = ?",
                    (filter_expression, existing["ID"]),
                )
            else:
                max_id += 1
                c.execute(
                    "INSERT INTO [TablePermission] (ID, RoleID, TableID, FilterExpression) VALUES (?, ?, ?, ?)",
                    (max_id, role_id, table_id, filter_expression),
                )

            # Update MAXID
            c.execute(
                "UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'",
                (str(max_id),)
            )
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_set)
        info["modified"] = True
        return ToolResponse.ok(f"RLS role '{role_name}' set on '{table_name}' with filter: {filter_expression}").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_evaluate_rls(
    alias: str,
    role_name: str,
    table_name: str,
    max_rows: int = 10,
) -> str:
    """Evaluate an RLS role's filter and show which rows would be visible.

    Uses the DAX engine to evaluate the role's filter expression against
    actual table data.

    Args:
        alias: The alias of the open file
        role_name: Name of the RLS role to evaluate
        table_name: Table to check visibility for
        max_rows: Maximum rows to show (default 10)
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found", DataModelCompressionError.code).to_text()

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm = f.read()
        abf = decompress_datamodel(dm)
        db_bytes = read_metadata_sqlite(abf)

        import sqlite3
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(db_bytes)
        tmp.close()
        try:
            conn = sqlite3.connect(tmp.name)
            conn.row_factory = sqlite3.Row

            role = conn.execute("SELECT ID FROM [Role] WHERE Name = ?", (role_name,)).fetchone()
            if not role:
                conn.close()
                return ToolResponse.error(f"Role '{role_name}' not found", PBIXMCPError.code).to_text()

            table_row = conn.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            if not table_row:
                conn.close()
                return ToolResponse.error(f"Table '{table_name}' not found", PBIXMCPError.code).to_text()

            perm = conn.execute(
                "SELECT FilterExpression FROM [TablePermission] WHERE RoleID = ? AND TableID = ?",
                (role["ID"], table_row["ID"])
            ).fetchone()
            conn.close()

            if not perm or not perm["FilterExpression"]:
                return ToolResponse.ok(f"Role '{role_name}' has no filter on table '{table_name}' — all rows visible.").to_text()

            filter_expr = perm["FilterExpression"]

            # Load table data and evaluate filter
            ctx = _get_dax_context(alias)
            tbl = ctx['tables'].get(table_name)
            if not tbl:
                return ToolResponse.error(f"Table '{table_name}' has no data", PBIXMCPError.code).to_text()

            from pbix_mcp.dax import engine as dax_engine
            eng = dax_engine.DAXEngine()

            total = len(tbl['rows'])
            visible = 0
            sample_rows: list = []

            for row in tbl['rows']:
                # Build row context
                row_expr = filter_expr
                for ci, cn in enumerate(tbl['columns']):
                    val = row[ci]
                    for pat in [f"'{table_name}'[{cn}]", f"{table_name}[{cn}]"]:
                        if pat in row_expr:
                            if isinstance(val, str):
                                row_expr = row_expr.replace(pat, f'"{val}"')
                            elif val is None:
                                row_expr = row_expr.replace(pat, 'BLANK()')
                            else:
                                row_expr = row_expr.replace(pat, str(val))

                eval_ctx = dax_engine.DAXContext(
                    ctx['tables'], ctx['measure_defs'], None, None, None,
                    ctx.get('relationships', [])
                )
                result = eng._eval_expr(row_expr, eval_ctx)
                if result is True or result == 1:
                    visible += 1
                    if len(sample_rows) < max_rows:
                        sample_rows.append({tbl['columns'][i]: row[i] for i in range(min(5, len(tbl['columns'])))})

            lines = [
                f"RLS evaluation for role '{role_name}' on '{table_name}':",
                f"  Filter: {filter_expr}",
                f"  Visible: {visible}/{total} rows ({visible/total*100:.1f}%)\n",
            ]
            if sample_rows:
                lines.append(f"  Sample visible rows (first {len(sample_rows)}):")
                for sr in sample_rows:
                    lines.append(f"    {sr}")

            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 10b: Perspectives ----


def _read_metadata_db(alias: str):
    """Helper: decompress DataModel and return a temp SQLite connection + path.

    Caller MUST close conn and os.unlink(tmp_path) when done.
    """
    info = _ensure_open(alias)
    dm_path = os.path.join(info["work_dir"], "DataModel")
    if not os.path.exists(dm_path):
        raise PBIXMCPError("No DataModel found", DataModelCompressionError.code)

    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    with open(dm_path, "rb") as f:
        dm = f.read()
    abf = decompress_datamodel(dm)
    db_bytes = read_metadata_sqlite(abf)

    fd, tmp_path = tempfile.mkstemp(suffix=".db")
    os.write(fd, db_bytes)
    os.close(fd)
    conn = sqlite3.connect(tmp_path)
    conn.row_factory = sqlite3.Row
    return info, conn, tmp_path


@mcp.tool()
def pbix_get_perspectives(alias: str) -> str:
    """Get all perspectives with their included tables, columns, and measures.

    Args:
        alias: The alias of the open file
    """
    try:
        info, conn, tmp_path = _read_metadata_db(alias)
        try:
            perspectives = conn.execute("SELECT ID, Name, Description FROM Perspective ORDER BY ID").fetchall()
            if not perspectives:
                return ToolResponse.ok("No perspectives defined in this file.").to_text()

            lines = [f"Perspectives ({len(perspectives)}):\n"]
            for p in perspectives:
                pid, pname, pdesc = p["ID"], p["Name"], p["Description"]
                lines.append(f"  {pname}" + (f" — {pdesc}" if pdesc else ""))

                ptables = conn.execute(
                    "SELECT pt.ID, pt.IncludeAll, t.Name FROM PerspectiveTable pt "
                    "JOIN [Table] t ON pt.TableID = t.ID "
                    "WHERE pt.PerspectiveID = ? ORDER BY t.Name", (pid,)
                ).fetchall()
                for pt in ptables:
                    ptid, include_all, tname = pt["ID"], pt["IncludeAll"], pt["Name"]
                    if include_all:
                        lines.append(f"    {tname} (all columns/measures)")
                    else:
                        cols = conn.execute(
                            # COALESCE: a perspective can include a
                            # calculated-table column, whose name lives in
                            # InferredName.
                            "SELECT COALESCE(c.ExplicitName, c.InferredName) "
                            "       AS ExplicitName FROM PerspectiveColumn pc "
                            "JOIN [Column] c ON pc.ColumnID = c.ID "
                            "WHERE pc.PerspectiveTableID = ? "
                            "ORDER BY 1", (ptid,)
                        ).fetchall()
                        measures = conn.execute(
                            "SELECT m.Name FROM PerspectiveMeasure pm "
                            "JOIN Measure m ON pm.MeasureID = m.ID "
                            "WHERE pm.PerspectiveTableID = ? ORDER BY m.Name", (ptid,)
                        ).fetchall()
                        col_names = [c["ExplicitName"] for c in cols]
                        meas_names = [m["Name"] for m in measures]
                        items = col_names + [f"[M] {m}" for m in meas_names]
                        lines.append(f"    {tname}: {', '.join(items) if items else '(no specific items)'}")
                lines.append("")

            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            conn.close()
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_perspective(
    alias: str,
    name: str,
    tables_json: str = "[]",
    description: str = "",
) -> str:
    """Add a perspective — a filtered view of the model for different user groups.

    Args:
        alias: The alias of the open file
        name: Name for the perspective (e.g. "Sales Analyst", "Executive View")
        tables_json: JSON array of tables to include, e.g.
            '[{"table": "Sales"}, {"table": "Product", "columns": ["Name", "Category"]}]'
            If columns/measures are omitted for a table, all are included.
            Optional per-table fields: "columns" (list), "measures" (list)
        description: Optional description
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        tables_spec = json.loads(tables_json) if tables_json else []

        def _do_add(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            # Check if perspective already exists
            existing = c.execute("SELECT ID FROM Perspective WHERE Name = ?", (name,)).fetchone()
            if existing:
                raise PBIXMCPError(f"Perspective '{name}' already exists", "DUPLICATE")

            # Get MAXID
            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0

            # Create Perspective
            max_id += 1
            persp_id = max_id
            c.execute(
                "INSERT INTO Perspective (ID, ModelID, Name, Description, ModifiedTime) "
                "VALUES (?, 1, ?, ?, ?)",
                (persp_id, name, description or None, int(datetime.now().timestamp() * 1e7)),
            )

            # Add tables
            for tspec in tables_spec:
                tname = tspec.get("table", "")
                trow = c.execute(
                    "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (tname,)
                ).fetchone()
                if not trow:
                    raise PBIXMCPError(f"Table '{tname}' not found", "TABLE_NOT_FOUND")
                table_id = trow["ID"]

                specific_cols = tspec.get("columns", [])
                specific_meas = tspec.get("measures", [])
                include_all = 1 if (not specific_cols and not specific_meas) else 0

                max_id += 1
                pt_id = max_id
                c.execute(
                    "INSERT INTO PerspectiveTable (ID, PerspectiveID, TableID, IncludeAll, ModifiedTime) "
                    "VALUES (?, ?, ?, ?, ?)",
                    (pt_id, persp_id, table_id, include_all, int(datetime.now().timestamp() * 1e7)),
                )

                if specific_cols:
                    for col_name in specific_cols:
                        crow = c.execute(
                            "SELECT ID FROM [Column] WHERE TableID = ? AND (ExplicitName = ? OR InferredName = ?)",
                            (table_id, col_name, col_name),
                        ).fetchone()
                        if crow:
                            max_id += 1
                            c.execute(
                                "INSERT INTO PerspectiveColumn (ID, PerspectiveTableID, ColumnID, ModifiedTime) "
                                "VALUES (?, ?, ?, ?)",
                                (max_id, pt_id, crow["ID"], int(datetime.now().timestamp() * 1e7)),
                            )

                if specific_meas:
                    for meas_name in specific_meas:
                        mrow = c.execute(
                            "SELECT ID FROM Measure WHERE TableID = ? AND Name = ?",
                            (table_id, meas_name),
                        ).fetchone()
                        if mrow:
                            max_id += 1
                            c.execute(
                                "INSERT INTO PerspectiveMeasure (ID, PerspectiveTableID, MeasureID, ModifiedTime) "
                                "VALUES (?, ?, ?, ?)",
                                (max_id, pt_id, mrow["ID"], int(datetime.now().timestamp() * 1e7)),
                            )

            # Update MAXID
            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        n_tables = len(tables_spec)
        return ToolResponse.ok(
            f"Perspective '{name}' created with {n_tables} table(s)."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_remove_perspective(alias: str, name: str) -> str:
    """Remove a perspective and all its included table/column/measure references.

    Args:
        alias: The alias of the open file
        name: Name of the perspective to remove
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        def _do_remove(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            row = c.execute("SELECT ID FROM Perspective WHERE Name = ?", (name,)).fetchone()
            if not row:
                raise PBIXMCPError(f"Perspective '{name}' not found", "NOT_FOUND")
            pid = row["ID"]

            # Get PerspectiveTable IDs for cascade delete
            pt_ids = [r["ID"] for r in c.execute(
                "SELECT ID FROM PerspectiveTable WHERE PerspectiveID = ?", (pid,)
            ).fetchall()]

            for pt_id in pt_ids:
                c.execute("DELETE FROM PerspectiveColumn WHERE PerspectiveTableID = ?", (pt_id,))
                c.execute("DELETE FROM PerspectiveMeasure WHERE PerspectiveTableID = ?", (pt_id,))
                c.execute("DELETE FROM PerspectiveHierarchy WHERE PerspectiveTableID = ?", (pt_id,))

            c.execute("DELETE FROM PerspectiveTable WHERE PerspectiveID = ?", (pid,))
            c.execute("DELETE FROM Perspective WHERE ID = ?", (pid,))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_remove)
        info["modified"] = True
        return ToolResponse.ok(f"Perspective '{name}' removed.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 10c: User Hierarchies ----


@mcp.tool()
def pbix_get_hierarchies(alias: str) -> str:
    """Get all user hierarchies with their levels and columns.

    Args:
        alias: The alias of the open file
    """
    try:
        info, conn, tmp_path = _read_metadata_db(alias)
        try:
            hierarchies = conn.execute(
                "SELECT h.ID, h.Name, h.IsHidden, h.Description, t.Name as TableName "
                "FROM Hierarchy h JOIN [Table] t ON h.TableID = t.ID "
                "ORDER BY t.Name, h.Name"
            ).fetchall()
            if not hierarchies:
                return ToolResponse.ok("No user hierarchies defined in this file.").to_text()

            lines = [f"Hierarchies ({len(hierarchies)}):\n"]
            for h in hierarchies:
                hid = h["ID"]
                hidden = " (hidden)" if h["IsHidden"] else ""
                desc = f" — {h['Description']}" if h["Description"] else ""
                lines.append(f"  {h['TableName']}.{h['Name']}{hidden}{desc}")

                levels = conn.execute(
                    # COALESCE, for the same reason the rebuild path already
                    # does (see _rebuild_datamodel's Level query): a
                    # calculated-table or auto-date column keeps its name in
                    # InferredName and leaves ExplicitName NULL, so reading only
                    # ExplicitName reported every auto-date hierarchy level as
                    # pointing at no column at all.
                    "SELECT l.Ordinal, l.Name, "
                    "       COALESCE(c.ExplicitName, c.InferredName) AS ColumnName "
                    "FROM Level l LEFT JOIN [Column] c ON l.ColumnID = c.ID "
                    "WHERE l.HierarchyID = ? ORDER BY l.Ordinal", (hid,)
                ).fetchall()
                for lv in levels:
                    lines.append(f"    {lv['Ordinal']}: {lv['Name']} → {lv['ColumnName']}")
                lines.append("")

            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            conn.close()
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_hierarchy(
    alias: str,
    table_name: str,
    hierarchy_name: str,
    levels_json: str,
) -> str:
    """Create a user hierarchy (drill-down path) on a table.

    Args:
        alias: The alias of the open file
        table_name: Table to add the hierarchy to
        hierarchy_name: Name for the hierarchy (e.g. "Geography", "Date Hierarchy")
        levels_json: JSON array of levels in drill-down order, e.g.
            '[{"name": "Country", "column": "Country"},
              {"name": "State", "column": "State-Province"},
              {"name": "City", "column": "City"}]'
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        levels = json.loads(levels_json)
        if not levels:
            return ToolResponse.error("levels_json must contain at least one level", "INVALID_INPUT").to_text()

        def _do_add(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            if not trow:
                raise PBIXMCPError(f"Table '{table_name}' not found", "TABLE_NOT_FOUND")
            table_id = trow["ID"]

            # Check duplicate
            existing = c.execute(
                "SELECT ID FROM Hierarchy WHERE TableID = ? AND Name = ?",
                (table_id, hierarchy_name),
            ).fetchone()
            if existing:
                raise PBIXMCPError(f"Hierarchy '{hierarchy_name}' already exists on '{table_name}'", "DUPLICATE")

            # Get MAXID
            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0

            # Create Hierarchy
            max_id += 1
            hier_id = max_id
            c.execute(
                "INSERT INTO Hierarchy (ID, TableID, Name, IsHidden, State, ModifiedTime, StructureModifiedTime) "
                "VALUES (?, ?, ?, 0, 1, ?, ?)",
                (hier_id, table_id, hierarchy_name,
                 int(datetime.now().timestamp() * 1e7),
                 int(datetime.now().timestamp() * 1e7)),
            )

            # Create Levels and build LevelDefinition string
            level_col_ids = []
            level_def_parts = []
            cumulative_offset = 0
            for ordinal, lspec in enumerate(levels):
                lname = lspec.get("name", f"Level {ordinal}")
                col_name = lspec.get("column", "")

                crow = c.execute(
                    "SELECT ID, ExplicitName FROM [Column] WHERE TableID = ? AND (ExplicitName = ? OR InferredName = ?)",
                    (table_id, col_name, col_name),
                ).fetchone()
                if not crow:
                    raise PBIXMCPError(
                        f"Column '{col_name}' not found in table '{table_name}'",
                        "COLUMN_NOT_FOUND",
                    )
                col_id = crow["ID"]
                col_explicit = crow["ExplicitName"]
                level_col_ids.append(col_id)

                # Build LevelDefinition: $ColumnName (ColumnID)$offset$
                level_def_parts.append(f"${col_explicit} ({col_id})${cumulative_offset}$")

                # Count distinct values for this column to compute next offset
                # Use sorted dictionary cardinality
                distinct = c.execute(
                    "SELECT COUNT(DISTINCT ExplicitName) FROM [Column] WHERE TableID = ? AND ID = ?",
                    (table_id, col_id),
                ).fetchone()[0]
                # Actually need to count distinct VALUES in data, not columns.
                # For now, we can't easily do this from metadata alone.
                # Use a placeholder - PBI Desktop may recompute this.
                # We'll set offset to ordinal position as placeholder.

                max_id += 1
                c.execute(
                    "INSERT INTO Level (ID, HierarchyID, Ordinal, Name, ColumnID, ModifiedTime) "
                    "VALUES (?, ?, ?, ?, ?, ?)",
                    (max_id, hier_id, ordinal, lname, col_id,
                     int(datetime.now().timestamp() * 1e7)),
                )

            # Create HierarchyStorage (unmaterialized — no U$ table needed)
            # PBI Desktop creates U$ tables on first data refresh
            level_def = "".join(
                f"${crow_name} ({cid})$-1"
                for crow_name, cid in zip(
                    [c.execute("SELECT ExplicitName FROM [Column] WHERE ID = ?", (cid,)).fetchone()[0]
                     for cid in level_col_ids],
                    level_col_ids,
                )
            ) + "$"
            max_id += 1
            hier_storage_id = max_id
            c.execute(
                "INSERT INTO HierarchyStorage (ID, HierarchyID, Name, LevelDefinition, "
                "MaterializationType, StructureType, SystemTableID) "
                "VALUES (?, ?, ?, ?, -1, 0, 0)",
                (hier_storage_id, hier_id, f"{hierarchy_name} ({hier_id})",
                 level_def),
            )

            # Update Hierarchy to point to storage, State=4 (unmaterialized)
            c.execute(
                "UPDATE Hierarchy SET HierarchyStorageID = ?, State = 4 WHERE ID = ?",
                (hier_storage_id, hier_id),
            )

            # Set IsAvailableInMDX=1 on columns referenced by hierarchy levels
            for cid in level_col_ids:
                c.execute("UPDATE [Column] SET IsAvailableInMDX = 1 WHERE ID = ?", (cid,))

            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        return ToolResponse.ok(
            f"Hierarchy '{hierarchy_name}' created on '{table_name}' with {len(levels)} levels."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_remove_hierarchy(alias: str, table_name: str, hierarchy_name: str) -> str:
    """Remove a user hierarchy and all its levels.

    Args:
        alias: The alias of the open file
        table_name: Table the hierarchy belongs to
        hierarchy_name: Name of the hierarchy to remove
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        def _do_remove(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            if not trow:
                raise PBIXMCPError(f"Table '{table_name}' not found", "TABLE_NOT_FOUND")

            hrow = c.execute(
                "SELECT ID FROM Hierarchy WHERE TableID = ? AND Name = ?",
                (trow["ID"], hierarchy_name),
            ).fetchone()
            if not hrow:
                raise PBIXMCPError(f"Hierarchy '{hierarchy_name}' not found on '{table_name}'", "NOT_FOUND")

            c.execute("DELETE FROM Level WHERE HierarchyID = ?", (hrow["ID"],))
            c.execute("DELETE FROM Hierarchy WHERE ID = ?", (hrow["ID"],))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_remove)
        info["modified"] = True
        return ToolResponse.ok(f"Hierarchy '{hierarchy_name}' removed from '{table_name}'.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 10d: Cultures & Translations ----


@mcp.tool()
def pbix_get_cultures(alias: str) -> str:
    """Get all cultures (languages) with translation counts.

    Args:
        alias: The alias of the open file
    """
    try:
        info, conn, tmp_path = _read_metadata_db(alias)
        try:
            cultures = conn.execute("SELECT ID, Name FROM Culture ORDER BY Name").fetchall()
            if not cultures:
                return ToolResponse.ok("No cultures defined in this file.").to_text()

            lines = [f"Cultures ({len(cultures)}):\n"]
            for cu in cultures:
                cid, cname = cu["ID"], cu["Name"]
                count = conn.execute(
                    "SELECT COUNT(*) as cnt FROM ObjectTranslation WHERE CultureID = ?", (cid,)
                ).fetchone()["cnt"]
                lines.append(f"  {cname} — {count} translation(s)")

                # Show sample translations
                samples = conn.execute(
                    "SELECT ot.ObjectType, ot.Property, ot.Value, "
                    "COALESCE(t.Name, c2.ExplicitName, m.Name, h.Name) as ObjName "
                    "FROM ObjectTranslation ot "
                    "LEFT JOIN [Table] t ON ot.ObjectID = t.ID AND ot.ObjectType = 3 "
                    "LEFT JOIN [Column] c2 ON ot.ObjectID = c2.ID AND ot.ObjectType = 4 "
                    "LEFT JOIN Measure m ON ot.ObjectID = m.ID AND ot.ObjectType = 8 "
                    "LEFT JOIN Hierarchy h ON ot.ObjectID = h.ID AND ot.ObjectType = 9 "
                    "WHERE ot.CultureID = ? LIMIT 5", (cid,)
                ).fetchall()
                type_map = {3: "Table", 4: "Column", 8: "Measure", 9: "Hierarchy", 10: "Level"}
                prop_map = {1: "Caption", 2: "Description", 3: "DisplayFolder"}
                for s in samples:
                    otype = type_map.get(s["ObjectType"], f"Type{s['ObjectType']}")
                    prop = prop_map.get(s["Property"], f"Prop{s['Property']}")
                    lines.append(f"    {otype} '{s['ObjName']}' {prop} = \"{s['Value']}\"")
                lines.append("")

            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            conn.close()
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_culture(alias: str, culture_name: str) -> str:
    """Add a culture (language) for translations.

    Args:
        alias: The alias of the open file
        culture_name: BCP-47 culture code (e.g. "nb-NO", "de-DE", "fr-FR", "ja-JP")
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        def _do_add(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            existing = c.execute("SELECT ID FROM Culture WHERE Name = ?", (culture_name,)).fetchone()
            if existing:
                raise PBIXMCPError(f"Culture '{culture_name}' already exists", "DUPLICATE")

            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0
            max_id += 1
            c.execute(
                "INSERT INTO Culture (ID, ModelID, Name, ModifiedTime, StructureModifiedTime) "
                "VALUES (?, 1, ?, ?, ?)",
                (max_id, culture_name,
                 int(datetime.now().timestamp() * 1e7),
                 int(datetime.now().timestamp() * 1e7)),
            )
            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        return ToolResponse.ok(f"Culture '{culture_name}' added.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_translations(alias: str, culture_name: str, translations_json: str) -> str:
    """Add translated names/descriptions for model objects in a culture.

    Args:
        alias: The alias of the open file
        culture_name: Target culture (e.g. "nb-NO")
        translations_json: JSON array of translations, e.g.
            '[{"object": "Sales", "type": "table", "property": "caption", "value": "Salg"},
              {"object": "Sales.Amount", "type": "column", "property": "caption", "value": "Beloep"}]'
            type: "table", "column", "measure", "hierarchy"
            property: "caption" (display name), "description", "displayFolder"
            For columns/measures: use "Table.Column" or "Table.Measure" dot notation
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        translations = json.loads(translations_json)

        # TOM ObjectTranslation.ObjectType: 3=Table, 4=Column, 8=Measure, 9=Hierarchy, 10=Level
        type_map = {"table": 3, "column": 4, "measure": 8, "hierarchy": 9, "level": 10}
        # TOM ObjectTranslation.Property enum: 1=Caption, 2=Description, 3=DisplayFolder
        prop_map = {"caption": 1, "description": 2, "displayfolder": 3}

        def _do_add(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            culture_row = c.execute("SELECT ID FROM Culture WHERE Name = ?", (culture_name,)).fetchone()
            if not culture_row:
                raise PBIXMCPError(f"Culture '{culture_name}' not found. Add it first with pbix_add_culture.", "NOT_FOUND")
            culture_id = culture_row["ID"]

            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0
            added = 0

            for tr in translations:
                obj_ref = tr.get("object", "")
                obj_type_str = tr.get("type", "").lower()
                prop_str = tr.get("property", "caption").lower()
                value = tr.get("value", "")

                obj_type = type_map.get(obj_type_str)
                prop_code = prop_map.get(prop_str, 0)
                if obj_type is None:
                    continue

                # Resolve object name to ID
                obj_id = None
                if obj_type == 3:  # Table
                    row = c.execute("SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (obj_ref,)).fetchone()
                    if row:
                        obj_id = row["ID"]
                elif obj_type == 4:  # Column (Table.Column)
                    parts = obj_ref.split(".", 1)
                    if len(parts) == 2:
                        trow = c.execute("SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (parts[0],)).fetchone()
                        if trow:
                            crow = c.execute(
                                "SELECT ID FROM [Column] WHERE TableID = ? AND (ExplicitName = ? OR InferredName = ?)",
                                (trow["ID"], parts[1], parts[1]),
                            ).fetchone()
                            if crow:
                                obj_id = crow["ID"]
                elif obj_type == 8:  # Measure (Table.Measure)
                    parts = obj_ref.split(".", 1)
                    if len(parts) == 2:
                        trow = c.execute("SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (parts[0],)).fetchone()
                        if trow:
                            mrow = c.execute(
                                "SELECT ID FROM Measure WHERE TableID = ? AND Name = ?",
                                (trow["ID"], parts[1]),
                            ).fetchone()
                            if mrow:
                                obj_id = mrow["ID"]
                elif obj_type == 9:  # Hierarchy (Table.Hierarchy)
                    parts = obj_ref.split(".", 1)
                    if len(parts) == 2:
                        trow = c.execute("SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (parts[0],)).fetchone()
                        if trow:
                            hrow = c.execute(
                                "SELECT ID FROM Hierarchy WHERE TableID = ? AND Name = ?",
                                (trow["ID"], parts[1]),
                            ).fetchone()
                            if hrow:
                                obj_id = hrow["ID"]

                if obj_id is None:
                    continue

                # Upsert: check if translation exists
                existing = c.execute(
                    "SELECT ID FROM ObjectTranslation WHERE CultureID = ? AND ObjectID = ? AND ObjectType = ? AND Property = ?",
                    (culture_id, obj_id, obj_type, prop_code),
                ).fetchone()
                if existing:
                    c.execute(
                        "UPDATE ObjectTranslation SET Value = ?, ModifiedTime = ? WHERE ID = ?",
                        (value, int(datetime.now().timestamp() * 1e7), existing["ID"]),
                    )
                else:
                    max_id += 1
                    c.execute(
                        "INSERT INTO ObjectTranslation (ID, CultureID, ObjectID, ObjectType, Property, Value, ModifiedTime) "
                        "VALUES (?, ?, ?, ?, ?, ?, ?)",
                        (max_id, culture_id, obj_id, obj_type, prop_code, value,
                         int(datetime.now().timestamp() * 1e7)),
                    )
                added += 1

            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        return ToolResponse.ok(f"Added/updated {len(translations)} translation(s) for culture '{culture_name}'.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_remove_culture(alias: str, culture_name: str) -> str:
    """Remove a culture and all its translations.

    Args:
        alias: The alias of the open file
        culture_name: Culture code to remove (e.g. "nb-NO")
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        def _do_remove(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()
            row = c.execute("SELECT ID FROM Culture WHERE Name = ?", (culture_name,)).fetchone()
            if not row:
                raise PBIXMCPError(f"Culture '{culture_name}' not found", "NOT_FOUND")
            c.execute("DELETE FROM ObjectTranslation WHERE CultureID = ?", (row["ID"],))
            c.execute("DELETE FROM Culture WHERE ID = ?", (row["ID"],))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_remove)
        info["modified"] = True
        return ToolResponse.ok(f"Culture '{culture_name}' and all its translations removed.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 10e: Partition Management ----


@mcp.tool()
def pbix_get_partitions(alias: str) -> str:
    """Get all partitions with table, type, mode, and M expression.

    Args:
        alias: The alias of the open file
    """
    try:
        info, conn, tmp_path = _read_metadata_db(alias)
        try:
            partitions = conn.execute(
                "SELECT p.Name, p.Type, p.Mode, p.QueryDefinition, t.Name as TableName "
                "FROM Partition p JOIN [Table] t ON p.TableID = t.ID "
                "WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%' AND t.Name NOT LIKE 'U$%' "
                "ORDER BY t.Name, p.Name"
            ).fetchall()
            if not partitions:
                return ToolResponse.ok("No partitions found.").to_text()

            type_map = {1: "Query", 2: "Calculated", 3: "None", 4: "M"}
            mode_map = {0: "Import", 1: "DirectQuery", 2: "Dual"}

            lines = [f"Partitions ({len(partitions)}):\n"]
            current_table = None
            for p in partitions:
                tname = p["TableName"]
                if tname != current_table:
                    current_table = tname
                    lines.append(f"  {tname}:")

                ptype = type_map.get(p["Type"], f"Type{p['Type']}")
                pmode = mode_map.get(p["Mode"], f"Mode{p['Mode']}")
                qd = p["QueryDefinition"]
                expr_preview = (qd[:80] + "...") if qd and len(qd) > 80 else (qd or "(none)")
                lines.append(f"    {p['Name']} [{ptype}/{pmode}]: {expr_preview}")

            return ToolResponse.ok("\n".join(lines)).to_text()
        finally:
            conn.close()
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_add_partition(
    alias: str,
    table_name: str,
    partition_name: str,
    expression: str,
    mode: str = "import",
) -> str:
    """Add a new M (Power Query) partition to a table.

    WARNING: This tool is blocked for PBIX files opened in PBI Desktop.
    Adding partitions requires PartitionStorage objects in VertiPaq which
    cannot be created via metadata-only modification. The partition metadata
    is written correctly but PBI Desktop will reject the file on open.
    Works for PBIP/TMDL export (pbix_export_pbip, pbix_export_tmdl).

    Args:
        alias: The alias of the open file
        table_name: Table to add the partition to
        partition_name: Name for the new partition
        expression: M (Power Query) expression for the partition
        mode: "import" (default) or "directQuery"
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        mode_code = 1 if mode.lower() == "directquery" else 0

        def _do_add(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            if not trow:
                raise PBIXMCPError(f"Table '{table_name}' not found", "TABLE_NOT_FOUND")

            existing = c.execute(
                "SELECT ID FROM Partition WHERE TableID = ? AND Name = ?",
                (trow["ID"], partition_name),
            ).fetchone()
            if existing:
                raise PBIXMCPError(f"Partition '{partition_name}' already exists on '{table_name}'", "DUPLICATE")

            maxid_row = c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
            max_id = int(maxid_row[0]) if maxid_row else 0
            max_id += 1

            c.execute(
                "INSERT INTO Partition (ID, TableID, Name, Type, Mode, State, "
                "ModifiedTime, RefreshedTime, QueryDefinition) "
                "VALUES (?, ?, ?, 4, ?, 1, ?, ?, ?)",
                (max_id, trow["ID"], partition_name, mode_code,
                 int(datetime.now().timestamp() * 1e7),
                 int(datetime.now().timestamp() * 1e7),
                 expression),
            )

            c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'", (str(max_id),))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_add)
        info["modified"] = True
        return ToolResponse.ok(f"Partition '{partition_name}' added to '{table_name}'.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_remove_partition(alias: str, table_name: str, partition_name: str) -> str:
    """Remove a partition from a table.

    Will not delete the last remaining partition of a table.

    Args:
        alias: The alias of the open file
        table_name: Table the partition belongs to
        partition_name: Name of the partition to remove
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")

        def _do_remove(conn: sqlite3.Connection):
            conn.row_factory = sqlite3.Row
            c = conn.cursor()

            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1", (table_name,)
            ).fetchone()
            if not trow:
                raise PBIXMCPError(f"Table '{table_name}' not found", "TABLE_NOT_FOUND")

            prow = c.execute(
                "SELECT ID FROM Partition WHERE TableID = ? AND Name = ?",
                (trow["ID"], partition_name),
            ).fetchone()
            if not prow:
                raise PBIXMCPError(f"Partition '{partition_name}' not found on '{table_name}'", "NOT_FOUND")

            # Guard: don't delete last partition
            count = c.execute(
                "SELECT COUNT(*) as cnt FROM Partition WHERE TableID = ?", (trow["ID"],)
            ).fetchone()["cnt"]
            if count <= 1:
                raise PBIXMCPError(
                    f"Cannot delete the last partition of table '{table_name}'",
                    "LAST_PARTITION",
                )

            c.execute("DELETE FROM Partition WHERE ID = ?", (prow["ID"],))
            conn.commit()

        old_size, new_size = _modify_metadata_only(dm_path, _do_remove)
        info["modified"] = True
        return ToolResponse.ok(f"Partition '{partition_name}' removed from '{table_name}'.").to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


# ---- Section 11: Diagnostics ----

@mcp.tool()
def pbix_get_password(alias: str) -> str:
    """Extract embedded passwords from a PBIX file.

    Scans the data model for password-like tables (tables with 'password'
    in the name) and DAX measures that reference them (ISFILTERED, SELECTEDVALUE).
    Extracts the expected password value from the DAX expression.

    This is useful for dashboards that use a password-slicer gate pattern
    where the report is locked until the correct password is entered.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=info.get("work_dir"))

        results = []

        # Strategy 1: Find tables with 'password' in the name and read their data
        schema = model.schema
        table_names = sorted(set(r['TableName'] for r in schema))
        for tname in table_names:
            if "password" in tname.lower():
                try:
                    td = model.get_table(tname)
                    if td and td.get('columns') and td.get('rows'):
                        # Get unique values per column
                        for ci, col in enumerate(td['columns']):
                            vals = sorted(set(
                                row[ci] for row in td['rows']
                                if ci < len(row) and row[ci] is not None
                            ), key=str)
                            if vals:
                                results.append(f"Table '{tname}', column '{col}': {len(vals)} values")
                                for v in vals[:10]:
                                    results.append(f"  {v}")
                                if len(vals) > 10:
                                    results.append(f"  ... and {len(vals) - 10} more")
                except Exception:
                    pass

        # Strategy 2: Find DAX measures that check passwords
        measures_list = model.dax_measures
        if measures_list:
            import re as _re
            for m_row in measures_list:
                expr = m_row.get("Expression", "")
                name = m_row.get("Name", "")
                if not expr:
                    continue
                # Look for SELECTEDVALUE(...[...]) = "value" patterns near password context
                for m in _re.finditer(
                    r"""SELECTEDVALUE\s*\(\s*'?([^')]+)'?\s*\[([^\]]+)\]\s*\)\s*=\s*["']([^"']+)["']""",
                    expr, _re.IGNORECASE
                ):
                    table = m.group(1).strip()
                    column = m.group(2).strip()
                    password = m.group(3)
                    if "password" in table.lower() or "password" in column.lower() or "password" in name.lower():
                        results.append(f"  >>> PASSWORD: \"{password}\"  (from SELECTEDVALUE('{table}'[{column}]) in measure '{name}')")

                # Also look for hardcoded password strings near password context
                skip_words = {"correct", "wrong", "true", "false", "password",
                              "enjoy", "dashboard", "filter", "warning", "error",
                              "selected", "value", "blank"}
                for m in _re.finditer(r'''["']([^"'\n]{3,30})["']''', expr):
                    candidate = m.group(1).strip()
                    if candidate.lower() in skip_words:
                        continue
                    # Only flag if near a password-related context
                    context_start = max(0, m.start() - 200)
                    context = expr[context_start:m.end()].lower()
                    if "password" in context:
                        # Skip obvious UI text
                        if any(w in candidate.lower() for w in ["correct", "wrong", "enjoy", "⚠", "✔"]):
                            continue
                        results.append(f"  Candidate in measure '{name}': \"{candidate}\"")

        if not results:
            return ToolResponse.ok("No password tables or password-checking measures found in this file.").to_text()

        return ToolResponse.ok("Password analysis:\n" + "\n".join(results)).to_text()

    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(str(e), "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_performance(alias: str) -> str:
    """Analyze report for performance issues and optimization opportunities.

    Checks for oversized tables, high column counts, complex measures,
    orphaned tables, inactive relationships, and wide schemas.

    Args:
        alias: The alias of the open file
    """
    import re as _re
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]

        from pbix_mcp.formats.model_reader import ModelReader
        model = ModelReader(info["path"], work_dir=work_dir)

        lines: list[str] = []
        warnings = 0
        infos = 0

        def warn(msg: str):
            nonlocal warnings
            lines.append(f"  WARNING: {msg}")
            warnings += 1

        def info_msg(msg: str):
            nonlocal infos
            lines.append(f"  INFO: {msg}")
            infos += 1

        lines.append("# Performance Analysis\n")

        # --- Table sizes ---
        lines.append("## Table Sizes")
        stats = model.statistics
        data_tables = [t for t in stats if not t["TableName"].startswith(
            ("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate"))]
        data_tables.sort(key=lambda t: t["RowCount"], reverse=True)

        total_rows = sum(t["RowCount"] for t in data_tables)
        total_cols = sum(t["ColumnCount"] for t in data_tables)
        lines.append(f"  Total: {len(data_tables)} tables, {total_rows:,} rows, {total_cols} columns")

        for t in data_tables:
            name, rows, cols = t["TableName"], t["RowCount"], t["ColumnCount"]
            flags = []
            if rows > 1_000_000:
                flags.append("LARGE (>1M rows)")
            elif rows > 100_000:
                flags.append("medium (>100K rows)")
            if cols > 30:
                flags.append(f"wide ({cols} columns)")
            elif cols > 20:
                flags.append(f"moderately wide ({cols} columns)")
            if rows == 0:
                flags.append("empty table")

            if flags:
                warn(f"{name}: {rows:,} rows, {cols} cols — {', '.join(flags)}")
            else:
                lines.append(f"  {name}: {rows:,} rows, {cols} cols")

        # --- Column analysis ---
        lines.append("\n## Column Analysis")
        schema = model.schema
        hidden_count = 0
        calc_count = 0
        by_type: dict[str, int] = {}
        for col in schema:
            if col["TableName"].startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate")):
                continue
            dt = col["DataType"]
            by_type[dt] = by_type.get(dt, 0) + 1
            if col.get("IsHidden"):
                hidden_count += 1
            if col.get("IsCalculated"):
                calc_count += 1

        lines.append(f"  Types: {', '.join(f'{k}={v}' for k, v in sorted(by_type.items()))}")
        if hidden_count:
            info_msg(f"{hidden_count} hidden columns (keys, internal)")
        if calc_count:
            info_msg(f"{calc_count} calculated columns (evaluated at refresh)")

        # String columns in large tables — high cardinality risk
        for t in data_tables:
            if t["RowCount"] > 50_000:
                str_cols = [c for c in schema
                            if c["TableName"] == t["TableName"]
                            and c["DataType"] == "String"
                            and not c.get("IsHidden")
                            and "RowNumber" not in c["ColumnName"]]
                if len(str_cols) > 5:
                    warn(f"{t['TableName']}: {len(str_cols)} string columns on {t['RowCount']:,} rows — potential high cardinality")

        # --- Measure complexity ---
        lines.append("\n## Measure Complexity")
        measures = model.dax_measures
        if measures:
            for m in measures:
                expr = m["Expression"]
                # Count table references
                table_refs = set(_re.findall(r"'([^']+)'\[", expr))
                table_refs |= set(_re.findall(r"\b([A-Za-z]\w+)\[", expr))
                # Count function calls
                func_calls = len(_re.findall(r"[A-Z]{2,}\s*\(", expr))
                # Count nesting depth (rough — count opening parens)
                max_depth = 0
                depth = 0
                for ch in expr:
                    if ch == '(':
                        depth += 1
                        max_depth = max(max_depth, depth)
                    elif ch == ')':
                        depth -= 1

                flags = []
                if len(table_refs) > 3:
                    flags.append(f"references {len(table_refs)} tables")
                if func_calls > 10:
                    flags.append(f"{func_calls} function calls")
                elif func_calls > 5:
                    flags.append(f"{func_calls} function calls")
                if max_depth > 5:
                    flags.append(f"nesting depth {max_depth}")
                if len(expr) > 500:
                    flags.append(f"{len(expr)} chars")

                if flags:
                    warn(f"Measure '{m['Name']}': {', '.join(flags)}")
                else:
                    lines.append(f"  {m['Name']}: {func_calls} functions, {len(table_refs)} table refs — OK")

            lines.append(f"  {len(measures)} measures total")
        else:
            lines.append("  No measures defined")

        # --- Relationships ---
        lines.append("\n## Relationships")
        rels = model.relationships
        inactive = [r for r in rels if not r.get("IsActive")]
        bidir = [r for r in rels if r.get("CrossFilteringBehavior") == 2]

        lines.append(f"  {len(rels)} relationships total")
        if inactive:
            for r in inactive:
                warn(f"Inactive: {r['FromTableName']}.{r['FromColumnName']} -> {r['ToTableName']}.{r['ToColumnName']}")
        if bidir:
            for r in bidir:
                warn(f"Bidirectional: {r['FromTableName']} <-> {r['ToTableName']} — can cause ambiguity")

        # Orphaned tables (no relationships)
        rel_tables = set()
        for r in rels:
            rel_tables.add(r["FromTableName"])
            rel_tables.add(r["ToTableName"])
        for t in data_tables:
            if t["TableName"] not in rel_tables and t["RowCount"] > 0:
                info_msg(f"Orphaned table '{t['TableName']}' — no relationships, {t['RowCount']:,} rows")

        # --- Summary ---
        lines.insert(1, f"Summary: {warnings} warnings, {infos} info items\n")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_diff(alias_a: str, alias_b: str) -> str:
    """Compare two open PBIX files and show what changed.

    Compares tables, columns, measures, relationships, pages, visuals,
    data sources, and theme colors between two files. Both files must
    be open.

    Args:
        alias_a: The alias of the first file (baseline / "old")
        alias_b: The alias of the second file (changed / "new")
    """
    try:
        info_a = _ensure_open(alias_a)
        info_b = _ensure_open(alias_b)

        from pbix_mcp.formats.model_reader import ModelReader
        model_a = ModelReader(info_a["path"], work_dir=info_a.get("work_dir"))
        model_b = ModelReader(info_b["path"], work_dir=info_b.get("work_dir"))

        lines: list[str] = []

        def section(title: str):
            lines.append(f"\n## {title}\n")

        def added(item: str):
            lines.append(f"  + {item}")

        def removed(item: str):
            lines.append(f"  - {item}")

        def changed(item: str, old_val: str, new_val: str):
            lines.append(f"  ~ {item}: {old_val} -> {new_val}")

        name_a = os.path.basename(info_a["path"])
        name_b = os.path.basename(info_b["path"])
        lines.append(f"# Diff: {name_a} vs {name_b}")

        # --- Tables ---
        section("Tables")
        try:
            stats_a = {t["TableName"]: t for t in model_a.statistics
                       if not t["TableName"].startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate"))}
            stats_b = {t["TableName"]: t for t in model_b.statistics
                       if not t["TableName"].startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate"))}

            for name in sorted(set(stats_b) - set(stats_a)):
                t = stats_b[name]
                added(f"{name} ({t['ColumnCount']} cols, {t['RowCount']:,} rows)")
            for name in sorted(set(stats_a) - set(stats_b)):
                t = stats_a[name]
                removed(f"{name} ({t['ColumnCount']} cols, {t['RowCount']:,} rows)")
            for name in sorted(set(stats_a) & set(stats_b)):
                a, b = stats_a[name], stats_b[name]
                changes = []
                if a["ColumnCount"] != b["ColumnCount"]:
                    changes.append(f"columns {a['ColumnCount']}->{b['ColumnCount']}")
                if a["RowCount"] != b["RowCount"]:
                    changes.append(f"rows {a['RowCount']:,}->{b['RowCount']:,}")
                if changes:
                    changed(name, "", ", ".join(changes))

            if not any(ln.startswith("  ") for ln in lines if lines.index(ln) > len(lines) - 10):
                lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Columns ---
        section("Columns")
        try:
            def _col_set(model):
                result = {}
                for c in model.schema:
                    tn = c["TableName"]
                    if tn.startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate")):
                        continue
                    if c.get("IsHidden") or "RowNumber" in c["ColumnName"]:
                        continue
                    result[f"{tn}.{c['ColumnName']}"] = c["DataType"]
                return result

            cols_a = _col_set(model_a)
            cols_b = _col_set(model_b)

            for col in sorted(set(cols_b) - set(cols_a)):
                added(f"{col} ({cols_b[col]})")
            for col in sorted(set(cols_a) - set(cols_b)):
                removed(f"{col} ({cols_a[col]})")
            for col in sorted(set(cols_a) & set(cols_b)):
                if cols_a[col] != cols_b[col]:
                    changed(col, cols_a[col], cols_b[col])

            added_count = len(set(cols_b) - set(cols_a))
            removed_count = len(set(cols_a) - set(cols_b))
            if added_count == 0 and removed_count == 0:
                lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Measures ---
        section("Measures")
        try:
            meas_a = {m["Name"]: m for m in model_a.dax_measures}
            meas_b = {m["Name"]: m for m in model_b.dax_measures}

            for name in sorted(set(meas_b) - set(meas_a)):
                m = meas_b[name]
                expr = m["Expression"].replace("\n", " ")[:60]
                added(f"{name} = {expr}")
            for name in sorted(set(meas_a) - set(meas_b)):
                removed(name)
            for name in sorted(set(meas_a) & set(meas_b)):
                if meas_a[name]["Expression"] != meas_b[name]["Expression"]:
                    old_expr = meas_a[name]["Expression"].replace("\n", " ")[:40]
                    new_expr = meas_b[name]["Expression"].replace("\n", " ")[:40]
                    changed(name, old_expr, new_expr)

            if not meas_a and not meas_b:
                lines.append("  (no measures in either file)")
            elif len(set(meas_b) - set(meas_a)) == 0 and len(set(meas_a) - set(meas_b)) == 0:
                has_expr_changes = any(meas_a[n]["Expression"] != meas_b[n]["Expression"]
                                       for n in set(meas_a) & set(meas_b))
                if not has_expr_changes:
                    lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Relationships ---
        section("Relationships")
        try:
            def _rel_key(r):
                return f"{r['FromTableName']}.{r['FromColumnName']}->{r['ToTableName']}.{r['ToColumnName']}"

            rels_a = {_rel_key(r): r for r in model_a.relationships}
            rels_b = {_rel_key(r): r for r in model_b.relationships}

            for key in sorted(set(rels_b) - set(rels_a)):
                added(key)
            for key in sorted(set(rels_a) - set(rels_b)):
                removed(key)
            for key in sorted(set(rels_a) & set(rels_b)):
                if rels_a[key].get("IsActive") != rels_b[key].get("IsActive"):
                    changed(key, f"active={rels_a[key].get('IsActive')}", f"active={rels_b[key].get('IsActive')}")

            if len(set(rels_b) - set(rels_a)) == 0 and len(set(rels_a) - set(rels_b)) == 0:
                lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Pages & Visuals ---
        section("Pages & Visuals")
        try:
            layout_a = _get_layout(info_a.get("work_dir", "")) or {}
            layout_b = _get_layout(info_b.get("work_dir", "")) or {}

            pages_a = {s.get("displayName", f"Page {i}"): s
                       for i, s in enumerate(layout_a.get("sections", []))}
            pages_b = {s.get("displayName", f"Page {i}"): s
                       for i, s in enumerate(layout_b.get("sections", []))}

            for pname in sorted(set(pages_b) - set(pages_a)):
                vc_count = len(pages_b[pname].get("visualContainers", []))
                added(f"Page '{pname}' ({vc_count} visuals)")
            for pname in sorted(set(pages_a) - set(pages_b)):
                removed(f"Page '{pname}'")
            for pname in sorted(set(pages_a) & set(pages_b)):
                vc_a = len(pages_a[pname].get("visualContainers", []))
                vc_b = len(pages_b[pname].get("visualContainers", []))
                if vc_a != vc_b:
                    changed(f"Page '{pname}'", f"{vc_a} visuals", f"{vc_b} visuals")

            if len(set(pages_b) - set(pages_a)) == 0 and len(set(pages_a) - set(pages_b)) == 0:
                has_visual_changes = any(
                    len(pages_a[p].get("visualContainers", [])) != len(pages_b[p].get("visualContainers", []))
                    for p in set(pages_a) & set(pages_b)
                )
                if not has_visual_changes:
                    lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Data Sources ---
        section("Data Sources")
        try:
            pq_a = {p["TableName"]: p.get("Expression", "") for p in model_a.power_query}
            pq_b = {p["TableName"]: p.get("Expression", "") for p in model_b.power_query}

            for tname in sorted(set(pq_b) - set(pq_a)):
                added(f"{tname} (new query)")
            for tname in sorted(set(pq_a) - set(pq_b)):
                removed(f"{tname}")
            for tname in sorted(set(pq_a) & set(pq_b)):
                if pq_a[tname] != pq_b[tname]:
                    changed(tname, "M expression modified", "")

            if len(set(pq_b) - set(pq_a)) == 0 and len(set(pq_a) - set(pq_b)) == 0:
                has_pq_changes = any(pq_a[t] != pq_b[t] for t in set(pq_a) & set(pq_b))
                if not has_pq_changes:
                    lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Theme ---
        section("Theme Colors")
        try:
            work_a = info_a.get("work_dir", "")
            work_b = info_b.get("work_dir", "")
            colors_a = set(_load_theme_data_colors(work_a))
            colors_b = set(_load_theme_data_colors(work_b))

            for c in sorted(colors_b - colors_a):
                added(c)
            for c in sorted(colors_a - colors_b):
                removed(c)
            if colors_a == colors_b:
                lines.append("  (no changes)")
        except Exception as e:
            lines.append(f"  Error: {e}")

        # --- Summary ---
        adds = sum(1 for ln in lines if ln.startswith("  + "))
        removes = sum(1 for ln in lines if ln.startswith("  - "))
        changes = sum(1 for ln in lines if ln.startswith("  ~ "))
        lines.insert(1, f"\nSummary: {adds} added, {removes} removed, {changes} changed")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_document(alias: str, output_path: str = "") -> str:
    """Auto-generate a comprehensive report documentation summary.

    Assembles all report metadata into a structured document: tables with
    row counts, column details, DAX measures with expressions, relationships,
    data sources, pages with visual inventories, RLS roles, and theme colors.

    Returns markdown in the response AND saves a .docx file to disk.

    Args:
        alias: The alias of the open file
        output_path: Where to save the .docx file. Default: next to the PBIX.
    """
    import re
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        pbix_path = info["path"]

        if not output_path:
            output_path = os.path.splitext(pbix_path)[0] + "_documentation.docx"

        md_lines: list[str] = []

        def md(line: str = ""):
            md_lines.append(line)

        # --- Header ---
        fname = os.path.basename(pbix_path)
        md(f"# Report Documentation: {fname}")
        md(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        md()

        # --- Tables ---
        md("## Data Model — Tables")
        md()
        try:
            from pbix_mcp.formats.model_reader import ModelReader
            model = ModelReader(pbix_path, work_dir=work_dir)
            stats = model.statistics
            data_tables = [t for t in stats if not t["TableName"].startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate"))]

            md("| Table | Columns | Rows |")
            md("|-------|---------|------|")
            for t in data_tables:
                md(f"| {t['TableName']} | {t['ColumnCount']} | {t['RowCount']:,} |")
            md()
            md(f"*{len(data_tables)} data tables, {sum(t['RowCount'] for t in data_tables):,} total rows*")
            md()
        except Exception as e:
            md(f"*Error reading tables: {e}*")
            md()

        # --- Schema ---
        md("## Column Details")
        md()
        try:
            schema = model.schema
            by_table: dict[str, list] = {}
            for col in schema:
                tname = col["TableName"]
                if tname.startswith(("H$", "R$", "U$", "LocalDateTable", "DateTableTemplate")):
                    continue
                by_table.setdefault(tname, []).append(col)

            for tname in sorted(by_table.keys()):
                cols = by_table[tname]
                visible = [c for c in cols if not c.get("IsHidden") and "RowNumber" not in c["ColumnName"]]
                if not visible:
                    continue
                md(f"### {tname}")
                md()
                md("| Column | Type |")
                md("|--------|------|")
                for c in visible:
                    md(f"| {c['ColumnName']} | {c['DataType']} |")
                md()
        except Exception as e:
            md(f"*Error reading schema: {e}*")
            md()

        # --- Measures ---
        md("## DAX Measures")
        md()
        try:
            measures = model.dax_measures
            if measures:
                md("| Table | Measure | Expression | Format |")
                md("|-------|---------|------------|--------|")
                for m in measures:
                    expr = m["Expression"].replace("\n", " ").replace("|", "\\|")
                    if len(expr) > 80:
                        expr = expr[:77] + "..."
                    fmt = m.get("FormatString", "") or ""
                    md(f"| {m['TableName']} | **{m['Name']}** | `{expr}` | {fmt} |")
                md()
                md(f"*{len(measures)} measures*")
            else:
                md("*No DAX measures defined*")
            md()
        except Exception as e:
            md(f"*Error reading measures: {e}*")
            md()

        # --- Relationships ---
        md("## Relationships")
        md()
        try:
            rels = model.relationships
            if rels:
                md("| From (Many) | | To (One) | Active |")
                md("|-------------|---|----------|--------|")
                for r in rels:
                    active = "Yes" if r.get("IsActive") else "No"
                    md(f"| {r['FromTableName']}.{r['FromColumnName']} | -> | {r['ToTableName']}.{r['ToColumnName']} | {active} |")
                md()
                md(f"*{len(rels)} relationships*")
            else:
                md("*No relationships*")
            md()
        except Exception as e:
            md(f"*Error reading relationships: {e}*")
            md()

        # --- Data Sources ---
        md("## Data Sources")
        md()
        try:
            pq = model.power_query
            if pq:
                md("| Table | M Expression (excerpt) |")
                md("|-------|----------------------|")
                for p in pq:
                    expr = p.get("Expression", "")
                    if expr:
                        # Find the Source = ... line (most informative)
                        lines = expr.split("\n")
                        source_line = next((ln.strip() for ln in lines if "Source" in ln and "=" in ln), lines[0].strip())
                        if len(source_line) > 80:
                            source_line = source_line[:77] + "..."
                        md(f"| {p['TableName']} | `{source_line}` |")
                md()
            else:
                md("*No Power Query expressions found*")
            md()
        except Exception:
            md("*No data source information available*")
            md()

        # --- Pages & Visuals ---
        md("## Report Pages & Visuals")
        md()
        layout = _get_layout(work_dir)
        if layout:
            for si, sec in enumerate(layout.get("sections", [])):
                page_name = sec.get("displayName", f"Page {si}")
                containers = sec.get("visualContainers", [])
                w = sec.get("width", 1280)
                h = sec.get("height", 720)
                md(f"### {page_name} ({w}x{h})")
                md()
                if containers:
                    md("| # | Type | Position | Size |")
                    md("|---|------|----------|------|")
                    for vi, vc in enumerate(containers):
                        config = _parse_visual_config(vc)
                        vtype = _get_visual_type(config)
                        x, y = int(vc.get("x", 0)), int(vc.get("y", 0))
                        vw, vh = int(vc.get("width", 0)), int(vc.get("height", 0))
                        md(f"| {vi} | {vtype} | ({x},{y}) | {vw}x{vh} |")
                else:
                    md("*No visuals on this page*")
                md()
        else:
            md("*No layout found*")
            md()

        # --- RLS Roles ---
        md("## Row-Level Security")
        md()
        try:
            meta = model._read_metadata()
            if meta:
                import sqlite3
                conn = sqlite3.connect(":memory:")
                conn.executescript("BEGIN;" if False else "")
                # Write meta to temp
                import tempfile
                fd, tmp = tempfile.mkstemp(suffix=".db")
                os.write(fd, meta)
                os.close(fd)
                conn = sqlite3.connect(tmp)
                conn.row_factory = sqlite3.Row
                roles = conn.execute("SELECT Name FROM Role").fetchall()
                if roles:
                    for role in roles:
                        rname = role["Name"]
                        perms = conn.execute(
                            "SELECT t.Name as TableName, tp.FilterExpression "
                            "FROM TablePermission tp JOIN Role r ON tp.RoleID=r.ID "
                            "JOIN [Table] t ON tp.TableID=t.ID "
                            "WHERE r.Name=?", (rname,)
                        ).fetchall()
                        md(f"**{rname}**")
                        for p in perms:
                            md(f"- `{p['TableName']}`: `{p['FilterExpression']}`")
                        md()
                else:
                    md("*No RLS roles defined*")
                conn.close()
                os.unlink(tmp)
            else:
                md("*No metadata available for RLS*")
        except Exception:
            md("*No RLS roles defined*")
        md()

        # --- Theme Colors ---
        md("## Theme Colors")
        md()
        data_colors = _load_theme_data_colors(work_dir)
        if data_colors:
            md("Data palette: " + ", ".join(f"`{c}`" for c in data_colors[:10]))
        else:
            md("*Default theme*")
        md()

        # --- Build markdown ---
        markdown = "\n".join(md_lines)

        # --- Build .docx ---
        docx_msg = ""
        try:
            from docx import Document
            from docx.shared import Pt

            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "Segoe UI"
            style.font.size = Pt(10)

            doc.add_heading(f"Report Documentation: {fname}", level=0)

            # Single-pass: process lines, collecting table rows inline
            i = 0
            while i < len(md_lines):
                line = md_lines[i]

                # Skip title (already added) and empty lines
                if (line.startswith("# ") and not line.startswith("## ")) or not line.strip():
                    i += 1
                    continue

                # Headings
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=2)
                    i += 1
                    continue
                if line.startswith("## "):
                    doc.add_heading(line[3:], level=1)
                    i += 1
                    continue

                # Table block: collect all consecutive | rows (including |--- separators)
                if line.startswith("|"):
                    rows = []
                    while i < len(md_lines) and md_lines[i].startswith("|"):
                        if "---|" not in md_lines[i]:
                            cells = [c.strip() for c in md_lines[i].split("|")[1:-1]]
                            rows.append(cells)
                        i += 1
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = "Light Grid Accent 1"
                        for ri, row in enumerate(rows):
                            for ci, cell in enumerate(row):
                                clean = re.sub(r'[`*]', '', cell)
                                table.rows[ri].cells[ci].text = clean
                        doc.add_paragraph()  # spacing after table
                    continue

                # Italic text
                if line.startswith("*") and line.endswith("*"):
                    p = doc.add_paragraph(line.strip("*"))
                    if p.runs:
                        p.runs[0].italic = True
                    i += 1
                    continue

                # Normal text
                doc.add_paragraph(line)
                i += 1

            doc.save(output_path)
            docx_msg = f"\n\nDocx saved to: {output_path}"
        except ImportError:
            docx_msg = "\n\n(python-docx not installed — skipping .docx generation. Install with: pip install python-docx)"
        except Exception as e:
            docx_msg = f"\n\n(Docx generation error: {e})"

        return ToolResponse.ok(markdown + docx_msg).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", "INTERNAL_ERROR").to_text()


@mcp.tool()
def pbix_doctor(alias: str) -> str:
    """Run comprehensive diagnostics on an open PBIX/PBIT file.

    Performs a full health check across every layer of the file:
    ZIP structure, report layout, DataModel compression, ABF archive,
    SQLite metadata, data source connections, storage modes,
    VertiPaq column data, relationships, measures, calculated tables,
    RLS roles, and slicer filters.

    Args:
        alias: The alias of the open file
    """
    checks = []

    def _check(name, fn):
        try:
            result = fn()
            checks.append(f"  ✅ {name}: {result}")
            return True
        except _DoctorWarning as w:
            # Valid but suspicious — Power BI tolerates it, a human may not
            # want it. Reported separately so a real defect is not buried.
            checks.append(f"  ⚠️  {name}: {w}")
            return True
        except Exception as e:
            checks.append(f"  ❌ {name}: {e}")
            return False

    try:
        info = _ensure_open(alias)
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        raise SessionError(f"Doctor failed: {e}")

    checks.append(f"Diagnostics for '{alias}':\n")

    # 1. File basics
    _check("File exists", lambda: f"{os.path.getsize(info['path']):,} bytes" if os.path.exists(info['path']) else "MISSING")
    _check("File type", lambda: "PBIT" if info.get("is_pbit") else "PBIX")

    # 2. Layout
    def check_layout():
        layout = _get_layout(info["work_dir"])
        if layout:
            pages = len(layout.get("sections", []))
            fmt = "PBIR" if layout.get("__pbir__") else "legacy"
            return f"{pages} pages ({fmt} format)"
        return "No layout found"
    _check("Report layout", check_layout)

    # --- Decompress DataModel ONCE for all subsequent checks ---
    dm_path = os.path.join(info["work_dir"], "DataModel")
    abf_data = None
    abf_files = None
    db_conn = None
    db_tmp_path = None

    def _init_datamodel():
        nonlocal abf_data, abf_files, db_conn, db_tmp_path
        if abf_data is not None:
            return
        import tempfile

        from pbix_mcp.formats.abf_rebuild import list_abf_files, read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
        with open(dm_path, "rb") as f:
            dm = f.read()
        abf_data = decompress_datamodel(dm)
        abf_files = list_abf_files(abf_data)
        db_bytes = read_metadata_sqlite(abf_data)
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(db_bytes)
        tmp.close()
        db_tmp_path = tmp.name
        db_conn = sqlite3.connect(db_tmp_path)

    try:
        # 3. DataModel
        def check_datamodel():
            if not os.path.exists(dm_path):
                return "NOT FOUND"
            size = os.path.getsize(dm_path)
            _init_datamodel()
            return f"{size:,} bytes compressed, {len(abf_data):,} bytes decompressed"
        _check("DataModel (XPress9)", check_datamodel)

        # 4. ABF contents
        def check_abf():
            _init_datamodel()
            return f"{len(abf_files)} internal files"
        _check("ABF archive", check_abf)

        # 5. SQLite metadata
        def check_sqlite():
            _init_datamodel()
            tables = db_conn.execute("SELECT COUNT(*) FROM [Table] WHERE ModelID=1").fetchone()[0]
            measures = db_conn.execute("SELECT COUNT(*) FROM [Measure]").fetchone()[0]
            rels = db_conn.execute("SELECT COUNT(*) FROM [Relationship]").fetchone()[0]
            return f"{tables} tables, {measures} measures, {rels} relationships"
        _check("Metadata SQLite", check_sqlite)

        # 6. Data sources & storage modes
        def check_data_sources():
            _init_datamodel()
            c = db_conn.cursor()
            mode_names = {0: "Import", 1: "DirectQuery", 2: "Dual"}
            results = []
            c.execute("""SELECT t.Name, p.Mode, SUBSTR(p.QueryDefinition, 1, 60)
                         FROM Partition p JOIN [Table] t ON p.TableID = t.ID
                         WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'
                         AND p.QueryDefinition IS NOT NULL
                         ORDER BY t.Name""")
            modes = set()
            sources = set()
            for row in c.fetchall():
                tname, mode, qd = row
                mode_str = mode_names.get(mode, f"Unknown({mode})")
                modes.add(mode_str)
                if qd:
                    if "PostgreSQL.Database" in qd:
                        sources.add("PostgreSQL")
                    elif "MySQL.Database" in qd:
                        sources.add("MySQL")
                    elif "Sql.Database" in qd:
                        sources.add("SQL Server")
                    elif "Odbc.DataSource" in qd:
                        sources.add("ODBC")
                    elif "Excel.Workbook" in qd:
                        sources.add("Excel")
                    elif "Web.Contents" in qd:
                        sources.add("Web/JSON")
                    elif "#table(" in qd:
                        sources.add("Embedded (Import)")
                    else:
                        sources.add("Other M expression")
                results.append(f"    {tname}: {mode_str}")
            summary = f"Modes: {', '.join(sorted(modes))} | Sources: {', '.join(sorted(sources))}"
            return summary + "\n" + "\n".join(results)
        _check("Data sources & storage modes", check_data_sources)

        # 7. Per-table column breakdown
        def check_columns():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT t.Name, COUNT(*) as cols,
                         GROUP_CONCAT(DISTINCT CASE col.ExplicitDataType
                             WHEN 2 THEN 'String' WHEN 6 THEN 'Int64' WHEN 8 THEN 'Double'
                             WHEN 9 THEN 'DateTime' WHEN 10 THEN 'Decimal' WHEN 11 THEN 'Boolean'
                             ELSE 'Type' || col.ExplicitDataType END)
                         FROM [Column] col JOIN [Table] t ON col.TableID = t.ID
                         WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'
                         AND col.Type = 1
                         GROUP BY t.Name ORDER BY t.Name""")
            lines = []
            total_cols = 0
            for row in c.fetchall():
                tname, ncols, types = row
                total_cols += ncols
                lines.append(f"    {tname}: {ncols} columns ({types})")
            return f"{total_cols} total data columns\n" + "\n".join(lines)
        _check("Column breakdown", check_columns)

        # 8. VertiPaq table data (row counts from ColumnStorage metadata)
        def check_tables():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT t.Name,
                         MAX(cs.Statistics_RowCount) as row_count
                         FROM [Table] t
                         JOIN [Column] col ON col.TableID = t.ID
                         LEFT JOIN ColumnStorage cs ON cs.ColumnID = col.ID
                         WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'
                         AND col.Type = 1
                         GROUP BY t.Name ORDER BY t.Name""")
            lines = []
            total_rows = 0
            table_count = 0
            for row in c.fetchall():
                tname, rcount = row
                rcount = rcount or 0
                total_rows += rcount
                table_count += 1
                lines.append(f"    {tname}: {rcount:,} rows")
            return f"{table_count} tables, {total_rows:,} total rows\n" + "\n".join(lines)
        _check("VertiPaq data (row counts)", check_tables)

        # 8b. String dictionary invariants (DBCC-style, issue #43): VertiPaq's
        # Unique Value store is CASE-INSENSITIVE, so two stored entries that
        # differ only by case make Desktop reject the whole model with
        # "A duplicate value has been detected in the Unique Value store".
        # Files written by pbix-mcp < 0.9.88 could carry this; flag it here
        # instead of letting Desktop's load failure be the first symptom.
        def check_string_dictionaries():
            _init_datamodel()
            from pbix_mcp.formats.vertipaq_decoder import decode_dictionary
            offenders = []
            scanned = 0
            for f in abf_files:
                path = f.get("Path", "")
                if not path.endswith(".dictionary") or ".tbl\\" not in path:
                    continue
                start, size = f.get("m_cbOffsetHeader"), f.get("Size")
                if start is None or not size:
                    continue
                try:
                    dict_type, values = decode_dictionary(
                        abf_data[start : start + size])
                except Exception:
                    continue  # numeric/undecodable dictionaries: not this check
                if not values or not isinstance(values[0], str):
                    continue
                scanned += 1
                folded: dict = {}
                for v in values:
                    if not isinstance(v, str):
                        continue
                    k = v.casefold()
                    if k in folded and folded[k] != v:
                        offenders.append(
                            f"{path}: {folded[k]!r} vs {v!r}")
                    else:
                        folded.setdefault(k, v)
            if offenders:
                raise Exception(
                    "case-colliding entries (Desktop WILL refuse to load): "
                    + "; ".join(offenders[:5])
                    + (f" (+{len(offenders) - 5} more)"
                       if len(offenders) > 5 else ""))
            return f"{scanned} string dictionaries, no case-folded duplicates"
        _check("String dictionary invariants", check_string_dictionaries)

        # 9. Relationships
        def check_relationships():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT ft.Name, COALESCE(fc.ExplicitName, fc.InferredName),
                                tt.Name, COALESCE(tc.ExplicitName, tc.InferredName),
                                r.IsActive
                         FROM Relationship r
                         JOIN [Table] ft ON r.FromTableID = ft.ID
                         JOIN [Column] fc ON r.FromColumnID = fc.ID
                         JOIN [Table] tt ON r.ToTableID = tt.ID
                         JOIN [Column] tc ON r.ToColumnID = tc.ID""")
            lines = []
            for row in c.fetchall():
                active = "active" if row[4] else "inactive"
                lines.append(f"    {row[0]}.{row[1]} → {row[2]}.{row[3]} ({active})")
            if not lines:
                return "None"
            return f"{len(lines)} relationships\n" + "\n".join(lines)
        _check("Relationships", check_relationships)

        # 10. Measures
        def check_measures():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT t.Name, m.Name, m.Expression
                         FROM Measure m JOIN [Table] t ON m.TableID = t.ID""")
            lines = []
            for row in c.fetchall():
                # Desktop writes placeholder measures with Expression NULL
                # (same class as the 0.9.53 builder fix). len(None) crashed
                # this check on MS_Life_Expectancy -- a working file.
                raw = row[2] or ""
                expr = raw[:40] + "..." if len(raw) > 40 else raw
                lines.append(f"    [{row[0]}] {row[1]} = {expr}")
            if not lines:
                return "None"
            return f"{len(lines)} measures\n" + "\n".join(lines)
        _check("DAX measures", check_measures)

        # 11. RLS roles
        def check_rls():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("SELECT COUNT(*) FROM Role WHERE ModelID=1")
            count = c.fetchone()[0]
            return f"{count} roles" if count else "None"
        _check("Row-Level Security (RLS)", check_rls)

        # 12. Calculated tables — a calculated (DAX) table is Partition.Type=2
        # (DATATABLE / GENERATESERIES / CALENDAR / etc.). Partition.Type=4 is a
        # plain M/import partition and must NOT be counted here — doing so
        # mislabels every imported table as "calculated". Calculation-group
        # source tables (Type=7) are reported separately and excluded too.
        def check_calc():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT t.Name FROM [Table] t
                         JOIN Partition p ON p.TableID = t.ID
                         WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'
                         AND p.Type = 2""")
            calc = [r[0] for r in c.fetchall()]
            return f"{len(calc)} calculated tables ({', '.join(calc)})" if calc else "None"
        _check("Calculated tables", check_calc)

        # 13. Default slicer filters
        def check_filters():
            filters = _get_all_default_filters(info["work_dir"])
            if filters:
                return f"{len(filters)} default slicer filters"
            return "None"
        _check("Default slicer filters", check_filters)

        # 14. Tables without VertiPaq storage (metadata exists, ABF files missing)
        def check_tables_have_storage():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("""SELECT t.ID, t.Name FROM [Table] t
                         WHERE t.Name NOT LIKE 'H$%' AND t.Name NOT LIKE 'R$%'
                         AND t.ModelID = 1""")
            abf_paths = [f.get("Path", "") for f in abf_files] if abf_files else []
            abf_str = "\n".join(abf_paths)
            missing = []
            for row in c.fetchall():
                tid, tname = row
                # Match by ID, never by name. Storage folder names are NOT the
                # display name: special characters are space-sanitized
                # (fct_Orders -> "fct Orders", "# Measures" -> "  Measures")
                # and a RENAMED table keeps its creation-time folder
                # ("PositiveYOY-NegativeYOY" lives in "Table (2542).tbl").
                # The old name-based marker flagged 10 of GeoSales_Dashboard's
                # 14 tables as storage-less on a file Desktop opens fine.
                marker = f"({tid}).tbl"
                if marker not in abf_str:
                    missing.append(tname)
            if missing:
                raise ValueError(
                    f"{len(missing)} table(s) in metadata have NO VertiPaq storage — "
                    f"PBI will crash (TMCacheManager): {', '.join(missing)}"
                )
            return "All metadata tables have VertiPaq storage"
        _check("Table/storage consistency", check_tables_have_storage)

        # 15. Orphaned foreign key references
        def check_orphaned_refs():
            _init_datamodel()
            c = db_conn.cursor()
            issues = []
            # Table.RefreshPolicyID → RefreshPolicy.ID
            c.execute("""SELECT t.Name, t.RefreshPolicyID FROM [Table] t
                         WHERE t.RefreshPolicyID IS NOT NULL AND t.RefreshPolicyID != 0
                         AND t.RefreshPolicyID NOT IN (SELECT ID FROM RefreshPolicy)""")
            for row in c.fetchall():
                issues.append(f"Table '{row[0]}' → missing RefreshPolicy ID {row[1]}")
            # Table.CalculationGroupID → CalculationGroup.ID
            c.execute("""SELECT t.Name, t.CalculationGroupID FROM [Table] t
                         WHERE t.CalculationGroupID IS NOT NULL AND t.CalculationGroupID != 0
                         AND t.CalculationGroupID NOT IN (SELECT ID FROM CalculationGroup)""")
            for row in c.fetchall():
                issues.append(f"Table '{row[0]}' → missing CalculationGroup ID {row[1]}")
            # CalculationGroup.TableID → Table.ID
            c.execute("""SELECT cg.ID, cg.TableID FROM CalculationGroup cg
                         WHERE cg.TableID NOT IN (SELECT ID FROM [Table])""")
            for row in c.fetchall():
                issues.append(f"CalculationGroup {row[0]} → missing Table ID {row[1]}")
            if issues:
                raise ValueError(
                    f"{len(issues)} orphaned reference(s) — PBI will reject file:\n    "
                    + "\n    ".join(issues)
                )
            return "No orphaned references"
        _check("Metadata referential integrity", check_orphaned_refs)

        # 16. Expression.Kind must be a valid enum value.
        #
        # This check used to demand a DataMashup whenever Expression rows
        # exist, which FALSE-POSITIVED on 9 of the 24 corpus files (37.5%):
        # V3 "enhanced metadata" files (Version >= 1.28, incl. every
        # Service-downloaded report) legitimately store shared M expressions
        # in the model with no DataMashup part at all. The REAL trigger of
        # PFE_TM_ENUM_VALUES_VALIDATION_FAILED is an out-of-range enum:
        # TOM's ExpressionKind has a single member, M = 0, and every one of
        # the 25 Desktop/Service-authored Expression rows across the corpus
        # is Kind=0 -- including parameter queries. Kind=1 is what the old
        # incremental-refresh writer inserted, and what PBI rejects.
        def check_expressions():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("SELECT COUNT(*) FROM Expression")
            expr_count = c.fetchone()[0]
            c.execute(
                "SELECT Name, Kind FROM Expression WHERE Kind IS NOT NULL "
                "AND Kind != 0")
            bad = c.fetchall()
            if bad:
                names = ", ".join(f"{n!r} (Kind={k})" for n, k in bad[:5])
                raise ValueError(
                    f"{len(bad)} Expression row(s) with an invalid Kind — "
                    f"TOM's ExpressionKind enum only defines M=0, so PBI "
                    f"rejects the file with "
                    f"PFE_TM_ENUM_VALUES_VALIDATION_FAILED: {names}"
                )
            return (f"{expr_count} expressions, all Kind=0"
                    if expr_count else "None")
        _check("Expression Kind validity", check_expressions)

        # 17. MAXID consistency
        def check_maxid():
            _init_datamodel()
            c = db_conn.cursor()
            c.execute("SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'")
            row = c.fetchone()
            if not row:
                raise ValueError("MAXID not found in DBPROPERTIES")
            maxid = int(row[0])
            # Find actual max ID across all object tables
            actual_max = 0
            for tbl in ("Table", "Column", "Measure", "Partition", "Relationship",
                        "Role", "CalculationGroup", "CalculationItem"):
                try:
                    c.execute(f"SELECT MAX(ID) FROM [{tbl}]")
                    r = c.fetchone()
                    if r and r[0]:
                        actual_max = max(actual_max, r[0])
                except Exception:
                    pass
            if maxid < actual_max:
                raise ValueError(
                    f"MAXID={maxid} but highest object ID is {actual_max} — "
                    f"PBI will crash with TMCCollectionObject::Add assertion"
                )
            return f"MAXID={maxid} (highest ID={actual_max})"
        _check("MAXID consistency", check_maxid)

        # ---- Report-definition integrity ----
        #
        # Each of these corresponds to a defect class found by auditing all 125
        # tools: state that a tool wrote but that never reached disk, a
        # reference that no longer resolves, or a classic-shaped value written
        # into a PBIR document. A report can be perfectly schema-valid and
        # still fail every one of them.
        for name, fn in _report_integrity_checks(info["work_dir"]):
            _check(name, fn)

    finally:
        # Clean up shared resources
        if db_conn:
            db_conn.close()
        if db_tmp_path and os.path.exists(db_tmp_path):
            os.unlink(db_tmp_path)

    return ToolResponse.ok("\n".join(checks)).to_text()


def _rebuild_path_dry_run(work_dir: str) -> tuple:
    """Would a rebuild-path edit succeed on this model?

    Returns ``(ok, reason, preserved_count)``. Runs the SAME planner the edit
    runs, so the answer cannot drift from reality — the only way to keep a
    diagnostic honest is to make it execute the real predicate rather than a
    restatement of it.
    """
    import sqlite3 as _sqlite3
    import tempfile

    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    dm = os.path.join(work_dir, "DataModel")
    try:
        with open(dm, "rb") as f:
            abf = decompress_datamodel(f.read())
        meta = read_metadata_sqlite(abf)
    except Exception as exc:
        return False, f"the DataModel could not be read ({exc}).", 0

    fd, tmp = tempfile.mkstemp(suffix=".db", dir=work_dir)
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(meta)
        conn = _sqlite3.connect(tmp)
        conn.row_factory = _sqlite3.Row
        try:
            planned, col_restamp, table_restamp = _plan_calc_preservation(
                conn, abf, meta, _relationships_from_metadata(conn))
        finally:
            conn.close()
    except ValueError as exc:
        return False, str(exc), 0
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}", 0
    finally:
        try:
            os.remove(tmp)
        except OSError:
            pass
    return True, "", len(col_restamp) + len(table_restamp)


class _DoctorWarning(Exception):
    """Something Power BI tolerates but a human probably wants to know about.

    Kept distinct from a hard failure so that a report which merely carries a
    stale reference is not reported as broken — Microsoft's own AI sample ships
    with bookmarks pointing at pages that were deleted.
    """


def _doctor_report_docs(work_dir: str) -> tuple:
    """(layout, report_config, is_pbir) for the integrity checks."""
    return (_get_layout(work_dir), _get_report_config(work_dir) or {},
            _is_pbir(work_dir))


def _report_integrity_checks(work_dir: str) -> list:
    """Report-definition checks, as (name, callable) pairs.

    Every callable returns a human-readable string on success and RAISES on a
    problem — matching the `_check` contract in pbix_doctor.
    """
    layout, rcfg, is_pbir = _doctor_report_docs(work_dir)
    sections = (layout or {}).get("sections", []) or []
    static_root = os.path.join(work_dir, "Report", "StaticResources",
                               "RegisteredResources")

    def _declared_resources() -> set:
        out = set()
        for pkg in (rcfg.get("resourcePackages") or []):
            inner = pkg.get("resourcePackage", pkg)
            if inner.get("name") != "RegisteredResources":
                continue
            for item in (inner.get("items") or []):
                out.add(item.get("path") or item.get("name") or "")
        return {x for x in out if x}

    def _on_disk_resources() -> set:
        if not os.path.isdir(static_root):
            return set()
        return {f for f in os.listdir(static_root)
                if os.path.isfile(os.path.join(static_root, f))}

    def check_resources():
        declared, on_disk = _declared_resources(), _on_disk_resources()
        # Microsoft: "Every resource file must have a corresponding entry in
        # the report.json file." An undeclared file simply never renders.
        undeclared = sorted(on_disk - declared)
        missing = sorted(declared - on_disk)
        if undeclared or missing:
            bits = []
            if undeclared:
                bits.append(f"{len(undeclared)} file(s) present but NOT "
                            f"declared (will not render): {undeclared[:4]}")
            if missing:
                bits.append(f"{len(missing)} declared but MISSING from the "
                            f"package: {missing[:4]}")
            raise ValueError("; ".join(bits))
        return f"{len(declared)} registered resource(s), all consistent"

    def check_custom_visuals():
        # A custom visual can be registered three ways, and a report is only
        # broken if it uses one that is registered by NONE of them:
        #   1. publicCustomVisuals  — AppSource visuals, by GUID
        #   2. Report/CustomVisuals/<type>/  — a private .pbiviz embedded here
        #   3. the tenant's organizational store — invisible in the file
        # Checking only (1) reports every embedded private visual as broken,
        # which is most of Microsoft's own samples.
        declared = set(rcfg.get("publicCustomVisuals") or [])
        cv_root = os.path.join(work_dir, "Report", "CustomVisuals")
        embedded = set(os.listdir(cv_root)) if os.path.isdir(cv_root) else set()
        registered = declared | embedded

        used = set()
        for sec in sections:
            for vc in (sec.get("visualContainers") or []):
                vt = _get_visual_type(_parse_visual_config(vc))
                if vt and vt != "unknown":
                    used.add(vt)
        # Built-ins are plain camelCase; a custom visual carries a GUID-ish or
        # timestamp suffix.
        custom_used = {v for v in used
                       if re.search(r"[0-9A-F]{16,}$", v, re.I)
                       or re.search(r"_[0-9A-F]{8}(_[0-9A-F]{4}){3}_[0-9A-F]{12}$", v, re.I)
                       or re.search(r"\d{10,}$", v)}
        unresolved = sorted(custom_used - registered)
        if unresolved:
            # Could still be an org-store visual, which no file inspection can
            # confirm — so this is a warning, not a failure.
            raise _DoctorWarning(
                f"{len(unresolved)} custom visual type(s) used but registered "
                f"neither in publicCustomVisuals nor Report/CustomVisuals/ — "
                f"they will only load if your tenant's organizational store "
                f"provides them: {unresolved[:3]}")
        return (f"{len(custom_used)} custom visual(s) in use, all registered "
                f"({len(declared)} AppSource, {len(embedded)} embedded)")

    def check_page_visual_names():
        # Page names must be globally unique; VISUAL names only need to be
        # unique within their page, because bookmarks address them as
        # sections[<page>].visualContainers[<visual>]. Microsoft's own AI
        # sample reuses 74 visual names across pages, which is fine.
        seen_pages, dup_pages = set(), []
        total_vis, dup_vis, unnamed = 0, [], 0
        for sec in sections:
            pn = sec.get("name")
            if pn in seen_pages:
                dup_pages.append(pn)
            seen_pages.add(pn)
            page_vis = set()
            for vc in (sec.get("visualContainers") or []):
                vn = _parse_visual_config(vc).get("name")
                total_vis += 1
                if not vn:
                    unnamed += 1
                    continue
                if vn in page_vis:
                    dup_vis.append(f"{sec.get('displayName')}:{vn}")
                page_vis.add(vn)
        problems = []
        if dup_pages:
            problems.append(f"{len(dup_pages)} duplicate page name(s) "
                            f"{dup_pages[:3]} — bookmarks and page navigation "
                            f"resolve by name")
        if dup_vis:
            problems.append(f"{len(dup_vis)} visual name(s) duplicated WITHIN "
                            f"a page {dup_vis[:3]}")
        if unnamed:
            problems.append(f"{unnamed} visual(s) with no name — they cannot "
                            f"be addressed by any tool")
        if problems:
            raise ValueError("; ".join(problems))
        return (f"{len(seen_pages)} page(s), {total_vis} visual(s), "
                f"names unique where they must be")

    def check_bookmarks():
        raw = (layout or {}).get("config") or "{}"
        try:
            cfg = json.loads(raw) if isinstance(raw, str) else (raw or {})
        except (json.JSONDecodeError, TypeError):
            cfg = {}
        bookmarks = cfg.get("bookmarks") or []
        if not bookmarks:
            return "none"
        page_names = {s.get("name") for s in sections}
        vis_names = {_parse_visual_config(vc).get("name")
                     for s in sections for vc in (s.get("visualContainers") or [])}
        bad = []

        def _walk(bm):
            for child in (bm.get("children") or []):
                _walk(child)
            state = bm.get("explorationState") or {}
            active = state.get("activeSection")
            if active and active not in page_names:
                bad.append(f"{bm.get('displayName')!r} -> missing page {active}")
            for sec_name, sec_state in (state.get("sections") or {}).items():
                if sec_name not in page_names:
                    bad.append(f"{bm.get('displayName')!r} -> missing page {sec_name}")
                for vn in (sec_state.get("visualContainers") or {}):
                    if vn not in vis_names:
                        bad.append(f"{bm.get('displayName')!r} -> missing visual {vn}")

        for bm in bookmarks:
            _walk(bm)
        if bad:
            # Power BI ignores a bookmark step that points at something gone,
            # so this is a stale reference rather than a broken report.
            raise _DoctorWarning(
                f"{len(bad)} bookmark reference(s) point at a page or visual "
                f"that no longer exists — those steps will do nothing: "
                f"{bad[:3]}")
        return f"{len(bookmarks)} bookmark(s), all references resolve"

    def check_rebuild_eligibility():
        """Whether the model can take a rebuild-path edit (add/replace a table,
        add/remove a relationship, remove a table).

        Worth knowing UP FRONT rather than from a failed call — but only if it
        is TRUE. This runs the real preservation planner as a dry run instead
        of approximating it. An earlier version inferred the answer from the
        presence of auto date/time tables and told two corpus reports
        "supported" when the edit then refused for an unrelated reason; a
        diagnostic that lies is worse than no diagnostic.
        """
        dm = os.path.join(work_dir, "DataModel")
        if not os.path.exists(dm):
            return "no DataModel (report-only file)"
        ok, reason, preserved = _rebuild_path_dry_run(work_dir)
        if not ok:
            raise _DoctorWarning(
                f"rebuild-path edits (add/replace a table, add/remove a "
                f"relationship, remove a table) will be REFUSED — {reason} "
                f"The surgical tools never rebuild and still work: "
                f"add_measure / modify_measure / remove_measure / "
                f"modify_column / set_sort_by_column.")
        return ("rebuild-path edits supported"
                + (f" ({preserved} calculated object(s) preserved)"
                   if preserved else ""))

    def check_rebuild_cost():
        """What a rebuild-path edit would COST this model, before making one.

        The edit reports what it could not carry AFTERWARDS; this says so
        beforehand, which is when the choice is still available. It lists the
        model features a rebuild has to re-create by name rather than copy —
        the ones whose loss a user would actually notice.
        """
        dm = os.path.join(work_dir, "DataModel")
        if not os.path.exists(dm):
            return "no DataModel (report-only file)"
        try:
            from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
            from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel
            with open(dm, "rb") as f:
                abf = decompress_datamodel(f.read())
            fd, tmp = tempfile.mkstemp(suffix=".db")
            os.write(fd, read_metadata_sqlite(abf))
            os.close(fd)
            try:
                conn = sqlite3.connect(tmp)
                counts = {}
                for tname, _fks, _ident in _CARRY_SPEC:
                    try:
                        n = conn.execute(
                            f"SELECT COUNT(*) FROM [{tname}]").fetchone()[0]
                    except sqlite3.Error:
                        continue
                    if n:
                        counts[tname] = n
                conn.close()
            finally:
                os.unlink(tmp)
        except Exception as exc:
            return f"could not inspect ({type(exc).__name__})"
        if not counts:
            return "a rebuild-path edit would carry nothing extra"
        by_meaning: dict[str, int] = {}
        for tname, n in counts.items():
            key = _CARRY_MEANING.get(tname, tname)
            by_meaning[key] = by_meaning.get(key, 0) + n
        listing = ", ".join(f"{k} ({v})" for k, v in
                            sorted(by_meaning.items(), key=lambda kv: -kv[1]))
        return (f"a rebuild-path edit re-creates by name: {listing}. Anything "
                f"it cannot re-attach is reported in the response warnings.")

    checks = [
        ("Registered resources", check_resources),
        ("Custom visual registration", check_custom_visuals),
        ("Page / visual naming", check_page_visual_names),
        ("Bookmark references", check_bookmarks),
        ("Rebuild-path eligibility", check_rebuild_eligibility),
        ("Rebuild-path cost", check_rebuild_cost),
    ]
    if is_pbir:
        checks += _pbir_integrity_checks(work_dir, layout, sections)
    return checks


def _pbir_integrity_checks(work_dir: str, layout: dict, sections: list) -> list:
    """Checks that only apply to the PBIR tree."""
    pages_dir = os.path.join(work_dir, "Report", "definition", "pages")

    def check_page_tree():
        meta_path = os.path.join(pages_dir, "pages.json")
        meta = {}
        if os.path.exists(meta_path):
            with open(meta_path, encoding="utf-8-sig") as f:
                meta = json.load(f)
        folders = {d for d in os.listdir(pages_dir)
                   if os.path.isdir(os.path.join(pages_dir, d))}
        order = list(meta.get("pageOrder") or [])
        problems = []
        missing = [p for p in order if p not in folders]
        extra = sorted(folders - set(order))
        if missing:
            problems.append(f"pageOrder names {len(missing)} folder(s) that "
                            f"do not exist: {missing[:3]}")
        if extra:
            problems.append(f"{len(extra)} page folder(s) absent from "
                            f"pageOrder (they will not appear): {extra[:3]}")
        active = meta.get("activePageName")
        if active and active not in folders:
            problems.append(f"activePageName {active!r} does not exist")
        for folder in folders:
            pj = os.path.join(pages_dir, folder, "page.json")
            if not os.path.exists(pj):
                problems.append(f"page folder {folder!r} has no page.json")
                continue
            with open(pj, encoding="utf-8-sig") as f:
                doc = json.load(f)
            if doc.get("name") != folder:
                problems.append(
                    f"page.json name {doc.get('name')!r} != folder {folder!r}")
        if problems:
            raise ValueError("; ".join(problems))
        return f"{len(folders)} page folder(s), order and names consistent"

    def check_visual_tree():
        problems = []
        count = 0
        for folder in os.listdir(pages_dir):
            vdir = os.path.join(pages_dir, folder, "visuals")
            if not os.path.isdir(vdir):
                continue
            for vid in os.listdir(vdir):
                vj = os.path.join(vdir, vid, "visual.json")
                if not os.path.exists(vj):
                    continue
                count += 1
                with open(vj, encoding="utf-8-sig") as f:
                    doc = json.load(f)
                if doc.get("name") != vid:
                    problems.append(
                        f"visual.json name {doc.get('name')!r} != folder {vid!r}")
        if problems:
            raise ValueError(
                f"{len(problems)} mismatch(es): {problems[:3]}")
        return f"{count} visual file(s), names match their folders"

    def check_naming_convention():
        # PBIR: "The object name and/or file/folder name must consist of one or
        # more word characters (letters, digits, underscores) or hyphens."
        bad = []
        for base, dirs, _files in os.walk(
                os.path.join(work_dir, "Report", "definition")):
            for d in dirs:
                if not re.fullmatch(r"[\w-]+", d):
                    bad.append(d)
        if bad:
            raise ValueError(
                f"{len(bad)} folder name(s) violate the PBIR naming rule "
                f"(Desktop ignores them as private files): {bad[:3]}")
        return "all folder names comply"

    def check_no_classic_leaks():
        # Classic spellings inside a PBIR page/visual mean the converter wrote
        # the legacy shape. Legitimate inside a bookmark's explorationState,
        # which models captured state in the classic vocabulary.
        leaks = []
        for base, _dirs, files in os.walk(
                os.path.join(work_dir, "Report", "definition")):
            for fn in files:
                if fn not in ("page.json", "visual.json"):
                    continue
                p = os.path.join(base, fn)
                with open(p, encoding="utf-8-sig") as f:
                    text = f.read()
                for token in ("singleVisual", "vcObjects", "prototypeQuery",
                              "__pbir_"):
                    if token in text:
                        leaks.append(f"{fn}:{token}")
        if leaks:
            raise ValueError(
                f"{len(leaks)} classic-shaped key(s) in PBIR files: "
                f"{sorted(set(leaks))[:4]}")
        return "no classic-shaped keys in page/visual files"

    def check_enum_fields():
        # The displayOption class: PBIR types these as string enums; writing
        # the classic int produces a file that imports and then will not open.
        bad = []
        for base, _dirs, files in os.walk(
                os.path.join(work_dir, "Report", "definition")):
            for fn in files:
                if fn != "page.json":
                    continue
                with open(os.path.join(base, fn), encoding="utf-8-sig") as f:
                    doc = json.load(f)
                for field, allowed in _PBIR_ENUM_FIELDS["page.json"].items():
                    if field in doc and not (
                            isinstance(doc[field], str) and doc[field] in allowed):
                        bad.append(f"{doc.get('displayName', '?')}.{field}"
                                   f"={doc[field]!r}")
        if bad:
            raise ValueError(
                f"{len(bad)} enum field(s) carrying a non-enum value "
                f"(Power BI will refuse to open the report): {bad[:3]}")
        return "enum fields all carry valid enum names"

    return [
        ("PBIR page tree", check_page_tree),
        ("PBIR visual tree", check_visual_tree),
        ("PBIR naming convention", check_naming_convention),
        ("PBIR classic-shape leaks", check_no_classic_leaks),
        ("PBIR enum fields", check_enum_fields),
    ]


# ---- Section 10b: TMDL Export ----


def _tmdl_escape(value: str) -> str:
    """Escape a quoted TMDL name: single quotes are doubled (TMDL's own
    escape), which parse_tmdl_document's reader reverses."""
    if not value:
        return ""
    return value.replace("'", "''")


def _tmdl_quote(name: str) -> str:
    """Quote a TMDL object name only when it needs it (non-identifier chars)."""
    import re as _re
    if _re.match(r"^[A-Za-z_][A-Za-z0-9_]*$", name or ""):
        return name
    return f"'{_tmdl_escape(name)}'"


def _tmdl_emit_expr(lines: list, indent: str, header: str, expr: str) -> None:
    """Emit ``header = expr`` inline for a single-line expression, or the
    TMDL block form (``header =`` + lines indented two levels deeper) for a
    multi-line one — matching Desktop's own TMDL writer, and what
    tmdl_reader.parse_tmdl_document reads back verbatim."""
    expr = (expr or "").strip()
    if "\n" in expr:
        lines.append(f"{indent}{header} =")
        for el in expr.split("\n"):
            lines.append(f"{indent}\t\t{el}" if el.strip() else "")
    else:
        lines.append(f"{indent}{header} = {expr}")


def _tmdl_extended_properties(c, object_type: int, object_id: int,
                              lines: list, indent: str) -> None:
    """Emit extendedProperty blocks for one object (field parameters live
    here: ParameterMetadata on the Fields column, ObjectType=4/Type=json)."""
    try:
        rows = c.execute(
            "SELECT Name, Type, Value FROM ExtendedProperty "
            "WHERE ObjectType = ? AND ObjectID = ? ORDER BY ID",
            (object_type, object_id)).fetchall()
    except sqlite3.OperationalError:
        return  # older metadata without the table
    for ep_name, _ep_type, ep_value in rows:
        _tmdl_emit_expr(lines, indent,
                        f"extendedProperty {_tmdl_quote(ep_name)}",
                        ep_value or "")


def _export_tmdl_from_sqlite(conn: sqlite3.Connection, output_dir: str) -> dict:
    """Export metadata SQLite to TMDL folder structure.

    Returns dict with counts of exported objects.
    """
    c = conn.cursor()
    stats = {"tables": 0, "columns": 0, "measures": 0, "relationships": 0,
             "roles": 0, "hierarchies": 0}

    # ---- database.tmdl ----
    c.execute("SELECT Name, Culture FROM Model LIMIT 1")
    model_row = c.fetchone()
    db_name = model_row[0] if model_row else "Model"
    compat = 1567  # Default PBI compatibility level

    with open(os.path.join(output_dir, "database.tmdl"), "w", encoding="utf-8") as f:
        f.write(f"database {db_name}\n")
        f.write(f"\tcompatibilityLevel: {compat}\n")

    # ---- model.tmdl ----
    culture = model_row[1] if model_row and model_row[1] else "en-US"
    c.execute(
        "SELECT DefaultPowerBIDataSourceVersion, DiscourageImplicitMeasures, "
        "SourceQueryCulture, DataAccessOptions FROM Model LIMIT 1"
    )
    model_props = c.fetchone()

    with open(os.path.join(output_dir, "model.tmdl"), "w", encoding="utf-8") as f:
        f.write("model Model\n")
        f.write(f"\tculture: {culture}\n")
        if model_props:
            dsv = model_props[0]
            # DefaultPowerBIDataSourceVersion: 2 = powerBI_V3
            dsv_map = {1: "powerBI_V1", 2: "powerBI_V3"}
            if dsv in dsv_map:
                f.write(f"\tdefaultPowerBIDataSourceVersion: {dsv_map[dsv]}\n")
            if model_props[1]:
                f.write("\tdiscourageImplicitMeasures\n")
            sqc = model_props[2]
            if sqc:
                f.write(f"\tsourceQueryCulture: {sqc}\n")
            dao = model_props[3]
            if dao:
                import json as _json
                try:
                    dao_obj = _json.loads(dao)
                    if dao_obj:
                        f.write("\tdataAccessOptions\n")
                        if dao_obj.get("legacyRedirects"):
                            f.write("\t\tlegacyRedirects\n")
                        if dao_obj.get("returnErrorValuesAsNull"):
                            f.write("\t\treturnErrorValuesAsNull\n")
                except Exception:
                    pass

    # ---- expressions.tmdl (shared M parameters) ----
    c.execute(
        "SELECT Name, Expression, Description, LineageTag FROM Expression "
        "WHERE ModelID = 1 ORDER BY ID"
    )
    exprs = c.fetchall()
    if exprs:
        lines = []
        for e_name, e_expr, e_desc, e_tag in exprs:
            if e_expr:
                _tmdl_emit_expr(lines, "",
                                f"expression {_tmdl_quote(e_name)}", e_expr)
                # Note: TMDL expression objects do not support 'description'
                if e_tag:
                    lines.append(f"\tlineageTag: {e_tag}")
                lines.append("")
        if lines:
            with open(os.path.join(output_dir, "expressions.tmdl"), "w", encoding="utf-8") as f:
                f.write("\n".join(lines) + "\n")

    # ---- tables/ ----
    tables_dir = os.path.join(output_dir, "tables")
    os.makedirs(tables_dir, exist_ok=True)

    c.execute("SELECT ID, Name, Description, IsHidden, LineageTag "
              "FROM [Table] ORDER BY ID")
    tables = c.fetchall()

    for table_id, table_name, table_desc, is_hidden, t_tag in tables:
        # Skip internal system tables (H$=hierarchy, R$=relationship, U$=user hierarchy)
        if table_name.startswith(("H$", "R$", "U$")):
            continue
        lines = [f"table '{_tmdl_escape(table_name)}'"]
        # Note: PBI Desktop's TMDL parser rejects 'description' on tables
        # even though the TOM model supports it — skip to avoid load errors
        if is_hidden:
            lines.append("\tisHidden")
        if t_tag:
            lines.append(f"\tlineageTag: {t_tag}")
        _tmdl_extended_properties(c, 3, table_id, lines, "\t")
        lines.append("")

        # Columns. SummarizeBy is emitted only when it differs from the
        # type-default (numeric -> default(1), other -> none(2)) — exactly the
        # default the import side's builder re-derives, which keeps
        # export -> import -> export byte-stable while preserving overrides.
        c.execute(
            "SELECT ID, ExplicitName, InferredName, ExplicitDataType, InferredDataType, "
            "IsHidden, IsKey, SourceColumn, Expression, FormatString, Description, Type, "
            "LineageTag, SummarizeBy, DataCategory, DisplayFolder, SortByColumnID "
            "FROM [Column] WHERE TableID = ? ORDER BY ID",
            (table_id,)
        )
        _dtype_map = {
            2: "string", 6: "int64", 8: "double", 9: "dateTime", 10: "decimal", 11: "boolean"
        }
        from pbix_mcp.formats.tmdl_reader import SUMMARIZE_BY_NAMES
        col_rows = c.fetchall()
        id2name = {r[0]: (r[1] or r[2] or "?") for r in col_rows}
        for col in col_rows:
            (col_id, expl_name, inf_name, expl_dt, inf_dt, is_col_hidden,
             is_key, source_col, expression, fmt_str, col_desc, col_type,
             c_tag, summarize_by, data_cat, disp_folder, sortby_id) = col
            col_name = expl_name or inf_name or "?"
            dtype_id = expl_dt if expl_dt else (inf_dt if inf_dt else 2)
            dtype = _dtype_map.get(dtype_id, "string")

            if col_type == 3:
                continue  # Skip RowNumber system columns

            if expression and col_type == 2:
                _tmdl_emit_expr(lines, "\t",
                                f"column '{_tmdl_escape(col_name)}'", expression)
            else:
                lines.append(f"\tcolumn '{_tmdl_escape(col_name)}'")

            lines.append(f"\t\tdataType: {dtype}")
            if source_col:
                lines.append(f"\t\tsourceColumn: {source_col}")
            if is_col_hidden:
                lines.append("\t\tisHidden")
            if is_key:
                lines.append("\t\tisKey")
            if fmt_str:
                lines.append(f"\t\tformatString: {fmt_str}")
            type_default_sb = 1 if dtype_id in (6, 8, 10) else 2
            if summarize_by and summarize_by != type_default_sb \
                    and summarize_by in SUMMARIZE_BY_NAMES:
                lines.append(f"\t\tsummarizeBy: {SUMMARIZE_BY_NAMES[summarize_by]}")
            if sortby_id and sortby_id in id2name:
                lines.append(
                    f"\t\tsortByColumn: '{_tmdl_escape(id2name[sortby_id])}'")
            if data_cat:
                lines.append(f"\t\tdataCategory: {data_cat}")
            if disp_folder:
                lines.append(f"\t\tdisplayFolder: {disp_folder}")
            if c_tag:
                lines.append(f"\t\tlineageTag: {c_tag}")
            _tmdl_extended_properties(c, 4, col_id, lines, "\t\t")
            # Note: PBI Desktop's TMDL parser rejects 'description' on columns
            lines.append("")
            stats["columns"] += 1

        # Hierarchies (user drill-down hierarchies with their levels)
        try:
            h_rows = c.execute(
                "SELECT ID, Name, IsHidden, DisplayFolder, LineageTag "
                "FROM Hierarchy WHERE TableID = ? ORDER BY ID",
                (table_id,)).fetchall()
        except sqlite3.OperationalError:
            h_rows = []
        for h_id, h_name, h_hidden, h_folder, h_tag in h_rows:
            lines.append(f"\thierarchy '{_tmdl_escape(h_name)}'")
            if h_hidden:
                lines.append("\t\tisHidden")
            if h_folder:
                lines.append(f"\t\tdisplayFolder: {h_folder}")
            if h_tag:
                lines.append(f"\t\tlineageTag: {h_tag}")
            lines.append("")
            for lv_name, lv_col_id, lv_tag in c.execute(
                    "SELECT Name, ColumnID, LineageTag FROM Level "
                    "WHERE HierarchyID = ? ORDER BY Ordinal", (h_id,)):
                lines.append(f"\t\tlevel {_tmdl_quote(lv_name)}")
                if lv_tag:
                    lines.append(f"\t\t\tlineageTag: {lv_tag}")
                lines.append(
                    f"\t\t\tcolumn: '{_tmdl_escape(id2name.get(lv_col_id, ''))}'")
                lines.append("")
            stats["hierarchies"] += 1

        # Measures
        c.execute(
            "SELECT ID, Name, Expression, FormatString, Description, IsHidden, "
            "DisplayFolder, DataCategory, LineageTag "
            "FROM Measure WHERE TableID = ? ORDER BY ID",
            (table_id,)
        )
        for meas in c.fetchall():
            (m_id, m_name, m_expr, m_fmt, m_desc, m_hidden, m_folder,
             m_cat, m_tag) = meas
            _tmdl_emit_expr(lines, "\t",
                            f"measure '{_tmdl_escape(m_name)}'", m_expr or "")
            if m_fmt:
                lines.append(f"\t\tformatString: {m_fmt}")
            # Note: PBI Desktop's TMDL parser rejects 'description' on measures
            if m_hidden:
                lines.append("\t\tisHidden")
            if m_folder:
                lines.append(f"\t\tdisplayFolder: {m_folder}")
            if m_cat:
                lines.append(f"\t\tdataCategory: {m_cat}")
            if m_tag:
                lines.append(f"\t\tlineageTag: {m_tag}")
            _tmdl_extended_properties(c, 8, m_id, lines, "\t\t")
            lines.append("")
            stats["measures"] += 1

        # Partitions
        c.execute(
            "SELECT Name, QueryDefinition, Mode, Type FROM [Partition] "
            "WHERE TableID = ? ORDER BY ID",
            (table_id,)
        )
        for part in c.fetchall():
            p_name, p_query, p_mode, p_type = part
            if p_query:
                mode_str = "directQuery" if p_mode == 1 else "import"
                if p_type == 2:
                    # Calculated partition (DAX)
                    lines.append(f"\tpartition '{_tmdl_escape(p_name)}' = calculated")
                else:
                    # Type 4 = M (Power Query), default
                    lines.append(f"\tpartition '{_tmdl_escape(p_name)}' = m")
                    lines.append(f"\t\tmode: {mode_str}")
                lines.append("\t\tsource =")
                for qline in p_query.strip().split("\n"):
                    lines.append(f"\t\t\t\t{qline}" if qline.strip() else "")
                lines.append("")

        # Write table TMDL
        safe_name = table_name.replace("/", "_").replace("\\", "_").replace(":", "_")
        with open(os.path.join(tables_dir, f"{safe_name}.tmdl"), "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
        stats["tables"] += 1

    # ---- relationships.tmdl ----
    c.execute(
        "SELECT r.Name, r.IsActive, r.CrossFilteringBehavior, "
        "ft.Name, COALESCE(fc.ExplicitName, fc.InferredName), "
        "tt.Name, COALESCE(tc.ExplicitName, tc.InferredName), "
        "r.FromCardinality, r.ToCardinality "
        "FROM [Relationship] r "
        "JOIN [Table] ft ON r.FromTableID = ft.ID "
        "JOIN [Column] fc ON r.FromColumnID = fc.ID "
        "JOIN [Table] tt ON r.ToTableID = tt.ID "
        "JOIN [Column] tc ON r.ToColumnID = tc.ID "
        "ORDER BY r.ID"
    )
    rels = c.fetchall()
    if rels:
        _card_names = {0: "none", 1: "one", 2: "many"}
        lines = []
        for rel in rels:
            (r_name, is_active, cross_filter, from_tbl, from_col, to_tbl,
             to_col, from_card, to_card) = rel
            lines.append(f"relationship {r_name or ''}")
            lines.append(f"\tfromColumn: '{_tmdl_escape(from_tbl)}'.'{_tmdl_escape(from_col)}'")
            lines.append(f"\ttoColumn: '{_tmdl_escape(to_tbl)}'.'{_tmdl_escape(to_col)}'")
            if not is_active:
                lines.append("\tisActive: false")
            # TOM CrossFilteringBehavior: 1=OneDirection (default), 2=BothDirections, 3=Automatic
            # Only emit non-default values in TMDL
            cfb_map = {2: "bothDirections", 3: "automatic"}
            if cross_filter in cfb_map:
                lines.append(f"\tcrossFilteringBehavior: {cfb_map[cross_filter]}")
            # Cardinality: many->one is TMDL's default; emit only deviations
            # (many-to-many, one-to-one) so plain star schemas stay terse.
            if from_card is not None and from_card != 2:
                lines.append(f"\tfromCardinality: {_card_names.get(from_card, 'many')}")
            if to_card is not None and to_card != 1:
                lines.append(f"\ttoCardinality: {_card_names.get(to_card, 'one')}")
            lines.append("")
            stats["relationships"] += 1

        with open(os.path.join(output_dir, "relationships.tmdl"), "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")

    # ---- roles/ ----
    c.execute("SELECT ID, Name, Description FROM Role ORDER BY ID")
    roles = c.fetchall()
    if roles:
        roles_dir = os.path.join(output_dir, "roles")
        os.makedirs(roles_dir, exist_ok=True)
        for role_id, role_name, role_desc in roles:
            lines = [f"role '{_tmdl_escape(role_name)}'"]
            # Note: PBI Desktop's TMDL parser rejects 'description' on roles

            c.execute(
                "SELECT t.Name, tp.FilterExpression FROM TablePermission tp "
                "JOIN [Table] t ON tp.TableID = t.ID "
                "WHERE tp.RoleID = ? ORDER BY tp.ID",
                (role_id,)
            )
            for tbl_name, filter_expr in c.fetchall():
                lines.append(f"\ttablePermission '{_tmdl_escape(tbl_name)}'")
                if filter_expr:
                    lines.append(f"\t\tfilterExpression: {filter_expr}")
            lines.append("")

            safe_name = role_name.replace("/", "_").replace("\\", "_")
            with open(os.path.join(roles_dir, f"{safe_name}.tmdl"), "w", encoding="utf-8") as f:
                f.write("\n".join(lines) + "\n")
            stats["roles"] += 1

    return stats


@mcp.tool()
def pbix_set_incremental_refresh(
    alias: str,
    table_name: str,
    archive_periods: int = 36,
    archive_granularity: str = "month",
    refresh_periods: int = 12,
    refresh_granularity: str = "month",
    detect_changes_column: str = "",
    mode: str = "import",
) -> str:
    """Configure incremental refresh policy for a table.

    Incremental refresh partitions a table by date range so only recent
    data is refreshed, dramatically reducing refresh time for large datasets.

    Requires the table's M expression to filter on RangeStart/RangeEnd parameters.
    These DateTime parameters are automatically created if they don't exist.

    Args:
        alias: The alias of the open file
        table_name: Table to apply the refresh policy to
        archive_periods: Number of periods to keep as historical (default 36)
        archive_granularity: Granularity for archive window — "day", "month",
                             "quarter", or "year" (default "month")
        refresh_periods: Number of periods to refresh each time (default 12)
        refresh_granularity: Granularity for refresh window — "day", "month",
                             "quarter", or "year" (default "month")
        detect_changes_column: Optional column name for change detection
                               (e.g. "ModifiedDate"). If set, only partitions
                               where this column changed will be refreshed.
        mode: "import" (default) or "hybrid". Hybrid adds a DirectQuery
              partition for real-time data on top of import partitions.
    """
    try:
        _GRAN_MAP = {"day": 1, "month": 2, "quarter": 3, "year": 4}
        _MODE_MAP = {"import": 0, "hybrid": 1}

        if archive_granularity not in _GRAN_MAP:
            raise ValueError(f"archive_granularity must be one of {list(_GRAN_MAP.keys())}")
        if refresh_granularity not in _GRAN_MAP:
            raise ValueError(f"refresh_granularity must be one of {list(_GRAN_MAP.keys())}")
        if mode not in _MODE_MAP:
            raise ValueError("mode must be 'import' or 'hybrid'")

        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Incremental refresh requires a DataMashup with M expressions that
        # filter on RangeStart/RangeEnd.  Without it, PBI rejects the file.
        mashup_path = os.path.join(info["work_dir"], "DataMashup")
        if not os.path.exists(mashup_path):
            return ToolResponse.error(
                "Incremental refresh requires a DataMashup section with M expressions "
                "that filter on RangeStart/RangeEnd parameters. This file has no "
                "DataMashup (it uses embedded data). Use source_csv or source_db when "
                "creating tables to enable incremental refresh.",
                "INVALID_OPERATION"
            ).to_text()

        policy_info = {}

        def _do_set(conn: sqlite3.Connection):
            c = conn.cursor()

            # Find table
            c.execute("SELECT ID FROM [Table] WHERE Name = ?", (table_name,))
            trow = c.fetchone()
            if not trow:
                raise ValueError(f"Table '{table_name}' not found")
            table_id = trow[0]

            # RangeStart/RangeEnd are M parameter expressions. Kind MUST be 0:
            # TOM's ExpressionKind enum defines a single member (M = 0), and
            # all 25 Desktop/Service-authored Expression rows across the
            # corpus -- parameter queries included -- carry Kind=0. This
            # writer used to insert Kind=1, which is exactly the out-of-range
            # enum PFE_TM_ENUM_VALUES_VALIDATION_FAILED complains about; the
            # missing-DataMashup theory the old comment blamed was disproven
            # by 9 corpus files holding Expression rows with no DataMashup.
            # The DataMashup gate itself is kept for now: the no-DataMashup
            # (V3) incremental-refresh flow has not been verified in Desktop.
            has_mashup = os.path.exists(os.path.join(info["work_dir"], "DataMashup"))
            if has_mashup:
                for param_name in ("RangeStart", "RangeEnd"):
                    c.execute("SELECT ID FROM Expression WHERE Name = ?", (param_name,))
                    if not c.fetchone():
                        c.execute("SELECT COALESCE(MAX(ID), 0) + 1 FROM Expression")
                        expr_id = c.fetchone()[0]
                        c.execute(
                            "INSERT INTO Expression (ID, ModelID, Name, Kind, "
                            "Expression, ModifiedTime) "
                            "VALUES (?, 1, ?, 0, ?, datetime('now'))",
                            (expr_id, param_name,
                             '#datetime(2020, 1, 1, 0, 0, 0) meta [IsParameterQuery=true, '
                             'Type="DateTime", IsParameterQueryRequired=true]')
                        )

            # Build polling expression for change detection
            polling_expr = ""
            if detect_changes_column:
                polling_expr = (
                    f"let\n"
                    f"    Source = {table_name},\n"
                    f"    MaxDate = List.Max(Source[{detect_changes_column}])\n"
                    f"in\n"
                    f"    MaxDate"
                )

            # Check for existing policy
            c.execute(
                "SELECT ID FROM RefreshPolicy WHERE TableID = ?",
                (table_id,)
            )
            existing = c.fetchone()

            if existing:
                # Update existing policy
                policy_id = existing[0]
                c.execute(
                    "UPDATE RefreshPolicy SET "
                    "PolicyType=1, RollingWindowGranularity=?, RollingWindowPeriods=?, "
                    "IncrementalGranularity=?, IncrementalPeriods=?, "
                    "IncrementalPeriodsOffset=?, PollingExpression=?, Mode=? "
                    "WHERE ID=?",
                    (_GRAN_MAP[archive_granularity], archive_periods,
                     _GRAN_MAP[refresh_granularity], refresh_periods,
                     -1 if mode == "hybrid" else 0,
                     polling_expr, _MODE_MAP[mode], policy_id)
                )
            else:
                # Create new policy
                c.execute("SELECT COALESCE(MAX(ID), 0) + 1 FROM RefreshPolicy")
                policy_id = c.fetchone()[0]
                c.execute(
                    "INSERT INTO RefreshPolicy (ID, TableID, PolicyType, "
                    "RollingWindowGranularity, RollingWindowPeriods, "
                    "IncrementalGranularity, IncrementalPeriods, "
                    "IncrementalPeriodsOffset, PollingExpression, Mode) "
                    "VALUES (?, ?, 1, ?, ?, ?, ?, ?, ?, ?)",
                    (policy_id, table_id,
                     _GRAN_MAP[archive_granularity], archive_periods,
                     _GRAN_MAP[refresh_granularity], refresh_periods,
                     -1 if mode == "hybrid" else 0,
                     polling_expr, _MODE_MAP[mode])
                )

            # Link table to policy
            c.execute(
                "UPDATE [Table] SET RefreshPolicyID = ? WHERE ID = ?",
                (policy_id, table_id)
            )

            conn.commit()
            policy_info["policy_id"] = policy_id
            policy_info["mode"] = mode

        old_size, new_size = _modify_metadata_only(dm_path, _do_set)
        info["modified"] = True

        detect_msg = f"\n  Change detection: {detect_changes_column}" if detect_changes_column else ""
        return ToolResponse.ok(
            f"Incremental refresh policy set on '{table_name}':\n"
            f"  Archive: {archive_periods} {archive_granularity}(s)\n"
            f"  Refresh: {refresh_periods} {refresh_granularity}(s)\n"
            f"  Mode: {mode}{detect_msg}\n"
            f"  DataModel: {old_size:,} → {new_size:,} bytes\n\n"
            f"The table's M expression must filter on RangeStart/RangeEnd parameters.\n"
            f"Power BI will automatically create date-based partitions on first refresh."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_get_incremental_refresh(alias: str) -> str:
    """Get incremental refresh policies for all tables.

    Args:
        alias: The alias of the open file
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        import tempfile

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        abf = decompress_datamodel(dm_bytes)
        db_bytes = read_metadata_sqlite(abf)

        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(db_bytes)
        tmp.close()

        try:
            conn = sqlite3.connect(tmp.name)
            c = conn.cursor()

            _GRAN_NAMES = {1: "day", 2: "month", 3: "quarter", 4: "year"}
            _MODE_NAMES = {0: "import", 1: "hybrid"}

            c.execute(
                "SELECT rp.ID, t.Name, rp.PolicyType, "
                "rp.RollingWindowGranularity, rp.RollingWindowPeriods, "
                "rp.IncrementalGranularity, rp.IncrementalPeriods, "
                "rp.IncrementalPeriodsOffset, rp.PollingExpression, rp.Mode "
                "FROM RefreshPolicy rp "
                "JOIN [Table] t ON rp.TableID = t.ID "
                "ORDER BY rp.ID"
            )
            policies = c.fetchall()
            conn.close()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        if not policies:
            return ToolResponse.ok("No incremental refresh policies configured.").to_text()

        lines = [f"Incremental refresh policies ({len(policies)}):\n"]
        for p in policies:
            pid, tbl, ptype, rw_gran, rw_periods, inc_gran, inc_periods, offset, polling, pmode = p
            lines.append(f"  Table: {tbl}")
            lines.append(f"    Archive: {rw_periods} {_GRAN_NAMES.get(rw_gran, '?')}(s)")
            lines.append(f"    Refresh: {inc_periods} {_GRAN_NAMES.get(inc_gran, '?')}(s)")
            lines.append(f"    Mode: {_MODE_NAMES.get(pmode, '?')}")
            if polling:
                lines.append("    Change detection: enabled")
            lines.append("")

        return ToolResponse.ok("\n".join(lines)).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


@mcp.tool()
def pbix_export_tmdl(alias: str, output_path: str = "") -> str:
    """Export the data model as TMDL (Tabular Model Definition Language) files.

    TMDL is a human-readable, Git-friendly text format for Power BI models.
    Creates a folder with .tmdl files for tables, relationships, roles, etc.

    Args:
        alias: The alias of the open file
        output_path: Output directory path. Defaults to <pbix_dir>/<alias>_tmdl/
    """
    try:
        info = _ensure_open(alias)
        dm_path = os.path.join(info["work_dir"], "DataModel")
        if not os.path.exists(dm_path):
            return ToolResponse.error("No DataModel found.", DataModelCompressionError.code).to_text()

        # Determine output directory
        if not output_path:
            pbix_dir = os.path.dirname(info.get("original_path", info["work_dir"]))
            output_path = os.path.join(pbix_dir, f"{alias}_tmdl")

        os.makedirs(output_path, exist_ok=True)

        from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
        from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

        with open(dm_path, "rb") as f:
            dm_bytes = f.read()

        abf = decompress_datamodel(dm_bytes)
        db_bytes = read_metadata_sqlite(abf)

        import tempfile
        tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        tmp.write(db_bytes)
        tmp.close()

        try:
            conn = sqlite3.connect(tmp.name)
            stats = _export_tmdl_from_sqlite(conn, output_path)
            conn.close()
        finally:
            # Close the SQLite handle BEFORE unlinking: Windows refuses to
            # delete a file that still has an open handle (WinError 32), which
            # made every calculated-column/table edit fail on the platform
            # nearly all Power BI users are on. POSIX allows it, so CI (ubuntu)
            # never saw this.
            try:
                conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp.name)
            except OSError:
                pass

        summary = (
            f"TMDL exported to: {output_path}\n"
            f"  Tables: {stats['tables']}\n"
            f"  Columns: {stats['columns']}\n"
            f"  Measures: {stats['measures']}\n"
            f"  Relationships: {stats['relationships']}\n"
            f"  Hierarchies: {stats.get('hierarchies', 0)}\n"
            f"  Roles: {stats['roles']}\n"
            f"Files are Git-friendly text — diff, merge, and version control your model."
        )
        return ToolResponse.ok(summary).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


# ---- Section 10c: TMDL Import + PBIP open/save (issue #34) ----


def _apply_tmdl_metadata(conn: sqlite3.Connection, model: dict) -> None:
    """Post-pass after the builder skeleton: write every TMDL property the
    builder cannot express straight into the metadata SQLite — partitions'
    real M/DAX sources, calculated-column expressions, lineage tags,
    format/summarize/sort/category overrides, shared M expressions,
    relationship names, roles, and extended properties (field parameters).
    Property-only updates + rows without storage, so it is safe under
    ``_modify_metadata_only``."""
    c = conn.cursor()
    c.execute(
        'CREATE TABLE IF NOT EXISTS [ExtendedProperty]( [ID] INTEGER, '
        '[ObjectID] INTEGER, [ObjectType] INTEGER, [Name] TEXT, '
        '[Type] INTEGER, [Value] TEXT, [ModifiedTime] INTEGER, '
        'PRIMARY KEY("ID" ASC) )')

    # Metadata object IDs are GLOBAL across tables — allocate from the same
    # DBPROPERTIES MAXID counter every other splice path uses.
    maxid_row = c.execute(
        "SELECT Value FROM DBPROPERTIES WHERE Name = 'MAXID'").fetchone()
    _ids = {"max": int(maxid_row[0]) if maxid_row else 0}

    def _next_id() -> int:
        _ids["max"] += 1
        return _ids["max"]

    def _add_ext_props(obj_type: int, obj_id: int, eps: list) -> None:
        for ep in eps or []:
            c.execute(
                "INSERT INTO [ExtendedProperty] (ID, ObjectID, ObjectType, "
                "Name, Type, Value) VALUES (?, ?, ?, ?, ?, ?)",
                (_next_id(), obj_id, obj_type,
                 ep["name"], ep["type"], ep["value"]))

    # --- Model properties ---
    sets, vals = [], []
    if model.get("culture"):
        sets.append("Culture = ?"); vals.append(model["culture"])
    if model.get("default_powerbi_data_source_version") is not None:
        sets.append("DefaultPowerBIDataSourceVersion = ?")
        vals.append(model["default_powerbi_data_source_version"])
    if model.get("discourage_implicit_measures"):
        sets.append("DiscourageImplicitMeasures = 1")
    if model.get("source_query_culture"):
        sets.append("SourceQueryCulture = ?")
        vals.append(model["source_query_culture"])
    dao = model.get("data_access_options")
    if dao is not None:
        sets.append("DataAccessOptions = ?")
        vals.append(json.dumps({k: True for k, v in dao.items() if v}))
    if sets:
        c.execute(f"UPDATE Model SET {', '.join(sets)} WHERE ID = 1", vals)

    # --- Shared M expressions (parameters) ---
    for e in model.get("expressions") or []:
        c.execute(
            "INSERT INTO Expression (ID, ModelID, Name, Kind, Expression, "
            "LineageTag) VALUES (?, 1, ?, 0, ?, ?)",
            (_next_id(), e["name"], e["expression"],
             e.get("lineage_tag")))

    # --- Per-table objects ---
    for t in model["tables"]:
        trow = c.execute(
            "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
            (t["name"],)).fetchone()
        if not trow:
            continue
        tid = trow[0]
        if t.get("lineage_tag"):
            c.execute("UPDATE [Table] SET LineageTag = ? WHERE ID = ?",
                      (t["lineage_tag"], tid))
        _add_ext_props(3, tid, t.get("extended_properties"))

        col_ids = {name: cid for cid, name in c.execute(
            "SELECT ID, COALESCE(ExplicitName, InferredName) "
            "FROM [Column] WHERE TableID = ?", (tid,))}
        for col in t["columns"]:
            cid = col_ids.get(col["name"])
            if cid is None:
                continue
            csets, cvals = [], []
            if col.get("format_string"):
                csets.append("FormatString = ?")
                cvals.append(col["format_string"])
            if col.get("is_key"):
                csets.append("IsKey = 1")
            if col.get("is_hidden"):
                csets.append("IsHidden = 1")
            if col.get("summarize_by") is not None:
                csets.append("SummarizeBy = ?")
                cvals.append(col["summarize_by"])
            if col.get("data_category"):
                csets.append("DataCategory = ?")
                cvals.append(col["data_category"])
            if col.get("display_folder"):
                csets.append("DisplayFolder = ?")
                cvals.append(col["display_folder"])
            if col.get("lineage_tag"):
                csets.append("LineageTag = ?")
                cvals.append(col["lineage_tag"])
            if col.get("expression") is not None:
                # Calculated column: DAX expression, no source column.
                csets.append("Type = 2")
                csets.append("Expression = ?")
                cvals.append(col["expression"])
                csets.append("SourceColumn = NULL")
            elif col.get("source_column") and col["source_column"] != col["name"]:
                csets.append("SourceColumn = ?")
                cvals.append(col["source_column"])
            if col.get("sort_by_column") and col["sort_by_column"] in col_ids:
                csets.append("SortByColumnID = ?")
                cvals.append(col_ids[col["sort_by_column"]])
            if csets:
                c.execute(
                    f"UPDATE [Column] SET {', '.join(csets)} WHERE ID = ?",
                    cvals + [cid])
            _add_ext_props(4, cid, col.get("extended_properties"))

        for m in t["measures"]:
            mrow = c.execute(
                "SELECT ID FROM Measure WHERE TableID = ? AND Name = ?",
                (tid, m["name"])).fetchone()
            if not mrow:
                continue
            msets, mvals = [], []
            if m.get("is_hidden"):
                msets.append("IsHidden = 1")
            if m.get("display_folder"):
                msets.append("DisplayFolder = ?")
                mvals.append(m["display_folder"])
            if m.get("lineage_tag"):
                msets.append("LineageTag = ?")
                mvals.append(m["lineage_tag"])
            if msets:
                c.execute(
                    f"UPDATE Measure SET {', '.join(msets)} WHERE ID = ?",
                    mvals + [mrow[0]])
            _add_ext_props(8, mrow[0], m.get("extended_properties"))

        # Partitions: replace the builder's enter-data source with the TMDL
        # one (in ID order); extra TMDL partitions become metadata-only rows.
        prows = [r[0] for r in c.execute(
            "SELECT ID FROM [Partition] WHERE TableID = ? ORDER BY ID",
            (tid,))]
        for i, part in enumerate(t["partitions"]):
            ptype = 2 if part["kind"] == "calculated" else 4
            pmode = 1 if part["mode"] == "directquery" else 0
            if i < len(prows):
                c.execute(
                    "UPDATE [Partition] SET Name = ?, QueryDefinition = ?, "
                    "Mode = ?, Type = ? WHERE ID = ?",
                    (part["name"], part["source"], pmode, ptype, prows[i]))
            else:
                c.execute(
                    "INSERT INTO [Partition] (ID, TableID, Name, "
                    "QueryDefinition, State, Type, Mode) "
                    "VALUES (?, ?, ?, ?, 1, ?, ?)",
                    (_next_id(), tid, part["name"],
                     part["source"], ptype, pmode))

        for h in t["hierarchies"]:
            hrow = c.execute(
                "SELECT ID FROM Hierarchy WHERE TableID = ? AND Name = ?",
                (tid, h["name"])).fetchone()
            if not hrow:
                continue
            if h.get("lineage_tag"):
                c.execute("UPDATE Hierarchy SET LineageTag = ? WHERE ID = ?",
                          (h["lineage_tag"], hrow[0]))
            if h.get("display_folder"):
                c.execute("UPDATE Hierarchy SET DisplayFolder = ? WHERE ID = ?",
                          (h["display_folder"], hrow[0]))
            if h.get("is_hidden"):
                c.execute("UPDATE Hierarchy SET IsHidden = 1 WHERE ID = ?",
                          (hrow[0],))
            for lv in h["levels"]:
                if lv.get("lineage_tag"):
                    c.execute(
                        "UPDATE Level SET LineageTag = ? "
                        "WHERE HierarchyID = ? AND Name = ?",
                        (lv["lineage_tag"], hrow[0], lv["name"]))

    # --- Relationship names (builder autogenerates; TMDL keeps originals) ---
    for r in model.get("relationships") or []:
        row = c.execute(
            "SELECT r.ID FROM [Relationship] r "
            "JOIN [Table] ft ON r.FromTableID = ft.ID "
            "JOIN [Column] fc ON r.FromColumnID = fc.ID "
            "JOIN [Table] tt ON r.ToTableID = tt.ID "
            "JOIN [Column] tc ON r.ToColumnID = tc.ID "
            "WHERE ft.Name = ? AND COALESCE(fc.ExplicitName, fc.InferredName) = ? "
            "AND tt.Name = ? AND COALESCE(tc.ExplicitName, tc.InferredName) = ?",
            (r["from_table"], r["from_column"], r["to_table"], r["to_column"]),
        ).fetchone()
        if row and r.get("name"):
            c.execute("UPDATE [Relationship] SET Name = ? WHERE ID = ?",
                      (r["name"], row[0]))

    # --- Roles + table permissions (same shape as the RLS splice path) ---
    for role in model.get("roles") or []:
        role_id = _next_id()
        c.execute(
            "INSERT INTO Role (ID, ModelID, Name, Description) "
            "VALUES (?, 1, ?, ?)",
            (role_id, role["name"], None))
        for tp in role.get("table_permissions") or []:
            trow = c.execute(
                "SELECT ID FROM [Table] WHERE Name = ? AND ModelID = 1",
                (tp["table"],)).fetchone()
            if trow:
                c.execute(
                    "INSERT INTO TablePermission (ID, RoleID, TableID, "
                    "FilterExpression) VALUES (?, ?, ?, ?)",
                    (_next_id(), role_id, trow[0],
                     tp.get("filter_expression") or None))

    if maxid_row:
        c.execute("UPDATE DBPROPERTIES SET Value = ? WHERE Name = 'MAXID'",
                  (str(_ids["max"]),))
    conn.commit()


def _build_pbix_from_tmdl_model(model: dict, output_path: str) -> dict:
    """Build a .pbix at ``output_path`` from a parsed TMDL model dict:
    builder skeleton (tables / columns / measures / relationships /
    hierarchies) + full-fidelity metadata post-pass. Returns stats."""
    import warnings as _warnings

    from pbix_mcp.builder import PBIXBuilder

    b = PBIXBuilder(model.get("name") or "Model")
    stats = {"tables": 0, "columns": 0, "measures": 0,
             "relationships": 0, "roles": len(model.get("roles") or []),
             "hierarchies": 0, "expressions": len(model.get("expressions") or [])}
    for t in model["tables"]:
        cols = [{"name": col["name"], "data_type": col["data_type"]}
                for col in t["columns"]]
        if not cols:
            raise ValueError(f"TMDL table '{t['name']}' declares no columns")
        calc = [col["name"] for col in t["columns"]
                if col.get("expression") is not None]
        b.add_table(t["name"], cols, rows=[], hidden=t.get("is_hidden", False),
                    calc_columns=calc)
        stats["tables"] += 1
        stats["columns"] += len(cols)
        for m in t["measures"]:
            b.add_measure(t["name"], m["name"], m["expression"],
                          format_string=m.get("format_string") or None,
                          data_category=m.get("data_category") or None)
            stats["measures"] += 1
        for h in t["hierarchies"]:
            levels = [{"name": lv["name"], "column": lv["column"]}
                      for lv in h["levels"] if lv.get("column")]
            if levels:
                b.add_user_hierarchy(t["name"], h["name"], levels)
                stats["hierarchies"] += 1
    for r in model.get("relationships") or []:
        b.add_relationship(
            r["from_table"], r["from_column"], r["to_table"], r["to_column"],
            is_active=r.get("is_active", True),
            cross_filter_behavior=r.get("cross_filtering_behavior", 1),
            from_cardinality=r.get("from_cardinality", 2),
            to_cardinality=r.get("to_cardinality", 1),
            auto_orient=False)
        stats["relationships"] += 1

    with _warnings.catch_warnings():
        # Schema-only import: every table is legitimately empty.
        _warnings.simplefilter("ignore", UserWarning)
        b.save(output_path)

    # Full-fidelity post-pass on the saved file's DataModel.
    tmp_dir = tempfile.mkdtemp(prefix="pbix_tmdl_import_")
    try:
        _extract_pbix(output_path, tmp_dir)
        dm_path = os.path.join(tmp_dir, "DataModel")
        _modify_metadata_only(dm_path, lambda conn: _apply_tmdl_metadata(conn, model))
        _repack_pbix(tmp_dir, output_path)
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)
    return stats


@mcp.tool()
def pbix_import_tmdl(tmdl_path: str, output_path: str = "", alias: str = "") -> str:
    """Import a TMDL model (folder or single .tmdl file) into a new PBIX.

    The inverse of pbix_export_tmdl: parses tables, columns (calculated
    included), measures, relationships, hierarchies, partitions (M/DAX
    sources), shared M expressions, roles, extended properties (field
    parameters) and lineage tags, and builds a working schema-only PBIX
    (TMDL carries no row data — refresh in Power BI Desktop repopulates
    from the partitions' M queries). Round-trip: export → import → export
    reproduces the same TMDL files.

    Args:
        tmdl_path: TMDL definition folder (with tables/*.tmdl), a folder
            containing a definition/ subfolder (SemanticModel layout), or a
            single .tmdl document file.
        output_path: Where to write the .pbix. Defaults to
            <tmdl_path>_imported.pbix next to the input.
        alias: When set, the imported file is left open under this alias
            (like pbix_open) for further editing.
    """
    try:
        from pbix_mcp.formats.tmdl_reader import (
            parse_tmdl_folder,
            parse_tmdl_string,
        )

        tmdl_path = os.path.abspath(tmdl_path)
        if not os.path.exists(tmdl_path):
            return ToolResponse.error(
                f"TMDL path not found: {tmdl_path}", "INVALID_INPUT").to_text()

        if os.path.isfile(tmdl_path):
            with open(tmdl_path, "r", encoding="utf-8-sig") as f:
                model = parse_tmdl_string(f.read())
            default_out = os.path.splitext(tmdl_path)[0] + "_imported.pbix"
        else:
            folder = tmdl_path
            # SemanticModel layout: the TMDL lives in definition/
            if not os.path.isdir(os.path.join(folder, "tables")) and \
                    os.path.isdir(os.path.join(folder, "definition")):
                folder = os.path.join(folder, "definition")
            model = parse_tmdl_folder(folder)
            default_out = tmdl_path.rstrip("\\/") + "_imported.pbix"

        output_path = os.path.abspath(output_path or default_out)
        if alias and alias in _open_files:
            raise FileAlreadyOpenError(
                f"Alias '{alias}' is already in use. Close it first or choose "
                f"a different alias.")

        stats = _build_pbix_from_tmdl_model(model, output_path)

        opened_note = ""
        if alias:
            # Reuse the normal open flow so the session behaves identically.
            result = json.loads(pbix_open(output_path, alias))
            if not result.get("success"):
                return ToolResponse.error(
                    f"Imported PBIX written to {output_path} but opening it "
                    f"failed: {result.get('message')}",
                    "IMPORT_OPEN_FAILED").to_text()
            opened_note = f"\nOpened as alias '{alias}'."

        return ToolResponse.ok(
            f"TMDL imported to: {output_path}\n"
            f"  Tables: {stats['tables']} ({stats['columns']} columns)\n"
            f"  Measures: {stats['measures']}\n"
            f"  Relationships: {stats['relationships']}\n"
            f"  Hierarchies: {stats['hierarchies']}\n"
            f"  Shared expressions: {stats['expressions']}\n"
            f"  Roles: {stats['roles']}\n"
            f"Schema-only import: tables are empty until refreshed from "
            f"their partition sources in Power BI Desktop.{opened_note}"
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


def _resolve_pbip_project(path: str) -> tuple[str, str, str, str]:
    """Resolve a PBIP reference (the .pbip file or its folder) to
    (project_root, base_name, report_dir, model_definition_dir)."""
    path = os.path.abspath(path)
    if os.path.isfile(path) and path.lower().endswith(".pbip"):
        root = os.path.dirname(path)
        pbip_file = path
    elif os.path.isdir(path):
        root = path
        candidates = [f for f in os.listdir(path) if f.lower().endswith(".pbip")]
        if not candidates:
            raise InvalidPBIXError(
                f"No .pbip file found in {path} — pass the project folder or "
                f"the .pbip file itself.")
        pbip_file = os.path.join(path, sorted(candidates)[0])
    else:
        raise InvalidPBIXError(f"Not a .pbip file or project folder: {path}")

    with open(pbip_file, "r", encoding="utf-8-sig") as f:
        pbip = json.load(f)
    report_rel = ""
    for art in pbip.get("artifacts") or []:
        if isinstance(art, dict) and "report" in art:
            report_rel = (art["report"] or {}).get("path", "")
            break
    report_dir = os.path.join(root, report_rel) if report_rel else ""
    if not report_dir or not os.path.isdir(report_dir):
        # fall back to the conventional folder name
        hits = [d for d in os.listdir(root) if d.endswith(".Report")]
        if not hits:
            raise InvalidPBIXError(f"No .Report folder found in {root}")
        report_dir = os.path.join(root, sorted(hits)[0])
    base_name = os.path.basename(report_dir)
    if base_name.endswith(".Report"):
        base_name = base_name[: -len(".Report")]

    # definition.pbir -> datasetReference.byPath -> SemanticModel folder
    model_dir = ""
    pbir_path = os.path.join(report_dir, "definition.pbir")
    if os.path.exists(pbir_path):
        with open(pbir_path, "r", encoding="utf-8-sig") as f:
            pbir = json.load(f)
        by_path = ((pbir.get("datasetReference") or {}).get("byPath") or {}).get("path", "")
        if by_path:
            model_dir = os.path.normpath(os.path.join(report_dir, by_path))
    if not model_dir or not os.path.isdir(model_dir):
        hits = [d for d in os.listdir(root) if d.endswith(".SemanticModel")]
        if not hits:
            raise InvalidPBIXError(
                f"No SemanticModel folder found for PBIP project {root}")
        model_dir = os.path.join(root, sorted(hits)[0])

    tmdl_def = os.path.join(model_dir, "definition")
    if not os.path.isdir(tmdl_def):
        if os.path.exists(os.path.join(model_dir, "model.bim")):
            raise UnsupportedFormatError(
                "This PBIP stores its model as TMSL (model.bim), not TMDL — "
                "only TMDL model folders are supported.")
        raise InvalidPBIXError(
            f"No TMDL definition/ folder in {model_dir}")
    return root, base_name, report_dir, tmdl_def


@mcp.tool()
def pbix_open_pbip(path: str, alias: str = "") -> str:
    """Open a PBIP project folder (TMDL model half + report half) as a live
    document — the same session every other tool operates on. pbix_save
    (with no output_path) writes edits back into the project folder;
    pbix_save with a .pbix output_path converts the project to a PBIX.

    Args:
        path: The .pbip file or the project folder containing it.
        alias: Short name for the session (defaults to the project name).
    """
    try:
        from pbix_mcp.formats.tmdl_reader import parse_tmdl_folder

        root, base_name, report_dir, tmdl_def = _resolve_pbip_project(path)
        if not alias:
            alias = base_name
        if alias in _open_files:
            raise FileAlreadyOpenError(
                f"Alias '{alias}' is already in use. Close it first or choose "
                f"a different alias.")

        model = parse_tmdl_folder(tmdl_def)

        # Build the model half into a scratch .pbix and extract it — this
        # becomes the live work_dir every existing tool already understands.
        gen_pbix = os.path.join(
            tempfile.gettempdir(),
            f"pbix_mcp_pbip_{alias}_{os.getpid()}_{uuid.uuid4().hex[:8]}.pbix")
        stats = _build_pbix_from_tmdl_model(model, gen_pbix)

        work_dir = os.path.join(
            tempfile.gettempdir(),
            f"pbix_mcp_{alias}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            f"_{os.getpid()}_{uuid.uuid4().hex[:8]}",
        )
        os.makedirs(work_dir, exist_ok=True)
        _work_dirs.add(work_dir)
        try:
            _extract_pbix(gen_pbix, work_dir)
        except Exception:
            shutil.rmtree(work_dir, ignore_errors=True)
            _work_dirs.discard(work_dir)
            raise

        # Attach the report half. Our own PBIP export writes the full legacy
        # layout as report.json; Desktop-authored projects carry a PBIR
        # definition/ tree, which the PBIR reader already understands.
        report_note = "none"
        report_json = os.path.join(report_dir, "report.json")
        pbir_tree = os.path.join(report_dir, "definition")
        if os.path.exists(report_json):
            with open(report_json, "r", encoding="utf-8-sig") as f:
                layout = json.load(f)
            layout_path = os.path.join(work_dir, "Report", "Layout")
            os.makedirs(os.path.dirname(layout_path), exist_ok=True)
            with open(layout_path, "wb") as f:
                f.write(json.dumps(layout, ensure_ascii=False).encode("utf-16-le"))
            report_note = "classic (report.json)"
        elif os.path.isdir(os.path.join(pbir_tree, "pages")):
            shutil.copytree(pbir_tree,
                            os.path.join(work_dir, "Report", "definition"),
                            dirs_exist_ok=True)
            # PBIR must not coexist with the generated classic Layout.
            gen_layout = os.path.join(work_dir, "Report", "Layout")
            if os.path.exists(gen_layout):
                os.unlink(gen_layout)
            report_note = "PBIR (definition/ tree)"
        static_src = os.path.join(report_dir, "StaticResources")
        if os.path.isdir(static_src):
            shutil.copytree(static_src,
                            os.path.join(work_dir, "Report", "StaticResources"),
                            dirs_exist_ok=True)

        _open_files[alias] = {
            "path": gen_pbix,
            "work_dir": work_dir,
            "is_pbit": False,
            "modified": False,
            "is_directquery": False,
            "pbip_dir": root,
            "pbip_base": base_name,
        }
        return ToolResponse.ok(
            f"Opened PBIP project '{root}' as '{alias}'\n"
            f"  Model: {stats['tables']} tables, {stats['measures']} measures, "
            f"{stats['relationships']} relationships (from TMDL)\n"
            f"  Report: {report_note}\n"
            f"Schema-only model: tables are empty until refreshed from their "
            f"partition sources.\n"
            f"pbix_save writes back into the project folder; "
            f"pbix_save with output_path='x.pbix' converts to PBIX."
        ).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", getattr(e, "code", None)).to_text()


def _save_pbip(info: dict) -> str:
    """Persist a PBIP-opened session back into its project folder: re-export
    the model half as TMDL into <base>.SemanticModel/definition/ and the
    report half as report.json (or patch the PBIR tree in place). Returns a
    human-readable summary."""
    from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
    from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

    root = info["pbip_dir"]
    base = info["pbip_base"]
    work_dir = info["work_dir"]
    report_dir = os.path.join(root, f"{base}.Report")
    model_dir = os.path.join(root, f"{base}.SemanticModel")
    tmdl_def = os.path.join(model_dir, "definition")
    os.makedirs(report_dir, exist_ok=True)

    # --- model half: wipe + re-export TMDL ---
    dm_path = os.path.join(work_dir, "DataModel")
    tmdl_stats = {"tables": 0, "measures": 0}
    if os.path.exists(dm_path):
        if os.path.isdir(tmdl_def):
            shutil.rmtree(tmdl_def)
        os.makedirs(tmdl_def, exist_ok=True)
        with open(dm_path, "rb") as f:
            abf = decompress_datamodel(f.read())
        db_bytes = read_metadata_sqlite(abf)
        fd, tmp_db = tempfile.mkstemp(suffix=".db")
        os.write(fd, db_bytes)
        os.close(fd)
        conn = None
        try:
            conn = sqlite3.connect(tmp_db)
            tmdl_stats = _export_tmdl_from_sqlite(conn, tmdl_def)
        finally:
            try:
                if conn is not None:
                    conn.close()
            except Exception:
                pass
            try:
                os.unlink(tmp_db)
            except OSError:
                pass

    # --- report half ---
    pbir_tree = os.path.join(work_dir, "Report", "definition")
    if os.path.isdir(pbir_tree):
        # PBIR: the work_dir tree IS the edited report — mirror it back.
        dest = os.path.join(report_dir, "definition")
        if os.path.isdir(dest):
            shutil.rmtree(dest)
        shutil.copytree(pbir_tree, dest)
        report_note = "definition/ (PBIR tree)"
    else:
        layout = _get_layout(work_dir)
        if layout is not None:
            with open(os.path.join(report_dir, "report.json"), "w",
                      encoding="utf-8") as f:
                json.dump(layout, f, indent=2, ensure_ascii=False)
            report_note = "report.json"
        else:
            report_note = "none"
    static_src = os.path.join(work_dir, "Report", "StaticResources")
    if os.path.isdir(static_src):
        dest_static = os.path.join(report_dir, "StaticResources")
        if os.path.isdir(dest_static):
            shutil.rmtree(dest_static)
        shutil.copytree(static_src, dest_static)

    # --- project scaffolding (create only when missing) ---
    pbip_file = os.path.join(root, f"{base}.pbip")
    if not os.path.exists(pbip_file):
        with open(pbip_file, "w", encoding="utf-8") as f:
            json.dump({
                "$schema": "https://developer.microsoft.com/json-schemas/fabric/pbip/pbipProperties/1.0.0/schema.json",
                "version": "1.0",
                "artifacts": [{"report": {"path": f"{base}.Report"}}],
                "settings": {"enableAutoRecovery": True},
            }, f, indent=2)
    pbism = os.path.join(model_dir, "definition.pbism")
    if not os.path.exists(pbism):
        with open(pbism, "w", encoding="utf-8") as f:
            json.dump({
                "$schema": "https://developer.microsoft.com/json-schemas/fabric/item/semanticModel/definitionProperties/1.0.0/schema.json",
                "version": "4.1", "settings": {},
            }, f, indent=2)
    pbir_desc = os.path.join(report_dir, "definition.pbir")
    if not os.path.exists(pbir_desc):
        with open(pbir_desc, "w", encoding="utf-8") as f:
            json.dump({
                "$schema": "https://developer.microsoft.com/json-schemas/fabric/item/report/definitionProperties/2.0.0/schema.json",
                "version": "1.0",
                "datasetReference": {"byPath": {"path": f"../{base}.SemanticModel"}},
            }, f, indent=2)

    return (f"Saved PBIP project: {root}\n"
            f"  Model: {tmdl_stats.get('tables', 0)} tables, "
            f"{tmdl_stats.get('measures', 0)} measures -> "
            f"{base}.SemanticModel/definition/\n"
            f"  Report: {report_note} -> {base}.Report/")


def _sanitize_pbir_name(name: str) -> str:
    """Make a name safe for PBIR folder/file naming (word chars or hyphens only)."""
    import re as _re
    sanitized = _re.sub(r'[^\w\-]', '_', name)
    if not sanitized:
        sanitized = "unnamed"
    return sanitized[:50]


def _pbix_config_to_pbir_visual(config: dict, x: float, y: float, w: float, h: float, z: float = 0) -> dict:
    """Convert a PBIX visualContainer config dict to a PBIR visual.json dict."""
    pbir_name = config.get("name", "visual")
    single_visual = config.get("singleVisual", {})
    visual_type = single_visual.get("visualType", "unknown")

    # PBIR visual structure
    visual_obj: dict = {
        "visualType": visual_type,
    }

    # Preserve drillFilterOtherVisuals if present
    if "drillFilterOtherVisuals" in single_visual:
        visual_obj["drillFilterOtherVisuals"] = single_visual["drillFilterOtherVisuals"]

    # Build query structure from prototypeQuery + projections
    proto = single_visual.get("prototypeQuery")
    projections = single_visual.get("projections")
    if proto or projections:
        query: dict = {"queryState": {}}
        if projections:
            # Translate projections to queryState role mappings
            for role_name, role_items in projections.items():
                query["queryState"][role_name] = {"projections": role_items}
        visual_obj["query"] = query
        if proto:
            # Translate an authored prototypeQuery.OrderBy into the PBIR
            # sortDefinition (alias SourceRefs become Entity refs); visuals
            # without one keep the default-sort marker.
            alias2entity = {f.get("Name"): f.get("Entity")
                            for f in (proto.get("From") or [])}

            def _entityize(node: dict) -> None:
                ref = (node.get("Expression") or {}).get("SourceRef") or {}
                if "Source" in ref:
                    node["Expression"]["SourceRef"] = {
                        "Entity": alias2entity.get(ref["Source"], ref["Source"])}

            sort_entries = []
            for ob in (proto.get("OrderBy") or []):
                expr = copy.deepcopy(ob.get("Expression") or {})
                inner = expr.get("Measure") or expr.get("Column")
                if inner is None and "Aggregation" in expr:
                    agg_inner = expr["Aggregation"].get("Expression") or {}
                    inner = agg_inner.get("Column") or agg_inner.get("Measure")
                if inner is None:
                    # Unknown expression shape (e.g. HierarchyLevel) — skip it
                    # rather than leak alias-based SourceRefs into the PBIR.
                    continue
                _entityize(inner)
                sort_entries.append({
                    "field": expr,
                    "direction": ("Ascending" if ob.get("Direction") == 1
                                  else "Descending"),
                })
            if sort_entries:
                visual_obj["query"]["sortDefinition"] = {
                    "sort": sort_entries, "isDefaultSort": False}
            else:
                visual_obj["query"]["sortDefinition"] = {"sort": [], "isDefaultSort": True}
            # Add prototypeQuery as dataViewMappings source
            visual_obj["query"]["queryRef"] = proto
    else:
        visual_obj["query"] = {"queryState": {}}

    # Copy data formatting (objects)
    if "objects" in single_visual:
        visual_obj["objects"] = single_visual["objects"]

    # vcObjects -> visualContainerObjects (title, background, border, ...).
    # It belongs INSIDE `visual`: visualContainer sets additionalProperties
    # false and does not permit it at the top level, and all 70 visuals in the
    # service-authored corpus put it under `visual`. Emitting it at the top
    # level produced PBIP that Power BI rejects.
    if "vcObjects" in single_visual:
        visual_obj["visualContainerObjects"] = single_visual["vcObjects"]

    result: dict = {
        "$schema": "https://developer.microsoft.com/json-schemas/fabric/item/report/definition/visualContainer/1.0.0/schema.json",
        "name": pbir_name,
        "position": {"x": x, "y": y, "z": z, "width": w, "height": h},
        "visual": visual_obj,
    }
    return result


@mcp.tool()
def pbix_export_pbip(alias: str, output_dir: str = "") -> str:
    """Convert a PBIX to PBIP (Power BI Project) folder structure.

    Creates a PBIP project with:
      - {name}.pbip            (root pointer JSON)
      - {name}.Report/         (report layout + static resources)
      - {name}.SemanticModel/  (semantic model as TMDL)
      - .gitignore             (standard PBIP ignores)

    PBIP is Microsoft's folder-based format for Git version control and CI/CD.

    Args:
        alias: The alias of the open file
        output_dir: Target directory. Defaults to <pbix_dir>/<name>_pbip/
    """
    try:
        info = _ensure_open(alias)
        work_dir = info["work_dir"]
        pbix_path = info["path"]

        base_name = os.path.splitext(os.path.basename(pbix_path))[0]
        # Strip spaces and special chars for safer folder names
        safe_base = "".join(c if c.isalnum() or c in "-_" else "_" for c in base_name)

        if not output_dir:
            pbix_dir = os.path.dirname(pbix_path)
            output_dir = os.path.join(pbix_dir, f"{safe_base}_pbip")

        # Clean or create output directory
        if os.path.exists(output_dir):
            import shutil
            shutil.rmtree(output_dir)
        os.makedirs(output_dir, exist_ok=True)

        report_dir = os.path.join(output_dir, f"{safe_base}.Report")
        model_dir = os.path.join(output_dir, f"{safe_base}.SemanticModel")
        os.makedirs(report_dir, exist_ok=True)
        os.makedirs(model_dir, exist_ok=True)

        # --- 1. Export TMDL to SemanticModel/definition/ ---
        dm_path = os.path.join(work_dir, "DataModel")
        tmdl_stats = {"tables": 0, "columns": 0, "measures": 0, "relationships": 0, "roles": 0}
        if os.path.exists(dm_path):
            from pbix_mcp.formats.abf_rebuild import read_metadata_sqlite
            from pbix_mcp.formats.datamodel_roundtrip import decompress_datamodel

            with open(dm_path, "rb") as f:
                dm_bytes = f.read()
            abf = decompress_datamodel(dm_bytes)
            db_bytes = read_metadata_sqlite(abf)

            tmdl_def_dir = os.path.join(model_dir, "definition")
            os.makedirs(tmdl_def_dir, exist_ok=True)

            fd, tmp_db = tempfile.mkstemp(suffix=".db")
            os.write(fd, db_bytes)
            os.close(fd)
            try:
                conn = sqlite3.connect(tmp_db)
                tmdl_stats = _export_tmdl_from_sqlite(conn, tmdl_def_dir)
                conn.close()
            finally:
                try:
                    os.unlink(tmp_db)
                except OSError:
                    pass

        # --- 2. Create definition.pbism (semantic model descriptor) ---
        pbism_content = {
            "$schema": "https://developer.microsoft.com/json-schemas/fabric/item/semanticModel/definitionProperties/1.0.0/schema.json",
            "version": "4.1",
            "settings": {}
        }
        with open(os.path.join(model_dir, "definition.pbism"), "w", encoding="utf-8") as f:
            json.dump(pbism_content, f, indent=2)

        # --- 3. Copy StaticResources to Report folder ---
        pbix_static = os.path.join(work_dir, "Report", "StaticResources")
        if os.path.isdir(pbix_static):
            import shutil
            dest_static = os.path.join(report_dir, "StaticResources")
            shutil.copytree(pbix_static, dest_static)

        # --- 4. Create definition.pbir (report descriptor with byPath ref to model) ---
        pbir_content = {
            "$schema": "https://developer.microsoft.com/json-schemas/fabric/item/report/definitionProperties/2.0.0/schema.json",
            "version": "1.0",
            "datasetReference": {
                "byPath": {
                    "path": f"../{safe_base}.SemanticModel"
                }
            }
        }
        with open(os.path.join(report_dir, "definition.pbir"), "w", encoding="utf-8") as f:
            json.dump(pbir_content, f, indent=2)

        # --- 5. Report layout (legacy format) ---
        # Use the original PBIX Layout JSON directly as report.json.
        # PBIR decomposed format (version 4.0) has rendering bugs in PBI Desktop,
        # so we use legacy format (version 1.0) with the full Layout JSON.
        layout = _get_layout(work_dir)
        if not layout:
            return ToolResponse.error("No layout found in PBIX", "LAYOUT_MISSING").to_text()

        sections = layout.get("sections", [])
        with open(os.path.join(report_dir, "report.json"), "w", encoding="utf-8") as f:
            json.dump(layout, f, indent=2, ensure_ascii=False)

        # --- 7. Root .pbip file ---
        pbip_content = {
            "$schema": "https://developer.microsoft.com/json-schemas/fabric/pbip/pbipProperties/1.0.0/schema.json",
            "version": "1.0",
            "artifacts": [
                {"report": {"path": f"{safe_base}.Report"}}
            ],
            "settings": {"enableAutoRecovery": True}
        }
        with open(os.path.join(output_dir, f"{safe_base}.pbip"), "w", encoding="utf-8") as f:
            json.dump(pbip_content, f, indent=2)

        # --- 8. .gitignore ---
        with open(os.path.join(output_dir, ".gitignore"), "w", encoding="utf-8") as f:
            f.write("**/.pbi/localSettings.json\n")
            f.write("**/.pbi/cache.abf\n")

        # Count output
        total_pages = len(sections)
        total_visuals = sum(len(s.get("visualContainers", [])) for s in sections)

        summary = (
            f"PBIP project exported to: {output_dir}\n\n"
            f"  {safe_base}.pbip                     (root)\n"
            f"  {safe_base}.Report/\n"
            f"    definition.pbir                    (report descriptor)\n"
            f"    report.json                        ({total_pages} pages, {total_visuals} visuals)\n"
            f"  {safe_base}.SemanticModel/\n"
            f"    definition.pbism                   (model descriptor)\n"
            f"    definition/                        (TMDL)\n"
            f"      {tmdl_stats['tables']} tables, {tmdl_stats['columns']} columns, "
            f"{tmdl_stats['measures']} measures, {tmdl_stats['relationships']} relationships, "
            f"{tmdl_stats['roles']} roles\n"
        )
        return ToolResponse.ok(summary).to_text()
    except PBIXMCPError as e:
        return ToolResponse.error(e.message, e.code).to_text()
    except Exception as e:
        return ToolResponse.error(f"{str(e)}\n{traceback.format_exc()}", "PBIP_EXPORT_ERROR").to_text()


# ---- Section 11: MCP main ----

if __name__ == "__main__":
    mcp.run()
